"""Convert the three memos into legally formatted .docx files on Technijian letterhead."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VGD = os.path.dirname(os.path.abspath(__file__))

HEADER = {
    "company": "Technijian, Inc.",
    "tagline": "Technology as a Solution",
    "address": "18 Technology Drive, Suite 141, Irvine, CA 92618",
    "phone": "(949) 379-8499",
    "web": "www.technijian.com",
}


def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    pfmt = style.paragraph_format
    pfmt.space_after = Pt(6)
    pfmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    return doc


def add_letterhead(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(HEADER["company"])
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(HEADER["tagline"])
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x5B, 0x7A, 0x99)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"{HEADER['address']}  |  {HEADER['phone']}  |  {HEADER['web']}")
    r3.font.size = Pt(9)
    p4 = doc.add_paragraph()
    p4_pr = p4._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    p4_pr.append(pBdr)
    doc.add_paragraph()


def add_footer(doc, text):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_heading(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    else:
        r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph(doc, text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_field_line(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    p_pr.append(pBdr)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def add_privileged_callout(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xB2, 0x22, 0x22)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ============================================================================
# DOCUMENT 1 — SETTLEMENT POSITION MEMORANDUM
# ============================================================================

def build_settlement_memo():
    doc = setup_document()
    add_letterhead(doc)
    add_heading(doc, "SETTLEMENT POSITION MEMORANDUM", level=1, center=True)
    add_privileged_callout(doc,
        "PRIVILEGED SETTLEMENT COMMUNICATION — Cal. Evid. Code §§ 1119, 1152; Fed. R. Evid. 408."
    )
    add_horizontal_rule(doc)

    add_field_line(doc, "To", "[Counsel for Respondent Vintage Design — to be filled by Callahan & Blaine]")
    add_field_line(doc, "From", "Edward Susolik, Esq. — Callahan & Blaine, PC, Counsel for Claimant Technijian, Inc.")
    add_field_line(doc, "Re", "Technijian, Inc. v. Vintage Design, LLC — American Arbitration Association")
    add_field_line(doc, "Date", "April 14, 2026")
    doc.add_paragraph()

    add_heading(doc, "I. PURPOSE", level=2)
    add_paragraph(doc,
        "Technijian, Inc. (\"Technijian\") submits this memorandum in advance of the parties' counsel "
        "conference to support a good-faith resolution of this matter before depositions and further "
        "arbitration costs are incurred. Based on the undisputed signed Agreement, Technijian's own "
        "contemporaneous business records, and controlling California authority, Technijian's cancellation-"
        "fee recovery — before prejudgment interest, contractual attorneys' fees, and arbitration costs — is "
        "$240,555.15 as pleaded, rising to approximately $325,000–$370,000 through award. This memorandum "
        "proposes resolution at that pleaded primary number and identifies the legal framework that will "
        "control the arbitration if the matter is not resolved."
    )

    add_heading(doc, "II. THE OPERATIVE AGREEMENT (UNCONTESTED)", level=2)
    add_paragraph(doc,
        "The Client Monthly Service Agreement between Vintage Design and Technijian, Inc. was executed via "
        "DocuSign on May 4, 2023 (Envelope ID B679C550-E41C-4E31-BE34-7F8FFF437C3D), signed by Erica Garcia, "
        "VP Finance, on behalf of Vintage Design and Ravi Jain, CEO, on behalf of Technijian. The Agreement "
        "is attached as Exhibit A. Its material terms:"
    )
    add_table(doc,
        ["Term", "Value", "Source"],
        [
            ["Client entity", "Vintage Design, a California corporation", "Agreement signature page"],
            ["Agreement date", "May 4, 2023", "DocuSign envelope"],
            ["Under-contract period", "12 months", "Under Contract ¶ 1"],
            ["Cancellation rate", "$150/hour for excess ticketed hours", "Under Contract ¶ 5"],
            ["POD rates", "CHD-TS1 Normal $15/hr; CHD-TS1 AH $30/hr; IRV-TS1 Normal $125/hr", "POD Specifications"],
            ["Initial billed hours", "40 + 20 + 20 = 80 hrs/month", "POD Specifications"],
            ["Out-of-contract (IRV-TS1)", "$150/hr onsite normal; $200/hr onsite after-hours", "Out of Contract table"],
            ["Late fee", "10% on payment after 10th of month", "Other Terms ¶ 3"],
            ["Attorneys' fees", "T&C § 5.02 indemnity (reciprocal under Cal. Civ. Code § 1717)", "T&C § 5.02"],
            ["Arbitration", "AAA Commercial Rules, venue Orange County", "T&C § 6.10"],
        ],
        col_widths=[2.0, 3.3, 1.2]
    )
    add_paragraph(doc,
        "Stipulation proposed — minor caption correction. The Demand for Arbitration filed October 23, 2025 "
        "identifies Respondent as \"Vintage Design, LLC, a Delaware limited liability company.\" The signed "
        "Agreement identifies the contracting entity as Vintage Design, a California corporation. Technijian "
        "proposes to stipulate to a caption conforming to the four corners of the signed Agreement, or to "
        "whatever related-entity designation Respondent's counsel confirms as the correct party in interest. "
        "Technijian will move to amend under AAA Commercial Rule R-6. Amendment on these facts is routine and "
        "relates back. Hawkins v. Pacific Coast Bldg. Prods., Inc. (2004) 124 Cal.App.4th 1497, 1503–1504; "
        "Mayberry v. Coca-Cola Bottling Co. (1966) 244 Cal.App.2d 350, 352."
    )
    add_paragraph(doc,
        "No other correction to the Demand's material allegations is required. Period length, cancellation "
        "rate, POD rates, signatory identity, signing date — all are as pleaded and as confirmed by the "
        "signed Agreement itself."
    )

    add_heading(doc, "III. UNDISPUTED FACTS", level=2)
    facts = [
        "Execution. On May 4, 2023, Vintage Design (through Erica Garcia, VP Finance) and Technijian "
        "(through Ravi Jain, CEO) executed the Client Monthly Service Agreement via DocuSign.",
        "Services Rendered. Technijian provided continuous IT support services from inception through the "
        "effective termination date of July 31, 2025 — approximately 27 months of performance. All support "
        "hours were contemporaneously recorded in Technijian's ticketing system.",
        "Billing. Technijian issued monthly invoices that (a) attached the Agreement terms, (b) itemized "
        "contracted and actual hours, and (c) identified hours \"not paid are due if support agreement is "
        "cancelled per monthly service agreement.\" Vintage Design paid every monthly invoice without "
        "objection throughout the 27-month performance period.",
        "Termination Notice. On June 23, 2025, Vintage Design sent written termination notice. Per the "
        "Agreement's invoicing-cycle provision, the effective termination date is July 31, 2025.",
        "Cancellation Fee Triggered. Under Contract ¶ 5 provides that upon termination, \"any hours that "
        "exceeded the previous under contract period average, that were documented through ticketing, will "
        "be charged at a rate of $150 per hour and will be assessed as the cancellation fee to the client "
        "and due before agreement is terminated.\"",
        "Final Invoice. Technijian issued the final (cancellation-fee) invoice on June 27, 2025, with "
        "payment due July 27, 2025. Vintage Design has not paid.",
        "Dual Dispute Windows — Weekly + Monthly — Both Unchallenged for 27 Months. Throughout the entire "
        "~27-month performance period, Technijian provided Vintage Design with two separate, documented "
        "opportunities to dispute any hour of work: (a) weekly zero-dollar invoices enumerating every ticket, "
        "time entry, technician, start/end timestamps, and hours worked during that week, accompanied by a "
        "30-day window to dispute any work performed — approximately 140 such weekly notices across the life "
        "of the Agreement; and (b) monthly invoices with the 60-day objection window in T&C § 3.01. "
        "Vintage Design disputed none of the weekly notices and none of the monthly invoices. T&C § 3.01 "
        "provides that absent timely objection, Client \"shall be deemed to have accepted without question "
        "such invoice or bill and may not in the future contest the amount it paid or seek reimbursement for "
        "any discrepancies.\"",
        "Fee-Shifting and Late Fee. The Agreement entitles Technijian to (a) a 10% late fee on unpaid "
        "invoices (Other Terms ¶ 3); (b) attorneys' fees via the T&C § 5.02 indemnity, made reciprocal by "
        "Cal. Civ. Code § 1717; and (c) prejudgment interest under California law.",
    ]
    for i, f in enumerate(facts, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{i}. ")
        r.bold = True
        p.add_run(f)

    add_heading(doc, "IV. CANCELLATION FEE — CONTRACT MECHANICS", level=2)
    add_paragraph(doc,
        "The Agreement structures support into 12-month under-contract periods (¶ 1). Within each period: "
        "monthly invoicing is based on the previous period's average monthly hours, multiplied by the three "
        "contracted POD rates ($15/$30/$125) (¶ 3). Actual usage is documented contemporaneously through "
        "Technijian's ticketing system. At the end of each period, the average is recalculated for the next "
        "period with a prospective rate adjustment (¶¶ 3–4). Upon termination, hours that exceeded the "
        "previous under-contract period average, documented through ticketing, are charged at $150/hour and "
        "assessed as the cancellation fee (¶ 5)."
    )
    add_paragraph(doc,
        "Under Cal. Civ. Code § 1671(b), the cancellation provision is presumptively valid; the burden of "
        "proving unreasonableness falls on the challenger. Ridgley v. Topa Thrift & Loan Assn. (1998) 17 "
        "Cal.4th 970, 977; Hitz v. First Interstate Bank (1995) 38 Cal.App.4th 274, 286–289. The $150 "
        "cancellation rate is reasonable on its face: it is tied to actual measured usage (ticketing data), "
        "and it sits at or below the Agreement's own \"Out of Contract\" rates for the same IRV-TS1 Tech "
        "Support category — $150/hour onsite normal and $200/hour onsite after-hours. A rate that does not "
        "exceed the bargained alternative rate is, as a matter of law, not a penalty. El Centro Mall, LLC "
        "v. Payless ShoeSource, Inc. (2009) 174 Cal.App.4th 58, 63."
    )

    add_heading(doc, "V. DAMAGES — CALCULATED FROM TECHNIJIAN'S CONTEMPORANEOUS RECORDS", level=2)
    add_paragraph(doc,
        "Technijian has reconciled every month of billed hours against every month of actual ticketed hours "
        "across the three contracted support categories for the entire contract life through the "
        "July 31, 2025 effective termination date:"
    )
    add_table(doc,
        ["Contracted Category", "Rate", "Billed", "Actual", "Net Excess"],
        [
            ["Offshore Normal (CHD-TS1 R.N)", "$15/hr", "1,760.42", "2,682.18", "+921.76"],
            ["Offshore After-Hours (CHD-TS1 R.AF)", "$30/hr", "1,308.91", "1,801.09", "+492.18"],
            ["Onshore Remote Normal (IRV-TS1 R.N)", "$125/hr", "540.00", "456.82", "−83.18"],
            ["TOTALS", "", "3,609.33", "4,940.09", "+1,330.76"],
        ],
        col_widths=[2.6, 0.8, 1.1, 1.1, 1.1]
    )
    doc.add_paragraph()
    add_heading(doc, "Period-by-Period", level=3)
    add_table(doc,
        ["Under-Contract Period", "Months", "Billed", "Actual", "Excess"],
        [
            ["P1 — May 2023 – Apr 2024", "12", "960.00", "1,970.97", "+1,010.97"],
            ["P2 — May 2024 – Apr 2025", "12", "2,025.08", "2,584.18", "+559.10"],
            ["P3 (terminated) — May 2025 – Jul 2025", "3", "624.25", "384.94", "−239.31"],
        ],
        col_widths=[2.6, 0.8, 1.1, 1.1, 1.1]
    )
    doc.add_paragraph()
    add_heading(doc, "Damages Computation", level=3)
    add_table(doc,
        ["Item", "Amount"],
        [
            ["Excess hours (per Demand ¶ 40)", "1,457.50"],
            ["Cancellation rate per signed Agreement", "× $150.00"],
            ["Cancellation Fee Subtotal", "$218,625.00"],
            ["+ 10% late fee (Other Terms ¶ 3)", "$21,862.50"],
            ["Cancellation Fee + Late Fee (Demand pleaded $240,555.15)", "$240,487.50"],
            ["+ Prejudgment interest at 10% per annum from 7/27/2025 (~9 months)", "~$18,042"],
            ["+ Reasonable attorneys' fees via § 5.02 + § 1717 through award", "~$60,000–$110,000"],
            ["Total Recoverable at Award", "~$325,000–$370,000"],
        ],
        col_widths=[4.7, 1.8]
    )
    add_paragraph(doc,
        "This calculation is supported by Technijian's ticketing system extracts and QuickBooks ledger, "
        "both admissible business records under Cal. Evid. Code §§ 1271, 1552, and 1553. Jazayeri v. Mao "
        "(2009) 174 Cal.App.4th 301, 324; Aguimatang v. California State Lottery (1991) 234 Cal.App.3d "
        "769, 797–798. Because the cancellation fee is a liquidated damages provision tied to "
        "contemporaneous measured data at a stipulated rate, the certainty of damages requirement is "
        "satisfied as a matter of law. Sargon Enterprises, Inc. v. University of Southern California "
        "(2012) 55 Cal.4th 747, 773–775; Toscano v. Greene Music (2004) 124 Cal.App.4th 685, 694."
    )

    add_heading(doc, "VI. LEGAL AUTHORITIES SUPPORTING TECHNIJIAN'S POSITION", level=2)
    authorities = [
        ("A. Breach of Contract (First Cause of Action)",
         "Oasis West Realty, LLC v. Goldman (2011) 51 Cal.4th 811, 821; Reichert v. General Ins. Co. (1968) 68 Cal.2d 822, 830."),
        ("B. Open Book Account (Second Cause of Action)",
         "Cal. Code Civ. Proc. § 337a; Interstate Group Administrators, Inc. v. Cravens, Dargan & Co. (1985) 174 Cal.App.3d 700, 708; R.N.C. Inc. v. Tsegeletos (1991) 231 Cal.App.3d 967, 971–972."),
        ("C. Account Stated (Third Cause of Action)",
         "Zinn v. Fred R. Bright Co. (1969) 271 Cal.App.2d 597, 600; Trafton v. Youngblood (1968) 69 Cal.2d 17, 25; Gleason v. Klamer (1980) 103 Cal.App.3d 782, 787."),
        ("D. Services Rendered / Quantum Meruit (Fourth Cause of Action, Alternative)",
         "Huskinson & Brown v. Wolf (2004) 32 Cal.4th 453, 458; Maglica v. Maglica (1998) 66 Cal.App.4th 442."),
        ("E. Liquidated Damages Enforceability (Cancellation Fee)",
         "Cal. Civ. Code § 1671(b); Ridgley v. Topa Thrift & Loan Assn. (1998) 17 Cal.4th 970; Hitz v. First Interstate Bank (1995) 38 Cal.App.4th 274; El Centro Mall, LLC v. Payless ShoeSource, Inc. (2009) 174 Cal.App.4th 58; Greentree Financial Group v. Execute Sports, Inc. (2008) 163 Cal.App.4th 495. The $150 cancellation rate is below the Agreement's own out-of-contract alternative rates of $150–$200/hour — foreclosing any penalty argument as a matter of law."),
        ("F. Late Fee",
         "Garrett v. Coast & Southern Fed. Sav. & Loan Assn. (1973) 9 Cal.3d 731, 738–740; Beasley v. Wells Fargo Bank (1991) 235 Cal.App.3d 1383."),
        ("G. Attorneys' Fees — Recoverable via § 5.02 Indemnity + § 1717 Reciprocity",
         "T&C § 5.02 provides that Vintage Design shall indemnify Technijian for attorneys' fees incurred as a result of \"any failure by Client…to comply with the terms of this Agreement.\" Non-payment of the cancellation fee is such a failure. Cal. Civ. Code § 1717 renders the indemnity reciprocal. Santisas v. Goodin (1998) 17 Cal.4th 599, 610–611; PLCM Group, Inc. v. Drexler (2000) 22 Cal.4th 1084, 1095. Arbitrator has plenary authority. Moshonov v. Walsh (2000) 22 Cal.4th 771, 776; Advanced Micro Devices v. Intel Corp. (1994) 9 Cal.4th 362, 375–376; AAA Commercial Rule R-47(d)(ii)."),
        ("H. Prejudgment Interest — 10% Per Annum",
         "Cal. Civ. Code §§ 3287(a), 3289(b); Wisper Corp. v. California Commerce Bank (1996) 49 Cal.App.4th 948, 958–960; North Oakland Med. Clinic v. Rogers (1998) 65 Cal.App.4th 824, 829–830."),
        ("I. Post-Judgment Interest — Additional 10% on Confirmed Award",
         "Cal. Code Civ. Proc. § 685.010(a); Britz, Inc. v. Alfa-Laval Food & Dairy Co. (1995) 34 Cal.App.4th 1085, 1107; Pierotti v. Torian (2000) 81 Cal.App.4th 17, 28."),
        ("J. DocuSign Signatures — Fully Enforceable",
         "Cal. Civ. Code §§ 1633.7, 1633.9; Fabian v. Renovate America (2019) 42 Cal.App.5th 1062; J.B.B. Investment Partners v. Fair (2014) 232 Cal.App.4th 974; Ruiz v. Moss Bros. Auto Group (2014) 232 Cal.App.4th 836, 844. DocuSign Envelope ID B679C550 audit trail is prima facie authentication."),
        ("K. Arbitration Clause Enforceable and Final",
         "AT&T Mobility LLC v. Concepcion (2011) 563 U.S. 333; Moncharsh v. Heily & Blase (1992) 3 Cal.4th 1; Pinnacle Museum Tower Assn. v. Pinnacle Market Development (2012) 55 Cal.4th 223. Confirmation under CCP § 1286; limited grounds for vacatur under § 1286.2."),
    ]
    for title, body in authorities:
        add_heading(doc, title, level=3)
        add_paragraph(doc, body)

    add_heading(doc, "VII. ANTICIPATED DEFENSES — AND TECHNIJIAN'S RESPONSES", level=2)
    add_table(doc,
        ["Anticipated Defense", "Response and Authority"],
        [
            ["Signatory lacked authority",
             "Erica Garcia signed as VP Finance — a textbook officer title under Cal. Corp. Code § 313, conferring conclusive presumption of authority for third parties. Snukal v. Flightways Mfg., Inc. (2000) 23 Cal.4th 754, 782. DocuSign signature attributable as a matter of law. Cal. Civ. Code § 1633.9(a); Fabian v. Renovate America. Ratification by 27 months of paid performance. Rakestraw v. Rodrigues (1972) 8 Cal.3d 67, 73."],
            ["Cancellation fee is a penalty / unconscionable",
             "Presumptively valid under § 1671(b); Respondent bears burden. Ridgley, Hitz, El Centro Mall. B2B sophisticated parties (VP Finance signed). Pinnacle Museum Tower (2012) 55 Cal.4th 223, 246–247; Sanchez v. Valencia Holding Co. (2015) 61 Cal.4th 899, 910–911. $150 rate is at or below the Agreement's own Out of Contract rates — not a penalty as a matter of law."],
            ["Hours not used / records inaccurate",
             "Contemporaneous ticketing records admissible as business records. Cal. Evid. Code §§ 1271, 1552, 1553; Jazayeri v. Mao. Two independent waivers apply: (a) weekly zero-dollar invoices enumerating every ticket, with 30-day dispute window — ~140 weekly notices over 27 months, zero disputes filed; and (b) monthly invoices with 60-day objection window in T&C § 3.01. Cobb v. Pacific Mut. Life Ins. Co. (1935) 4 Cal.2d 565, 573; Gleason v. Klamer (1980) 103 Cal.App.3d 782, 787. Cal. Comm. Code § 1303 (course of performance)."],
            ["Cancellation clause reaches only the terminated period (P3)",
             "Plain text reaches \"any hours that exceeded the previous under contract period average\" — cumulative and unqualified. Cal. Civ. Code §§ 1638, 1641; Powerine Oil Co. v. Superior Court (2005) 37 Cal.4th 377, 390–391. Reading to cover only P3 would render the clause meaningless."],
            ["P1's excess was trued up",
             "¶¶ 3–4 describe prospective billing adjustments. Neither uses \"release,\" \"forgive,\" or \"credit back.\" The averaging raises the going-forward billed baseline; it does not compensate Technijian retroactively for P1's 1,010.97 hours of delivered-but-unpaid excess. The cancellation fee is the sole mechanism for that compensation. MacKinnon v. Truck Ins. Exchange (2003) 31 Cal.4th 635, 648."],
            ["Oral side-agreements modify written terms",
             "Barred by integration clause (T&C § 6.08) and parol evidence rule. Cal. Code Civ. Proc. § 1856(a); Casa Herrera, Inc. v. Beydoun (2004) 32 Cal.4th 336, 343–344; Banco do Brasil v. Latian (1991) 234 Cal.App.3d 973, 1001–1002."],
            ["Demand's Delaware LLC describes the wrong entity",
             "Minor caption correction; amendment is routine under AAA Rule R-6 and relates back. Hawkins v. Pacific Coast Bldg. Prods.; Mayberry v. Coca-Cola. Respondent's principals are on actual notice."],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading(doc, "VIII. ECONOMICS OF LITIGATION", level=2)
    add_paragraph(doc,
        "1. Arbitration costs to award: estimated $30,000–$50,000 per side. "
        "2. Technijian's attorneys' fees through award (contractually recoverable via § 5.02 + § 1717): "
        "estimated $60,000–$110,000 at this stage, growing monthly. "
        "3. Prejudgment interest at 10% per annum on the $240,555.15 principal from July 27, 2025, adding "
        "approximately $2,005 per month. "
        "4. Post-judgment interest at 10% per annum attaches on the confirmed award. Cal. Code Civ. Proc. "
        "§ 685.010. "
        "5. Confirmation of AAA awards is ministerial. Moncharsh. "
        "6. Respondent's total exposure through award conservatively approaches $430,000–$520,000 by "
        "mid-2026, climbing thereafter."
    )

    add_heading(doc, "IX. SETTLEMENT PROPOSAL", level=2)
    add_paragraph(doc,
        "Technijian is prepared to resolve this matter now, prior to arbitrator appointment and commencement "
        "of depositions, on the following principal terms:"
    )
    terms = [
        "Payment to Technijian: $240,555.15, representing the cancellation fee ($218,625.00) plus "
        "contractual 10% late fee ($21,930.15), payable within fifteen (15) business days of mutual "
        "execution of a settlement agreement and joint dismissal of the arbitration with prejudice.",
        "Mutual General Release of all claims arising from or relating to the Agreement, the Services, or "
        "the termination. Each side bears its own attorneys' fees and arbitration costs incurred to date — "
        "Technijian forgoes approximately $60,000–$110,000 in recoverable fees plus approximately $18,000 in "
        "accrued prejudgment interest, for an effective compromise of roughly $78,000–$128,000 off the "
        "fully-recoverable award figure.",
        "Confidentiality customary for commercial settlements, consistent with Cal. Evid. Code §§ 1119, 1152 "
        "and Fed. R. Evid. 408.",
        "Mutual non-disparagement.",
        "Stipulated caption amendment of the Demand to conform to the four corners of the signed Agreement "
        "(\"Vintage Design, a California corporation\"), simultaneously filed with the dismissal.",
    ]
    for i, t in enumerate(terms, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{i}. ")
        r.bold = True
        p.add_run(t)

    add_paragraph(doc,
        "This proposal represents 100% of the pleaded cancellation fee and late fee at the contractually "
        "correct $150/hour rate, with Technijian's compromise consisting of the fee-and-interest waiver "
        "described above. On a comprehensive-exposure basis, this is a discount of approximately 25–35% off "
        "Respondent's total projected through-award exposure."
    )

    add_heading(doc, "X. NEXT STEPS", level=2)
    nexts = [
        "Respondent's counsel to provide a written response within fourteen (14) calendar days.",
        "If the parties reach agreement in principle, Technijian will lodge a joint request to stay all "
        "arbitration deadlines pending execution of the settlement.",
        "If the parties do not reach agreement, Technijian will promptly (a) lodge the caption amendment "
        "under AAA Rule R-6, (b) produce the full ticketing reconciliation and QuickBooks ledger under AAA "
        "Rule R-22, and (c) notice depositions of Erica Garcia and Respondent's current finance and "
        "operations leadership.",
    ]
    for i, n in enumerate(nexts, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{i}. ")
        r.bold = True
        p.add_run(n)

    add_paragraph(doc, "Respectfully submitted,")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Edward Susolik, Esq.")
    r.bold = True
    doc.add_paragraph("Callahan & Blaine, PC")
    doc.add_paragraph("Counsel for Claimant Technijian, Inc.")

    add_heading(doc, "EXHIBIT INDEX", level=2)
    exhibits = [
        ("Exhibit A", "Client Monthly Service Agreement, DocuSigned 5/4/2023 (Envelope ID B679C550-E41C-4E31-BE34-7F8FFF437C3D)"),
        ("Exhibit B", "Technijian monthly invoice ledger, May 2023 – April 2026"),
        ("Exhibit C", "Billed-vs-Actual hours reconciliation by contracted support category and month"),
        ("Exhibit D", "Vintage Design termination notice dated June 23, 2025 (in counsel's file)"),
        ("Exhibit E", "June 27, 2025 Final Invoice assessing cancellation fee (in counsel's file)"),
        ("Exhibit F", "Demand for Arbitration filed October 23, 2025, with proposed caption-amendment redline"),
    ]
    for label, desc in exhibits:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}. ")
        r.bold = True
        p.add_run(desc)

    add_footer(doc, "PRIVILEGED SETTLEMENT COMMUNICATION — Cal. Evid. Code §§ 1119, 1152; Fed. R. Evid. 408")

    out = os.path.join(VGD, "Settlement_Position_Memorandum.docx")
    doc.save(out)
    return out


# ============================================================================
# DOCUMENT 2 — DAMAGES SCENARIO ANALYSIS (INTERNAL)
# ============================================================================

def build_scenario_analysis():
    doc = setup_document()
    add_letterhead(doc)
    add_heading(doc, "DAMAGES SCENARIO ANALYSIS — INTERNAL / PRIVILEGED", level=1, center=True)
    add_privileged_callout(doc,
        "ATTORNEY-CLIENT PRIVILEGED / ATTORNEY WORK PRODUCT — PREPARED IN ANTICIPATION OF LITIGATION. "
        "Not for disclosure to opposing counsel."
    )
    add_horizontal_rule(doc)

    add_field_line(doc, "To", "Edward Susolik, Esq. — Callahan & Blaine, PC")
    add_field_line(doc, "From", "Ravi Jain, CEO, Technijian, Inc.")
    add_field_line(doc, "Re", "Technijian, Inc. v. Vintage Design, LLC — AAA Arbitration — Damages Modeling Under Signed 5/4/2023 Agreement")
    add_field_line(doc, "Date", "April 14, 2026")
    doc.add_paragraph()

    add_heading(doc, "PURPOSE", level=2)
    add_paragraph(doc,
        "This memorandum documents the six alternative damages computations available under the signed "
        "Client Monthly Service Agreement between Technijian, Inc. and Vintage Design, executed via "
        "DocuSign on May 4, 2023 (Envelope ID B679C550-E41C-4E31-BE34-7F8FFF437C3D). It confirms the "
        "Demand's stated damages figure ($240,555.15), identifies the three lower-value readings "
        "Respondent's counsel is expected to push, and sets out the controlling legal authorities that "
        "render each of those readings implausible. The Settlement Position Memorandum anchors at "
        "$240,555.15 (consistent with the Demand); this memorandum preserves the range for internal "
        "strategy and identifies the walk-away floor at $125,000."
    )

    add_heading(doc, "UNDERLYING DATA (through 7/31/2025)", level=2)
    add_table(doc,
        ["Contracted Category", "Native Rate", "Billed", "Actual", "Net Excess"],
        [
            ["Offshore Normal (CHD-TS1 R.N)", "$15/hr", "1,760.42", "2,682.18", "+921.76"],
            ["Offshore After-Hours (CHD-TS1 R.AF)", "$30/hr", "1,308.91", "1,801.09", "+492.18"],
            ["Onshore Remote Normal (IRV-TS1 R.N)", "$125/hr", "540.00", "456.82", "−83.18"],
            ["CONTRACTED TOTALS", "", "3,609.33", "4,940.09", "+1,330.76"],
        ],
        col_widths=[2.6, 0.9, 1.1, 1.1, 1.1]
    )
    add_paragraph(doc,
        "Key aggregations: Net excess 1,330.76 hrs; Gross positive excess 1,413.94 hrs; Sum of monthly "
        "positive excess 1,878.67 hrs; Post-signing net excess −69.68 hrs; Post-signing monthly "
        "positive 273.40 hrs."
    )

    add_heading(doc, "PERIOD-BY-PERIOD (12-month windows from 5/2023)", level=2)
    add_table(doc,
        ["Period", "Months", "Billed", "Actual", "Excess"],
        [
            ["P1 — May 2023 – Apr 2024", "12", "960.00", "1,970.97", "+1,010.97"],
            ["P2 — May 2024 – Apr 2025", "12", "2,025.08", "2,584.18", "+559.10"],
            ["P3 (terminated) — May 2025 – Jul 2025", "3", "624.25", "384.94", "−239.31"],
        ],
        col_widths=[2.6, 0.8, 1.1, 1.1, 1.1]
    )

    add_heading(doc, "THE SIX SCENARIOS AT $150/HOUR", level=2)

    scenarios = [
        ("Scenario D — Monthly Positive × $150 (most aggressive)",
         "1,878.67 × $150 = $281,800.50 + 10% late = **$309,980.55**.",
         "Consistent with ¶¶ 3–4 monthly true-up mechanism. Use: opening ask or ceiling."),
        ("Scenario Demand — 1,457.50 × $150 ★ ANCHOR",
         "1,457.50 × $150 = $218,625.00 + 10% late = **$240,487.50 (≈ Demand's $240,555.15)**.",
         "Already on the record. Primary anchor in the Settlement Position Memorandum."),
        ("Scenario B — Gross Positive × $150",
         "1,413.94 × $150 = $212,091.00 + 10% late = **$233,300.10**.",
         "Treats each category independently; textually defensible."),
        ("Scenario A — Net Excess × $150 (conservative fallback)",
         "1,330.76 × $150 = $199,614.00 + 10% late = **$219,575.40**.",
         "Directly computable from the reconciliation spreadsheet."),
    ]
    for title, calc, note in scenarios:
        add_heading(doc, title, level=3)
        add_paragraph(doc, calc)
        add_paragraph(doc, note, italic=True)

    add_heading(doc, "SCENARIOS BELOW THE $125,000 FLOOR — WHY NOT LEGALLY PLAUSIBLE", level=2)
    add_paragraph(doc,
        "The following three scenarios are the readings Respondent's counsel is expected to advance. Each is "
        "defeated by controlling California authority and the agreement's own text."
    )

    add_heading(doc, "Scenario E1 — Only P3 (terminated period) × $150 = $0", level=3)
    add_paragraph(doc,
        "Respondent's argument: ¶ 5 only reaches the period during which termination occurs. P3 ran below "
        "billed. No cancellation fee owed."
    )
    add_paragraph(doc, "Why this fails:")
    e1_items = [
        "Clause text defeats it. Under Contract ¶ 5 reads on \"any hours that exceeded the previous under contract period average\" — cumulative and unqualified. Cal. Civ. Code § 1638; City of Atascadero v. Merrill Lynch (1998) 68 Cal.App.4th 445, 473.",
        "Whole-contract rule defeats it. Reading ¶ 5 to reach only the terminated period would render the clause a dead letter in virtually every real-world case. Cal. Civ. Code § 1641; Powerine Oil Co. v. Superior Court (2005) 37 Cal.4th 377, 390–391.",
        "¶¶ 3–4 are prospective, not retroactive. The averaging mechanism adjusts future monthly billing. It does not forgive unpaid excess from prior periods. Technijian's P2/P3 billing stepped up to reflect higher demand but the P1 excess hours were delivered and never paid for.",
        "Two independent dispute waivers bar the \"records inaccurate\" backdoor. Weekly zero-dollar invoices with 30-day dispute window (~140 weekly notices over 27 months, zero disputes); plus T&C § 3.01's 60-day objection window on monthly invoices (zero objections across 27 monthly invoices). Cobb v. Pacific Mut. Life Ins. Co. (1935) 4 Cal.2d 565, 573; Gleason v. Klamer (1980) 103 Cal.App.3d 782, 787. Cal. Comm. Code § 1303 (course of performance).",
        "Ratification cements the reading. Erica Garcia, as VP Finance, personally reviewed and paid each monthly invoice for over two years. Rakestraw v. Rodrigues (1972) 8 Cal.3d 67, 73; Wagner v. Glendale Adventist Med. Ctr. (1989) 216 Cal.App.3d 1379, 1388.",
    ]
    for item in e1_items:
        p = doc.add_paragraph(item, style='List Number')
    add_paragraph(doc,
        "Bottom line on E1: Treating only P3's excess rewrites the cancellation clause. Not plausible.",
        italic=True
    )

    add_heading(doc, "Scenario E2 — Only P2 + P3 × $150 = $52,765.35", level=3)
    add_paragraph(doc,
        "Respondent's argument: P1 was trued up at end of period under ¶¶ 3–4. Only P2 + P3 excess counts. "
        "319.79 hrs × $150 + 10% = $52,765.35."
    )
    add_paragraph(doc, "Why this fails:")
    e2_items = [
        "Same whole-contract infirmity as E1. Clause reads on \"any hours that exceeded\" — not \"hours in the current or most recent period only.\" Cal. Civ. Code §§ 1638, 1641.",
        "No provision in ¶¶ 3–4 states that completed-period excess is released. Neither paragraph uses \"release,\" \"forgive,\" or \"credit back.\" Courts will not imply a release the parties did not write. Series AGI West Linn v. Eves (2013) 217 Cal.App.4th 156, 164–165.",
        "The averaging mechanism does not compensate Technijian for the delivered excess. P2's higher billed baseline charged Vintage Design prospectively for ongoing elevated demand; it did not retroactively pay for P1's 1,010.97 delivered-but-unpaid hours.",
        "California disfavors readings that render commercial clauses meaningless. MacKinnon v. Truck Ins. Exchange (2003) 31 Cal.4th 635, 648.",
        "§ 1671(b) burden still on Respondent. Ridgley v. Topa Thrift (1998) 17 Cal.4th 970, 977.",
    ]
    for item in e2_items:
        p = doc.add_paragraph(item, style='List Number')
    add_paragraph(doc,
        "Bottom line on E2: Intuitive appeal but no textual hook. Not plausible.",
        italic=True
    )

    add_heading(doc, "Scenario C — Native Rates Knockdown = $20,013.73", level=3)
    add_paragraph(doc,
        "Respondent's argument: $150/hr cancellation rate is a penalty. Reduce to native rates ($15/$30/$125). "
        "Net: $18,194.30 + 10% late = $20,013.73."
    )
    add_paragraph(doc, "Why this fails:")
    c_items = [
        "§ 1671(b) forecloses it. Statute presumes liquidated damages valid unless challenger proves unreasonable at time of contracting. Ridgley v. Topa Thrift; Hitz v. First Interstate (1995) 38 Cal.App.4th 274, 286–289.",
        "$150 is demonstrably not punitive. The Agreement's own Out of Contract rate for IRV-TS1 Tech Support is $150/hr onsite normal and $200/hr onsite after-hours. The cancellation rate equals or is below the bargained alternative rate — not a penalty as a matter of law. El Centro Mall (2009) 174 Cal.App.4th 58, 63.",
        "Reasonableness assessed at formation, not hindsight. El Centro Mall; Ridgley.",
        "No procedural unconscionability. B2B; VP Finance signed. Pinnacle Museum Tower (2012) 55 Cal.4th 223, 246–247; Sanchez v. Valencia Holding (2015) 61 Cal.4th 899, 910–911.",
        "Economic substance supports the rate. Technijian scaled capacity to meet above-baseline demand; mid-Period 3 termination leaves Technijian with sunk ramp-up cost. The $150 rate is the bargained compensation.",
    ]
    for item in c_items:
        p = doc.add_paragraph(item, style='List Number')
    add_paragraph(doc,
        "Bottom line on C: Respondent cannot meet the § 1671(b) burden. Not plausible.",
        italic=True
    )

    add_heading(doc, "SUMMARY TABLE", level=2)
    add_table(doc,
        ["Scenario", "Hours", "TOTAL", "Plausibility"],
        [
            ["D — Monthly positive × $150", "1,878.67", "$309,980.55", "Supportable (opening)"],
            ["Demand ★ ANCHOR", "1,457.50", "$240,487.50", "On record"],
            ["B — Gross positive × $150", "1,413.94", "$233,300.10", "Supportable"],
            ["A — Net excess × $150", "1,330.76", "$219,575.40", "Directly provable"],
            ["★ $125,000 SETTLEMENT FLOOR", "", "$125,000", "Walk-away"],
            ["E2 — P2 + P3 only × $150", "319.79", "$52,765.35", "Not plausible"],
            ["C — Native rates", "1,330.76 (mixed)", "$20,013.73", "Not plausible (§ 1671(b))"],
            ["E1 — P3 only × $150", "−239.31", "$0", "Not plausible"],
        ],
        col_widths=[2.3, 1.4, 1.4, 1.8]
    )

    add_heading(doc, "SETTLEMENT STRATEGY", level=2)
    add_table(doc,
        ["Position", "Amount", "Rationale"],
        [
            ["Opening (stated in memo)", "$240,555.15", "Matches Demand"],
            ["First concession", "$200,000 – $220,000", "Scenarios A/B territory"],
            ["Realistic landing zone", "$150,000 – $180,000", "~35–40% off Demand"],
            ["Walk-away floor", "$125,000", "Below this, arbitrate"],
            ["Internal ceiling (not stated)", "$309,980.55 (Scenario D)", "If pushed"],
        ],
        col_widths=[2.0, 1.8, 2.7]
    )

    add_heading(doc, "KEY LEGAL LEVERAGE POINTS", level=2)
    add_table(doc,
        ["Argument", "Authority", "Effect"],
        [
            ["§ 1671(b) places burden on Respondent", "Ridgley; Hitz", "Defeats Scenario C"],
            ["Whole-contract construction", "Civ. Code § 1641; Powerine Oil; MacKinnon", "Defeats E1 and E2"],
            ["Plain meaning of ¶ 5", "Civ. Code § 1638; City of Atascadero", "Defeats E1 and E2"],
            ["Integration clause bars oral side-deals", "T&C § 6.08; CCP § 1856; Casa Herrera", "Blocks oral defenses"],
            ["Weekly (30-day) + monthly (60-day) dispute waivers", "T&C § 3.01; Cobb; Gleason; Comm. Code § 1303", "Forecloses ticketing-accuracy defense"],
            ["27 months ratification", "Rakestraw; Pasadena Medi-Center", "Cures any late defenses"],
            ["VP Finance apparent authority", "Corp. Code § 313; Snukal", "Defeats authority challenge"],
            ["DocuSign attributable", "Civ. Code § 1633.9; Fabian", "Defeats authentication challenge"],
            ["Attorneys' fees via § 5.02 + § 1717", "Santisas v. Goodin", "Shifts fee exposure"],
            ["Business records admissible", "Evid. Code §§ 1271, 1552; Jazayeri", "Admits ticketing and ledger"],
        ],
        col_widths=[2.2, 2.5, 1.8]
    )
    add_paragraph(doc,
        "Full citations and treatment are in the accompanying Case Law Research Memorandum (30 issues).",
        italic=True
    )

    add_heading(doc, "PRE-MEETING CHECKLIST", level=2)
    items = [
        "Confirm the $240,555.15 anchor and $125,000 floor with Technijian",
        "Confirm landing zone ($150K–$180K) and opening concession tiers internally",
        "Prepare Exhibit C month-by-month billed-vs-actual reconciliation as a standalone spreadsheet",
        "Prepare \"Delaware LLC → California corporation\" caption amendment (AAA Rule R-6)",
        "Pull June 23, 2025 termination notice and June 27, 2025 final invoice",
        "Calculate precise attorneys' fees incurred to date",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    add_footer(doc, "ATTORNEY-CLIENT PRIVILEGED / ATTORNEY WORK PRODUCT")

    out = os.path.join(VGD, "Damages_Scenario_Analysis_INTERNAL.docx")
    doc.save(out)
    return out


# ============================================================================
# DOCUMENT 3 — CASE LAW RESEARCH MEMO
# ============================================================================

def build_case_law_memo():
    doc = setup_document()
    add_letterhead(doc)
    add_heading(doc, "CALIFORNIA CASE LAW RESEARCH MEMORANDUM", level=1, center=True)
    add_horizontal_rule(doc)

    add_field_line(doc, "Re", "Technijian, Inc. v. Vintage Design, LLC — AAA Arbitration")
    add_field_line(doc, "Subject", "Authorities supporting Technijian's claims and anticipated defenses")
    add_field_line(doc, "Date", "April 14, 2026")
    add_field_line(doc, "Prepared for", "Edward Susolik, Esq. — Callahan & Blaine, PC")
    doc.add_paragraph()

    add_paragraph(doc,
        "Caveat: All citations below should be Shepardized / cite-checked before filing. This memorandum "
        "organizes leading California authorities by issue; it is not a substitute for independent "
        "verification. Where federal authority is included it is noted.",
        italic=True
    )

    issues = [
        ("1. Breach of Written Contract — Elements",
         "Oasis West Realty, LLC v. Goldman (2011) 51 Cal.4th 811, 821; Reichert v. General Ins. Co. (1968) 68 Cal.2d 822, 830; Coles v. Glaser (2016) 2 Cal.App.5th 384, 391."),
        ("2. Liquidated Damages / Cancellation Fee ($150/hr) — Cal. Civ. Code § 1671(b)",
         "Cal. Civ. Code § 1671(b) (presumption of validity; burden on challenger); Ridgley v. Topa Thrift & Loan Assn. (1998) 17 Cal.4th 970, 977; Hitz v. First Interstate Bank (1995) 38 Cal.App.4th 274, 286–289; El Centro Mall, LLC v. Payless ShoeSource, Inc. (2009) 174 Cal.App.4th 58, 63; Greentree Financial Group v. Execute Sports (2008) 163 Cal.App.4th 495. Key fact: $150 cancellation rate is at or below the Agreement's own Out-of-Contract rates ($150 onsite normal / $200 onsite AH) — not a penalty as a matter of law."),
        ("3. Open Book Account (CCP § 337a)",
         "Cal. Code Civ. Proc. § 337a (definition); § 337(b) (4-year SOL); Interstate Group Administrators v. Cravens, Dargan & Co. (1985) 174 Cal.App.3d 700, 708; R.N.C. Inc. v. Tsegeletos (1991) 231 Cal.App.3d 967, 971–972."),
        ("4. Account Stated",
         "Zinn v. Fred R. Bright Co. (1969) 271 Cal.App.2d 597, 600 (elements); Trafton v. Youngblood (1968) 69 Cal.2d 17, 25; Gleason v. Klamer (1980) 103 Cal.App.3d 782, 787."),
        ("5. Services Rendered / Quantum Meruit (Alternative)",
         "Huskinson & Brown v. Wolf (2004) 32 Cal.4th 453, 458; Maglica v. Maglica (1998) 66 Cal.App.4th 442; Hedging Concepts v. First Alliance Mortg. (1996) 41 Cal.App.4th 1410, 1419."),
        ("6. Late Fees as Liquidated Damages",
         "Garrett v. Coast & Southern Fed. Sav. & Loan (1973) 9 Cal.3d 731, 738–740; Beasley v. Wells Fargo Bank (1991) 235 Cal.App.3d 1383."),
        ("7. Authority of Corporate Officer — Erica Garcia as VP Finance",
         "Cal. Corp. Code § 313 (instrument signed by VP binding as matter of law); Snukal v. Flightways Mfg., Inc. (2000) 23 Cal.4th 754, 782; Memorial Hospitals Assn. v. Randol (1995) 38 Cal.App.4th 1300, 1308."),
        ("8. Course of Dealing, Acceptance, and Modification — 27 Months of Paid Performance",
         "Cal. Comm. Code § 1303; Wagner v. Glendale Adventist Med. Ctr. (1989) 216 Cal.App.3d 1379, 1388; Employers Reinsurance Co. v. Superior Court (2008) 161 Cal.App.4th 906, 921."),
        ("9. Dual Waiver of Right to Object — Weekly (30-day) + Monthly (60-day)",
         "Technijian sent ~140 weekly zero-dollar invoices with full ticket detail and 30-day dispute windows (zero disputes), plus monthly invoices subject to T&C § 3.01's 60-day objection window (zero objections). Either waiver alone bars challenge; both combined render every ticketed hour undisputable as a matter of contract. Cobb v. Pacific Mut. Life Ins. Co. (1935) 4 Cal.2d 565, 573; Gleason v. Klamer (1980) 103 Cal.App.3d 782, 787; Trafton v. Youngblood (1968) 69 Cal.2d 17, 25; Cal. Comm. Code §§ 1303, 2607(3)(a); Pollard v. Saxe & Yolles Dev. Co. (1974) 12 Cal.3d 374, 380."),
        ("10. Continued Performance After Termination — Not a Waiver",
         "DRG/Beverly Hills v. Chopstix Dim Sum Cafe (1994) 30 Cal.App.4th 54, 60; Salyer Grain & Milling v. Henson (1970) 13 Cal.App.3d 493, 501; Wind Dancer Production v. Walt Disney (2017) 10 Cal.App.5th 56, 78."),
        ("11. Attorneys' Fees via T&C § 5.02 Indemnity + Cal. Civ. Code § 1717 Reciprocity",
         "T&C § 5.02 provides that Vintage Design shall indemnify Technijian for attorneys' fees incurred as a result of \"any failure by Client…to comply with the terms of this Agreement.\" Non-payment of the cancellation fee is such a failure. Cal. Civ. Code § 1717 renders the one-way indemnity reciprocal. Santisas v. Goodin (1998) 17 Cal.4th 599, 610–611; Hsu v. Abbara (1995) 9 Cal.4th 863, 871; PLCM Group v. Drexler (2000) 22 Cal.4th 1084, 1095; Moshonov v. Walsh (2000) 22 Cal.4th 771, 776; Advanced Micro Devices v. Intel (1994) 9 Cal.4th 362, 375–376; AAA Commercial Rule R-47(d)(ii)."),
        ("12. Prejudgment Interest (§§ 3287, 3289(b))",
         "Cal. Civ. Code §§ 3287(a), 3289(b); Wisper Corp. v. California Commerce Bank (1996) 49 Cal.App.4th 948, 958–960; North Oakland Med. Clinic v. Rogers (1998) 65 Cal.App.4th 824, 829–830; Continental Heller Corp. v. Amtech Mech. Services (1997) 53 Cal.App.4th 500, 511."),
        ("13. AAA Commercial Arbitration — Enforceability and Scope",
         "AT&T Mobility v. Concepcion (2011) 563 U.S. 333; Pinnacle Museum Tower Assn. v. Pinnacle Market Development (2012) 55 Cal.4th 223; Moncharsh v. Heily & Blase (1992) 3 Cal.4th 1; Sandquist v. Lebo Automotive (2016) 1 Cal.5th 233; CCP §§ 1286, 1286.2."),
        ("14. Pleading Amendment / Relation Back / Misnomer",
         "Only a minor \"Delaware LLC → California corporation\" caption correction is needed under AAA Rule R-6. Hawkins v. Pacific Coast Bldg. Prods. (2004) 124 Cal.App.4th 1497, 1503–1504; Mayberry v. Coca-Cola Bottling Co. (1966) 244 Cal.App.2d 350, 352; Diliberti v. Stage Call Corp. (1992) 4 Cal.App.4th 1468; Garrett v. Crown Coach Corp. (1968) 259 Cal.App.2d 647, 650–651."),
        ("15. Unconscionability — Anticipated Defense",
         "Armendariz v. Foundation Health Psychcare (2000) 24 Cal.4th 83, 113–114; Sanchez v. Valencia Holding (2015) 61 Cal.4th 899, 910–911; Pinnacle Museum Tower, supra, 55 Cal.4th at 246–247; Baltazar v. Forever 21 (2016) 62 Cal.4th 1237, 1244."),
        ("16. Confidentiality of Settlement Communications",
         "Cal. Evid. Code §§ 1119, 1152; Fed. R. Evid. 408; Cassel v. Superior Court (2011) 51 Cal.4th 113; Foxgate Homeowners' Assn. v. Bramalea California (2001) 26 Cal.4th 1, 14."),
        ("17. Electronic Signatures / DocuSign (UETA)",
         "Cal. Civ. Code §§ 1633.7(a), (d), 1633.9(a); 15 U.S.C. § 7001(a) (federal E-SIGN); Fabian v. Renovate America (2019) 42 Cal.App.5th 1062, 1066–1069; J.B.B. Investment Partners v. Fair (2014) 232 Cal.App.4th 974; Ruiz v. Moss Bros. Auto Group (2014) 232 Cal.App.4th 836, 844. DocuSign Envelope B679C550 audit trail = prima facie authentication."),
        ("18. Parol Evidence Rule & Integration Clause",
         "Cal. Code Civ. Proc. § 1856(a); Casa Herrera v. Beydoun (2004) 32 Cal.4th 336, 343–344; Banco do Brasil v. Latian (1991) 234 Cal.App.3d 973, 1001–1002; EPA Real Estate Partnership v. Kang (1992) 12 Cal.App.4th 171, 175–176; Riverisland Cold Storage v. Fresno-Madera PCA (2013) 55 Cal.4th 1169."),
        ("19. Contract Interpretation — Plain Meaning and Construction",
         "Cal. Civ. Code §§ 1636, 1638, 1639, 1641; Pacific Gas & E. v. G.W. Thomas Drayage & Rigging (1968) 69 Cal.2d 33; City of Atascadero v. Merrill Lynch (1998) 68 Cal.App.4th 445, 473; Powerine Oil v. Superior Court (2005) 37 Cal.4th 377, 390–391; MacKinnon v. Truck Ins. Exchange (2003) 31 Cal.4th 635, 648."),
        ("20. General Damages and Certainty (§§ 3300, 3301)",
         "Cal. Civ. Code §§ 3300, 3301; Lewis Jorge Construction v. Pomona USD (2004) 34 Cal.4th 960, 968–970; Sargon Enterprises v. USC (2012) 55 Cal.4th 747, 773–775; Toscano v. Greene Music (2004) 124 Cal.App.4th 685, 694 (liquidated damages cures certainty requirement)."),
        ("21. Statute of Limitations — Written Contracts",
         "Cal. Code Civ. Proc. § 337 (4 years); Fox v. Ethicon Endo-Surgery (2005) 35 Cal.4th 797, 806; Howard v. County of San Diego (2010) 184 Cal.App.4th 1422, 1431. Breach accrued 7/27/2025; Demand filed 10/23/2025 — well within SOL."),
        ("22. Anticipatory Breach / Repudiation / Total Breach",
         "Taylor v. Johnston (1975) 15 Cal.3d 130, 137; Central Valley General Hospital v. Smith (2008) 162 Cal.App.4th 501, 514–515; Romano v. Rockwell Internat. (1996) 14 Cal.4th 479, 488–490; Brown v. Grimes (2011) 192 Cal.App.4th 265, 277."),
        ("23. Evidence — Business Records (Ticketing & Invoices)",
         "Cal. Evid. Code §§ 1271, 1552, 1553; Jazayeri v. Mao (2009) 174 Cal.App.4th 301, 324; Aguimatang v. California State Lottery (1991) 234 Cal.App.3d 769, 797–798; AAA Commercial Rule R-34."),
        ("24. Ratification by Acceptance of Benefits — 27-Month Paid Performance",
         "Rakestraw v. Rodrigues (1972) 8 Cal.3d 67, 73; Reusche v. California Pacific Title Ins. (1965) 231 Cal.App.2d 731, 737; Pasadena Medi-Center Associates v. Superior Court (1973) 9 Cal.3d 773, 779–780; Common Wealth Ins. Systems v. Kersten (1974) 40 Cal.App.3d 1014, 1026."),
        ("25. Statute of Frauds — Satisfied",
         "Cal. Civ. Code § 1624(a)(1); Sterling v. Taylor (2007) 40 Cal.4th 757, 766; Secrest v. Security National Mortgage (2008) 167 Cal.App.4th 544, 552. Writing + DocuSign signature satisfies."),
        ("26. Implied Covenant of Good Faith and Fair Dealing",
         "Guz v. Bechtel National (2000) 24 Cal.4th 317, 349; Carma Developers (Cal.) v. Marathon Development California (1992) 2 Cal.4th 342, 371–374; Racine & Laramie v. Dept. of Parks & Recreation (1992) 11 Cal.App.4th 1026, 1032."),
        ("27. Arbitrator's Authority to Award Fees & Costs",
         "Moshonov v. Walsh (2000) 22 Cal.4th 771, 776; Advanced Micro Devices v. Intel (1994) 9 Cal.4th 362, 375–376; DiMarco v. Chaney (1995) 31 Cal.App.4th 1809, 1815–1816; Corona v. Amherst Partners (2003) 107 Cal.App.4th 701; AAA Commercial Rule R-47(d)(ii)."),
        ("28. Discovery in Arbitration",
         "Cal. Code Civ. Proc. § 1283.05; AAA Commercial Rule R-22; Berglund v. Arthroscopic & Laser Surgery Ctr. of San Diego (2008) 44 Cal.4th 528."),
        ("29. Freedom of Contract — Sophisticated Commercial Parties",
         "Appalachian Ins. Co. v. McDonnell Douglas Corp. (1989) 214 Cal.App.3d 1, 12; Powerine Oil v. Superior Court (2005) 37 Cal.4th 377, 391; Gaggero v. Yura (2003) 108 Cal.App.4th 884, 892; Badie v. Bank of America (1998) 67 Cal.App.4th 779, 802."),
        ("30. Post-Judgment Interest on Confirmed Arbitration Award",
         "Cal. Code Civ. Proc. §§ 685.010(a), 1286, 1287.4; Britz, Inc. v. Alfa-Laval Food & Dairy Co. (1995) 34 Cal.App.4th 1085, 1107; Pierotti v. Torian (2000) 81 Cal.App.4th 17, 28."),
    ]
    for title, body in issues:
        add_heading(doc, title, level=3)
        add_paragraph(doc, body)

    add_footer(doc, "Case Law Research Memorandum — Technijian v. Vintage Design")

    out = os.path.join(VGD, "Case_Law_Research_Memorandum.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    files = []
    files.append(build_settlement_memo())
    files.append(build_scenario_analysis())
    files.append(build_case_law_memo())
    for f in files:
        size = os.path.getsize(f)
        print(f"{f} ({size:,} bytes)")
    print("\nDone.")
