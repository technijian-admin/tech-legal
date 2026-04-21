"""Generate BST legal deliverable DOCX files from MD sources.

Matches the VTD folder pattern (professional headings, tables for markdown tables,
preserves italics/bold inline, indents bulleted lists).
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BST_DIR = os.path.dirname(os.path.abspath(__file__))

TASKS = [
    ("BST_DAMAGES_ANALYSIS.md", "BST_Damages_Scenario_Analysis_INTERNAL.docx"),
    ("BST_SETTLEMENT_POSITION_MEMORANDUM.md", "BST_Settlement_Position_Memorandum.docx"),
    ("BST_CASE_LAW_RESEARCH_MEMORANDUM.md", "BST_Case_Law_Research_Memorandum.docx"),
    ("EMAIL_TO_SUSOLIK.md", "BST_Email_to_Susolik.docx"),
]


def add_inline(paragraph, text):
    """Handle bold (**x**) and italics (*x*) inline."""
    # Simple tokenizer — alternating bold/italic/plain
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = paragraph.add_run(p[2:-2])
            r.bold = True
        elif p.startswith("*") and p.endswith("*") and not p.startswith("**"):
            r = paragraph.add_run(p[1:-1])
            r.italic = True
        elif p.startswith("`") and p.endswith("`"):
            r = paragraph.add_run(p[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            paragraph.add_run(p)


def parse_table_rows(lines, i):
    """Parse a markdown table starting at lines[i]. Return (table_rows, new_i)."""
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        line = lines[i].strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        rows.append(cells)
        i += 1
    # Filter out the separator row (|---|---|)
    rows = [r for r in rows if not all(re.match(r'^[-:\s]+$', c) for c in r)]
    return rows, i


def render_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()
    doc = Document()
    # base style
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.split("\n")
    i = 0
    in_code = False
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Fenced code
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(raw)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("_" * 70)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            if level == 1:
                p = doc.add_heading(level=0)
            else:
                p = doc.add_heading(level=min(level - 1, 4))
            add_inline(p, heading_text)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]:
            rows, i = parse_table_rows(lines, i)
            render_table(doc, rows)
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            content = stripped[1:].lstrip()
            r = p.add_run(content)
            r.italic = True
            r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            i += 1
            continue

        # Bulleted list
        if re.match(r"^[-*+]\s+", stripped):
            content = re.sub(r"^[-*+]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, content)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, raw)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    for md, dx in TASKS:
        md_path = os.path.join(BST_DIR, md)
        dx_path = os.path.join(BST_DIR, dx)
        if not os.path.exists(md_path):
            print(f"SKIP {md} (not found)")
            continue
        convert(md_path, dx_path)
    print("Done.")
