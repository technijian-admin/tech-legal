"""
Brandywine Homes (BWH) Executive Summary Generator
Migration from old T&C to new MSA — existing client upgrade
Uses Technijian Brand Guide 2026 formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Brand Colors (from Technijian_Brand_Guide_2026) --
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex versions for shading
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_GREEN = "28A745"

FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")

doc = Document()

# -- Page Setup --
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = BRAND_GREY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# -- Helper Functions --

def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def accent_bar(color_hex, height_pt=4):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pf = p._element.get_or_add_pPr()
    spacing = pf.makeelement(qn("w:spacing"), {
        qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
    })
    pf.append(spacing)
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def run_styled(paragraph, text, size=11, bold=False, color=None, italic=False, font=FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def section_header(title):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    bar_cell = tbl.rows[0].cells[0]
    set_cell_shading(bar_cell, HEX_BLUE)
    bar_cell.width = Inches(0.08)
    bar_cell.text = ""
    bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    title_cell = tbl.rows[0].cells[1]
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run_styled(p, title, size=14, bold=True, color=CORE_BLUE)
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def styled_table(headers, rows, col_widths=None, total_row_indices=None):
    total_row_indices = total_row_indices or []
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        run_styled(p, h, size=9, bold=True, color=WHITE)
    for row_idx, row_data in enumerate(rows):
        row = tbl.add_row()
        is_total = row_idx in total_row_indices
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if is_total:
                set_cell_shading(cell, HEX_OFF_WHITE)
                run_styled(p, str(val), size=9, bold=True, color=DARK_CHARCOAL)
            else:
                run_styled(p, str(val), size=9, color=BRAND_GREY)
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return tbl


def green_table(headers, rows, col_widths=None, total_row_indices=None):
    """Table with green header for savings/comparison sections."""
    total_row_indices = total_row_indices or []
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEX_GREEN)
        cell.text = ""
        p = cell.paragraphs[0]
        run_styled(p, h, size=9, bold=True, color=WHITE)
    for row_idx, row_data in enumerate(rows):
        row = tbl.add_row()
        is_total = row_idx in total_row_indices
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if is_total:
                set_cell_shading(cell, HEX_OFF_WHITE)
                run_styled(p, str(val), size=9, bold=True, color=DARK_CHARCOAL)
            else:
                run_styled(p, str(val), size=9, color=BRAND_GREY)
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return tbl


def bullet_item(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_styled(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)
    return p


def body_text(text, space_after=8):
    p = doc.add_paragraph()
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ================================================================
# CLIENT DATA
# ================================================================

CLIENT_NAME = "Brandywine Homes"
CLIENT_SHORT = "BWH"
CLIENT_CODE = "BWH"
CLIENT_ADDRESS = "2355 Main St #220, Irvine, CA 92614"
CLIENT_INDUSTRY = "Residential Real Estate / Homebuilder"

# April (current) totals — COMBINED from both invoices
APRIL_MONTHLY_INVOICE = 9041.05       # Monthly Invoice #28148
APRIL_RECURRING_INVOICE = 1234.70     # Recurring Invoice #28116
APRIL_TOTAL = 10275.75                # Combined old T&C total

# May (proposed) totals — after changes
MAY_DESKTOPS = 32
MAY_SERVERS = 12

# May Online Services (32 desktops, no Network Assessment)
# Desktop services (32 units): Patch $4, SI $6, MR $2, OPS-NET $3.25, AVH $6, AV $8.50 = $29.75/desktop
# 32 * $29.75 = $952.00
# Server services (12 units): Patch $4, SI $6, MR $2, OPS-NET $3.25, IB $15, AVH $6, AV $10.50 = $46.75/server
# 12 * $46.75 = $561.00
# Cloud: Backup Storage 13*$50=$650, Veeam One 11*$3=$33 = $683.00
# Network: OPS-Config 3*$6=$18, RTPT 6*$7=$42, OPS-Traffic 5*$14=$70, OPS-Wifi 6*$1=$6,
#          OPS-Storage 2*$4.75=$9.50, Sophos FW 1*$270=$270, Edge Appliance 1*$100=$100, OPS-Traffic 1*$14=$14 = $459.50
# Email: SA 1*$50=$50, DMARC 1*$20=$20, Anti-Spam 83*$6.25=$518.75, Phishing 85*$6=$510,
#        Veeam 365 110*$2.50=$275, My Disk 1*$16=$16 = $1,389.75
MAY_ONLINE_SERVICES = 952.00 + 561.00 + 683.00 + 459.50 + 1389.75  # = $4,045.25

# May Schedule B — Subscriptions
MAY_SCHEDULE_B = 7.20  # Microsoft Entra ID P1

# May Virtual Staff (15% discount)
MAY_VS_ARCHITECT = 5.00 * 170.00    # $850.00
MAY_VS_US_TECH = 15.26 * 106.25     # $1,621.38
MAY_VS_INDIA_NORMAL = 58.13 * 12.75 # $741.16
MAY_VS_INDIA_AFTER = 42.82 * 25.50  # $1,091.91
MAY_VIRTUAL_STAFF = MAY_VS_ARCHITECT + MAY_VS_US_TECH + MAY_VS_INDIA_NORMAL + MAY_VS_INDIA_AFTER  # $4,304.44

MAY_TOTAL = MAY_ONLINE_SERVICES + MAY_SCHEDULE_B + MAY_VIRTUAL_STAFF  # $8,356.89

MONTHLY_SAVINGS = APRIL_TOTAL - MAY_TOTAL  # $1,918.86
SAVINGS_PCT = (MONTHLY_SAVINGS / APRIL_TOTAL) * 100  # 18.7%


# ================================================================
# COVER PAGE
# ================================================================

accent_bar(HEX_BLUE, height_pt=6)

for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO_PATH):
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(3.0))

# Orange divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
div_tbl = doc.add_table(rows=1, cols=3)
div_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in div_tbl.rows[0].cells:
    cell.text = ""
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
mid_cell = div_tbl.rows[0].cells[1]
mid_cell.width = Inches(2.5)
set_cell_shading(mid_cell, HEX_ORANGE)
div_tbl.rows[0].cells[0].width = Inches(2.25)
div_tbl.rows[0].cells[2].width = Inches(2.25)
div_pr = div_tbl._element.tblPr
div_borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    "</w:tblBorders>"
)
div_pr.append(div_borders)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(8)
run_styled(p, "Executive Summary", size=26, bold=True, color=DARK_CHARCOAL)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run_styled(p, "MSA Migration \u2014 Managed IT Services, Cybersecurity & Virtual Staff", size=14, color=BRAND_GREY)

# Prepared for
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run_styled(p, "Prepared for ", size=12, color=BRAND_GREY)
run_styled(p, CLIENT_NAME, size=12, bold=True, color=CORE_BLUE)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run_styled(p, "April 2026", size=11, color=BRAND_GREY)

for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

accent_bar(HEX_ORANGE, height_pt=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(
    p,
    "Technijian  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)


# ================================================================
# PAGE 2 - MIGRATION OVERVIEW
# ================================================================
doc.add_page_break()

section_header("Migration Overview")
body_text(
    f"Technijian is pleased to present this proposal to migrate {CLIENT_NAME} from the existing "
    "Terms & Conditions (T&C) agreement to a new Master Service Agreement (MSA). This migration "
    "formalizes the service relationship under a comprehensive legal framework while delivering "
    "immediate cost savings through service optimization and rate consolidation."
)
body_text(
    f"As an existing Technijian client, {CLIENT_NAME} has been receiving managed IT services, "
    "cybersecurity, and virtual staff support under the legacy T&C structure. The new MSA provides "
    "enhanced protections for both parties, clearer service definitions, structured billing with "
    "full transparency, and negotiated rates that reflect the strength of this ongoing partnership."
)

# Key Changes callout
p = doc.add_paragraph()
run_styled(p, "Key Changes in This Migration:", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)

bullet_item("Desktop count reduced from 41 to 32 (decommissioned/consolidated endpoints)", bold_prefix="Endpoint Optimization: ")
bullet_item("Network Assessment service removed (redundant with existing Site Assessment coverage)", bold_prefix="Service Streamlining: ")
bullet_item("15% reduction across all Virtual Staff roles, effective immediately under the new MSA", bold_prefix="Rate Reduction: ")
bullet_item("From legacy T&C to full MSA with Schedule A, Schedule B, and Schedule C (Rate Card)", bold_prefix="Agreement Upgrade: ")
bullet_item("Virtual Staff services billed on a 12-month cycle with weekly tracking invoices", bold_prefix="Billing Cycle: ")
bullet_item("Technijian will target 20% under billed hours monthly to reduce the accumulated unpaid hour balance (~1,007 hrs) over the 12-month cycle", bold_prefix="Unpaid Hours Recovery: ")
bullet_item("Quota right-sized from 24 TB to 13 TB with deduplication enabled to keep storage usage down", bold_prefix="Backup Storage Optimization: ")
bullet_item("Two separate invoices (monthly + recurring) merged into single monthly MSA invoice", bold_prefix="Invoice Consolidation: ")
bullet_item("Veeam 365 backup consolidated to 110 users under the new MSA", bold_prefix="Veeam 365 Consolidated: ")

# Current Environment
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Current Environment")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Details"],
    [
        ("Managed Desktops", "32 (reduced from 41)"),
        ("Managed Servers", "12"),
        ("Email Users", "83 (Anti-Spam) / 85 (Phishing Training) / 110 (Veeam 365)"),
        ("Firewalls", "1 (Sophos 2C-4G)"),
        ("Switches", "3"),
        ("Access Points", "6"),
        ("Backup Storage", "13 TB"),
        ("Cloud Platform", "Microsoft 365 (Exchange Online, SharePoint, OneDrive, Teams)"),
        ("Industry", CLIENT_INDUSTRY),
        ("Address", CLIENT_ADDRESS),
    ],
    col_widths=[1.8, 4.2],
)


# ================================================================
# PAGE 3 - APRIL vs MAY COMPARISON
# ================================================================
doc.add_page_break()

section_header("April 2026 vs May 2026 \u2014 Cost Comparison")
body_text(
    "The table below shows a side-by-side comparison of the current April 2026 combined invoices "
    "(old T&C) versus the proposed May 2026 consolidated invoice (new MSA). All figures are monthly costs."
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph()
run_styled(p, "April (Old T&C \u2014 Combined): ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(p, "Monthly Invoice #28148: $9,041.05 + Recurring Invoice #28116: $1,234.70 = $10,275.75", size=10, color=BRAND_GREY)
p.paragraph_format.space_after = Pt(4)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
green_table(
    ["Category", "April (Combined)", "May (Proposed)", "Savings", "Change"],
    [
        ("Online Services", "$10,275.75", "$4,045.25", "", "Consolidated"),
        ("Schedule B Subscriptions", "", "$7.20", "", "Entra ID P1"),
        ("Virtual Staff", "(incl. above)", "$4,304.44", "", "15% rate reduction"),
        ("TOTAL", "$10,275.75", "$8,356.89", "$1,918.86", "18.7% savings"),
    ],
    col_widths=[1.4, 1.2, 1.2, 1.0, 1.3],
    total_row_indices=[3],
)

# What Changed detail
doc.add_paragraph().paragraph_format.space_after = Pt(6)
p = doc.add_paragraph()
run_styled(p, "What Changed \u2014 Detail", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run_styled(p, "1. Desktop Reduction (41 to 32): ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(p, "9 desktops decommissioned or consolidated. This reduces costs across 6 per-desktop services: "
    "Patch Management, My Secure Internet, My Remote, My Ops-Net, AVH Protection, and AV Protection. "
    "Savings: $267.75/mo.", size=10, color=BRAND_GREY)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run_styled(p, "2. Network Assessment Removed: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(p, "The Network Assessment service has been consolidated into the "
    "Site Assessment already included in the service stack. "
    "Savings: $78.00/mo.", size=10, color=BRAND_GREY)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run_styled(p, "3. Virtual Staff Rate Reduction (15%): ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(p, "All Virtual Staff roles receive a 15% rate discount under the new MSA, recognizing the "
    "long-term partnership. Systems Architect: $200 to $170/hr. USA Tech: $125 to $106.25/hr. "
    "India Normal: $15 to $12.75/hr. India After Hours: $30 to $25.50/hr. "
    "Savings: $759.61/mo.", size=10, color=BRAND_GREY)
p.paragraph_format.space_after = Pt(4)


# ================================================================
# PAGE 4 - DETAILED MAY SERVICE BREAKDOWN
# ================================================================
doc.add_page_break()

section_header("Online Services \u2014 Infrastructure (May)")
body_text(
    f"Enterprise infrastructure management for {CLIENT_NAME}\u2019s 32 desktops, 12 servers, "
    "and full network stack. All pricing reflects May 2026 rates under the new MSA."
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Desktop services
p = doc.add_paragraph()
run_styled(p, "Desktop Services (32 units)", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)

styled_table(
    ["Service", "Qty", "Unit Price", "Monthly"],
    [
        ("Patch Management", "32", "$4.00", "$128.00"),
        ("My Secure Internet", "32", "$6.00", "$192.00"),
        ("My Remote", "32", "$2.00", "$64.00"),
        ("My Ops - Net", "32", "$3.25", "$104.00"),
        ("AVH Protection - Desktop", "32", "$6.00", "$192.00"),
        ("AV Protection - Desktop", "32", "$8.50", "$272.00"),
        ("SUBTOTAL \u2014 DESKTOP", "", "", "$952.00"),
    ],
    col_widths=[2.5, 0.6, 0.9, 1.0],
    total_row_indices=[6],
)

# Server services
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
run_styled(p, "Server Services (12 units)", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)

styled_table(
    ["Service", "Qty", "Unit Price", "Monthly"],
    [
        ("Patch Management", "12", "$4.00", "$48.00"),
        ("My Secure Internet", "12", "$6.00", "$72.00"),
        ("My Remote", "12", "$2.00", "$24.00"),
        ("My Ops - Net", "12", "$3.25", "$39.00"),
        ("Image Backup", "12", "$15.00", "$180.00"),
        ("AVH Protection - Server", "12", "$6.00", "$72.00"),
        ("AV Protection - Server", "12", "$10.50", "$126.00"),
        ("SUBTOTAL \u2014 SERVER", "", "", "$561.00"),
    ],
    col_widths=[2.5, 0.6, 0.9, 1.0],
    total_row_indices=[7],
)

# Other infrastructure
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
run_styled(p, "Other Infrastructure", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)

styled_table(
    ["Service", "Qty", "Unit Price", "Monthly"],
    [
        ("Backup Storage (TB)", "13", "$50.00", "$650.00"),
        ("Veeam One", "11", "$3.00", "$33.00"),
        ("My Ops - Config (Switches)", "3", "$6.00", "$18.00"),
        ("Real Time Pen Testing", "6", "$7.00", "$42.00"),
        ("My Ops - Traffic (Firewalls)", "1", "$14.00", "$14.00"),
        ("Site Assessment", "1", "$50.00", "$50.00"),
        ("DMARC/DKIM", "1", "$20.00", "$20.00"),
        ("My Ops - Storage (Disks)", "2", "$4.75", "$9.50"),
        ("My Ops - Wifi (APs)", "6", "$1.00", "$6.00"),
        ("Sophos Firewall Subscription 2C-4G", "1", "$270.00", "$270.00"),
        ("MPC Edge Appliance (16GB/8 cores/512GB)", "1", "$100.00", "$100.00"),
        ("SUBTOTAL \u2014 OTHER", "", "", "$1,112.50"),
    ],
    col_widths=[2.5, 0.6, 0.9, 1.0],
    total_row_indices=[11],
)

# Online Services Total
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["", "Monthly"],
    [
        ("TOTAL ONLINE SERVICES", "$4,045.25"),
    ],
    col_widths=[4.0, 2.0],
    total_row_indices=[0],
)


# ================================================================
# PAGE 5 - EMAIL & USERS + VIRTUAL STAFF
# ================================================================
doc.add_page_break()

section_header("Email & User Services (May)")
body_text(
    "Email protection, compliance training, and backup services for all users. "
    "These services remain unchanged from the April invoice."
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Qty", "Unit Price", "Monthly"],
    [
        ("Anti-Spam Standard", "83", "$6.25", "$518.75"),
        ("Phishing Training", "85", "$6.00", "$510.00"),
        ("Veeam 365", "110", "$2.50", "$275.00"),
        ("My Disk", "1", "$16.00", "$16.00"),
        ("SUBTOTAL \u2014 EMAIL & USERS", "", "", "$1,319.75"),
    ],
    col_widths=[2.5, 0.6, 0.9, 1.0],
    total_row_indices=[4],
)

# Virtual Staff section
doc.add_paragraph().paragraph_format.space_after = Pt(6)
section_header("Virtual Staff \u2014 Contracted Support (May)")
body_text(
    "All Virtual Staff roles receive a 15% rate reduction under the new MSA. "
    "Hours remain based on current utilization. Virtual Staff services are billed on a "
    "12-month billing cycle with weekly tracking invoices for full transparency."
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Role", "Hours/Mo", "Old Rate", "New Rate", "Monthly"],
    [
        ("Systems Architect (US)", "5.00", "$200.00/hr", "$170.00/hr", "$850.00"),
        ("USA Tech Normal", "15.26", "$125.00/hr", "$106.25/hr", "$1,621.38"),
        ("India Tech Normal", "58.13", "$15.00/hr", "$12.75/hr", "$741.16"),
        ("India Tech After Hours", "42.82", "$30.00/hr", "$25.50/hr", "$1,091.91"),
        ("SUBTOTAL \u2014 VIRTUAL STAFF", "121.21", "", "", "$4,304.44"),
    ],
    col_widths=[1.6, 0.7, 0.9, 0.9, 1.0],
    total_row_indices=[4],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run_styled(p, "Billing Model: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "12-month cycle-based billing with reconciliation. Initial cycle based on estimated "
    "hours per the Service Order. Actual usage tracked from day one with weekly in-contract "
    "invoices showing ticket-level detail. Reconciliation occurs at the end of each 12-month cycle.",
    size=10, color=BRAND_GREY,
)

# Unpaid Hours Recovery Plan
doc.add_paragraph().paragraph_format.space_after = Pt(8)
section_header("Unpaid Hours Recovery Plan")

body_text(
    "The current under-contract period has accumulated unpaid hours (actual hours consumed "
    "exceeding billed hours) across Virtual Staff roles. Under the new MSA, Technijian will "
    "implement a structured recovery plan to reduce these balances over the 12-month billing cycle."
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Role", "Current Unpaid Hours", "New Rate", "Value at New Rate"],
    [
        ("India Tech Normal", "573.81 hrs", "$12.75/hr", "$7,316.08"),
        ("India Tech After Hours", "290.23 hrs", "$25.50/hr", "$7,400.87"),
        ("USA Tech Normal", "143.26 hrs", "$106.25/hr", "$15,221.38"),
        ("Systems Architect (US)", "0.00 hrs", "$170.00/hr", "$0.00"),
        ("TOTAL UNPAID BALANCE", "1,007.30 hrs", "", "$29,938.33"),
    ],
    col_widths=[1.4, 1.0, 0.8, 1.0],
    total_row_indices=[4],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run_styled(p, "Recovery Strategy: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "Technijian will target utilizing approximately 20% fewer hours than billed each month. "
    "This means actual work will be managed to stay below the monthly billed amount, allowing "
    "the unpaid balance to decrease each month. Over the 12-month billing cycle, this disciplined "
    "approach is designed to bring the unpaid hours balance down significantly.",
    size=10, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run_styled(p, "How It Works: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "Each month, BWH is billed a fixed amount based on the 12-month cycle average. "
    "If Technijian uses fewer hours than billed (targeting 20% under), the running balance "
    "decreases. Weekly in-contract invoices will show the current balance so both parties "
    "can track progress. At the end of the 12-month cycle, reconciliation will reflect "
    "the reduced balance.",
    size=10, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run_styled(p, "12-Month Cycle Rationale: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "The current 12-month billing cycle is maintained specifically to provide a longer "
    "runway for the unpaid hours to come down. A shorter cycle would require immediate "
    "reconciliation of the balance, whereas the 12-month cycle allows gradual recovery "
    "while maintaining consistent monthly billing for BWH.",
    size=10, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run_styled(p, "Projected Recovery (20% Under Target): ", size=10, bold=True, color=CORE_BLUE)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Role", "Billed/Mo", "Target Actual (80%)", "Monthly Paydown", "12-Mo Paydown"],
    [
        ("India Tech Normal", "58.13 hrs", "46.50 hrs", "11.63 hrs", "139.56 hrs"),
        ("India Tech After Hours", "42.82 hrs", "34.26 hrs", "8.56 hrs", "102.72 hrs"),
        ("USA Tech Normal", "15.26 hrs", "12.21 hrs", "3.05 hrs", "36.60 hrs"),
        ("TOTAL MONTHLY PAYDOWN", "116.21 hrs", "92.97 hrs", "23.24 hrs", "278.88 hrs"),
    ],
    col_widths=[1.3, 0.8, 0.9, 0.9, 0.9],
    total_row_indices=[3],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run_styled(
    p,
    "At the 20% under-utilization target, approximately 278.88 unpaid hours will be recovered "
    "over the 12-month cycle, reducing the balance from 1,007.30 hours to approximately 728.42 hours. "
    "Technijian commits to transparency through weekly invoices showing the current running balance "
    "for each role.",
    size=10, color=BRAND_GREY, italic=True,
)


# Total Investment Summary
doc.add_paragraph().paragraph_format.space_after = Pt(8)
section_header("Total Monthly Investment Summary (May 2026)")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Monthly"],
    [
        ("Online Services", "$4,045.25"),
        ("Schedule B Subscriptions", "$7.20"),
        ("Virtual Staff Support", "$4,304.44"),
        ("TOTAL MONTHLY (NEW MSA)", "$8,356.89"),
        ("", ""),
        ("Previous Combined (Old T&C)", "$10,275.75"),
        ("MONTHLY SAVINGS", "$1,918.86 (18.7%)"),
    ],
    col_widths=[3.5, 2.5],
    total_row_indices=[3, 6],
)


# ================================================================
# PAGE 6 - WHY EACH SERVICE
# ================================================================
doc.add_page_break()

section_header("Why Each Service")
body_text(
    f"Every component in this agreement addresses a specific risk or operational need for "
    f"{CLIENT_NAME}\u2019s environment. Below is the justification for each service, tailored "
    f"to the residential real estate and homebuilder industry."
)

services_justification = [
    ("CrowdStrike Falcon EDR (AV Protection)",
     "Next-generation endpoint detection and response deployed on all managed desktops and servers.",
     f"With {MAY_DESKTOPS} desktops and 12 servers handling buyer contracts, architectural plans, "
     f"financial records, and construction schedules, {CLIENT_NAME} needs real-time threat detection "
     "that goes beyond traditional antivirus. CrowdStrike uses AI-driven behavioral analysis to detect "
     "ransomware, fileless attacks, and zero-day exploits. Homebuilders are increasingly targeted by "
     "ransomware groups due to the volume of personal and financial data processed \u2014 buyer SSNs, "
     "bank details, escrow records, and construction loan documents."),

    ("Huntress Managed Detection (AVH Protection)",
     "24/7 SOC-backed threat hunting and persistent foothold detection with human-led analysis.",
     "CrowdStrike stops known and behavioral threats in real time, but sophisticated attackers "
     "establish persistence mechanisms that evade automated tools. Huntress provides a dedicated "
     "Security Operations Center that actively hunts for footholds, backdoors, and living-off-the-land "
     "techniques. This layered approach is critical for an environment handling buyer PII, escrow "
     "documents, and high-value real estate transactions."),

    ("Cisco Umbrella (My Secure Internet)",
     "DNS-layer security that blocks malicious domains, phishing sites, and command-and-control "
     "callbacks before a connection is ever established.",
     "Most cyberattacks begin with a DNS query \u2014 whether a phishing link in an email, a "
     "malicious ad, or a compromised website. Cisco Umbrella intercepts these requests at the "
     f"DNS layer, blocking threats before they reach the endpoint. With {MAY_DESKTOPS} desktops and "
     "12 servers, this provides consistent protection across office workstations and construction "
     "site laptops accessing cloud applications, project management tools, and financial systems."),

    ("Patch Management (ManageEngine Endpoint Central Plus)",
     "Automated OS and third-party application patching for all desktops and servers.",
     f"Unpatched vulnerabilities are the #1 attack vector for ransomware and malware. With "
     f"{MAY_DESKTOPS} desktops and 12 servers running Windows plus construction and real estate "
     "applications, keeping software current is critical. Endpoint Central Plus provides automated "
     "vulnerability scanning, supports 850+ third-party application updates, and includes software "
     "deployment and configuration management from a single console."),

    ("Image Backup (Veeam)",
     "Server-level image backup for all 12 managed servers with 13 TB of backup storage.",
     "Construction project files, architectural plans, contracts, and financial records are "
     "business-critical data that must be recoverable in the event of ransomware, hardware failure, "
     "or accidental deletion. Veeam provides image-level backup with rapid recovery capabilities, "
     "ensuring business continuity for active development projects and escrow transactions."),

    ("Anti-Spam Standard (Email Security)",
     "Email filtering and encryption for 83 mailboxes, protecting against phishing, BEC, and spam.",
     "Email is the primary attack vector for business email compromise, and with 83 users "
     "on Microsoft 365, every mailbox is a potential entry point. Real estate transactions involve "
     "frequent email exchanges with buyers, title companies, lenders, and subcontractors \u2014 "
     "making wire transfer fraud and vendor impersonation particularly dangerous. Anti-spam filtering "
     "catches fraudulent wire instructions, fake lien releases, and impersonation emails."),

    ("Phishing Training",
     "Security awareness training for 85 users with simulated phishing campaigns.",
     "Employees at construction sites and in the office regularly receive emails that appear to come "
     "from subcontractors, material suppliers, title companies, and buyers. Phishing training builds "
     "a human firewall by teaching staff to recognize and report social engineering attempts before "
     "they result in credential theft or unauthorized wire transfers."),

    ("Veeam 365 Backup",
     "Cloud backup for 110 Microsoft 365 mailboxes, SharePoint, and OneDrive.",
     "Microsoft 365 does not provide comprehensive backup by default. If a user accidentally deletes "
     "an email thread containing contract negotiations, or if an account is compromised and data is "
     "purged, Veeam 365 provides independent backup and point-in-time recovery for all mailbox, "
     "SharePoint, and OneDrive content \u2014 critical for maintaining project records and legal "
     "documentation throughout the construction lifecycle."),

    ("DMARC/DKIM Monitoring",
     f"Email authentication monitoring to prevent domain spoofing of {CLIENT_NAME}\u2019s domain.",
     "Without DMARC and DKIM enforcement, attackers can send emails that appear to come from "
     f"{CLIENT_NAME}\u2019s domain \u2014 targeting buyers, subcontractors, and financial partners "
     "with fraudulent invoices or wire instructions. Ongoing monitoring ensures email authentication "
     "policies are correctly configured and alerts on spoofing attempts."),

    ("Monthly Site Assessment",
     f"Comprehensive security and compliance assessment for {CLIENT_NAME}\u2019s infrastructure.",
     "Monthly automated deep-scan assessments covering network vulnerabilities, Active Directory "
     "health, M365 Secure Score analysis, and external attack surface exposure. Produces executive-ready "
     "risk scorecards that track security posture over time \u2014 providing leadership with clear "
     "visibility into risk trends across all locations."),

    ("Virtual Staff Support",
     "121.21 hours/month of blended US and offshore technical support across four specialized roles.",
     "With 32 desktops, 12 servers, and a complex infrastructure spanning office and construction "
     "environments, ongoing support demand is significant. The blended model provides a US-based "
     "Systems Architect for strategic planning, US technicians for complex escalations, and India-based "
     "staff for cost-effective daytime coverage and after-hours support \u2014 ensuring issues are "
     "resolved around the clock without waiting for the next business day."),
]

for svc_name, svc_desc, svc_why in services_justification:
    p = doc.add_paragraph()
    run_styled(p, svc_name, size=11, bold=True, color=CORE_BLUE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run_styled(p, svc_desc, size=10, italic=True, color=DARK_CHARCOAL)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run_styled(p, "Why: ", size=10, bold=True, color=CORE_ORANGE)
    run_styled(p, svc_why, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(6)


# ================================================================
# PAGE 7 - RATE CARD + SLA + AGREEMENT STRUCTURE + NEXT STEPS
# ================================================================
doc.add_page_break()

section_header("Full Rate Card (Schedule C)")
body_text(
    "The following rates apply to all services provided under the new MSA. "
    "BWH Contracted Rates reflect the negotiated 15% discount for Virtual Staff services."
)

# US Staff Rates
doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph()
run_styled(p, "United States \u2014 Based Staff", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)
styled_table(
    ["Role", "Standard Rate", "After-Hours", "BWH Contracted", "Description"],
    [
        ("Systems Architect", "$250/hr", "$350/hr", "$170.00/hr", "Strategic architecture and design"),
        ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr", "Technology leadership and advisory"),
        ("Developer", "$150/hr", "N/A", "$125/hr", "Software development and engineering"),
        ("Tech Support", "$150/hr", "$250/hr", "$106.25/hr", "Technical support and sysadmin"),
    ],
    col_widths=[1.2, 0.8, 0.8, 0.9, 2.0],
)

# Offshore Staff Rates
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
run_styled(p, "Offshore \u2014 Based Staff", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)
styled_table(
    ["Role", "Standard Rate", "After-Hours", "BWH Contracted", "Description"],
    [
        ("Developer", "$45/hr", "N/A", "$30/hr", "Software development and engineering"),
        ("SEO Specialist", "$45/hr", "N/A", "$30/hr", "SEO and digital marketing"),
        ("Tech Support (Normal)", "$15/hr", "$30/hr", "$12.75/hr", "India business hours support"),
        ("Tech Support (After Hrs)", "$30/hr", "N/A", "$25.50/hr", "India night shift (US biz hours)"),
    ],
    col_widths=[1.4, 0.8, 0.8, 0.9, 1.8],
)

# Project Rates
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
run_styled(p, "Project & Ad Hoc Rates", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)
styled_table(
    ["Service", "Rate", "Minimum", "Notes"],
    [
        ("On-Site Support (US)", "$150/hr", "2-hour minimum", "No trip charges"),
        ("Remote Support (ad hoc)", "$150/hr", "15-min increments", "Non-contracted work"),
        ("Emergency / Critical", "$250/hr", "1-hour minimum", "After-hours and critical incidents"),
        ("Project Management", "$150/hr", "N/A", "For SOW-based engagements"),
    ],
    col_widths=[2.0, 0.8, 1.1, 2.0],
)

# Rate definitions
doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph()
run_styled(p, "Normal Business Hours: ", size=9, bold=True, color=DARK_CHARCOAL)
run_styled(p, "Monday through Friday, 8:00 AM to 6:00 PM Pacific Time, excluding US federal holidays.", size=9, color=BRAND_GREY)
p = doc.add_paragraph()
run_styled(p, "After-Hours: ", size=9, bold=True, color=DARK_CHARCOAL)
run_styled(p, "All hours outside of Normal Business Hours, including weekends and US federal holidays.", size=9, color=BRAND_GREY)
p = doc.add_paragraph()
run_styled(p, "BWH Contracted Rate: ", size=9, bold=True, color=DARK_CHARCOAL)
run_styled(p, "15% discounted rate under the BWH MSA, applied to Virtual Staff services billed through the 12-month Cycle-Based Billing Model (Schedule A, Part 3).", size=9, color=BRAND_GREY)


# SLA
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Service Level Commitments")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service Level", "Target"],
    [
        ("Infrastructure Uptime", "99.9% for monitored services"),
        ("Critical Incident Response", "Within 1 hour of notification"),
        ("Standard Support Response", "Within 4 business hours"),
        ("Scheduled Maintenance", "Tuesday evenings and Saturdays (with advance notice)"),
        ("24/7 Monitoring", "All desktops, servers, and email security continuously monitored"),
        ("Monthly Reporting", "Service reports summarizing incidents, uptime, and support activity"),
        ("Quarterly Reviews", f"Scheduled service reviews with {CLIENT_NAME}\u2019s designated representative"),
    ],
    col_widths=[2.2, 3.8],
)

# Agreement Structure
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Agreement Structure")
agreements = [
    ("Master Service Agreement (MSA): ",
     "Governs the overall relationship, payment terms, confidentiality, liability, and dispute "
     "resolution for a 12-month initial term with automatic annual renewal. Replaces the existing T&C."),
    ("Schedule A \u2014 Monthly Services: ",
     "Details the managed desktop security, server management, email security, monitoring, "
     "and virtual staff services with SLA commitments."),
    ("Schedule B \u2014 Subscription Services: ",
     "Covers any third-party software subscriptions and licenses managed by Technijian on "
     f"{CLIENT_NAME}\u2019s behalf."),
    ("Schedule C \u2014 Rate Card: ",
     "Establishes hourly rates for all labor, including the BWH contracted rates with 15% discount, "
     "on-site support, emergency response, and ad hoc services."),
]
for bold_text, desc in agreements:
    bullet_item(desc, bold_prefix=bold_text)

# Next Steps
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Next Steps")
steps = [
    "Review this executive summary and the attached MSA",
    "Confirm any questions with your Technijian account representative",
    "Sign the MSA, Schedule A, Schedule B, and Schedule C",
    "New rates take effect May 1, 2026",
    "Weekly tracking invoices begin immediately under the new billing structure",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run_styled(p, f"{i}. ", size=10, bold=True, color=CORE_BLUE)
    run_styled(p, step, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)

# Footer divider
doc.add_paragraph().paragraph_format.space_after = Pt(8)
accent_bar(HEX_BLUE, height_pt=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run_styled(
    p,
    "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(
    p,
    f"CONFIDENTIAL \u2014 Prepared exclusively for {CLIENT_NAME}",
    size=8, italic=True, color=BRAND_GREY,
)

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "02_MSA", "BWH-Executive-Summary.docx")
doc.save(OUTPUT)
print(f"Created {OUTPUT}")
