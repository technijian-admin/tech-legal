"""
Brandywine Homes (BWH) — Complete MSA + Schedules A/B/C Generator
Single DOCX: Cover + MSA (Sections 1-11) + Schedule A + Schedule B + Schedule C
Uses Technijian Brand Guide 2026 formatting (BrandedDoc pattern from OKL)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors (Technijian Brand Guide 2026) ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "technijian-logo.png")

CLIENT_NAME = "Brandywine Homes"
CLIENT_SHORT = "BWH"
CLIENT_CODE = "BWH"


def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


def set_grey_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        return r

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_borders(tbl)

    def cover_page(self, title, subtitle, date="May 1, 2026"):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(LOGO_PATH, width=Inches(3.0))
        # Orange divider
        div = self.doc.add_table(rows=1, cols=3)
        div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in div.rows[0].cells:
            c.text = ""
            c.paragraphs[0].paragraph_format.space_before = Pt(0)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div.rows[0].cells[1], HEX_ORANGE)
        div.rows[0].cells[0].width = Inches(2.25)
        div.rows[0].cells[1].width = Inches(2.5)
        div.rows[0].cells[2].width = Inches(2.25)
        remove_borders(div)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Prepared for ", size=12, color=BRAND_GREY)
        self.run(p, CLIENT_NAME, size=12, bold=True, color=CORE_BLUE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.accent_bar(HEX_ORANGE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar = tbl.rows[0].cells[0]
        set_cell_shading(bar, HEX_BLUE)
        bar.width = Inches(0.08)
        bar.text = ""
        bar.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar.paragraphs[0].paragraph_format.space_after = Pt(0)
        tc = tbl.rows[0].cells[1]
        tc.text = ""
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_borders(tbl)

    def part_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_borders(tbl)

    def styled_table(self, headers, rows, total_indices=None):
        total_indices = total_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            set_cell_shading(c, HEX_BLUE)
            c.text = ""
            self.run(c.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for ri, rd in enumerate(rows):
            row = tbl.add_row()
            is_total = ri in total_indices
            for i, val in enumerate(rd):
                c = row.cells[i]
                c.text = ""
                if is_total:
                    set_cell_shading(c, HEX_OFF_WHITE)
                    self.run(c.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(c.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        set_grey_borders(tbl)
        return tbl

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            self.run(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, num, text, indent=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent
        self.run(p, f"{prefix}{num} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self, client_name=CLIENT_NAME.upper()):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)
        # Technijian signature block
        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label, tag in [("By: ___________________________________", "/tSign/"),
                           ("Name: _________________________________", "/tName/"),
                           ("Title: _________________________________", "/tTitle/"),
                           ("Date: _________________________________", "/tDate/")]:
            p = self.doc.add_paragraph()
            self._sig_line(p, label, tag)
            p.paragraph_format.space_after = Pt(2)
        self.spacer(12)
        # Client signature block
        p = self.doc.add_paragraph()
        self.run(p, client_name, size=11, bold=True, color=DARK_CHARCOAL)
        for label, tag in [("By: ___________________________________", "/cSign/"),
                           ("Name: _________________________________", "/cName/"),
                           ("Title: _________________________________", "/cTitle/"),
                           ("Date: _________________________________", "/cDate/")]:
            p = self.doc.add_paragraph()
            self._sig_line(p, label, tag)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "02_MSA", filename)
        self.doc.save(path)
        sz = os.path.getsize(path)
        print(f"Created {path}  ({sz:,} bytes)")

    # ── DocuSign hidden anchor for signature fields ──
    def _sig_line(self, p, label, anchor_tag):
        r = p.add_run(anchor_tag)
        r.font.size = Pt(1)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = FONT
        self.run(p, label, size=10, color=BRAND_GREY)

    # ── BWH-specific helpers (used by content below) ──
    def body_bold_prefix(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=9, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=9, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def info_box(self, title, text):
        """Shaded info/disclosure box (e.g. ARL notice)."""
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_OFF_WHITE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        self.run(p, title, size=10, bold=True, color=DARK_CHARCOAL)
        r = p.add_run("\n")
        r.font.size = Pt(4)
        self.run(p, text, size=9, color=BRAND_GREY)
        set_grey_borders(tbl)


# ════════════════════════════════════════════════════════════════════
#  PART 1 — COVER PAGE
# ════════════════════════════════════════════════════════════════════
def build_cover(d):
    d.cover_page(
        "Master Service Agreement",
        "Managed IT Services, Cybersecurity & Virtual Staff",
    )


# ════════════════════════════════════════════════════════════════════
#  PART 2 — MSA BODY (Sections 1-11)
# ════════════════════════════════════════════════════════════════════
def build_msa_body(d):
    d.page_break()
    d.part_header("MASTER SERVICE AGREEMENT")
    d.spacer()
    d.body("Agreement Number: MSA-BWH")
    d.body("Effective Date: May 1, 2026")
    d.spacer()
    d.body("This Master Service Agreement (\u201cAgreement\u201d) is entered into by and between:")
    d.spacer()
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", "")
    d.body("18 Technology Drive, Suite 141")
    d.body("Irvine, California 92618")
    d.spacer()
    d.body("and")
    d.spacer()
    d.body_bold(f"{CLIENT_NAME} (\u201cClient\u201d)", "")
    d.body("2355 Main St #220")
    d.body("Irvine, CA 92614")
    d.spacer()
    d.body("(collectively, the \u201cParties\u201d)")

    # RECITALS
    d.spacer()
    d.section_header("RECITALS")
    d.body("WHEREAS, Technijian provides managed IT services, cybersecurity, cloud infrastructure, telephony, and related technology solutions; and")
    d.body("WHEREAS, Client desires to engage Technijian to provide certain services as described in the Schedules attached hereto;")
    d.body("NOW, THEREFORE, for good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:")

    # ── SECTION 1 ──
    d.spacer()
    d.section_header("SECTION 1 \u2014 SCOPE OF SERVICES")

    d.numbered("1.01.", "Services. Technijian shall provide the services described in the Schedules attached to this Agreement, which are incorporated herein by reference:")
    d.bullet("Schedule A \u2014 Monthly Managed Services (Online Services, SIP Trunk, Virtual Staff)")
    d.bullet("Schedule B \u2014 Subscription and License Services")
    d.bullet("Schedule C \u2014 Rate Card")
    d.body("Additional services may be provided through Statements of Work (\u201cSOWs\u201d) executed under this Agreement. Each SOW shall be signed by authorized representatives of both Parties and shall reference this Agreement by number. Upon execution, each SOW is incorporated into and governed by this Agreement.")

    d.numbered("1.02.", "Standard of Care. Technijian shall perform all services in a professional and workmanlike manner, consistent with industry standards for managed IT service providers.")
    d.numbered("1.03.", "Service Level Agreement. The service levels applicable to the services are set forth in Schedule A. Technijian shall use commercially reasonable efforts to meet the service levels described therein.")

    d.numbered("1.04.", "Client Responsibilities. Client shall:")
    d.numbered("(a)", "Provide Technijian with reasonable access to Client\u2019s systems, facilities, and personnel as necessary for Technijian to perform the services;", indent=1)
    d.numbered("(b)", "Designate a primary point of contact for communications with Technijian;", indent=1)
    d.numbered("(c)", "Maintain current and accurate information regarding Client\u2019s systems and infrastructure;", indent=1)
    d.numbered("(d)", "Comply with all applicable laws and regulations in connection with its use of the services; and", indent=1)
    d.numbered("(e)", "Be solely responsible for the security and management of Client\u2019s account credentials and passwords.", indent=1)

    d.numbered("1.05.", "Independent Contractor. Technijian is an independent contractor. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the Parties.")
    d.numbered("1.06.", "Client Data. \u201cClient Data\u201d means any data, information, content, records, or files that belong to Client or Client\u2019s customers, employees, or agents, that are stored, processed, or transmitted using the services, including personal information as defined under the California Consumer Privacy Act.")
    d.numbered("1.07.", "Currency. All fees under this Agreement are stated and payable in United States Dollars (USD).")
    d.numbered("1.08.", "Order of Precedence. In the event of a conflict between this Agreement and any Schedule, SOW, or Service Order, the following order of precedence shall apply (highest to lowest): (a) the applicable SOW or Service Order (but only for the specific services described therein); (b) the applicable Schedule; (c) this Master Service Agreement. A more specific provision in a lower-priority document shall not be deemed to conflict with a general provision in a higher-priority document unless it expressly states an intent to override.")
    d.numbered("1.09.", "Representations and Warranties. Each Party represents and warrants that: (a) it has the legal power and authority to enter into this Agreement; (b) the execution of this Agreement has been duly authorized by all necessary corporate action; (c) this Agreement constitutes a valid and binding obligation enforceable against it in accordance with its terms; and (d) its performance under this Agreement will not violate any applicable law, regulation, or existing contractual obligation.")
    d.numbered("1.10.", "Subcontractors. Technijian may engage subcontractors, including offshore personnel, to perform services under this Agreement, provided that: (a) Technijian shall remain fully responsible for all services performed by its subcontractors; (b) all subcontractors shall be bound by confidentiality and data protection obligations at least as protective as those in this Agreement; and (c) Technijian shall be liable for the acts and omissions of its subcontractors as if they were Technijian\u2019s own.")

    # ── SECTION 2 ──
    d.spacer()
    d.section_header("SECTION 2 \u2014 TERM AND RENEWAL")

    d.numbered("2.01.", "Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the \u201cInitial Term\u201d).")

    # ARL Auto-Renewal Disclosure Box
    d.spacer()
    d.info_box(
        "AUTOMATIC RENEWAL NOTICE:",
        "This Agreement will automatically renew for successive twelve (12) month periods unless you cancel at least sixty (60) days before the end of the current term. You may cancel by sending written notice to Technijian at the address above or by email to contracts@technijian.com. Technijian will send a renewal reminder at least thirty (30) days before each renewal date."
    )
    d.spacer()

    d.numbered("2.02.", "Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a \u201cRenewal Term\u201d), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term. Technijian shall send Client a written renewal reminder at least thirty (30) days prior to each renewal date, which shall restate the auto-renewal terms and cancellation method.")
    d.numbered("2.03.", "Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party. If Client terminates for convenience during the Initial Term or any Renewal Term, Client shall pay an early termination fee equal to: (a) any unrecoverable third-party costs committed by Technijian on Client\u2019s behalf (including prepaid licenses, committed hosting, and contracted offshore resources); plus (b) a wind-down fee calculated as follows: 75% of the average monthly recurring fees for the three (3) months preceding the termination notice, multiplied by the number of months remaining in the current term, up to a maximum of three (3) months\u2019 average recurring fees. The early termination fee constitutes liquidated damages and represents a reasonable estimate of Technijian\u2019s anticipated damages from early termination, including committed capacity, staffing, and unrecoverable vendor obligations, and is not a penalty. The Parties acknowledge that actual damages from early termination would be difficult or impracticable to calculate at the time of contracting.")
    d.numbered("2.04.", "Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party:")
    d.numbered("(a)", "Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days after receiving written notice of the breach; or", indent=1)
    d.numbered("(b)", "Becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.", indent=1)

    d.numbered("2.05.", "Effect of Termination.")
    d.numbered("(a)", "Upon termination for any reason, all fees and charges for services rendered through the date of termination shall become immediately due and payable, including: (i) any remaining obligations for annual licenses and subscriptions procured on Client\u2019s behalf; (ii) for Virtual Staff services, any unpaid hour balance invoiced at the applicable Hourly Rate from the Rate Card (Schedule C) as set forth in Schedule A, Section 3.3(e); (iii) all accrued and unpaid late fees under Section 3.04; and (iv) any early termination fees as calculated under Section 2.03. Termination shall not relieve Client of any payment obligation that accrued prior to or as a result of termination.", indent=1)
    d.numbered("(b)", "Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, provided that Client has paid all amounts owed under this Agreement in full, including any accrued late fees and collection costs. If Client has any outstanding balance at the time of termination, Technijian may withhold transition assistance until payment is received, or condition transition assistance upon Client\u2019s execution of a payment plan acceptable to Technijian.", indent=1)
    d.numbered("(c)", "Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format. If Client has outstanding unpaid invoices at the time of termination, Technijian may require Client to execute a payment plan acceptable to Technijian as a condition of data return, but shall not withhold Client Data beyond sixty (60) days following termination regardless of payment status. Notwithstanding any other provision of this Agreement, Technijian shall not withhold Client Data to the extent such withholding would prevent Client from complying with applicable law, including data breach notification obligations under California Civil Code Section 1798.82, consumer rights requests under the CCPA, HIPAA obligations, or other regulatory requirements. Nothing in this subsection shall be construed to grant Technijian any ownership interest in Client Data, nor shall the return of Client Data relieve Client of any payment obligation.", indent=1)
    d.numbered("(d)", "The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 9.03 (Severability), Section 9.04 (Waiver), Section 9.05 (Assignment), Section 9.08 (Governing Law), Section 9.09 (Personnel Transition Fee), Section 10 (Data Protection), and Section 11 (Insurance, which shall remain in effect through the end of any transition assistance period under Section 2.05(b)), and any other provision that by its nature is intended to survive termination.", indent=1)

    d.page_break()

    # ── SECTION 3 ──
    d.section_header("SECTION 3 \u2014 PAYMENT")

    d.numbered("3.01.", "Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.")
    d.numbered("3.02.", "Invoice Types. Client may receive the following types of invoices from Technijian during the term of this Agreement. Each invoice will clearly identify its type, the applicable Schedule or SOW, and the billing period or delivery event.")
    d.numbered("(a)", "Monthly Service Invoice. Issued on the first business day of each month for recurring managed services under Schedule A (Online Services, infrastructure, monitoring, desktop/server management). Billed in advance for the upcoming month.", indent=1)
    d.numbered("(b)", "Monthly Recurring Subscription Invoice. Issued on the first business day of each month for subscription and license services under Schedule B (software licenses, SaaS subscriptions, SIP trunk services). Billed in advance for the upcoming month. Subscription quantities and pricing are as specified in the applicable Service Order.", indent=1)
    d.numbered("(c)", "Weekly In-Contract Invoice. Issued every Friday for Virtual Staff (contracted support) services performed under Schedule A, Part 3, during the preceding week (Monday through Friday). Each weekly in-contract invoice shall include: (i) a listing of each support ticket addressed during the period; (ii) the assigned resource, role, and hours spent per ticket; (iii) a description of the work performed per ticket; (iv) whether work was performed during normal business hours or after-hours; and (v) the current running balance for each contracted role. The weekly in-contract invoice is issued for transparency and tracking purposes; the actual billed amount is governed by the cycle-based billing model described in Schedule A, Section 3.3.", indent=1)
    d.numbered("(d)", "Weekly Out-of-Contract Invoice. Issued every Friday for labor services performed outside the scope of any active Schedule or SOW \u2014 including ad-hoc support requests, emergency work, and services performed under a SOW with hourly billing (such as CTO Advisory engagements). Each weekly out-of-contract invoice shall include: (i) a listing of each support ticket or task performed during the period; (ii) the assigned resource, role, and applicable hourly rate from the Rate Card (Schedule C); (iii) time entries with hours billed per activity (in 15-minute increments); (iv) whether work was performed during normal business hours or after-hours; and (v) the total hours and total amount for the week. Out-of-contract work is billed in arrears at the applicable rates from Schedule C.", indent=1)
    d.numbered("(e)", "Equipment and Materials Invoice. Issued upon delivery or procurement of hardware, software licenses (perpetual), or other tangible goods on Client\u2019s behalf. Each equipment invoice shall include: (i) item description, manufacturer, and model/part number; (ii) quantity and unit price; (iii) applicable sales tax; (iv) shipping and handling charges, if any; and (v) total amount due. Title to equipment shall not pass to Client until payment is received in full, as set forth in Section 3.09.", indent=1)
    d.numbered("(f)", "Project Milestone Invoice. Issued upon completion of a project milestone as defined in an applicable SOW. The milestone, deliverables, and invoiced amount are as specified in the payment schedule of the SOW. Milestone invoices are billed in arrears upon acceptance of the deliverables or deemed acceptance under the SOW\u2019s acceptance provisions.", indent=1)
    d.numbered("3.03.", "Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified in the applicable Schedule or SOW.")
    d.numbered("3.04.", "Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated as simple interest from the date payment was due until the date payment is received in full. Late fees are payable in addition to the outstanding principal amount. The Parties acknowledge that the late fee represents a reasonable estimate of Technijian\u2019s administrative costs and damages resulting from late payment, including cash-flow disruption, collection overhead, and the cost of carrying accounts receivable, and is not intended as a penalty.")
    d.numbered("3.05.", "Disputed Invoices.")
    d.numbered("(a)", "Weekly Invoices (In-Contract and Out-of-Contract). Because weekly invoices include detailed ticket descriptions and time entries, Client shall have thirty (30) days from the invoice date to review and dispute any portion of a weekly invoice. Client shall notify Technijian in writing, specifying the ticket number(s) and nature of the dispute with reasonable particularity. Undisputed tickets and time entries on the same invoice shall remain payable by the due date. Failure to provide a timely written dispute notice within the thirty (30) day period shall constitute acceptance of all tickets and time entries on the invoice.", indent=1)
    d.numbered("(b)", "All Other Invoices. For monthly service invoices, monthly recurring subscription invoices, equipment invoices, and project milestone invoices, Client shall notify Technijian in writing within fifteen (15) days of the invoice date if Client disputes any portion, specifying the nature and basis of the dispute with reasonable particularity. Client shall pay all undisputed amounts by the due date. Failure to provide a timely written dispute notice shall constitute acceptance of the invoice.", indent=1)
    d.numbered("(c)", "Resolution. The Parties shall work in good faith to resolve any invoice dispute within thirty (30) days of the dispute notice. If the dispute results in an adjustment, Technijian shall issue a credit memo or revised invoice within ten (10) business days of resolution.", indent=1)
    d.numbered("3.06.", "Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend services under the Schedule or SOW associated with the unpaid invoice until payment is received in full, including all accrued late fees. If Client fails to pay any undisputed invoice within sixty (60) days of the due date, Technijian may suspend all services under this Agreement and any related Schedules or SOWs. Recurring fees for suspended services shall continue to accrue for a period not to exceed thirty (30) days following the date of suspension, after which Technijian may terminate the affected Schedule, SOW, or this Agreement upon written notice. Technijian continues to maintain infrastructure, licenses, and reserved capacity during any suspension period, which justifies the continued accrual. Suspension of services shall not relieve Client of its payment obligations.")
    d.numbered("3.07.", "Acceleration. Upon the occurrence of any of the following events, all fees, charges, and amounts owing under this Agreement, all Schedules, and all SOWs shall become immediately due and payable without further notice or demand: (a) Client fails to pay any undisputed invoice within forty-five (45) days of the due date; (b) Client terminates this Agreement or any Schedule while any invoices remain unpaid; (c) Client becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets; or (d) Client is the subject of a material adverse change in its financial condition that, in Technijian\u2019s reasonable judgment, impairs Client\u2019s ability to perform its payment obligations.")
    d.numbered("3.08.", "Collection Costs and Attorney\u2019s Fees. This Section applies exclusively to the collection of fees, invoices, and other amounts owed under this Agreement and is separate from and does not apply to disputes regarding service quality, professional performance, errors and omissions, or any other non-payment claims. In any Collection Effort (as defined in Section 8.04), the prevailing Party shall be entitled to recover from the non-prevailing Party all reasonable costs of collection, including but not limited to: (a) reasonable attorney\u2019s fees and legal costs (including fees for in-house counsel at market rates); (b) collection agency fees and commissions; (c) court costs, arbitration filing fees, and administrative costs; (d) costs of investigation, skip tracing, and asset searches; and (e) all costs of appeal. This obligation applies regardless of whether a lawsuit or arbitration is commenced, and such costs shall be in addition to all other amounts owed. Pursuant to California Civil Code Section 1717, the Parties acknowledge that this attorney\u2019s fees provision is reciprocal and shall be enforced as such. For avoidance of doubt, this Section does not entitle either Party to recover attorney\u2019s fees or costs in connection with any counterclaim, cross-claim, or separate claim arising from alleged service deficiencies, professional negligence, or other non-payment matters \u2014 such claims are governed by Section 8.05.")
    d.numbered("3.09.", "Right of Setoff and Lien. (a) Technijian shall have the right to set off any amounts owed by Client under this Agreement against any amounts Technijian may owe to Client under this or any other agreement between the Parties. (b) Technijian shall retain a lien on all work product, deliverables, custom development, documentation, and materials (excluding Client Data as defined in Section 1.06) in its possession until all amounts owed by Client are paid in full. Technijian shall not be required to deliver, transfer, or release any work product or grant any license until all outstanding invoices, including accrued late fees and collection costs, are satisfied. (c) In the event of non-payment, Technijian may withhold transition assistance and credential transfers described in Section 2.05 until all amounts owed are paid in full, subject to the Client Data return obligations in Section 2.05(c) and the regulatory carve-outs therein.")
    d.numbered("3.10.", "Grant of Security Interest (UCC).")
    d.numbered("(a)", "Grant. To secure the full and timely payment of all fees, charges, late fees, collection costs, and any other amounts now or hereafter owing by Client to Technijian under this Agreement, any Schedule, any SOW, or any other agreement between the Parties (collectively, the \u201cSecured Obligations\u201d), Client hereby grants to Technijian a continuing security interest in the following property of Client, whether now owned or hereafter acquired (collectively, the \u201cCollateral\u201d): (i) all equipment, hardware, and fixtures procured by Technijian on Client\u2019s behalf or used in connection with the services; (ii) all work product, deliverables, custom development, and documentation produced by Technijian under this Agreement or any SOW; (iii) all proceeds of the foregoing; and (iv) all books and records relating to the foregoing. For avoidance of doubt, the Collateral does not include Client Data (as defined in Section 1.06), Client\u2019s pre-existing intellectual property, or Client\u2019s general accounts receivable or general intangibles unrelated to the services. This security interest shall be subordinate to any prior perfected security interest held by Client\u2019s primary lender(s) of record as of the Effective Date of this Agreement.", indent=1)
    d.numbered("(b)", "Filing of UCC-1 Financing Statement. Client authorizes Technijian to file a UCC-1 Financing Statement (and any amendments, continuations, or renewals thereof) with the California Secretary of State or any other applicable filing office to perfect the security interest granted herein. Technijian may file such UCC-1 Financing Statement at any time after execution of this Agreement; provided, however, that Technijian shall provide Client with ten (10) days written notice before filing, except that no prior notice shall be required if Client is more than forty-five (45) days past due on any undisputed invoice.", indent=1)
    d.numbered("(c)", "Client Cooperation. Client shall: (i) execute and deliver any financing statements, amendments, or other documents reasonably requested by Technijian to perfect, maintain, or enforce the security interest; (ii) not grant any security interest in the Collateral that would be senior to Technijian\u2019s security interest without Technijian\u2019s prior written consent; and (iii) promptly notify Technijian of any change in Client\u2019s legal name, state of organization, or organizational identification number.", indent=1)
    d.numbered("(d)", "Remedies Upon Default. If Client fails to pay any Secured Obligation within forty-five (45) days of the due date (a \u201cPayment Default\u201d), Technijian shall have, in addition to all other rights and remedies under this Agreement and applicable law, all the rights and remedies of a secured party under the California Uniform Commercial Code (Cal. Com. Code \u00a7 9101 et seq.), including the right to take possession of Collateral that is in Technijian\u2019s physical or constructive possession and to dispose of it in a commercially reasonable manner. Technijian shall provide Client with at least fifteen (15) days prior written notice before exercising any disposition remedy. All self-help remedies shall be exercised without breach of the peace as required by Cal. Com. Code \u00a7 9609. For avoidance of doubt, nothing in this Section authorizes Technijian to access Client\u2019s computer systems, networks, or accounts to enforce the security interest; enforcement against intangible Collateral shall be conducted through lawful judicial process or notification to account debtors pursuant to Cal. Com. Code \u00a7 9607.", indent=1)
    d.numbered("(e)", "Termination and Release. Within thirty (30) days after all Secured Obligations have been paid in full and this Agreement has been terminated or expired with no further obligations outstanding, Technijian shall, at its own expense, file a UCC-3 Termination Statement to release the security interest and provide Client with written confirmation of the release. If Technijian fails to file a termination statement within such period after Client\u2019s written request (and all Secured Obligations are paid in full), Client shall be entitled to file a UCC-3 Termination Statement on its own behalf, and Technijian hereby authorizes such filing.", indent=1)
    d.numbered("3.11.", "Credit Reporting and Collections. Technijian reserves the right to report delinquent accounts to commercial credit reporting agencies and to assign delinquent accounts to third-party collection agencies, in each case after sixty (60) days of non-payment and ten (10) days written notice to Client.")
    d.numbered("3.12.", "Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian\u2019s income.")

    # ── SECTION 4 ──
    d.spacer()
    d.section_header("SECTION 4 \u2014 CONFIDENTIALITY")

    d.numbered("4.01.", "Definition. \u201cConfidential Information\u201d means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information. If the Parties have executed a separate Non-Disclosure Agreement, its terms are incorporated herein by reference.")
    d.numbered("4.02.", "Obligations. Each Party shall:")
    d.numbered("(a)", "Hold the other Party\u2019s Confidential Information in confidence using at least the same degree of care it uses for its own confidential information, but not less than reasonable care;", indent=1)
    d.numbered("(b)", "Not disclose Confidential Information to third parties without prior written consent, except to employees, agents, and subcontractors who have a need to know and are bound by equivalent obligations; and", indent=1)
    d.numbered("(c)", "Not use Confidential Information for any purpose other than performing obligations under this Agreement.", indent=1)
    d.numbered("4.03.", "Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known to the receiving Party prior to disclosure, is independently developed, or is received from a third party without restriction.")
    d.numbered("4.04.", "Compelled Disclosure. If required by law or court order to disclose Confidential Information, the receiving Party shall provide prompt written notice to the disclosing Party (to the extent legally permitted) and cooperate in seeking a protective order.")
    d.numbered("4.05.", "Duration. Confidentiality obligations shall survive termination for a period of three (3) years.")

    # ── SECTION 5 ──
    d.spacer()
    d.section_header("SECTION 5 \u2014 LIMITATION OF LIABILITY")

    d.numbered("5.01.", "Limitation. EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, NEITHER PARTY\u2019S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM (THE \u201cSTANDARD CAP\u201d).")
    d.numbered("5.02.", "Exclusion of Consequential Damages. EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL, REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE OR WHETHER EITHER PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.")
    d.numbered("5.03.", "Enhanced Cap for Certain Claims. FOR CLAIMS ARISING FROM BREACHES OF SECTION 4 (CONFIDENTIALITY), SECTION 10 (DATA PROTECTION), INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, THE TOTAL AGGREGATE LIABILITY OF THE RESPONSIBLE PARTY SHALL NOT EXCEED THREE (3) TIMES THE STANDARD CAP DEFINED IN SECTION 5.01 (THE \u201cENHANCED CAP\u201d). THE ENHANCED CAP SHALL NOT APPLY TO LIABILITY ARISING FROM A PARTY\u2019S WILLFUL AND INTENTIONAL MISAPPROPRIATION OF THE OTHER PARTY\u2019S CONFIDENTIAL INFORMATION OR CLIENT DATA, FOR WHICH LIABILITY SHALL BE UNCAPPED.")
    d.numbered("5.04.", "Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data in its possession, Client acknowledges that: (a) Client is solely responsible for maintaining backup copies of its data; (b) Technijian\u2019s liability for data loss shall be limited to using commercially reasonable efforts to restore data from available backups; and (c) Technijian shall not be liable for data loss caused by Client\u2019s actions, third-party attacks that Technijian could not have reasonably prevented through the safeguards required under Section 10.03, or events beyond Technijian\u2019s reasonable control.")

    d.page_break()

    # ── SECTION 6 ──
    d.section_header("SECTION 6 \u2014 INDEMNIFICATION")

    d.numbered("6.01.", "By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from: (a) Technijian\u2019s gross negligence or willful misconduct in performing the services; or (b) any claim that Technijian IP (as defined in Section 7.01) used in performing the services infringes such third party\u2019s United States intellectual property rights, provided the infringement does not arise from Client\u2019s modifications to the Technijian IP, Client\u2019s combination of the Technijian IP with non-Technijian materials, or Client\u2019s use of the services in a manner not authorized by this Agreement. If any Technijian IP becomes the subject of an infringement claim, Technijian may, at its option and expense: (i) procure for Client the right to continue using the affected IP; (ii) modify the affected IP to make it non-infringing while maintaining substantially equivalent functionality; or (iii) replace the affected IP with a non-infringing alternative. If none of these options are commercially practicable, either Party may terminate the affected services upon thirty (30) days\u2019 notice, and Technijian shall refund any prepaid fees for the terminated services covering the period after termination.")
    d.numbered("6.02.", "By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from: (a) Client\u2019s use of the services in violation of applicable law; (b) Client\u2019s breach of this Agreement; or (c) any data, content, or materials provided by Client.")
    d.numbered("6.03.", "Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party\u2019s prior written consent.")

    # ── SECTION 7 ──
    d.spacer()
    d.section_header("SECTION 7 \u2014 INTELLECTUAL PROPERTY")

    d.numbered("7.01.", "Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes used in providing the services (\u201cTechnijian IP\u201d). Client receives no rights to Technijian IP except as expressly set forth in this Agreement.")
    d.numbered("7.02.", "Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property (\u201cClient IP\u201d).")
    d.numbered("7.03.", "Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW. Unless otherwise specified, Technijian shall retain ownership of any general-purpose tools, frameworks, or methodologies developed during the engagement.")

    # ── SECTION 8 ──
    d.spacer()
    d.section_header("SECTION 8 \u2014 DISPUTE RESOLUTION")

    d.numbered("8.01.", "Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations between their respective designated representatives for a period of thirty (30) days.")
    d.numbered("8.02.", "Mediation. If the dispute is not resolved through negotiation, the Parties shall submit the dispute to mediation administered by a mutually agreed-upon mediator in Orange County, California, for a period not to exceed sixty (60) days.")
    d.numbered("8.03.", "Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association under its Commercial Arbitration Rules. The arbitration shall take place in Orange County, California, before a single arbitrator. The arbitrator shall have the authority to award any remedy that would be available in a court of competent jurisdiction, including injunctive relief and specific performance. The Parties shall equally share arbitrator fees and AAA administrative costs; each Party shall bear its own attorney\u2019s fees except as provided in Sections 8.04 and 8.05. The arbitrator shall issue a reasoned written award.")
    d.numbered("8.04.", "Fees \u2014 Payment Collection Actions. Notwithstanding Section 8.05, in any dispute, arbitration, mediation, litigation, or other proceeding in which the primary relief sought is the collection of fees, invoices, or other amounts owed under this Agreement, any Schedule, or any SOW (a \u201cCollection Action\u201d), the prevailing Party shall be entitled to recover from the non-prevailing Party all reasonable costs and expenses incurred in connection with such Collection Action, including but not limited to: (a) attorney\u2019s fees and legal costs (including fees for in-house counsel at market rates); (b) expert witness fees; (c) arbitration and mediation filing fees, administrative costs, and arbitrator compensation; (d) court costs and filing fees; (e) costs of discovery, depositions, and document production; and (f) all costs of appeal. For purposes of this Section, a \u201cCollection Action\u201d includes any proceeding initiated to recover unpaid invoices, late fees, accelerated amounts under Section 3.07, or amounts due upon termination under Section 2.05, as well as any action to enforce a judgment or arbitration award arising from such a proceeding. The \u201cprevailing Party\u201d means the Party that substantially obtains the relief sought, whether by settlement, judgment, or award, as determined by the arbitrator or court. This provision is in addition to and does not limit the pre-litigation collection costs recoverable under Section 3.08.")
    d.numbered("8.05.", "Fees \u2014 All Other Disputes. Except as expressly provided in Section 8.04 (Collection Actions) and Section 3.08 (Collection Costs), in any dispute, arbitration, mediation, litigation, or other proceeding arising out of or relating to this Agreement \u2014 including but not limited to claims relating to service quality, professional negligence, errors and omissions, breach of warranty, indemnification, data protection, or any counterclaim asserted in response to a Collection Action \u2014 each Party shall bear its own attorney\u2019s fees and costs. Neither Party shall be entitled to recover attorney\u2019s fees or costs from the other Party in connection with such non-collection disputes, regardless of which Party prevails. For avoidance of doubt, if a Collection Action under Section 8.04 and a non-collection dispute under this Section 8.05 are joined or heard in the same proceeding, the arbitrator or court shall apportion fees accordingly, awarding fees to the prevailing Party only with respect to the Collection Action claims.")
    d.numbered("8.06.", "Injunctive and Provisional Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief in a court of competent jurisdiction to prevent irreparable harm. Either Party may seek provisional remedies from a court of competent jurisdiction pursuant to California Code of Civil Procedure Section 1281.8 pending appointment of the arbitrator or resolution of the arbitration, without waiving the right to arbitrate.")

    # ── SECTION 9 ──
    d.spacer()
    d.section_header("SECTION 9 \u2014 GENERAL PROVISIONS")

    d.numbered("9.01.", "Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties and supersedes all prior agreements, whether written or oral, relating to the subject matter hereof.")
    d.numbered("9.02.", "Amendment. This Agreement may only be amended by a written instrument signed by both Parties. Technijian may update its Rate Card (Schedule C) upon sixty (60) days written notice to Client, effective at the start of the next Renewal Term.")
    d.numbered("9.03.", "Severability. If any provision is found to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.")
    d.numbered("9.04.", "Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party. A waiver of any breach shall not constitute a waiver of any subsequent breach.")
    d.numbered("9.05.", "Assignment. Neither Party may assign this Agreement without the prior written consent of the other Party, except that either Party may assign this Agreement in connection with a merger, acquisition, or sale of substantially all of its assets.")
    d.numbered("9.06.", "Force Majeure.")
    d.numbered("(a)", "Neither Party shall be liable for delays or failures in performance caused by events beyond its reasonable control, including natural disasters, acts of government, labor disputes, pandemics, cyberattacks on critical infrastructure, or failures of major third-party infrastructure services (such as cloud platform outages) that are not attributable to the affected Party\u2019s failure to implement reasonable redundancy (\u201cForce Majeure Event\u201d).", indent=1)
    d.numbered("(b)", "The affected Party shall notify the other Party in writing within five (5) business days of becoming aware of a Force Majeure Event and shall use commercially reasonable efforts to mitigate the impact and resume performance.", indent=1)
    d.numbered("(c)", "Payment obligations are not excused by a Force Majeure Event.", indent=1)
    d.numbered("(d)", "If a Force Majeure Event prevents performance of a material portion of the services for more than ninety (90) consecutive days, either Party may terminate the affected Schedule, SOW, or this Agreement upon fifteen (15) days written notice without liability for early termination fees, and Client shall pay only for services actually rendered through the date of termination.", indent=1)
    d.numbered("9.07.", "Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier to the addresses set forth above (or as updated in writing).")
    d.numbered("9.08.", "Governing Law. This Agreement shall be governed by and construed in accordance with the laws of the State of California without regard to conflict of law principles.")
    d.numbered("9.09.", "Personnel Transition Fee. The Parties acknowledge that each invests significant resources in recruiting, training, and retaining skilled personnel. If either Party hires (whether as an employee or independent contractor) any individual who was an employee of the other Party and who was directly involved in performing or receiving services under this Agreement, and such hiring occurs during the term of this Agreement or within twelve (12) months following termination, the hiring Party shall pay the other Party a personnel transition fee equal to 25% of the hired individual\u2019s first-year annual compensation (base salary or annualized contractor fees). This fee represents a reasonable estimate of the non-hiring Party\u2019s recruiting and training costs and is not intended as a restraint on trade or employment. This Section does not restrict any individual\u2019s right to seek or obtain employment, and shall not apply to individuals who: (a) respond to general public job postings or advertisements not specifically targeted at the other Party\u2019s employees; or (b) are referred by a third-party recruiting firm without the hiring Party\u2019s direction to target the other Party\u2019s employees.")
    d.numbered("9.10.", "Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.")

    # ── SECTION 10 ──
    d.spacer()
    d.section_header("SECTION 10 \u2014 DATA PROTECTION")

    d.numbered("10.01.", "CCPA/CPRA Compliance. To the extent Technijian processes, stores, or has access to personal information (as defined under the California Consumer Privacy Act, as amended by the California Privacy Rights Act, Cal. Civ. Code \u00a7 1798.100 et seq., collectively \u201cCCPA\u201d) on behalf of Client, Technijian acts as a \u201cservice provider\u201d as defined under Cal. Civ. Code \u00a7 1798.140(ag) and shall:")
    d.numbered("(a)", "Process such personal information only as necessary to perform the services and in accordance with Client\u2019s documented instructions and this Agreement;", indent=1)
    d.numbered("(b)", "Not sell, share, retain, use, or disclose personal information for any purpose other than performing the services, including not using personal information for targeted advertising or cross-context behavioral advertising;", indent=1)
    d.numbered("(c)", "Not combine personal information received from Client with personal information received from other sources or collected from Technijian\u2019s own interactions with individuals, except as expressly permitted by the CCPA to perform the services;", indent=1)
    d.numbered("(d)", "Implement reasonable security measures appropriate to the nature of the personal information, consistent with the requirements of Section 10.03;", indent=1)
    d.numbered("(e)", "Cooperate with Client in responding to verifiable consumer rights requests under the CCPA (including access, deletion, correction, and opt-out requests) within ten (10) business days of Client\u2019s request;", indent=1)
    d.numbered("(f)", "Notify Client within five (5) business days if Technijian determines that it can no longer meet its obligations as a service provider under the CCPA;", indent=1)
    d.numbered("(g)", "Ensure that all subcontractors (including offshore personnel engaged under Section 1.10) who process personal information on behalf of Client are bound by written agreements containing data protection obligations at least as protective as those in this Section 10, and Technijian shall remain liable for the acts and omissions of its subcontractors with respect to personal information;", indent=1)
    d.numbered("(h)", "Permit Client, upon thirty (30) days written notice and no more than once per twelve (12) month period, to audit or inspect Technijian\u2019s data processing practices to verify compliance with this Section 10, or, at Technijian\u2019s option, provide Client with a summary of a recent independent third-party audit (such as SOC 2 Type II) covering the relevant controls; and", indent=1)
    d.numbered("(i)", "Certify that Technijian understands the restrictions in this Section 10 and will comply with them, including that Technijian shall not sell or share personal information as those terms are defined under the CCPA.", indent=1)
    d.body("If the Parties require more detailed data processing terms (for example, to address GDPR, HIPAA, or industry-specific requirements), the Parties shall execute a separate Data Processing Addendum, which shall be incorporated into this Agreement by reference.")

    d.numbered("10.02.", "Security Incident Notification. If Technijian becomes aware of a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Client Data (\u201cSecurity Incident\u201d), Technijian shall: (a) notify Client in writing without unreasonable delay and in no event later than forty-eight (48) hours after becoming aware of the Security Incident; (b) provide Client with sufficient information to enable Client to comply with its obligations under California Civil Code \u00a7 1798.82 (data breach notification), including the categories and approximate number of records affected, the nature of the incident, and the measures taken or proposed to address it, and any other applicable data breach notification laws; (c) cooperate with Client\u2019s investigation of the Security Incident; and (d) take reasonable steps to contain and remediate the Security Incident. If Client Data includes protected health information subject to HIPAA, notification shall also comply with 45 CFR \u00a7 164.410.")
    d.numbered("10.03.", "Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards designed to protect Client Data from unauthorized access, use, or disclosure, consistent with industry standards for managed IT service providers. Such safeguards shall include, at a minimum: (a) encryption of Client Data in transit and at rest; (b) access controls limiting access to authorized personnel; (c) regular security assessments and vulnerability testing; and (d) employee security awareness training.")
    d.numbered("10.04.", "Regulatory Compliance. If Client is subject to the Health Insurance Portability and Accountability Act (\u201cHIPAA\u201d), the Payment Card Industry Data Security Standard (\u201cPCI DSS\u201d), the General Data Protection Regulation (\u201cGDPR\u201d), or other industry-specific data protection requirements, the Parties shall execute a separate addendum addressing the additional obligations applicable to the regulated data. Technijian shall cooperate with Client in implementing controls necessary to meet such requirements.")
    d.numbered("10.05.", "Data Return and Deletion. Subject to Section 2.05(c), upon termination of this Agreement or upon Client\u2019s written request, Technijian shall securely delete or return all Client Data in its possession within thirty (30) days, using methods consistent with NIST SP 800-88 (Guidelines for Media Sanitization) or equivalent standards, and shall certify such deletion in writing upon request. Technijian may retain copies of Client Data only to the extent required by applicable law, provided such retained data remains subject to the confidentiality and data protection obligations of this Agreement.")

    # ── SECTION 11 ──
    d.spacer()
    d.section_header("SECTION 11 \u2014 INSURANCE")

    d.numbered("11.01.", "Required Coverage. During the term of this Agreement, Technijian shall maintain the following insurance coverage with carriers rated A- VII or better by A.M. Best: (a) Commercial General Liability insurance on an occurrence basis with limits of not less than $1,000,000 per occurrence and $2,000,000 in the aggregate; (b) Professional Liability (Errors and Omissions) insurance on a claims-made basis with limits of not less than $1,000,000 per claim and $2,000,000 in the aggregate, with a retroactive date no later than the Effective Date of this Agreement; (c) Cyber Liability insurance on a claims-made basis with limits of not less than $1,000,000 per claim, covering data breaches, network security failures, and privacy violations; and (d) Workers\u2019 Compensation insurance as required by the laws of the State of California.")
    d.numbered("11.02.", "Additional Requirements. (a) Client shall be named as an additional insured on Technijian\u2019s Commercial General Liability policy with respect to the services provided under this Agreement. (b) Technijian\u2019s insurance shall be primary and non-contributory with respect to any insurance maintained by Client. (c) For claims-made policies (Professional Liability and Cyber Liability), Technijian shall maintain tail coverage (extended reporting period) for a minimum of two (2) years following termination of this Agreement. (d) Technijian shall include a waiver of subrogation in favor of Client on the Commercial General Liability and Workers\u2019 Compensation policies.")
    d.numbered("11.03.", "Certificates of Insurance. Upon Client\u2019s written request, Technijian shall provide certificates of insurance evidencing the coverage required under this Section, including evidence of additional insured status and waiver of subrogation. Technijian shall provide Client with at least thirty (30) days\u2019 prior written notice of any material change to or cancellation of such coverage.")

    # ── SIGNATURES ──
    d.signatures()
    d.spacer()
    d.body("Schedules:")
    d.bullet("Schedule A \u2014 Monthly Managed Services")
    d.bullet("Schedule B \u2014 Subscription and License Services")
    d.bullet("Schedule C \u2014 Rate Card")
    d.footer_bar()


# ════════════════════════════════════════════════════════════════════
#  PART 3 — SCHEDULE A: Monthly Managed Services
# ════════════════════════════════════════════════════════════════════
def build_schedule_a(d):
    d.page_break()
    d.part_header("SCHEDULE A \u2014 MONTHLY MANAGED SERVICES")
    d.body("Attached to Master Service Agreement MSA-BWH")
    d.body("Effective Date: May 1, 2026")
    d.body(
        "This Schedule describes the Monthly Managed Services provided by Technijian, Inc. "
        "(\u201cTechnijian\u201d) to Brandywine Homes (\u201cClient\u201d) under the Master Service Agreement."
    )
    d.spacer()

    # ── PART 1: ONLINE SERVICES ──
    d.part_header("PART 1 \u2014 ONLINE SERVICES")
    d.spacer()

    d.section_header("1.1 Description")
    d.body(
        "Online Services include managed infrastructure, security, monitoring, and related IT services "
        "delivered on a recurring monthly basis. Services are selected from the Technijian Price List "
        "and itemized in the Service Order below."
    )

    d.section_header("1.2 Service Categories")
    d.body(
        "Online Services are organized into the following categories. Specific services, quantities, "
        "and pricing are detailed in the Service Order attached to this Schedule."
    )
    d.styled_table(
        ["Category", "Description"],
        [
            ("Desktop Management", "Patch management, antivirus (CrowdStrike, Huntress), DNS filtering, remote access, monitoring"),
            ("Server Management", "Patch management, image backup, antivirus (CrowdStrike, Huntress), remote access, monitoring"),
            ("Cloud Infrastructure", "Backup storage, Veeam management"),
            ("Network & Security", "Switch config backup, firewall traffic monitoring, WiFi monitoring, pen testing, storage monitoring"),
            ("Email & Compliance", "Site assessment, DMARC/DKIM, anti-spam, phishing training, Veeam 365, file sharing"),
        ],
    )

    d.spacer()
    d.section_header("1.3 Service Order")
    d.body(
        "The specific services, quantities, and monthly pricing for Client are detailed below. "
        "The Service Order may be updated by mutual written agreement."
    )

    # Desktop Management (32 desktops)
    d.body_bold("Desktop Management (32 desktops)", "")
    d.styled_table(
        ["Service", "Code", "Unit Price", "Qty", "Monthly"],
        [
            ("Patch Management", "PMW", "$4.00", "32", "$128.00"),
            ("My Secure Internet", "SI", "$6.00", "32", "$192.00"),
            ("My Remote", "MR", "$2.00", "32", "$64.00"),
            ("My Ops \u2014 Net", "OPS-NET", "$3.25", "32", "$104.00"),
            ("AVH Protection \u2014 Desktop (Huntress)", "AVMH", "$6.00", "32", "$192.00"),
            ("AV Protection \u2014 Desktop (CrowdStrike)", "AVD", "$8.50", "32", "$272.00"),
            ("Desktop Subtotal", "", "", "", "$952.00"),
        ],
        total_indices=[6],
    )

    d.spacer()
    # Server Management (12 servers)
    d.body_bold("Server Management (12 servers)", "")
    d.styled_table(
        ["Service", "Code", "Unit Price", "Qty", "Monthly"],
        [
            ("Patch Management", "PMW", "$4.00", "12", "$48.00"),
            ("My Secure Internet", "SI", "$6.00", "12", "$72.00"),
            ("My Remote", "MR", "$2.00", "12", "$24.00"),
            ("My Ops \u2014 Net", "OPS-NET", "$3.25", "12", "$39.00"),
            ("Image Backup (Veeam)", "IB", "$15.00", "12", "$180.00"),
            ("AVH Protection \u2014 Server (Huntress)", "AVHS", "$6.00", "12", "$72.00"),
            ("AV Protection \u2014 Server (CrowdStrike)", "AVS", "$10.50", "12", "$126.00"),
            ("Server Subtotal", "", "", "", "$561.00"),
        ],
        total_indices=[7],
    )

    d.spacer()
    # Cloud Infrastructure
    d.body_bold("Cloud Infrastructure", "")
    d.styled_table(
        ["Service", "Code", "Unit Price", "Qty", "Monthly"],
        [
            ("Backup Storage (TB)", "TB-BSTR", "$50.00", "13", "$650.00"),
            ("Veeam One", "VONE", "$3.00", "11", "$33.00"),
            ("Cloud Subtotal", "", "", "", "$683.00"),
        ],
        total_indices=[2],
    )

    d.spacer()
    # Network & Security
    d.body_bold("Network & Security", "")
    d.styled_table(
        ["Service", "Code", "Unit Price", "Qty", "Monthly"],
        [
            ("My Ops \u2014 Config (Switches)", "OPS-BKP", "$6.00", "3", "$18.00"),
            ("Real Time Pen Testing", "RTPT", "$7.00", "6", "$42.00"),
            ("My Ops \u2014 Traffic (Firewalls)", "OPS-TR", "$14.00", "1", "$14.00"),
            ("My Ops \u2014 Wifi (APs)", "OPS-WF", "$1.00", "6", "$6.00"),
            ("My Ops \u2014 Storage (Disks)", "OPS-ST", "$4.75", "2", "$9.50"),
            ("Sophos Firewall Subscription 2C-4G", "SO-2C4G", "$270.00", "1", "$270.00"),
            ("MPC Edge Appliance (16GB/8 cores/512GB)", "Edge-16M", "$100.00", "1", "$100.00"),
            ("Network Subtotal", "", "", "", "$459.50"),
        ],
        total_indices=[7],
    )

    d.spacer()
    # Email & Compliance
    d.body_bold("Email & Compliance", "")
    d.styled_table(
        ["Service", "Code", "Unit Price", "Qty", "Monthly"],
        [
            ("Site Assessment", "SA", "$50.00", "1", "$50.00"),
            ("DMARC/DKIM", "DKIM", "$20.00", "1", "$20.00"),
            ("Anti-Spam Standard", "ASA", "$6.25", "83", "$518.75"),
            ("Phishing Training", "PHT", "$6.00", "85", "$510.00"),
            ("Veeam 365", "V365", "$2.50", "110", "$275.00"),
            ("My Disk", "MDU", "$16.00", "1", "$16.00"),
            ("Email Subtotal", "", "", "", "$1,389.75"),
        ],
        total_indices=[6],
    )

    d.spacer()
    # Grand total
    d.styled_table(
        ["", "", "", "", ""],
        [
            ("TOTAL MONTHLY ONLINE SERVICES", "", "", "", "$4,045.25"),
        ],
        total_indices=[0],
    )

    d.spacer()
    d.section_header("1.4 Service Levels")
    d.styled_table(
        ["Service Level", "Target"],
        [
            ("Infrastructure Uptime", "99.9% monthly (excluding scheduled maintenance)"),
            ("Scheduled Maintenance", "Tuesday evenings and Saturdays (with advance notice)"),
            ("Critical Incident Response", "Within 1 hour of notification"),
            ("Standard Support Response", "Within 4 business hours"),
            ("Emergency Maintenance", "As needed with reasonable notice"),
        ],
    )

    d.spacer()
    d.section_header("1.4a Service Level Remedies")
    d.numbered("(a)", "Service Credits. If Technijian fails to meet the Infrastructure Uptime target of 99.9% in any calendar month, Client shall be entitled to a service credit equal to 5% of the monthly recurring charges for the affected service for each full 0.1% below the target, up to a maximum credit of 25% of the monthly recurring charges for that service.")
    d.numbered("(b)", "Chronic Failure. If Technijian fails to meet the Infrastructure Uptime target for three (3) or more consecutive months, Client shall have the right to terminate the affected service without penalty upon thirty (30) days written notice.")
    d.numbered("(c)", "Credit Requests. To receive a service credit, Client must submit a written request within thirty (30) days of the end of the affected month. Service credits shall be applied against future invoices and shall not be paid as cash refunds.")
    d.numbered("(d)", "Exclusions. Service level targets and remedies do not apply during scheduled maintenance windows, force majeure events, or outages caused by Client\u2019s actions, third-party services not managed by Technijian, or factors outside Technijian\u2019s reasonable control.")
    d.numbered("(e)", "Sole Remedy. Service credits under this Section 1.4a are Client\u2019s sole and exclusive remedy for Technijian\u2019s failure to meet the applicable service levels.")

    d.spacer()
    d.section_header("1.5 Monitoring and Reporting")
    d.body("Technijian shall provide:")
    d.numbered("(a)", "24/7 monitoring of Client\u2019s infrastructure included in the Service Order;")
    d.numbered("(b)", "Monthly service reports summarizing uptime, incidents, and support activity; and")
    d.numbered("(c)", "Quarterly service reviews with Client\u2019s designated representative.")
    d.numbered("(d)", "Escalation Path. Support requests shall be escalated as follows: Tier 1 (initial response) \u2014 assigned technician; Tier 2 (within 4 hours if unresolved) \u2014 senior engineer or team lead; Tier 3 (within 8 hours if unresolved) \u2014 department manager or CTO Advisory.")

    d.page_break()

    # ── PART 2: SIP TRUNK SERVICES ──
    d.part_header("PART 2 \u2014 SIP TRUNK SERVICES")
    d.spacer()

    d.section_header("2.1 Description")
    d.body("SIP Trunk Services include voice-over-IP telephony, SIP trunking, and related telecommunications services.")

    d.section_header("2.2 Service Components")
    d.styled_table(
        ["Component", "Description"],
        [
            ("SIP Trunk", "Primary voice connectivity with failover"),
            ("Voice Package", "Bundled calling plans (domestic, long distance)"),
            ("DID Numbers", "Direct inward dialing numbers"),
            ("E911", "Emergency services routing"),
        ],
    )

    d.spacer()
    d.section_header("2.3 Service Levels")
    d.styled_table(
        ["Service Level", "Target"],
        [
            ("Voice Uptime", "99.9% monthly"),
            ("Call Quality (MOS)", "4.0 or higher"),
            ("Number Porting", "Completed within 10 business days of request"),
        ],
    )

    d.spacer()
    d.section_header("2.4 SIP Trunk Service Level Remedies")
    d.body(
        "The service level remedies set forth in Section 1.4a (Service Credits, Chronic Failure, "
        "Credit Requests, Exclusions, and Sole Remedy) shall apply equally to SIP Trunk Service Levels, "
        "with \u201cVoice Uptime\u201d substituted for \u201cInfrastructure Uptime\u201d where applicable."
    )
    d.body(
        "Note: SIP Trunk Services are not included in the current Service Order. This section is included "
        "for reference should Client elect to add SIP Trunk Services in the future.",
        space_after=4,
    )

    d.page_break()

    # ── PART 3: VIRTUAL STAFF ──
    d.part_header("PART 3 \u2014 VIRTUAL STAFF (CONTRACTED SUPPORT)")
    d.spacer()

    d.section_header("3.1 Description")
    d.body(
        "Virtual Staff services provide Client with dedicated technology support personnel on a contracted "
        "basis. This service operates on a 12-month cycle-based billing model as described below."
    )

    d.section_header("3.2 Support Roles")
    d.styled_table(
        ["Role", "Location", "Hours/Mo", "BWH Rate"],
        [
            ("Systems Architect", "US (IRV-AD1)", "5.00", "$170.00/hr"),
            ("USA Tech Normal", "US (IRV-TS1)", "15.26", "$106.25/hr"),
            ("India Tech Normal", "India (CHD-TS1)", "58.13", "$12.75/hr"),
            ("India Tech After Hours", "India (CHD-TS1)", "42.82", "$25.50/hr"),
        ],
    )

    d.spacer()
    d.section_header("3.3 Cycle-Based Billing Model")

    d.numbered("(a)", "Billing Cycle. Client has selected a billing cycle of 12 months (the \u201cCycle\u201d). The purpose of the Cycle is to provide Client with a structured path to eliminate any unpaid hour balance by the end of each Cycle, thereby avoiding cancellation fees.")
    d.numbered("(b)", "Monthly Billed Amount Calculation. The fixed monthly billing rate for each role is calculated as follows:")
    d.numbered("1.", "At the start of each new Cycle, Technijian calculates the average monthly hours consumed per role during the previous Cycle, excluding the final month of that Cycle.", indent=1)
    d.numbered("2.", "This average is then adjusted to account for any unpaid hour balance carried forward from the previous Cycle, so that the new monthly billed amount is set at a level designed to bring the unpaid balance to zero by the end of the current Cycle.", indent=1)
    d.numbered("3.", "The monthly billed amount for each role equals the adjusted monthly billed hours multiplied by the applicable Contracted Rate from the Rate Card (Schedule C).", indent=1)

    d.numbered("(c)", "Running Balance. Technijian maintains a running balance for each role:")
    d.numbered("1.", "At the start of each month, the running balance is adjusted by adding the actual hours used during the previous month and subtracting the monthly billed hours for that month.", indent=1)
    d.numbered("2.", "A positive running balance indicates hours consumed in excess of billed amounts (an unpaid hour balance that the Cycle is designed to resolve).", indent=1)
    d.numbered("3.", "A negative running balance indicates hours billed in excess of consumption. A negative balance does not entitle Client to a credit or refund; it simply means Client has no unpaid hour balance (the unpaid balance is zero). The underlying usage data from months producing a negative balance is used in calculating the average for the next Cycle per Section 3.3(b).", indent=1)

    d.numbered("(d)", "Cycle Reconciliation. At the end of each Cycle:")
    d.numbered("1.", "The running balance for each role is reconciled.", indent=1)
    d.numbered("2.", "Any net positive balance (hours consumed but not yet billed) is carried forward into the next Cycle. The next Cycle\u2019s monthly billed amount will be recalculated per Section 3.3(b) to absorb this balance and target a zero unpaid balance by the end of the next Cycle.", indent=1)
    d.numbered("3.", "Any net negative balance (hours billed but not consumed) resets to zero. No credit or refund is issued. The actual usage data from the completed Cycle is used to calculate the monthly billed amount for the next Cycle per Section 3.3(b).", indent=1)

    d.numbered("(e)", "Cancellation. If Client terminates Virtual Staff services or this Agreement:")
    d.numbered("1.", "Any positive running balance (actual hours exceeding billed hours) shall become immediately due and payable. The unpaid hours shall be invoiced at the applicable Hourly Rate from the Rate Card (Schedule C), not the Contracted Rate. This reflects the rate that would have applied had Client engaged Technijian on an ad hoc hourly basis without a Cycle commitment.", indent=1)
    d.numbered("2.", "Any negative running balance (billed hours exceeding actual hours) does not entitle Client to a credit, refund, or offset against the final invoice. The unpaid balance is simply zero and no further amounts are due from either Party with respect to that role.", indent=1)
    d.numbered("3.", "The Cycle-Based Billing Model provides Client with Contracted Rates that are lower than the standard Hourly Rates as consideration for Client\u2019s commitment to the Cycle. The application of Hourly Rates upon cancellation reflects the removal of that commitment and the corresponding rate benefit.", indent=1)

    d.spacer()
    d.section_header("3.4 Weekly Service Reports")
    d.body("Technijian shall provide Client with weekly service reports that detail:")
    d.numbered("(a)", "Each support ticket addressed during the period;")
    d.numbered("(b)", "The role, resource name, and hours spent per ticket;")
    d.numbered("(c)", "Whether the work was performed during normal or after-hours;")
    d.numbered("(d)", "A description of the work performed; and")
    d.numbered("(e)", "The current running balance for each role.")

    d.spacer()
    d.section_header("3.5 New Client Onboarding")
    d.body("For new Clients without a previous Cycle history, the initial Cycle billing shall be based on:")
    d.numbered("(a)", "A mutually agreed-upon estimated monthly hours per role, documented in the Service Order; and")
    d.numbered("(b)", "Actual usage tracking beginning immediately, with the first Cycle reconciliation occurring at the end of the initial Cycle period.")

    d.spacer()
    d.section_header("3.6 Acceptable Use")
    d.body(
        "Client shall direct Virtual Staff only to perform work within the scope of the applicable role "
        "description and the services contemplated by this Agreement. Client shall not direct or request Virtual Staff to:"
    )
    d.numbered("(a)", "Perform any activity that violates applicable law, regulation, or third-party rights;")
    d.numbered("(b)", "Access, process, or transmit data in violation of data protection laws or Client\u2019s own data handling policies;")
    d.numbered("(c)", "Perform work for any entity other than Client without Technijian\u2019s prior written consent; or")
    d.numbered("(d)", "Perform work outside the scope of the role for which they are contracted without a corresponding change to the Service Order.")
    d.body(
        "Technijian reserves the right to reassign or remove Virtual Staff if Client directs work that falls "
        "outside these guidelines, without affecting Client\u2019s billing obligations under Section 3.3."
    )

    d.page_break()

    # ── GENERAL TERMS ──
    d.part_header("GENERAL TERMS FOR THIS SCHEDULE")
    d.spacer()

    d.section_header("Payment and Collection")
    d.body(
        "All payment terms, late fees, dispute procedures, and collection remedies for invoices arising under "
        "this Schedule are governed by Section 3 of the Master Service Agreement."
    )

    d.section_header("Changes to Services")
    d.body(
        "Either Party may request changes to the services described in this Schedule by providing thirty (30) "
        "days written notice. Changes to quantities, roles, or service levels shall be documented in an updated "
        "Service Order signed by both Parties."
    )

    d.section_header("Pricing Adjustments")
    d.body(
        "Pricing for services under this Schedule is subject to the Rate Card (Schedule C). Technijian may "
        "adjust rates upon sixty (60) days written notice, effective at the start of the next Renewal Term of the Agreement."
    )


# ════════════════════════════════════════════════════════════════════
#  PART 4 — SCHEDULE B: Subscription and License Services
# ════════════════════════════════════════════════════════════════════
def build_schedule_b(d):
    d.page_break()
    d.part_header("SCHEDULE B \u2014 SUBSCRIPTION AND LICENSE SERVICES")
    d.body("Attached to Master Service Agreement MSA-BWH")
    d.body("Effective Date: May 1, 2026")
    d.body(
        "This Schedule describes the Subscription and License Services provided by Technijian, Inc. "
        "(\u201cTechnijian\u201d) to Brandywine Homes (\u201cClient\u201d) under the Master Service Agreement."
    )
    d.spacer()

    d.section_header("Current Subscriptions Billed Through Technijian")
    d.styled_table(
        ["Service", "Code", "Unit Price", "Qty", "Monthly"],
        [
            ("Microsoft Entra ID P1", "MS-ADP1", "$7.20/user", "1", "$7.20"),
            ("TOTAL", "", "", "", "$7.20"),
        ],
        total_indices=[1],
    )
    d.spacer()
    d.body(
        "Note: Other subscriptions (Microsoft 365, etc.) are billed directly by the vendor and are not "
        "included in Technijian invoices. If Client elects to procure additional subscriptions through "
        "Technijian, they will be added to this Schedule via an updated Service Order."
    )

    d.spacer()
    d.section_header("Subscription Terms")
    d.numbered("(a)", "Subscriptions are billed monthly in advance per MSA Section 3.02(b).")
    d.numbered("(b)", "Quantity changes require thirty (30) days written notice and take effect on the first day of the following month.")
    d.numbered("(c)", "Annual subscriptions procured on Client\u2019s behalf are non-refundable for the committed term.")


# ════════════════════════════════════════════════════════════════════
#  PART 5 — SCHEDULE C: Rate Card
# ════════════════════════════════════════════════════════════════════
def build_schedule_c(d):
    d.page_break()
    d.part_header("SCHEDULE C \u2014 RATE CARD")
    d.spacer()
    d.body("Attached to Master Service Agreement MSA-BWH")
    d.body("Effective Date: May 1, 2026")
    d.spacer()

    # ── SECTION 1: VIRTUAL STAFF RATES ──
    d.part_header("1. VIRTUAL STAFF RATES")
    d.spacer()
    d.body("The following rates reflect BWH contracted rates with a 15% discount from standard pricing. Virtual Staff billing operates on a 12-month billing cycle.")

    d.section_header("1.1 United States \u2014 Based Staff")
    d.styled_table(
        ["Role", "Standard Rate", "After-Hours Rate", "BWH Contracted Rate"],
        [
            ("Systems Architect", "$250/hr", "$350/hr", "$170.00/hr"),
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr"),
            ("Developer", "$150/hr", "N/A", "$125/hr"),
            ("Tech Support", "$150/hr", "$250/hr", "$106.25/hr"),
        ],
    )

    d.spacer()
    d.section_header("1.2 Offshore \u2014 Based Staff")
    d.styled_table(
        ["Role", "Standard Rate", "After-Hours Rate", "BWH Contracted Rate"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr"),
            ("Tech Support (Normal)", "$15/hr", "$30/hr", "$12.75/hr"),
            ("Tech Support (After Hours)", "$30/hr", "N/A", "$25.50/hr"),
        ],
    )

    d.spacer()
    d.body_bold_prefix("Normal Business Hours: ", "Monday through Friday, 8:00 AM to 6:00 PM Pacific Time, excluding US federal holidays.")
    d.body_bold_prefix("After-Hours: ", "All hours outside of Normal Business Hours, including weekends and US federal holidays.")
    d.body_bold_prefix("BWH Contracted Rate: ", "The 15% discounted hourly rate applied under the BWH Master Service Agreement. These rates apply to Virtual Staff services billed through the 12-month Cycle-Based Billing Model described in Schedule A, Part 3.")
    d.body_bold_prefix("Standard Rate: ", "The standard rate for ad hoc (non-contracted) services. This rate is also applied to calculate cancellation fees on any unpaid hour balance upon termination of Virtual Staff services, as described in Schedule A, Section 3.3(e).")
    d.body_bold_prefix("Billing Cycle: ", "Virtual Staff services are billed on a 12-month cycle. Initial cycle based on estimated monthly hours per the Service Order. Actual usage tracked from day one with reconciliation at end of each cycle period.")

    # ── SECTION 2: PROJECT RATES ──
    d.spacer()
    d.part_header("2. PROJECT RATES")
    d.spacer()
    d.styled_table(
        ["Service", "Rate", "Minimum", "Notes"],
        [
            ("On-Site Support (US)", "$150/hr", "2-hour minimum", "No trip charges"),
            ("Remote Support (ad hoc, non-contracted)", "$150/hr", "15-min increments", "Billed in 15-minute increments"),
            ("Emergency / Critical Response", "$250/hr", "1-hour minimum", "After-hours and critical incidents"),
            ("Project Management", "$150/hr", "N/A", "For SOW-based engagements"),
        ],
    )

    d.page_break()

    # ── SECTION 3: ONLINE SERVICES INFRASTRUCTURE ──
    d.part_header("3. ONLINE SERVICES \u2014 INFRASTRUCTURE")
    d.spacer()

    d.section_header("3.1 Cloud Infrastructure")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Cloud VM \u2014 vCore", "CL-VC", "Per vCore", "M", "$6.25"),
            ("Cloud VM \u2014 Memory", "CL-GB", "Per GB", "M", "$0.63"),
            ("Cloud VM \u2014 Shared Bandwidth", "CL-SBW", "Per connection", "M", "$15.00"),
            ("Hosting \u2014 CDN", "HC", "Per domain", "M", "$30.00"),
            ("Production Storage", "TB-PSTR", "Per TB", "M", "$200.00"),
            ("Replicated Storage", "TB-RSTR", "Per TB", "M", "$100.00"),
            ("Backup Storage", "TB-BSTR", "Per TB", "M", "$50.00"),
        ],
    )

    d.spacer()
    d.section_header("3.2 Server Management")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Patch Management", "PMW", "Per server", "M", "$4.00"),
            ("Image Backup (Veeam)", "IB", "Per server", "M", "$15.00"),
            ("AV Protection \u2014 Server (CrowdStrike)", "AVS", "Per server", "M", "$10.50"),
            ("AVM Protection \u2014 Server (MalwareBytes)", "AVMS", "Per server", "M", "$16.00"),
            ("AVH Protection \u2014 Server (Huntress)", "AVHS", "Per server", "M", "$6.00"),
            ("My Secure Internet (DNS Filtering)", "SI", "Per server", "M", "$6.00"),
            ("My Remote", "MR", "Per server", "M", "$2.00"),
            ("Health Monitoring (SNMP)", "SHM", "Per device", "M", "$2.00"),
            ("Syslog Monitoring (SNMP)", "SSM", "Per device", "M", "$2.00"),
            ("My Ops \u2014 Net", "OPS-NET", "Per server", "M", "$3.25"),
            ("My Ops \u2014 Config", "OPS-BKP", "Per switch", "M", "$6.00"),
            ("My Ops \u2014 Traffic", "OPS-TR", "Per firewall", "M", "$14.00"),
            ("My Ops \u2014 Port", "OPS-PRT", "Per port", "M", "$0.25"),
            ("My Ops \u2014 Storage", "OPS-ST", "Per disk", "M", "$4.75"),
            ("My Ops \u2014 Wifi", "OPS-WF", "Per AP", "M", "$1.00"),
            ("MyAudit Server", "AMS", "Per server", "M", "$252.00"),
        ],
    )

    d.spacer()
    d.section_header("3.3 Desktop Management")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Patch Management", "PMW", "Per desktop", "M", "$4.00"),
            ("Patch Management (Mac OS)", "PMMAC", "Per desktop", "M", "$11.00"),
            ("AV Protection \u2014 Desktop (CrowdStrike)", "AVD", "Per desktop", "M", "$8.50"),
            ("AVM Protection \u2014 Desktop (MalwareBytes)", "AVMD", "Per desktop", "M", "$5.00"),
            ("AVH Protection \u2014 Desktop (Huntress)", "AVMH", "Per desktop", "M", "$6.00"),
            ("My Secure Internet (DNS Filtering)", "SI", "Per desktop", "M", "$6.00"),
            ("My Remote", "MR", "Per desktop", "M", "$2.00"),
            ("My Ops \u2014 Net", "OPS-NET", "Per desktop", "M", "$3.25"),
            ("Audit Monitoring (Base)", "AM", "Per desktop", "M", "$9.50"),
            ("Audit Monitoring (30 Day)", "AM30", "Per desktop", "M", "$12.50"),
            ("Audit Monitoring (90 Day)", "AM90", "Per desktop", "M", "$15.00"),
            ("Audit Monitoring (6 Months)", "AM6M", "Per desktop", "M", "$25.00"),
            ("Audit Monitoring (1 Year)", "AM1Y", "Per desktop", "M", "$45.00"),
        ],
    )

    d.spacer()
    d.section_header("3.3a Advanced Audit Monitoring")
    d.body("Advanced audit monitoring tiers include User Activity Monitoring (UAM) and Data Loss Prevention (DLP).")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Audit Monitoring UAM (Base + UAM)", "AMUAM", "Per desktop", "M", "$42.00"),
            ("Audit Monitoring UAM (30 Day)", "AMUAM30", "Per desktop", "M", "$44.90"),
            ("Audit Monitoring UAM (90 Day)", "AMUAM90", "Per desktop", "M", "$47.40"),
            ("Audit Monitoring UAM (6 Months)", "AMUAM6M", "Per desktop", "M", "$57.40"),
            ("Audit Monitoring UAM (1 Year)", "AMUAM1Y", "Per desktop", "M", "$77.40"),
            ("Audit Monitoring DLP (UAM + DLP)", "AMDLP", "Per desktop", "M", "$49.00"),
            ("Audit Monitoring DLP (30 Day)", "AMDLP30", "Per desktop", "M", "$51.90"),
            ("Audit Monitoring DLP (90 Day)", "AMDLP90", "Per desktop", "M", "$57.30"),
            ("Audit Monitoring DLP (6 Months)", "AMDLP6M", "Per desktop", "M", "$72.70"),
            ("Audit Monitoring DLP (1 Year)", "AMDLP1Y", "Per desktop", "M", "$108.10"),
        ],
    )

    d.page_break()

    d.section_header("3.4 Network and Security")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Sophos 1C-4G (Up to 100 Mbps)", "SO-1C4G", "Per appliance", "Y", "$145.00"),
            ("Sophos 2C-4G (Up to 300 Mbps)", "SO-2C4G", "Per appliance", "Y", "$350.00"),
            ("Sophos 4C-6G (Up to 500 Mbps)", "SO-4C6G", "Per appliance", "Y", "$510.00"),
            ("Sophos 6C-8G (Up to 750 Mbps)", "SO-6C8G", "Per appliance", "Y", "$855.00"),
            ("Sophos 8C-16G (Up to 1 Gbps)", "SO-8C16G", "Per appliance", "Y", "$1,330.00"),
            ("Sophos 16C-24G (Up to 1.5 Gbps)", "SO-16C24G", "Per appliance", "Y", "$2,435.00"),
            ("Sophos UC-UG (Over 1.5 Gbps)", "SO-UCG", "Per appliance", "Y", "$3,320.00"),
            ("VeloCloud SD-WAN 30M", "VC-30M", "Per appliance", "Y", "$200.00"),
            ("VeloCloud SD-WAN 50M", "VC-50M", "Per appliance", "Y", "$250.00"),
            ("VeloCloud SD-WAN 100M", "VC-100M", "Per appliance", "Y", "$310.00"),
            ("VeloCloud SD-WAN 200M", "VC-200M", "Per appliance", "Y", "$385.00"),
            ("Edge Appliance 16GB", "Edge-16M", "Per appliance", "Y", "$100.00"),
            ("Real-Time Penetration Testing", "RTPT", "Per IP", "M", "$7.00"),
            ("Vulnerability Assessment", "PT", "Per device", "M", "$1.00"),
            ("Network Assessment", "NA", "Per device", "M", "$1.00"),
        ],
    )

    d.spacer()
    d.section_header("3.5 Email, Compliance, and Identity")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("E-Mail Archiving", "EA", "Per user", "M", "$4.00"),
            ("DMARC/DKIM Management", "DKIM", "Per domain", "M", "$20.00"),
            ("Site Assessment", "SA", "Per domain", "M", "$50.00"),
            ("Veeam 365 Backup", "V365", "Per user", "M", "$2.50"),
            ("Anti-Spam Basic", "ASB", "Per user", "M", "$4.25"),
            ("Anti-Spam Standard (+ Encryption)", "ASA", "Per user", "M", "$6.25"),
            ("Anti-Spam Professional", "ASP", "Per user", "M", "$10.25"),
            ("Anti-Spam GSuite", "ASG", "Per user", "M", "$6.25"),
            ("Phishing Training", "PHT", "Per user", "M", "$6.00"),
            ("SSO / 2FA Login", "SSO-2FA", "Per user", "M", "$12.00"),
            ("SSO / Desktop", "SSO-DSK", "Per user", "M", "$6.00"),
            ("Credential Manager", "CRM", "Per user", "M", "$5.00"),
        ],
    )

    # ── SECTION 4: SIP TRUNK, HOSTING, LICENSING ──
    d.spacer()
    d.part_header("4. SIP TRUNK, HOSTING, AND LICENSING SERVICES")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("SIP Trunk", "\u2014", "Per trunk", "M", "Per Service Order"),
            ("Voice Package", "\u2014", "Per package", "M", "Per Service Order"),
            ("DID Number", "\u2014", "Per number", "M", "Per Service Order"),
            ("Hosting \u2014 Web IIS", "HIIS", "Per domain", "M", "$10.00"),
            ("Hosting \u2014 Web WordPress", "HWP", "Per domain", "M", "$10.00"),
            ("Hosting \u2014 Web Magento", "HMG", "Per domain", "M", "$10.00"),
            ("My VDI", "MYVDI", "Per user", "M", "$10.00"),
            ("MDM (iOS)", "MDMIOS", "Per device", "M", "$4.00"),
            ("My Disk (File Sharing & Backup)", "MDU", "Per user", "M", "$16.00"),
            ("Veeam VBR (Backup & Replication)", "VBR", "Per VM", "M", "$13.00"),
            ("Veeam One (Backup Management)", "VONE", "Per VM", "M", "$3.00"),
            ("365 Azure Directory P2", "MS-ADP2", "Per user", "M", "$9.00"),
            ("MSFT \u2014 RDP SAL", "MS-WRDS", "Per license", "M", "$9.20"),
            ("MSFT \u2014 SQL Std SAL", "MS-SSSE", "Per license", "M", "$18.50"),
            ("MSFT \u2014 SQL Std Core", "MS-SSSC", "Per license", "M", "$175.00"),
            ("MSFT \u2014 Sharepoint SAL", "MS-SPSE", "Per license", "M", "$4.00"),
            ("MSFT \u2014 Exchange Std SAL", "MS-ESS", "Per license", "M", "$3.50"),
            ("MSFT \u2014 Exchange Ent SAL", "MS-EESS", "Per license", "M", "$4.00"),
            ("MSFT \u2014 Std Core 2 LIC", "MS-STD", "Per license", "M", "$5.25"),
            ("Foxit PDF Business", "FX-BSPDF", "Per license", "Y", "$14.00"),
        ],
    )

    d.page_break()

    # ── SECTION 5: TERMS ──
    d.part_header("5. TERMS")
    d.spacer()

    d.section_header("5.01 Rate Adjustments")
    d.body(
        "Technijian may adjust the rates in this Schedule upon sixty (60) days written notice to Client. "
        "Adjusted rates shall take effect at the start of the next Renewal Term of the Agreement."
    )

    d.section_header("5.02 Volume Discounts")
    d.body(
        "Volume-based pricing may be negotiated and documented in the applicable Service Order or Subscription Order."
    )

    d.section_header("5.03 Minimum Billing")
    d.body(
        "Contracted Virtual Staff hours are billed per the Cycle-Based Billing Model described in Schedule A, "
        "Part 3, at the BWH Contracted Rate. Ad hoc (non-contracted) support is billed at the Standard Rate in "
        "15-minute increments. Upon cancellation or termination, any unpaid hour balance shall be invoiced at "
        "the applicable Standard Rate as set forth in Schedule A, Section 3.3(e)."
    )

    d.section_header("5.04 Relationship to Price List")
    d.body(
        "Technijian maintains a separate Services Price List that provides the current catalog of available "
        "services and list prices. The Price List may be updated by Technijian from time to time and shall be "
        "made available to Client upon request. In the event of a conflict between the Price List and this Rate "
        "Card, the rates documented in this Rate Card (Schedule C) shall govern for the duration of the "
        "then-current term. Any rate change shall be subject to the sixty (60) day written notice and Renewal "
        "Term effective date requirements set forth in Section 5.01 above and MSA Section 9.02."
    )

    # Rate Card Signatures
    d.signatures()
    d.footer_bar()


# ════════════════════════════════════════════════════════════════════
#  MAIN — Build complete document
# ════════════════════════════════════════════════════════════════════
def build_all():
    d = BrandedDoc()
    build_cover(d)
    build_msa_body(d)
    build_schedule_a(d)
    build_schedule_b(d)
    build_schedule_c(d)
    d.save("MSA-BWH.docx")


if __name__ == "__main__":
    build_all()
