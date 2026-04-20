"""
Generate SOW-AFFG-004-Rev1 matching the EXACT format of SOW-AFFG-004.

Title block (page 1): deepcopied from original — logo, orange rules, date, company.
Content (page 2+):    rebuilt with exact XML matching original's typography.

Original structure (first 14 body elements):
  [0]  empty para
  [1]  logo (inline image, centered, spacing after=200)
  [2]  empty para
  [3]  orange horizontal rule (pBdr bottom F67D4B sz=2)
  [4]  "Statement of Work" (24pt bold #1A1A2E centered)
  [5]  subtitle (12pt #59595B centered)
  [6]  orange horizontal rule
  [7]  empty para
  [8]  date (11pt #59595B centered)  <-- updated to "Rev 1 – May 1, 2026"
  [9]  empty para
  [10] "Technijian, Inc." (11pt bold #1A1A2E centered)
  [11] address (10pt #59595B centered)
  [12] page break
  [13] section properties para (defines page 1 margins)
  [14] Heading1 "STATEMENT OF WORK"  <-- content starts here
"""
import copy
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = r'c:\vscode\tech-legal\tech-legal\clients\AFFG\03_SOW\SOW-AFFG-004-Managed-Device-Migration.docx'
DEST = r'c:\vscode\tech-legal\tech-legal\clients\AFFG\03_SOW\SOW-AFFG-004-Rev1-Managed-Device-Migration.docx'

DARK  = '1A1A2E'
GREY  = '59595B'
BLUE  = '006DB6'
WHITE = 'FFFFFF'
ROW   = 'F8F9FA'
BRDR  = 'DDDDDD'
TOTAL = 'E8EFF5'
FONT  = 'Open Sans'

# ── Step 1: deepcopy title block from original ────────────────────────────────
src_doc = Document(SRC)
src_body_children = list(src_doc.element.body)
title_block = [copy.deepcopy(src_body_children[i]) for i in range(14)]

# Update date text in element [8]
date_elem = title_block[8]
for t_el in date_elem.findall('.//' + qn('w:t')):
    if t_el.text and '2026' in t_el.text:
        t_el.text = 'Rev 1 \u2013 May 1, 2026'

# ── Step 2: clone original → dest, clear body ────────────────────────────────
shutil.copy2(SRC, DEST)
doc = Document(DEST)
body = doc.element.body
final_sectPr = body.find(qn('w:sectPr'))
if final_sectPr is not None:
    final_sectPr = copy.deepcopy(final_sectPr)
for child in list(body):
    body.remove(child)

# ── Step 3: reinsert title block ──────────────────────────────────────────────
for elem in title_block:
    body.append(elem)
if final_sectPr is not None:
    body.append(final_sectPr)

# ── XML helpers ───────────────────────────────────────────────────────────────

def make_rPr(size_pt=11, bold=False, color=GREY):
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    for attr in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        fonts.set(qn(f'w:{attr}'), FONT)
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    szC = OxmlElement('w:szCs')
    szC.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szC)
    return rPr

def add_run(p_elem, text, size_pt=11, bold=False, color=GREY):
    r = OxmlElement('w:r')
    r.append(make_rPr(size_pt, bold, color))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_elem.append(r)

def insert_para(elems=None):
    """Create a bare w:p and insert before final sectPr."""
    p = OxmlElement('w:p')
    if elems:
        for e in elems:
            p.append(e)
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        body.insert(list(body).index(sectPr), p)
    else:
        body.append(p)
    return p

def heading(level, text):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading{level}')
    pPr.append(pStyle)
    p.append(pPr)
    add_run(p, text, size_pt=11, color=GREY)
    return p

def body_para(text, before=0, after=6, bold=False, size_pt=11, color=GREY):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(before * 20)))
    sp.set(qn('w:after'), str(int(after * 20)))
    pPr.append(sp)
    p.append(pPr)
    add_run(p, text, size_pt=size_pt, bold=bold, color=color)
    return p

def kv_para(label, value):
    """Matches original exactly: no pPr, two runs (bold dark + grey)."""
    p = insert_para()
    add_run(p, label, size_pt=11, bold=True, color=DARK)
    add_run(p, value, size_pt=11, color=GREY)
    return p

def bullet_para(text):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'ListParagraph')
    pPr.append(pStyle)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '2')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
    pPr.append(sp)
    p.append(pPr)
    add_run(p, text, size_pt=11, color=GREY)
    return p

def clause_para(num, text):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
    pPr.append(sp); p.append(pPr)
    add_run(p, num + '  ', size_pt=11, bold=True, color=DARK)
    add_run(p, text, size_pt=11, color=GREY)
    return p

# ── Table builder ─────────────────────────────────────────────────────────────
# Ticket cols match original exactly: [1600, 4200, 1600, 960] DXA
TICKET_COLS = [1600, 4200, 1600, 960]
TICKET_HDR  = ['Ticket', 'Title', 'Assigned To', 'Est. Hours']

def make_cell(text, fill, bold, size_pt=10, width_dxa=None):
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    if width_dxa:
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_dxa)); tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:color'), BRDR); b.set(qn('w:sz'), '1')
        tcB.append(b)
    tcPr.append(tcB)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top','80'),('left','120'),('bottom','80'),('right','120')):
        m = OxmlElement(f'w:{side}'); m.set(qn('w:w'), val); m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp); p.append(pPr)
    r = OxmlElement('w:r')
    r.append(make_rPr(size_pt=size_pt, bold=bold,
                      color=WHITE if fill == BLUE else GREY))
    t_el = OxmlElement('w:t')
    t_el.set(qn('xml:space'), 'preserve'); t_el.text = str(text)
    r.append(t_el); p.append(r); tc.append(p)
    return tc

def add_table(headers, rows, col_dxa):
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(col_dxa))); tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_dxa:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tbl.append(tblGrid)
    # header row
    hdr_tr = OxmlElement('w:tr')
    for i, h in enumerate(headers):
        hdr_tr.append(make_cell(h, BLUE, bold=True, width_dxa=col_dxa[i]))
    tbl.append(hdr_tr)
    # data rows
    for row_data in rows:
        is_total = (str(row_data[0]) == '' and 'Total' in str(row_data[1])) \
                   or str(row_data[0]) in ('TOTAL', 'Total', 'Subtotal')
        fill = TOTAL if is_total else ROW
        tr = OxmlElement('w:tr')
        for i, val in enumerate(row_data):
            tr.append(make_cell(val, fill, bold=is_total, width_dxa=col_dxa[i]))
        tbl.append(tr)
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        body.insert(list(body).index(sectPr), tbl)
    else:
        body.append(tbl)
    insert_para()  # spacing after table

# ── CONTENT ───────────────────────────────────────────────────────────────────

heading(1, 'STATEMENT OF WORK')

for label, value in [
    ('SOW Number: ', 'SOW-AFFG-004 Rev 1'),
    ('Title: ', 'Fleet Right-Sizing and Compliance Configuration'),
    ('Effective Date: ', 'May 1, 2026'),
    ('Revision: ', 'Rev 1 \u2013 replaces SOW-AFFG-004 (April 2026)'),
    ('Parent Agreement: ', 'MSA-AFFG-2026'),
    ('Primary Contact: ', 'Iris Liu  |  iris.liu@americanfundstars.com  |  949-439-2392'),
    ('Regulatory Scope: ', 'SEC Reg S-P (2024), FINRA 3110, FINRA 4370, SEC 17a-4, NIST SP 800-53 Rev 5'),
]:
    kv_para(label, value)

insert_para()

body_para(
    'This Statement of Work (\u201cSOW\u201d) is entered into by and between Technijian, Inc. (\u201cTechnijian\u201d), '
    '18 Technology Drive, Suite 141, Irvine, California 92618 and American Fundstars Financial Group LLC '
    '(\u201cClient\u201d or \u201cAFFG\u201d), 1 Park Plaza, Suite 210, Irvine, California 92618. '
    'Primary Contact: Iris Liu.'
)
body_para(
    'Revision Notes (Rev 1): (1) BYOD MAM-only adopted for personal mobile devices \u2014 no company phones; '
    '(2) managed device fleet confirmed as 4 Mac Mini + 6 Windows laptops (10 endpoints); '
    '(3) SSO/2FA Gateway product removed \u2014 Azure Entra ID SSO (M365 E3/E5 included) used for identity and MFA; '
    '(4) CloudBrink ZTNA replaces Cisco Umbrella for DNS filtering and zero-trust remote egress; '
    '(5) custodian portal access controlled via office WAN IP whitelist + CloudBrink egress IP for traveling laptops.'
)

# ── Section 1 ─────────────────────────────────────────────────────────────────
heading(1, '1. PROJECT OVERVIEW')
body_para(
    'AFFG currently operates 16 physical desktops under the Technijian standard managed services program. '
    'This SOW covers two concurrent workstreams: (1) right-sizing the managed device fleet from 16 desktops '
    'to 10 company-owned endpoints (4 Mac Mini + 6 Windows laptops), and (2) deploying the compliance controls '
    'required under SEC Regulation S-P (2024 amendments) and FINRA Rule 3110 that are not present in the '
    'current environment. Compliance controls include Microsoft Intune enrollment with Conditional Access, '
    'CloudBrink Zero Trust Network Access (ZTNA) replacing Cisco Umbrella, MyAudit UAM+DLP on all 10 '
    'managed endpoints, and Intune MAM-only for employee personal mobile devices (BYOD).'
)
heading(2, '1.1 Objectives')
for obj in [
    'Enroll 4 Mac Mini (Apple Business Manager) and 6 Windows laptops (Autopilot) in Microsoft Intune with full endpoint security stack',
    'Deploy CloudBrink ZTNA on all 10 endpoints to replace Cisco Umbrella; whitelist CloudBrink egress IP at Schwab Advisor Services and IBKR',
    'Configure Entra ID Conditional Access: named location policy (office WAN IP) + Intune-compliant device requirement for M365; block legacy auth',
    'Deploy MyAudit UAM+DLP (AMDLP1Y) on all 10 managed endpoints \u2014 block USB, clipboard, print, and local download of financial data',
    'Configure BYOD Intune MAM (Outlook App) for employee personal mobile devices \u2014 no MDM enrollment of personal devices',
    'Offboard 16 legacy desktops from Technijian managed services and amend Schedule A accordingly',
    'Conduct user training and deliver post-implementation compliance documentation',
]:
    bullet_para(obj)

heading(2, '1.2 Out of Scope')
for exc in [
    'CloudBrink per-user subscription licensing (procured directly by AFFG)',
    'Hardware procurement (AFFG provides Mac Mini desktops and Windows laptops)',
    'Microsoft 365 license costs (AFFG procures directly \u2014 E3 or E5 required)',
    'Technijian My Archive (AFFG operates own email-archiving platform)',
    'Company-owned mobile phones and MDM enrollment thereof (BYOD model adopted)',
    'Physical removal or disposal of retired desktop hardware (AFFG responsibility)',
]:
    bullet_para(exc)

# ── Section 2 ─────────────────────────────────────────────────────────────────
heading(1, '2. IMPLEMENTATION SCOPE')

heading(2, 'Phase 1: Endpoint Enrollment & Foundation (Weeks 1\u20134)')
body_para('7 tickets  |  24 hours')
body_para(
    'Configure Microsoft Intune Autopilot (Windows) and Apple Business Manager (Mac Mini). Define device '
    'compliance policies and configuration profiles for both platforms. Enroll and image all 10 endpoints. '
    'Deploy full security stack on each endpoint: CrowdStrike Falcon EDR, Huntress MDR, Patch Management, '
    'My Remote (ScreenConnect), and CloudBrink agent. Deploy MyAudit UAM+DLP on all 10 endpoints.'
)
heading(3, 'Phase 1 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-001', 'Intune Autopilot Configuration & Enrollment Profiles \u2013 Windows', 'CHD-TS1', '3'),
    ('AFFG-004-002', 'Apple Business Manager (ABM) Setup & MDM Configuration \u2013 Mac Mini', 'CHD-TS1', '3'),
    ('AFFG-004-003', 'Intune Device Compliance & Configuration Policies (Windows + macOS)', 'CHD-TS1', '2'),
    ('AFFG-004-004', 'Windows Laptop Autopilot Enrollment & Imaging \u2013 6 Devices', 'CHD-TS1', '4'),
    ('AFFG-004-005', 'Mac Mini ABM Enrollment & Imaging \u2013 4 Devices', 'CHD-TS1', '4'),
    ('AFFG-004-006', 'Endpoint Security Stack Deployment \u2013 All 10 Endpoints', 'CHD-TS1', '4'),
    ('AFFG-004-007', 'MyAudit UAM+DLP Deployment & Policy Configuration \u2013 All 10 Endpoints', 'CHD-TS1', '4'),
    ('', 'Phase Total', '', '24'),
], TICKET_COLS)

heading(2, 'Phase 2: Access Control & ZTNA (Weeks 3\u20136)')
body_para('5 tickets  |  16 hours')
body_para(
    'Configure Entra ID Conditional Access policies: office WAN IP named location, Intune-compliant device '
    'requirement for M365, legacy authentication block, and MFA enforcement. Deploy CloudBrink ZTNA '
    '(tenant provisioning, connector, agent on all 10 endpoints, per-application micro-segmentation). '
    'Whitelist CloudBrink egress IP at Schwab Advisor Services and IBKR as the sole permitted remote '
    'access origin. Remove Cisco Umbrella from all 10 endpoints.'
)
heading(3, 'Phase 2 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-008', 'Conditional Access Policy Suite \u2013 Named Location + Compliant Device + MFA', 'IRV-TS1', '4'),
    ('AFFG-004-009', 'CloudBrink ZTNA Tenant & Connector Provisioning', 'IRV-TS1', '3'),
    ('AFFG-004-010', 'CloudBrink Agent Deployment \u2013 All 10 Endpoints (replaces Cisco Umbrella)', 'CHD-TS1', '3'),
    ('AFFG-004-011', 'CloudBrink Access Policies & Per-App Micro-Segmentation', 'CHD-TS1', '3'),
    ('AFFG-004-012', 'Schwab & IBKR IP Whitelist Update \u2013 Office WAN IP + CloudBrink Egress IP', 'IRV-TS1', '3'),
    ('', 'Phase Total', '', '16'),
], TICKET_COLS)

heading(2, 'Phase 3: BYOD Mobile MAM (Weeks 5\u20137)')
body_para('2 tickets  |  4 hours')
body_para(
    'Configure Intune App Protection (MAM) policies for employee personal mobile devices. No MDM enrollment '
    'of personal devices. Employees install the Outlook App from their personal app store; MAM policies '
    'containerize corporate data, enforce app PIN, and enable selective corporate data wipe on separation '
    'without accessing personal content. Full M365 browser access from mobile is blocked via Conditional Access '
    'unless on an Intune-managed company endpoint.'
)
heading(3, 'Phase 3 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-013', 'Intune MAM Policy Configuration \u2013 BYOD Outlook App (iOS & Android)', 'CHD-TS1', '2'),
    ('AFFG-004-014', 'MAM Policy Testing & Employee BYOD Setup Guide', 'CHD-TS1', '2'),
    ('', 'Phase Total', '', '4'),
], TICKET_COLS)

heading(2, 'Phase 4: Legacy Desktop Offboarding (Weeks 5\u20138)')
body_para('3 tickets  |  8 hours')
body_para(
    'Decommission the 16 legacy physical desktops from Technijian managed services. Remove all Technijian '
    'agents (CrowdStrike, Huntress, Patch Management, My Remote, Cisco Umbrella) from retiring devices. '
    'User data transfer assistance for any users moving from a retiring desktop to a new managed endpoint. '
    'Draft and deliver Schedule A amendment to reflect the reduced fleet and updated service stack.'
)
heading(3, 'Phase 4 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-015', 'Agent Removal from 16 Legacy Desktops & Managed Services Offboarding', 'CHD-TS1', '4'),
    ('AFFG-004-016', 'User Data Transfer Assistance \u2013 Legacy Desktop to New Endpoint', 'CHD-TS1', '2'),
    ('AFFG-004-017', 'Schedule A Amendment \u2013 Fleet Right-Sizing & Compliance Stack Update', 'IRV-TS1', '2'),
    ('', 'Phase Total', '', '8'),
], TICKET_COLS)

heading(2, 'Phase 5: Validation & Training (Weeks 7\u201310)')
body_para('4 tickets  |  12 hours')
body_para(
    'End-to-end security and compliance validation across all 10 managed endpoints. Verify MyAudit DLP '
    'policy enforcement, CloudBrink ZTNA posture checks, Conditional Access enforcement, and MAM containment. '
    'User training on managed-device workflows and CloudBrink remote access. Compliance gap assessment '
    'documenting SEC Reg S-P and FINRA control coverage.'
)
heading(3, 'Phase 5 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-018', 'End-to-End Security & Compliance Validation \u2013 All 10 Endpoints', 'IRV-TS1', '4'),
    ('AFFG-004-019', 'MyAudit DLP & CloudBrink Policy Verification', 'CHD-TS1', '3'),
    ('AFFG-004-020', 'User Training \u2013 Managed Endpoints & CloudBrink Remote Access', 'CHD-TS1', '3'),
    ('AFFG-004-021', 'Compliance Gap Assessment & Control Coverage Report (SEC/FINRA)', 'IRV-TS1', '2'),
    ('', 'Phase Total', '', '12'),
], TICKET_COLS)

heading(2, 'Phase 6: Documentation (Weeks 9\u201312)')
body_para('2 tickets  |  8 hours')
body_para(
    'Update compliance documentation suite to reflect the managed-device architecture: endpoint security '
    'policy, ZTNA access control runbook, incident response procedure, and data handling policy. '
    'All documents delivered in Technijian standard format.'
)
heading(3, 'Phase 6 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-022', 'Update Compliance Documentation Suite \u2013 Managed Device Architecture', 'CHD-TS1', '6'),
    ('AFFG-004-023', 'Final Deliverable Package & Client Acceptance Sign-Off', 'IRV-TS1', '2'),
    ('', 'Phase Total', '', '8'),
], TICKET_COLS)

# ── Section 3 ─────────────────────────────────────────────────────────────────
heading(1, '3. CONTROL-TO-CITATION MAP')
body_para('The following table maps each design component to its regulatory control objective:')
add_table(
    ['Design Component', 'Control Objective', 'Citation(s)'],
    [
        ('Intune-managed endpoints \u2013 4 Mac Mini + 6 Windows laptops',
         'Company-owned, encrypted, policy-enforced devices with centralized MDM control',
         'Reg S-P 248.30(a)(1)\u2013(3); FINRA 3110(a)'),
        ('CloudBrink ZTNA on all 10 endpoints (replaces Cisco Umbrella)',
         'Zero-trust egress; device posture verified per session; DNS filtering; blocks unmanaged device access to portals',
         'Reg S-P 248.30(a)(3); NIST AC-17, SC-7'),
        ('Office WAN IP + CloudBrink egress IP whitelist at Schwab & IBKR',
         'Custodian portals accessible only from office or CloudBrink-tunneled managed device \u2014 MFA alone does not restrict device origin',
         'Reg S-P 248.30(a)(3); NIST AC-3, SC-7'),
        ('Entra ID Conditional Access \u2013 Named Location + Compliant Device',
         'M365 access restricted to office WAN IP or Intune-enrolled device; legacy authentication blocked',
         'Reg S-P 248.30(a)(3); NIST AC-3, IA-2(1)'),
        ('Intune MAM \u2013 BYOD Outlook App (personal mobile)',
         'Corporate email containerized on personal device; selective corporate wipe on separation; no personal data access',
         'Reg S-P 248.30(a)(1)\u2013(3); NIST AC-19, MP-6'),
        ('MyAudit UAM+DLP (AMDLP1Y) \u2013 all 10 managed endpoints',
         'Block USB exfiltration, local downloads, print, clipboard copy, and screen capture of financial data',
         'Reg S-P 248.30(a)(3); NIST AC-19, MP-7'),
        ('Technijian Veeam Backup for M365',
         'Immutable, non-rewriteable backup of full M365 tenant (Exchange, Teams, SharePoint, OneDrive)',
         'SEC Rule 17a-4(b),(f); FINRA 4511'),
        ('Technijian monthly network assessment (SA)',
         'Detect configuration drift; provide attested evidence of ongoing control operation',
         'FINRA 3110(c); FINRA 3120'),
    ],
    [2100, 2900, 2360]
)

# ── Section 4 ─────────────────────────────────────────────────────────────────
heading(1, '4. DELIVERABLES')
for d in [
    '10 Intune-managed company endpoints (4 Mac Mini via ABM, 6 Windows laptops via Autopilot) with full security stack deployed and verified',
    'CloudBrink ZTNA deployed on all 10 endpoints with per-application micro-segmented access and device posture enforcement',
    'Schwab Advisor Services and IBKR whitelisted for office WAN IP and CloudBrink egress IP; all other remote access blocked',
    'Entra ID Conditional Access enforcing M365 access: office WAN IP or Intune-compliant device required; legacy auth blocked',
    'MyAudit UAM+DLP active on all 10 managed endpoints with DLP policy covering USB, clipboard, print, and download',
    'Intune MAM policies for BYOD personal mobile devices (iOS and Android, Outlook App)',
    '16 legacy desktops offboarded from Technijian managed services with all agents removed',
    'Schedule A amendment reflecting updated fleet (10 endpoints) and compliance service stack',
    'Updated compliance documentation suite (endpoint security policy, ZTNA runbook, IRP, data handling policy)',
    'Compliance gap assessment confirming SEC Reg S-P and FINRA control coverage on managed-device architecture',
]:
    bullet_para(d)

# ── Section 5 ─────────────────────────────────────────────────────────────────
heading(1, '5. PRICING AND PAYMENT')
heading(2, '5.1 Labor Summary')
add_table(
    ['Role', 'Rate', 'Hours', 'Labor'],
    [
        ('US Tech Support (IRV-TS1)', '$150/hr', '20', '$3,000.00'),
        ('India Tech Support (CHD-TS1)', '$45/hr', '52', '$2,340.00'),
        ('TOTAL', '', '72 hrs', '$5,340.00'),
    ],
    [4000, 1400, 1200, 1760]
)

heading(2, '5.2 Payment Schedule')
add_table(
    ['Milestone', 'Trigger', 'Amount'],
    [
        ('50% upon SOW execution', 'SOW signed by both parties', '$2,670.00'),
        ('25% upon Phase 4 completion', '16 legacy desktops offboarded; Schedule A amendment delivered', '$1,335.00'),
        ('25% upon Phase 6 delivery', 'Compliance documentation + client acceptance sign-off', '$1,335.00'),
        ('Total', '', '$5,340.00'),
    ],
    [3000, 4200, 1160]
)

heading(2, '5.3 Payment Terms')
body_para(
    'All invoices are due and payable within thirty (30) days of the invoice date. '
    'Late payments are subject to a late fee of 1.5% per month on the unpaid balance.'
)

# ── Section 6 ─────────────────────────────────────────────────────────────────
heading(1, '6. ASSUMPTIONS')
for a in [
    'AFFG has procured or will procure 4 Mac Mini desktops and 6 Windows 11 Pro/Enterprise laptops prior to Phase 1 commencement',
    'AFFG will enroll Mac Minis in Apple Business Manager (ABM) prior to Phase 1; Technijian will provide setup assistance',
    'AFFG has procured CloudBrink ZTNA per-user subscription licenses for all 10 managed users',
    "AFFG\u2019s M365 tenant is licensed at E3 or E5 (Intune, Conditional Access, and Entra ID P1/P2 included)",
    'AFFG employees will use personal mobile devices (BYOD) for corporate email via Outlook App; no company mobile phones',
    'AFFG provides the office static WAN IP address to Technijian prior to Phase 2 commencement',
    'Schwab Advisor Services and IBKR can update their IP whitelist within 5\u201310 business days of Technijian\u2019s request',
    'The 16 legacy desktops will be available for agent removal during Phase 4; physical disposal is AFFG\u2019s responsibility',
    'Users transitioning from legacy desktops to new managed endpoints are available during Weeks 5\u20138 for data transfer',
]:
    bullet_para(a)

# ── Section 7 ─────────────────────────────────────────────────────────────────
heading(1, '7. EXCLUSIONS')
for e in [
    'CloudBrink per-user subscription licensing (AFFG-procured)',
    'Hardware procurement (Mac Mini desktops and Windows laptops)',
    'Microsoft 365 license costs (AFFG procures directly)',
    'Technijian My Archive (AFFG operates own archiving platform)',
    'Company-owned mobile phones and MDM enrollment thereof',
    'Physical removal, disposal, or resale of the 16 legacy desktop devices',
    'Remediation beyond defined scope; additional work will be scoped as a separate Change Order',
]:
    bullet_para(e)

# ── Section 8 ─────────────────────────────────────────────────────────────────
heading(1, '8. MONTHLY COST IMPACT')
body_para(
    'The figures below reflect AFFG\u2019s actual signed invoice baseline (\u201cCurrent\u201d) and the '
    'projected recurring charges after full SOW-004 implementation (\u201cProposed\u201d). '
    'The net increase is a compliance investment \u2014 not a cost-reduction initiative. '
    'AFFG currently operates 16 endpoints with no endpoint DLP, no ZTNA, and no device-posture '
    'enforcement at custodian portals. The MyAudit UAM+DLP stack and CloudBrink ZTNA close gaps '
    'required under SEC Reg S-P (2024 amendments) and FINRA Rule 3110. The regulatory exposure '
    'associated with non-compliance materially exceeds the incremental monthly investment shown here.'
)
heading(2, '8.1 Current vs. Proposed Monthly Recurring')
body_para('Current Monthly (actual signed invoice \u2013 no VDI, no device compliance) = $2,794.50/mo')
add_table(
    ['Category', 'Current Monthly', 'Proposed Monthly', 'Change'],
    [
        ('Endpoint security \u2013 16 desktops (AVD/AVMH/MR/PMW/SI)', '$424.00', '$0.00', '\u2212$424.00'),
        ('Endpoint security \u2013 6 Windows laptops (AVD/AVMH/MR/PMW, no Umbrella)', '$0.00', '$123.00', '+$123.00'),
        ('Endpoint security \u2013 4 Mac Mini (AVD/AVMH/MR/PMMAC, no Umbrella)', '$0.00', '$110.00', '+$110.00'),
        ('CloudBrink ZTNA \u2013 10 endpoints @ $8.00 (replaces Cisco Umbrella)', '$0.00', '$80.00', '+$80.00'),
        ('Compliance stack add-ons \u2013 10 endpoints (see 8.2)', '$0.00', '$1,131.00', '+$1,131.00'),
        ('All-user services \u2013 31 users (V365 + PHT, unchanged)', '$310.00', '$310.00', '$0.00'),
        ('Domain / DKIM / RTPT (unchanged)', '$62.00', '$62.00', '$0.00'),
        ('Site Assessment \u2013 1 site (unchanged)', '$50.00', '$50.00', '$0.00'),
        ('Virtual Staff Support (unchanged)', '$1,948.50', '$1,948.50', '$0.00'),
        ('TOTAL MONTHLY', '$2,794.50', '$3,764.50', '+$970.00'),
    ],
    [3800, 1440, 1440, 1680]
)

heading(2, '8.2 Compliance Stack Add-Ons Detail (NEW \u2013 10 endpoints)')
add_table(
    ['Service', 'Code', 'Qty', 'Unit Price', 'Monthly'],
    [
        ('MyAudit UAM+DLP / 1-Year', 'AMDLP1Y', '10', '$108.10', '$1,081.00'),
        ('Site Assessment (Network Detective)', 'SA', '1 site', '$50.00', '$50.00'),
        ('Subtotal', '', '', '', '$1,131.00'),
    ],
    [3200, 1200, 640, 1440, 1880]
)
body_para(
    'Notes: (1) SSO/2FA Gateway product removed \u2014 Azure Entra ID SSO (included in M365 E3) provides '
    'identity federation and MFA natively. (2) CloudBrink ZTNA subscription is AFFG-procured; $80/mo shown '
    'above is the Technijian-billed component only. (3) Cisco Umbrella (SI) removed from security stack; '
    'CloudBrink ZTNA provides DNS filtering, zero-trust egress, and replaces the former VDI egress IP. '
    '(4) Mac Mini devices use PMMAC ($11.00/device) in place of PMW ($4.00/device).'
)

# ── Sections 9\u201311 ─────────────────────────────────────────────────────────────
heading(1, '9. CHANGE MANAGEMENT')
for num, text in [
    ('9.01.', 'Any changes to the scope, schedule, or pricing of this SOW must be documented in a written Change Order signed by both Parties before work on the change begins.'),
    ('9.02.', 'If Client requests work outside the defined scope, Technijian shall provide a Change Order detailing the additional work, estimated hours, and cost impact.'),
    ('9.03.', "Technijian shall not proceed with out-of-scope work without an approved Change Order, except in cases where delay would result in harm to Client\u2019s systems."),
]:
    clause_para(num, text)

heading(1, '10. ACCEPTANCE')
for num, text in [
    ('10.01.', 'Upon completion of each phase, Technijian shall notify Client in writing that the deliverables are ready for review.'),
    ('10.02.', 'Client shall review the deliverables and provide written acceptance or a detailed description of deficiencies within five (5) business days.'),
    ('10.03.', 'If Client does not respond within the review period, the deliverables shall be deemed accepted.'),
    ('10.04.', 'If deficiencies are identified, Technijian shall correct them and resubmit for review. This process shall repeat until acceptance is achieved.'),
]:
    clause_para(num, text)

heading(1, '11. GOVERNING TERMS')
clause_para('11.01.', 'This SOW is governed by the Master Service Agreement MSA-AFFG-2026 between the Parties. '
    'In the event of a conflict between this SOW and the MSA, the MSA shall prevail unless this SOW expressly states otherwise.')

# ── Signatures ────────────────────────────────────────────────────────────────
heading(1, 'SIGNATURES')
body_para('TECHNIJIAN, INC.', bold=True, color=DARK, before=6, after=0)
for line in ['By: ___________________________________',
             'Name: _________________________________',
             'Title: _________________________________',
             'Date: _________________________________']:
    body_para(line, before=12, after=0)

insert_para()
body_para('AMERICAN FUNDSTARS FINANCIAL GROUP LLC', bold=True, color=DARK, before=6, after=0)
for line in ['By: ___________________________________',
             'Name: _________________________________',
             'Title: _________________________________',
             'Date: _________________________________']:
    body_para(line, before=12, after=0)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DEST)
print('Saved:', DEST)
