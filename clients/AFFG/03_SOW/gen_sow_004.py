"""
SOW-AFFG-004 MASTER generator (single source of truth, no rev suffix).

Architecture (consolidated April 2026):
  - 10 Apple endpoints: 4 Mac Mini + 6 MacBook (model TBD by AFFG, all macOS, all enrolled via Apple Business Manager)
  - 10 personal mobile devices via Intune BYOD MAM (no company phones)
  - Endpoint stack per device: PMMAC + CrowdStrike (AVD) + Huntress (AVMH) + My Remote (MR) + CloudBrink ZTNA
  - DLP layer: MyAudit UAM+DLP (AMDLP1Y) on all 10 endpoints
  - Cisco Umbrella REMOVED (replaced by CloudBrink ZTNA for DNS + zero-trust egress)
  - SSO/2FA Gateway and Credential Manager REMOVED (Azure Entra ID provides identity + MFA via M365 E3/E5)
  - CloudBrink ZTNA per-user licensing is Technijian-billed at $8.00/user/month (NOT AFFG-procured)

Pricing model:
  - One-time SOW labor: 72 hrs / $5,340 (20 US @ $150 + 52 India @ $45)
  - Current monthly: $2,794.50 (signed MSA baseline, 16 desktops standard stack)
  - Proposed monthly: $3,806.50 (+$1,012/mo net compliance investment)

§12 of this SOW amends the prior Client Monthly Service Agreement (DocuSign envelope
F3CDFC05) to incorporate the 2026 MSA Framework (Sections 2-10) attached as Exhibit A.

Output: SOW-AFFG-004-Managed-Device-Migration.docx  (in-place master regeneration)

DocuSign anchors: 1pt white text /tSign/ /tName/ /tTitle/ /tDate/ /cSign/ /cName/ /cTitle/ /cDate/
embedded on each signature line (invisible to readers, used by send-docusign.ps1).
"""
import copy
import os
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'c:\vscode\tech-legal\tech-legal\clients\AFFG\03_SOW'
SRC  = os.path.join(BASE, 'SOW-AFFG-004-Managed-Device-Migration.docx')
TMP  = os.path.join(BASE, 'SOW-AFFG-004-Managed-Device-Migration.master.tmp.docx')
DEST = os.path.join(BASE, 'SOW-AFFG-004-Managed-Device-Migration.docx')

DARK  = '1A1A2E'
GREY  = '59595B'
BLUE  = '006DB6'
WHITE = 'FFFFFF'
ROW   = 'F8F9FA'
BRDR  = 'DDDDDD'
TOTAL = 'E8EFF5'
FONT  = 'Open Sans'

# ── Step 1: deepcopy title block from existing source ────────────────────────
src_doc = Document(SRC)
src_body_children = list(src_doc.element.body)
title_block = [copy.deepcopy(src_body_children[i]) for i in range(14)]

# Update date text in title block element [8]
date_elem = title_block[8]
for t_el in date_elem.findall('.//' + qn('w:t')):
    if t_el.text and '2026' in t_el.text:
        t_el.text = 'May 1, 2026'

# ── Step 2: clone source → temp, clear body ──────────────────────────────────
shutil.copy2(SRC, TMP)
doc = Document(TMP)
body = doc.element.body
final_sectPr = body.find(qn('w:sectPr'))
if final_sectPr is not None:
    final_sectPr = copy.deepcopy(final_sectPr)
for child in list(body):
    body.remove(child)

# ── Step 3: reinsert title block ─────────────────────────────────────────────
for elem in title_block:
    body.append(elem)
if final_sectPr is not None:
    body.append(final_sectPr)

# ── XML helpers ──────────────────────────────────────────────────────────────

def make_rPr(size_pt=11, bold=False, color=GREY, caps=False):
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    for attr in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        fonts.set(qn(f'w:{attr}'), FONT)
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    if caps:
        rPr.append(OxmlElement('w:caps'))
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    szC = OxmlElement('w:szCs')
    szC.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szC)
    return rPr

def add_run(p_elem, text, size_pt=11, bold=False, color=GREY):
    r = OxmlElement('w:r')
    r.append(make_rPr(size_pt, bold, color))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_elem.append(r)

def insert_para(elems=None):
    p = OxmlElement('w:p')
    if elems:
        for e in elems:
            p.append(e)
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        body.insert(list(body).index(sectPr), p)
    else:
        body.append(p)
    return p

def heading(level, text):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading{level}')
    pPr.append(pStyle)
    p.append(pPr)
    add_run(p, text, size_pt=11, color=GREY)
    return p

def body_para(text, before=0, after=6, bold=False, size_pt=11, color=GREY):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(before * 20)))
    sp.set(qn('w:after'), str(int(after * 20)))
    pPr.append(sp)
    p.append(pPr)
    add_run(p, text, size_pt=size_pt, bold=bold, color=color)
    return p

def body_para_with_anchor(visible_text, anchor, before=0, after=0):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(before * 20)))
    sp.set(qn('w:after'), str(int(after * 20)))
    pPr.append(sp)
    p.append(pPr)
    add_run(p, anchor, size_pt=1, color=WHITE)
    add_run(p, visible_text, size_pt=11, color=GREY)
    return p

def kv_para(label, value):
    p = insert_para()
    add_run(p, label, size_pt=11, bold=True, color=DARK)
    add_run(p, value, size_pt=11, color=GREY)
    return p

def bullet_para(text):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'ListParagraph')
    pPr.append(pStyle)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '2')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
    pPr.append(sp)
    p.append(pPr)
    add_run(p, text, size_pt=11, color=GREY)
    return p

def clause_para(num, text):
    p = insert_para()
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
    pPr.append(sp); p.append(pPr)
    add_run(p, num + '  ', size_pt=11, bold=True, color=DARK)
    add_run(p, text, size_pt=11, color=GREY)
    return p

# ── Table builder ────────────────────────────────────────────────────────────
TICKET_COLS = [1600, 4200, 1600, 960]
TICKET_HDR  = ['Ticket', 'Title', 'Assigned To', 'Est. Hours']

def make_cell(text, fill, bold, size_pt=10, width_dxa=None):
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    if width_dxa:
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_dxa)); tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:color'), BRDR); b.set(qn('w:sz'), '1')
        tcB.append(b)
    tcPr.append(tcB)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top','80'),('left','120'),('bottom','80'),('right','120')):
        m = OxmlElement(f'w:{side}'); m.set(qn('w:w'), val); m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp); p.append(pPr)
    r = OxmlElement('w:r')
    r.append(make_rPr(size_pt=size_pt, bold=bold,
                      color=WHITE if fill == BLUE else GREY))
    t_el = OxmlElement('w:t')
    t_el.set(qn('xml:space'), 'preserve'); t_el.text = str(text)
    r.append(t_el); p.append(r); tc.append(p)
    return tc

def add_table(headers, rows, col_dxa):
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(col_dxa))); tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_dxa:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tbl.append(tblGrid)
    hdr_tr = OxmlElement('w:tr')
    for i, h in enumerate(headers):
        hdr_tr.append(make_cell(h, BLUE, bold=True, width_dxa=col_dxa[i]))
    tbl.append(hdr_tr)
    for row_data in rows:
        is_total = (str(row_data[0]) == '' and 'Total' in str(row_data[1])) \
                   or str(row_data[0]) in ('TOTAL', 'Total', 'Subtotal')
        fill = TOTAL if is_total else ROW
        tr = OxmlElement('w:tr')
        for i, val in enumerate(row_data):
            tr.append(make_cell(val, fill, bold=is_total, width_dxa=col_dxa[i]))
        tbl.append(tr)
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        body.insert(list(body).index(sectPr), tbl)
    else:
        body.append(tbl)
    insert_para()

# ── CONTENT ──────────────────────────────────────────────────────────────────

heading(1, 'STATEMENT OF WORK')

for label, value in [
    ('SOW Number: ', 'SOW-AFFG-004'),
    ('Title: ', 'Managed Device Control Deployment'),
    ('Effective Date: ', 'May 1, 2026'),
    ('Parent Agreement: ', 'Client Monthly Service Agreement (DocuSign Envelope F3CDFC05-B156-8A6F-8020-13A164F4E3F1, effective March 11, 2026), as amended by Section 12 of this SOW'),
    ('Source: ', 'AFFG Managed Device Strategy (REF: AFFG-MDS-2026-04, April 2026)'),
    ('Primary Contact: ', 'Iris Liu  |  iris.liu@americanfundstars.com  |  949-439-2392'),
    ('Regulatory Scope: ', 'SEC Reg S-P (2024 Amendments), FINRA Rule 3110, FINRA Rule 4370, SEC Rule 17a-4, NIST SP 800-53 Rev 5'),
]:
    kv_para(label, value)

insert_para()

body_para(
    'This Statement of Work (“SOW”) is entered into by and between Technijian, Inc. (“Technijian”), '
    '18 Technology Drive, Suite 141, Irvine, California 92618 and American Fundstars Financial Group LLC '
    '(“Client” or “AFFG”), 1 Park Plaza, Suite 210, Irvine, California 92618. '
    'Primary Contact: Iris Liu.'
)

# ── Section 1 ────────────────────────────────────────────────────────────────
heading(1, '1. PROJECT OVERVIEW')
body_para(
    'AFFG currently operates 16 physical desktops under the Technijian standard managed services program. '
    'This SOW covers two concurrent workstreams: (1) right-sizing the managed device fleet from 16 desktops '
    'to 10 company-owned Apple endpoints (4 Mac Mini + 6 MacBook, all running macOS), and (2) deploying '
    'the compliance controls required under SEC Regulation S-P (2024 amendments) and FINRA Rule 3110 that are '
    'not present in the current environment. Compliance controls include Apple Business Manager (ABM) '
    'enrollment in Microsoft Intune with Conditional Access, CloudBrink Zero Trust Network Access (ZTNA) '
    'replacing Cisco Umbrella, MyAudit UAM+DLP on all 10 managed endpoints, and Intune MAM-only for employee '
    'personal mobile devices (BYOD).'
)

heading(2, '1.1 Objectives')
for obj in [
    'Enroll all 10 Apple endpoints (4 Mac Mini + 6 MacBook) in Microsoft Intune via Apple Business Manager (ABM) with full endpoint security stack',
    'Deploy CloudBrink ZTNA on all 10 endpoints (Technijian-billed per-user license) to replace Cisco Umbrella; whitelist CloudBrink egress IP at Schwab Advisor Services and IBKR',
    'Configure Entra ID Conditional Access: named location policy (office WAN IP) + Intune-compliant device requirement for M365; block legacy auth',
    'Deploy MyAudit UAM+DLP (AMDLP1Y) on all 10 managed macOS endpoints — block USB, clipboard, print, and local download of financial data',
    'Configure BYOD Intune MAM (Outlook App) for employee personal mobile devices — no MDM enrollment of personal devices',
    'Offboard 16 legacy desktops from Technijian managed services and amend Schedule A accordingly',
    'Conduct user training and deliver post-implementation compliance documentation',
]:
    bullet_para(obj)

heading(2, '1.2 Out of Scope')
for exc in [
    'Hardware procurement (AFFG provides Mac Mini desktops and MacBook laptops)',
    'Microsoft 365 license costs (AFFG procures directly — E3 or E5 required)',
    'Technijian My Archive (AFFG operates own email-archiving platform)',
    'Company-owned mobile phones and MDM enrollment thereof (BYOD model adopted)',
    'Physical removal or disposal of retired desktop hardware (AFFG responsibility)',
]:
    bullet_para(exc)

# ── Section 2 ────────────────────────────────────────────────────────────────
heading(1, '2. IMPLEMENTATION SCOPE')

heading(2, 'Phase 1: Endpoint Enrollment & Foundation (Weeks 1–4)')
body_para('7 tickets  |  24 hours')
body_para(
    'Configure Apple Business Manager (ABM) integration with Microsoft Intune. Define device compliance '
    'policies and configuration profiles for macOS. Enroll and image all 10 Apple endpoints (4 Mac Mini + '
    '6 MacBook). Deploy full security stack on each endpoint: CrowdStrike Falcon EDR, Huntress MDR, '
    'macOS Patch Management (PMMAC), My Remote (ScreenConnect), and CloudBrink ZTNA agent. Deploy '
    'MyAudit UAM+DLP on all 10 endpoints with Full Disk Access and Screen Recording permissions provisioned '
    'silently via Intune configuration profile.'
)
heading(3, 'Phase 1 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-001', 'Apple Business Manager (ABM) Setup & Intune Integration', 'CHD-TS1', '3'),
    ('AFFG-004-002', 'Intune macOS Compliance & Configuration Profiles (FDA + Screen Recording)', 'CHD-TS1', '3'),
    ('AFFG-004-003', 'ABM Enrollment & Imaging – 4 Mac Mini', 'CHD-TS1', '4'),
    ('AFFG-004-004', 'ABM Enrollment & Imaging – 6 MacBook', 'CHD-TS1', '4'),
    ('AFFG-004-005', 'Endpoint Security Stack Deployment (CrowdStrike, Huntress, PMMAC, MR) – 10 Endpoints', 'CHD-TS1', '4'),
    ('AFFG-004-006', 'CloudBrink ZTNA Agent Deployment – 10 Endpoints', 'CHD-TS1', '2'),
    ('AFFG-004-007', 'MyAudit UAM+DLP Deployment & Policy Configuration – 10 macOS Endpoints', 'CHD-TS1', '4'),
    ('', 'Phase Total', '', '24'),
], TICKET_COLS)

heading(2, 'Phase 2: Access Control & ZTNA (Weeks 3–6)')
body_para('5 tickets  |  16 hours')
body_para(
    'Configure Entra ID Conditional Access policies: office WAN IP named location, Intune-compliant device '
    'requirement for M365, legacy authentication block, and MFA enforcement. Provision CloudBrink ZTNA tenant '
    'and connector. Author per-application micro-segmentation policies. Whitelist CloudBrink egress IP at '
    'Schwab Advisor Services and IBKR as the sole permitted remote access origin. Remove Cisco Umbrella from '
    'all retiring desktops.'
)
heading(3, 'Phase 2 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-008', 'Conditional Access Policy Suite – Named Location + Compliant Device + MFA', 'IRV-TS1', '4'),
    ('AFFG-004-009', 'CloudBrink ZTNA Tenant & Connector Provisioning', 'IRV-TS1', '3'),
    ('AFFG-004-010', 'CloudBrink Access Policies & Per-App Micro-Segmentation', 'CHD-TS1', '3'),
    ('AFFG-004-011', 'Schwab & IBKR IP Whitelist Update – Office WAN IP + CloudBrink Egress IP', 'IRV-TS1', '3'),
    ('AFFG-004-012', 'Cisco Umbrella Removal from Legacy Desktops', 'CHD-TS1', '3'),
    ('', 'Phase Total', '', '16'),
], TICKET_COLS)

heading(2, 'Phase 3: BYOD Mobile MAM (Weeks 5–7)')
body_para('2 tickets  |  4 hours')
body_para(
    'Configure Intune App Protection (MAM) policies for employee personal mobile devices. No MDM enrollment '
    'of personal devices. Employees install the Outlook App from their personal app store; MAM policies '
    'containerize corporate data, enforce app PIN, and enable selective corporate data wipe on separation '
    'without accessing personal content. Full M365 browser access from mobile is blocked via Conditional '
    'Access unless on an Intune-managed company endpoint.'
)
heading(3, 'Phase 3 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-013', 'Intune MAM Policy Configuration – BYOD Outlook App (iOS & Android)', 'CHD-TS1', '2'),
    ('AFFG-004-014', 'MAM Policy Testing & Employee BYOD Setup Guide', 'CHD-TS1', '2'),
    ('', 'Phase Total', '', '4'),
], TICKET_COLS)

heading(2, 'Phase 4: Legacy Desktop Offboarding (Weeks 5–8)')
body_para('3 tickets  |  8 hours')
body_para(
    'Decommission the 16 legacy physical desktops from Technijian managed services. Remove all Technijian '
    'agents (CrowdStrike, Huntress, Patch Management, My Remote) from retiring devices. User data transfer '
    'assistance for any users moving from a retiring desktop to a new managed Apple endpoint. Draft and '
    'deliver Schedule A amendment to reflect the reduced fleet and updated service stack.'
)
heading(3, 'Phase 4 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-015', 'Agent Removal from 16 Legacy Desktops & Managed Services Offboarding', 'CHD-TS1', '4'),
    ('AFFG-004-016', 'User Data Transfer Assistance – Legacy Desktop to Apple Endpoint', 'CHD-TS1', '2'),
    ('AFFG-004-017', 'Schedule A Amendment – Fleet Right-Sizing & Compliance Stack Update', 'IRV-TS1', '2'),
    ('', 'Phase Total', '', '8'),
], TICKET_COLS)

heading(2, 'Phase 5: Validation & Training (Weeks 7–10)')
body_para('4 tickets  |  12 hours')
body_para(
    'End-to-end security and compliance validation across all 10 managed endpoints. Verify MyAudit DLP '
    'policy enforcement on macOS, CloudBrink ZTNA posture checks, Conditional Access enforcement, and MAM '
    'containment. User training on managed-device workflows and CloudBrink remote access. Compliance gap '
    'assessment documenting SEC Reg S-P and FINRA control coverage.'
)
heading(3, 'Phase 5 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-018', 'End-to-End Security & Compliance Validation – All 10 Endpoints', 'IRV-TS1', '4'),
    ('AFFG-004-019', 'MyAudit DLP & CloudBrink Policy Verification (macOS)', 'CHD-TS1', '3'),
    ('AFFG-004-020', 'User Training – Managed Endpoints & CloudBrink Remote Access', 'CHD-TS1', '3'),
    ('AFFG-004-021', 'Compliance Gap Assessment & Control Coverage Report (SEC/FINRA)', 'IRV-TS1', '2'),
    ('', 'Phase Total', '', '12'),
], TICKET_COLS)

heading(2, 'Phase 6: Documentation (Weeks 9–12)')
body_para('2 tickets  |  8 hours')
body_para(
    'Update compliance documentation suite to reflect the managed-device architecture: endpoint security '
    'policy, ZTNA access control runbook, incident response procedure, and data handling policy. '
    'All documents delivered in Technijian standard format.'
)
heading(3, 'Phase 6 Tickets')
add_table(TICKET_HDR, [
    ('AFFG-004-022', 'Update Compliance Documentation Suite – Managed Device Architecture', 'CHD-TS1', '6'),
    ('AFFG-004-023', 'Final Deliverable Package & Client Acceptance Sign-Off', 'IRV-TS1', '2'),
    ('', 'Phase Total', '', '8'),
], TICKET_COLS)

# ── Section 3 ────────────────────────────────────────────────────────────────
heading(1, '3. CONTROL-TO-CITATION MAP')
body_para('The following table maps each design component to its regulatory control objective:')
add_table(
    ['Design Component', 'Control Objective', 'Citation(s)'],
    [
        ('Intune-managed Apple endpoints – 4 Mac Mini + 6 MacBook',
         'Company-owned, encrypted (FileVault), policy-enforced devices with centralized MDM control via ABM',
         'Reg S-P 248.30(a)(1)–(3); FINRA 3110(a)'),
        ('CloudBrink ZTNA on all 10 endpoints (replaces Cisco Umbrella)',
         'Zero-trust egress; device posture verified per session; DNS filtering; blocks unmanaged device access to portals',
         'Reg S-P 248.30(a)(3); NIST AC-17, SC-7'),
        ('Office WAN IP + CloudBrink egress IP whitelist at Schwab & IBKR',
         'Custodian portals accessible only from office or CloudBrink-tunneled managed device — MFA alone does not restrict device origin',
         'Reg S-P 248.30(a)(3); NIST AC-3, SC-7'),
        ('Entra ID Conditional Access – Named Location + Compliant Device',
         'M365 access restricted to office WAN IP or Intune-enrolled device; legacy authentication blocked',
         'Reg S-P 248.30(a)(3); NIST AC-3, IA-2(1)'),
        ('Intune MAM – BYOD Outlook App (personal mobile)',
         'Corporate email containerized on personal device; selective corporate wipe on separation; no personal data access',
         'Reg S-P 248.30(a)(1)–(3); NIST AC-19, MP-6'),
        ('MyAudit UAM+DLP (AMDLP1Y) – all 10 managed macOS endpoints',
         'Block USB exfiltration, local downloads, print, clipboard copy, and screen capture of financial data',
         'Reg S-P 248.30(a)(3); NIST AC-19, MP-7'),
        ('Technijian Veeam Backup for M365',
         'Immutable, non-rewriteable backup of full M365 tenant (Exchange, Teams, SharePoint, OneDrive)',
         'SEC Rule 17a-4(b),(f); FINRA 4511'),
        ('Technijian monthly network assessment (SA)',
         'Detect configuration drift; provide attested evidence of ongoing control operation',
         'FINRA 3110(c); FINRA 3120'),
    ],
    [2100, 2900, 2360]
)

# ── Section 4 ────────────────────────────────────────────────────────────────
heading(1, '4. DELIVERABLES')
for d in [
    '10 Intune-managed company Apple endpoints (4 Mac Mini + 6 MacBook, all enrolled via ABM) with full security stack deployed and verified',
    'CloudBrink ZTNA deployed on all 10 endpoints (Technijian-billed per-user licensing) with per-application micro-segmented access and device posture enforcement',
    'Schwab Advisor Services and IBKR whitelisted for office WAN IP and CloudBrink egress IP; all other remote access blocked',
    'Entra ID Conditional Access enforcing M365 access: office WAN IP or Intune-compliant device required; legacy auth blocked',
    'MyAudit UAM+DLP active on all 10 managed macOS endpoints with DLP policy covering USB, clipboard, print, screen capture, and download',
    'Intune MAM policies for BYOD personal mobile devices (iOS and Android, Outlook App)',
    '16 legacy desktops offboarded from Technijian managed services with all agents removed',
    'Schedule A amendment reflecting updated fleet (10 Apple endpoints) and compliance service stack',
    'Updated compliance documentation suite (endpoint security policy, ZTNA runbook, IRP, data handling policy)',
    'Compliance gap assessment confirming SEC Reg S-P and FINRA control coverage on managed-device architecture',
]:
    bullet_para(d)

# ── Section 5 ────────────────────────────────────────────────────────────────
heading(1, '5. PRICING AND PAYMENT')
heading(2, '5.1 Labor Summary')
add_table(
    ['Role', 'Rate', 'Hours', 'Labor'],
    [
        ('US Tech Support (IRV-TS1)', '$150/hr', '20', '$3,000.00'),
        ('India Tech Support (CHD-TS1)', '$45/hr', '52', '$2,340.00'),
        ('TOTAL', '', '72 hrs', '$5,340.00'),
    ],
    [4000, 1400, 1200, 1760]
)

heading(2, '5.2 Payment Schedule')
add_table(
    ['Milestone', 'Trigger', 'Amount'],
    [
        ('50% upon SOW execution', 'SOW signed by both parties', '$2,670.00'),
        ('25% upon Phase 4 completion', '16 legacy desktops offboarded; Schedule A amendment delivered', '$1,335.00'),
        ('25% upon Phase 6 delivery', 'Compliance documentation + client acceptance sign-off', '$1,335.00'),
        ('Total', '', '$5,340.00'),
    ],
    [3000, 4200, 1160]
)

heading(2, '5.3 Payment Terms')
body_para(
    'All invoices are due and payable within thirty (30) days of the invoice date. '
    'Late payments are subject to a late fee of 1.5% per month on the unpaid balance.'
)

# ── Section 6 ────────────────────────────────────────────────────────────────
heading(1, '6. ASSUMPTIONS')
for a in [
    'AFFG has procured or will procure 4 Mac Mini desktops and 6 MacBook laptops prior to Phase 1 commencement (subject to Apple delivery timeline)',
    'AFFG will enroll all Apple endpoints in Apple Business Manager (ABM) prior to Phase 1; Technijian will provide setup assistance',
    'CloudBrink ZTNA per-user licensing is provided by Technijian as part of the monthly endpoint stack (no separate AFFG procurement required)',
    "AFFG’s M365 tenant is licensed at E3 or E5 (Intune, Conditional Access, and Entra ID P1/P2 included)",
    'AFFG employees will use personal mobile devices (BYOD) for corporate email via Outlook App; no company mobile phones',
    'AFFG provides the office static WAN IP address to Technijian prior to Phase 2 commencement',
    'Schwab Advisor Services and IBKR can update their IP whitelist within 5–10 business days of Technijian’s request',
    'The 16 legacy desktops will be available for agent removal during Phase 4; physical disposal is AFFG’s responsibility',
    'Users transitioning from legacy desktops to new managed Apple endpoints are available during Weeks 5–8 for data transfer',
]:
    bullet_para(a)

# ── Section 7 ────────────────────────────────────────────────────────────────
heading(1, '7. EXCLUSIONS')
for e in [
    'Hardware procurement (Mac Mini desktops and MacBook laptops)',
    'Microsoft 365 license costs (AFFG procures directly)',
    'Technijian My Archive (AFFG operates own archiving platform)',
    'Company-owned mobile phones and MDM enrollment thereof',
    'Physical removal, disposal, or resale of the 16 legacy desktop devices',
    'Remediation beyond defined scope; additional work will be scoped as a separate Change Order',
]:
    bullet_para(e)

# ── Section 8 ────────────────────────────────────────────────────────────────
heading(1, '8. MONTHLY COST IMPACT')
body_para(
    'The figures below reflect AFFG’s actual signed invoice baseline (“Current”) and the '
    'projected recurring charges after full SOW-004 implementation (“Proposed”). '
    'The net increase is a compliance investment — not a cost-reduction initiative. '
    'AFFG currently operates 16 endpoints with no endpoint DLP, no ZTNA, and no device-posture '
    'enforcement at custodian portals. The MyAudit UAM+DLP stack and CloudBrink ZTNA close gaps '
    'required under SEC Reg S-P (2024 amendments) and FINRA Rule 3110. The regulatory exposure '
    'associated with non-compliance materially exceeds the incremental monthly investment shown here.'
)
heading(2, '8.1 Current vs. Proposed Monthly Recurring')
body_para('Current Monthly (actual signed invoice – no VDI, no device compliance) = $2,794.50/mo')
add_table(
    ['Category', 'Current Monthly', 'Proposed Monthly', 'Change'],
    [
        ('Endpoint security – 16 desktops (AVD/AVMH/MR/PMW/SI)', '$424.00', '$0.00', '−$424.00'),
        ('Apple endpoint stack – 10 endpoints (PMMAC/AVD/AVMH/MR + CloudBrink)', '$0.00', '$355.00', '+$355.00'),
        ('Compliance stack – 10 endpoints (MyAudit AMDLP1Y)', '$0.00', '$1,081.00', '+$1,081.00'),
        ('All-user services – 31 users (V365 + PHT, unchanged)', '$310.00', '$310.00', '$0.00'),
        ('Domain / DKIM / RTPT (unchanged)', '$62.00', '$62.00', '$0.00'),
        ('Site Assessment – 1 site (unchanged)', '$50.00', '$50.00', '$0.00'),
        ('Virtual Staff Support (unchanged)', '$1,948.50', '$1,948.50', '$0.00'),
        ('TOTAL MONTHLY', '$2,794.50', '$3,806.50', '+$1,012.00'),
    ],
    [3800, 1440, 1440, 1680]
)

heading(2, '8.2 Apple Endpoint Stack Detail (10 endpoints)')
add_table(
    ['Service', 'Code', 'Qty', 'Unit Price', 'Monthly'],
    [
        ('macOS Patch Management', 'PMMAC', '10', '$11.00', '$110.00'),
        ('CrowdStrike Falcon EDR', 'AVD', '10', '$8.50', '$85.00'),
        ('Huntress Managed Detection', 'AVMH', '10', '$6.00', '$60.00'),
        ('My Remote (ScreenConnect)', 'MR', '10', '$2.00', '$20.00'),
        ('CloudBrink ZTNA (Technijian-billed)', 'CB-ZTNA', '10', '$8.00', '$80.00'),
        ('Subtotal – Apple endpoint stack', '', '10', '$35.50', '$355.00'),
        ('MyAudit UAM+DLP / 1-Year', 'AMDLP1Y', '10', '$108.10', '$1,081.00'),
        ('Subtotal – with DLP', '', '', '', '$1,436.00'),
    ],
    [3200, 1200, 640, 1440, 1880]
)
body_para(
    'Notes: (1) SSO/2FA Gateway and Credential Manager removed — Azure Entra ID SSO (included in M365 E3) '
    'provides identity federation, MFA, and credential vault natively. (2) CloudBrink ZTNA per-user licensing '
    'is Technijian-billed at $8.00/user/month and is included in the proposed monthly recurring above; AFFG '
    'does not need to procure CloudBrink subscriptions separately. (3) Cisco Umbrella (SI) is removed from the '
    'security stack — CloudBrink ZTNA replaces it for DNS filtering and zero-trust egress. (4) macOS endpoints '
    'use PMMAC ($11.00/device) in place of PMW ($4.00/device).'
)

# ── Sections 9–11 ───────────────────────────────────────────────────────────
heading(1, '9. CHANGE MANAGEMENT')
for num, text in [
    ('9.01.', 'Any changes to the scope, schedule, or pricing of this SOW must be documented in a written Change Order signed by both Parties before work on the change begins.'),
    ('9.02.', 'If Client requests work outside the defined scope, Technijian shall provide a Change Order detailing the additional work, estimated hours, and cost impact.'),
    ('9.03.', "Technijian shall not proceed with out-of-scope work without an approved Change Order, except in cases where delay would result in harm to Client’s systems."),
]:
    clause_para(num, text)

heading(1, '10. ACCEPTANCE')
for num, text in [
    ('10.01.', 'Upon completion of each phase, Technijian shall notify Client in writing that the deliverables are ready for review.'),
    ('10.02.', 'Client shall review the deliverables and provide written acceptance or a detailed description of deficiencies within five (5) business days.'),
    ('10.03.', 'If Client does not respond within the review period, the deliverables shall be deemed accepted.'),
    ('10.04.', 'If deficiencies are identified, Technijian shall correct them and resubmit for review. This process shall repeat until acceptance is achieved.'),
]:
    clause_para(num, text)

heading(1, '11. GOVERNING TERMS')
clause_para('11.01.', 'This SOW is governed by the Client Monthly Service Agreement between the Parties executed via DocuSign on March 11, 2026 '
    '(DocuSign Envelope F3CDFC05-B156-8A6F-8020-13A164F4E3F1) (the “Original Agreement”), as amended by Section 12 of this SOW. '
    'In the event of a conflict between this SOW (including Section 12 and Exhibit A) and the Original Agreement, this SOW shall prevail.')

# ── Section 12 ───────────────────────────────────────────────────────────────
heading(1, '12. INCORPORATION OF 2026 MSA FRAMEWORK (AMENDMENT TO ORIGINAL AGREEMENT)')
for num, text in [
    ('12.01.', 'Purpose. This Section 12 is a mutual written amendment to the Original Agreement, executed pursuant to California Civil Code § 1698(a) (“A contract in writing may be modified by a contract in writing”). The Parties enter into this amendment to update the Original Agreement to reflect Technijian’s 2026 Master Service Agreement framework, which contains additional legal and data-protection provisions appropriate to AFFG’s status as a dually registered Investment Adviser and Broker-Dealer subject to SEC Regulation S-P (including the 2024 Amendments), FINRA Rule 3110, FINRA Rule 4370, and SEC Rule 17a-4.'),
    ('12.02.', 'Incorporation by Reference. Upon execution of this SOW by both Parties, Sections 2 through 10 of Technijian’s 2026 Master Service Agreement framework (the “MSA Framework Provisions”), the full text of which is set forth in Exhibit A attached hereto and incorporated herein by this reference, shall apply to the Original Agreement and to this SOW as if fully set forth in the Original Agreement.'),
    ('12.03.', 'Addition and Substitution. The MSA Framework Provisions supplement the Original Agreement and, to the extent of any conflict with a corresponding provision of the Original Agreement, supersede such corresponding provision. The late-fee rate of 1.5% per month stated in the Original Agreement is consistent with Section 3.04 of the MSA Framework Provisions and remains unchanged. The Virtual Staff Support hours and rates under the Original Agreement (CTO Advisory 3 hrs @ $225, US Tech Support 6 hrs @ $125, India Tech Normal 22 hrs @ $15, India Tech After-Hours 8 hrs @ $30) remain UNCHANGED and carry forward. Section 9.01 (Entire Agreement) of the MSA Framework Provisions is hereby modified, for purposes of this amendment only, to provide that the Original Agreement is amended (not superseded) by this Section 12, and that the monthly service charges under the Original Agreement shall be updated only as expressly provided in Section 8 of this SOW.'),
    ('12.04.', 'Effect on Services. Upon execution, all services provided by Technijian to Client — including the recurring services under Schedule A of the Original Agreement, the services described in this SOW, and any future Statement of Work executed between the Parties — shall be governed by the Original Agreement as amended by this Section 12. Future Statements of Work may reference the “Parent Agreement” as “Client Monthly Service Agreement as amended by SOW-AFFG-004 Section 12.”'),
    ('12.05.', 'No Other Changes. Except as expressly modified by this Section 12 and Section 8 of this SOW, the Original Agreement remains in full force and effect. Services already rendered, fees already paid, and rights already accrued under the Original Agreement are not affected.'),
    ('12.06.', 'Severability. If any provision of this Section 12 or of Exhibit A is held invalid or unenforceable by a court of competent jurisdiction, the remaining provisions of the Original Agreement (as amended) and this SOW shall continue in full force and effect.'),
    ('12.07.', 'Effective Date of Amendment. This amendment is effective upon execution of this SOW by authorized representatives of both Parties, and the MSA Framework Provisions shall govern services from and after the Effective Date stated in the SOW header.'),
]:
    clause_para(num, text)

# ── Signatures ───────────────────────────────────────────────────────────────
heading(1, 'SIGNATURES')
body_para('TECHNIJIAN, INC.', bold=True, color=DARK, before=6, after=0)
body_para_with_anchor('By: ___________________________________',   '/tSign/',  before=12, after=0)
body_para_with_anchor('Name: _________________________________',   '/tName/',  before=12, after=0)
body_para_with_anchor('Title: _________________________________',  '/tTitle/', before=12, after=0)
body_para_with_anchor('Date: _________________________________',   '/tDate/',  before=12, after=0)

insert_para()
body_para('AMERICAN FUNDSTARS FINANCIAL GROUP LLC', bold=True, color=DARK, before=6, after=0)
body_para_with_anchor('By: ___________________________________',   '/cSign/',  before=12, after=0)
body_para_with_anchor('Name: _________________________________',   '/cName/',  before=12, after=0)
body_para_with_anchor('Title: _________________________________',  '/cTitle/', before=12, after=0)
body_para_with_anchor('Date: _________________________________',   '/cDate/',  before=12, after=0)

# ── Exhibit A: 2026 MSA Framework Provisions ─────────────────────────────────
insert_para()
heading(1, 'EXHIBIT A — 2026 MSA FRAMEWORK PROVISIONS INCORPORATED BY REFERENCE')
body_para(
    'The following provisions (Sections 2 through 10) are the “MSA Framework Provisions” incorporated '
    'into the Original Agreement by Section 12 of this SOW. This Exhibit A is an integral part of this SOW '
    'and of the Original Agreement as amended. Capitalized terms not defined herein have the meanings given '
    'in the Original Agreement or the SOW body.'
)

heading(2, 'SECTION 2 — TERM AND RENEWAL')
for num, text in [
    ('2.01.', 'Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the “Initial Term”).'),
    ('2.02.', 'Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a “Renewal Term”), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term. Technijian shall send Client a written renewal reminder at least thirty (30) days prior to each renewal date.'),
    ('2.03.', 'Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party.'),
    ('2.04.', 'Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party (a) commits a material breach of this Agreement and fails to cure within thirty (30) days after receiving written notice; or (b) becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.'),
    ('2.05.', 'Effect of Termination. (a) Client shall pay all fees and charges for services rendered through the date of termination, including any remaining obligations for annual licenses and subscriptions procured on Client’s behalf, and any unpaid balance of contracted Virtual Staff Support hours actually worked in excess of contracted levels. (b) Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, subject to payment of applicable fees. (c) Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format, provided Client is not in breach of this Agreement. (d) The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 9.09 (Personnel Transition Fee), and Section 10 (Data Protection).'),
]:
    clause_para(num, text)

heading(2, 'SECTION 3 — PAYMENT')
for num, text in [
    ('3.01.', 'Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.'),
    ('3.02.', 'Invoice Types. Client may receive Monthly Service, Weekly In-Contract, Weekly Out-of-Contract, Equipment and Materials, and Project Milestone invoices. Each invoice will identify its type, the applicable Schedule or SOW, and the billing period or delivery event.'),
    ('3.03.', 'Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified in the applicable Schedule or SOW.'),
    ('3.04.', 'Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated as simple interest from the date payment was due until the date payment is received in full.'),
    ('3.05.', 'Disputed Invoices. Weekly invoices may be disputed within thirty (30) days of the invoice date with specific ticket numbers and basis. All other invoices may be disputed within fifteen (15) days. Failure to provide timely written dispute notice shall constitute acceptance of the invoice. Undisputed amounts remain payable by the due date.'),
    ('3.06.', 'Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend services until payment is received. Recurring fees for suspended services shall continue to accrue for a period not to exceed thirty (30) days following suspension.'),
    ('3.07.', 'Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian’s income.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 4 — CONFIDENTIALITY')
for num, text in [
    ('4.01.', 'Definition. “Confidential Information” means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information.'),
    ('4.02.', 'Obligations. Each Party shall hold the other Party’s Confidential Information in confidence using at least reasonable care; not disclose to third parties without prior written consent except to employees, agents, and subcontractors bound by equivalent obligations; and not use Confidential Information for any purpose other than performing obligations under this Agreement.'),
    ('4.03.', 'Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known prior to disclosure, is independently developed, or is received from a third party without restriction.'),
    ('4.04.', 'Compelled Disclosure. If required by law to disclose, the receiving Party shall provide prompt written notice and cooperate in seeking a protective order.'),
    ('4.05.', 'Duration. Confidentiality obligations shall survive termination for a period of three (3) years.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 5 — LIMITATION OF LIABILITY')
for num, text in [
    ('5.01.', 'EXCEPT AS PROVIDED IN SECTION 5.03, NEITHER PARTY’S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM (THE “STANDARD CAP”).'),
    ('5.02.', 'EXCEPT AS PROVIDED IN SECTION 5.03, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL.'),
    ('5.03.', 'ENHANCED CAP. FOR CLAIMS ARISING FROM BREACHES OF SECTION 4 (CONFIDENTIALITY), SECTION 10 (DATA PROTECTION), INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, THE TOTAL AGGREGATE LIABILITY OF THE RESPONSIBLE PARTY SHALL NOT EXCEED THREE (3) TIMES THE STANDARD CAP (THE “ENHANCED CAP”). THE ENHANCED CAP SHALL NOT APPLY TO LIABILITY ARISING FROM A PARTY’S WILLFUL AND INTENTIONAL MISAPPROPRIATION OF THE OTHER PARTY’S CONFIDENTIAL INFORMATION OR CLIENT DATA, FOR WHICH LIABILITY SHALL BE UNCAPPED.'),
    ('5.04.', 'Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data, Client acknowledges that Client is solely responsible for maintaining backup copies of its data.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 6 — INDEMNIFICATION')
for num, text in [
    ('6.01.', 'By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from Technijian’s gross negligence or willful misconduct in performing the services.'),
    ('6.02.', 'By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from Client’s use of the services in violation of applicable law, Client’s breach of this Agreement, or any data, content, or materials provided by Client.'),
    ('6.03.', 'Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party’s prior written consent.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 7 — INTELLECTUAL PROPERTY')
for num, text in [
    ('7.01.', 'Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes.'),
    ('7.02.', 'Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property.'),
    ('7.03.', 'Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 8 — DISPUTE RESOLUTION')
for num, text in [
    ('8.01.', 'Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations for a period of thirty (30) days.'),
    ('8.02.', 'Mediation. If not resolved, the Parties shall submit the dispute to mediation in Orange County, California, for a period not to exceed sixty (60) days.'),
    ('8.03.', 'Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association in Orange County, California.'),
    ('8.04.', 'Injunctive Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief to prevent irreparable harm.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 9 — GENERAL PROVISIONS')
for num, text in [
    ('9.01.', 'Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties. As modified by Section 12.03 of SOW-AFFG-004, this Section 9.01 does not supersede the Original Agreement; rather, the Original Agreement is amended by SOW-AFFG-004 Section 12 to incorporate the MSA Framework Provisions.'),
    ('9.02.', 'Amendment. This Agreement may only be amended by a written instrument signed by both Parties.'),
    ('9.03.', 'Severability. If any provision is found invalid or unenforceable, the remaining provisions shall continue in full force.'),
    ('9.04.', 'Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party.'),
    ('9.05.', 'Assignment. Neither Party may assign this Agreement without prior written consent, except in connection with a merger, acquisition, or sale of substantially all assets.'),
    ('9.06.', 'Force Majeure. Neither Party shall be liable for delays or failures in performance caused by events beyond its reasonable control. The affected Party shall notify the other Party within five (5) business days of becoming aware of a Force Majeure Event. Payment obligations are not excused. If a Force Majeure Event prevents performance for more than ninety (90) consecutive days, either Party may terminate the affected SOW upon fifteen (15) days written notice without early termination fees.'),
    ('9.07.', 'Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier.'),
    ('9.08.', 'Governing Law. This Agreement shall be governed by the laws of the State of California.'),
    ('9.09.', 'Personnel Transition Fee. If either Party hires (whether as an employee or independent contractor) any individual who was an employee of the other Party and who was directly involved in performing or receiving services under this Agreement, and such hiring occurs during the term of this Agreement or within twelve (12) months following termination, the hiring Party shall pay the other Party a personnel transition fee equal to 25% of the hired individual’s first-year annual compensation. This fee represents a reasonable estimate of recruiting and training costs and is not a restraint on trade or employment. This Section does not restrict any individual’s right to seek or obtain employment, and shall not apply to individuals who respond to general public job postings or are referred by a third-party recruiting firm without the hiring Party’s direction to target the other Party’s employees.'),
    ('9.10.', 'Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.'),
]:
    clause_para(num, text)

heading(2, 'SECTION 10 — DATA PROTECTION')
for num, text in [
    ('10.01.', 'CCPA/CPRA Compliance. To the extent Technijian processes, stores, or has access to personal information (as defined under the CCPA/CPRA, Cal. Civ. Code § 1798.100 et seq.) on behalf of Client, Technijian acts as a “service provider” as defined under Cal. Civ. Code § 1798.140(ag) and shall: (a) process personal information only as necessary to perform the services and per Client’s documented instructions; (b) not sell, share, retain, use, or disclose personal information for any purpose other than performing the services; (c) not combine personal information received from Client with information from other sources; (d) implement reasonable security measures per Section 10.03; (e) cooperate with Client in responding to verifiable consumer rights requests within ten (10) business days; (f) notify Client within five (5) business days if Technijian can no longer meet its service-provider obligations; (g) ensure subcontractors are bound by equivalent obligations; (h) permit Client to audit data processing practices once per year, or provide a SOC 2 Type II summary; and (i) certify compliance with this Section 10.'),
    ('10.02.', 'Security Incident Notification. If Technijian becomes aware of a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Client Data (“Security Incident”), Technijian shall: (a) notify Client in writing without unreasonable delay and in no event later than forty-eight (48) hours after becoming aware of the Security Incident; (b) provide Client with sufficient information to enable Client to comply with California Civil Code § 1798.82, applicable SEC Regulation S-P customer notification requirements, and FINRA Rule 4530 reporting obligations; (c) cooperate with Client’s investigation; and (d) take reasonable steps to contain and remediate the Security Incident.'),
    ('10.03.', 'Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards designed to protect Client Data from unauthorized access, use, or disclosure, consistent with industry standards for managed IT service providers. Such safeguards shall include, at a minimum: (a) encryption of Client Data in transit and at rest; (b) access controls limiting access to authorized personnel; (c) regular security assessments and vulnerability testing; and (d) employee security awareness training.'),
    ('10.04.', 'Regulatory Compliance. Client is a dually registered Investment Adviser and Broker-Dealer subject to SEC Regulation S-P (including the 2024 Amendments), FINRA Rule 4370, FINRA Rule 3110, and SEC Rule 17a-4. The Parties shall cooperate in good faith to implement the controls described in the applicable Schedules and SOWs to support Client’s compliance with these requirements.'),
    ('10.05.', 'Data Return and Deletion. Upon termination of this Agreement or upon Client’s written request, Technijian shall securely delete or return all Client Data in its possession within thirty (30) days, using methods consistent with NIST SP 800-88 or equivalent standards, and shall certify such deletion in writing upon request.'),
]:
    clause_para(num, text)

body_para('END OF EXHIBIT A', bold=True, color=DARK, before=12)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(TMP)

# Atomic move temp → master
if os.path.exists(DEST):
    os.remove(DEST)
os.rename(TMP, DEST)
print('Saved:', DEST)
