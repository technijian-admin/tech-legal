"""
AFFG Document Generator
Generates the full MSA document suite for American Fundstars Financial Group LLC:
  1. MSA-AFFG.docx                          — Master Service Agreement
  2. Schedule-A-AFFG.docx                   — Monthly Managed Services ($6,922.35/mo)
  3. Schedule-B-AFFG.docx                   — Subscription and License Services (placeholder)
  4. Schedule-C-AFFG.docx                   — Rate Card
  5. SOW-AFFG-003-Compliance-Strategy.docx  — 7-phase VDI + compliance implementation

The new MSA supersedes the prior Client Monthly Service Agreement dated 3/11/2026
(DocuSign Envelope F3CDFC05-B156-8A6F-8020-13A164F4E3F1).

Virtual Staff Support hours and rates from the prior MSA carry forward UNCHANGED
into Schedule A, Part 5 (CTO Advisory 3 hrs, US Tech 6 hrs, India Normal 22 hrs,
India After-Hours 8 hrs).
"""
import csv
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# -- Brand Colors (Technijian Brand Guide 2026) --
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "technijian-logo.png")
TICKETS_CSV = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SOW-AFFG-003-Tickets.csv")
OUT_DIR = os.path.dirname(os.path.abspath(__file__))

CLIENT_FULL = "American Fundstars Financial Group LLC"
CLIENT_SHORT = "AFFG"
CLIENT_UPPER = "AMERICAN FUNDSTARS FINANCIAL GROUP LLC"
CLIENT_ADDR_1 = "1 Park Plaza, Suite 210"
CLIENT_ADDR_2 = "Irvine, California 92618"
AGREEMENT_NO = "MSA-AFFG-2026"
EFFECTIVE_DATE = "May 1, 2026"
SIGNER = "[AUTHORIZED SIGNER]"
PRIOR_ENVELOPE = "F3CDFC05-B156-8A6F-8020-13A164F4E3F1"

SUPERSEDE_CLAUSE = (
    "This Agreement supersedes the prior Client Monthly Service Agreement between "
    "the Parties (DocuSign Envelope " + PRIOR_ENVELOPE + ", effective 3/11/2026). "
    "Upon the Effective Date of this Agreement, the monthly charges for workstation "
    "services, endpoint security, user services, domain/site/IP services, and shared "
    "storage set forth in Schedule A shall REPLACE (not be billed in addition to) the "
    "corresponding charges under the prior Monthly Service Agreement. The Virtual "
    "Staff Support hours and rates under the prior Monthly Service Agreement "
    "(CTO Advisory 3 hrs @ $225, US Tech Support 6 hrs @ $125, India Tech Normal "
    "22 hrs @ $15, India Tech After-Hours 8 hrs @ $30) remain UNCHANGED and carry "
    "forward into Schedule A, Part 5 of this Agreement at the same rates and "
    "contracted hour levels."
)


# ================================================================
# Low-level helpers
# ================================================================
def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


def set_grey_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


# ================================================================
# BrandedDoc — copied verbatim from AAVA/OKL pattern
# ================================================================
class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        return r

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_borders(tbl)

    def cover_page(self, title, subtitle, client_name=CLIENT_FULL, date=EFFECTIVE_DATE):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(LOGO_PATH, width=Inches(3.0))
        div = self.doc.add_table(rows=1, cols=3)
        div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in div.rows[0].cells:
            c.text = ""
            c.paragraphs[0].paragraph_format.space_before = Pt(0)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div.rows[0].cells[1], HEX_ORANGE)
        div.rows[0].cells[0].width = Inches(2.25)
        div.rows[0].cells[1].width = Inches(2.5)
        div.rows[0].cells[2].width = Inches(2.25)
        remove_borders(div)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Prepared for ", size=12, color=BRAND_GREY)
        self.run(p, client_name, size=12, bold=True, color=CORE_BLUE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.accent_bar(HEX_ORANGE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar = tbl.rows[0].cells[0]
        set_cell_shading(bar, HEX_BLUE)
        bar.width = Inches(0.08)
        bar.text = ""
        bar.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar.paragraphs[0].paragraph_format.space_after = Pt(0)
        tc = tbl.rows[0].cells[1]
        tc.text = ""
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_borders(tbl)

    def part_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_borders(tbl)

    def styled_table(self, headers, rows, total_indices=None):
        total_indices = total_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            set_cell_shading(c, HEX_BLUE)
            c.text = ""
            self.run(c.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for ri, rd in enumerate(rows):
            row = tbl.add_row()
            is_total = ri in total_indices
            for i, val in enumerate(rd):
                c = row.cells[i]
                c.text = ""
                if is_total:
                    set_cell_shading(c, HEX_OFF_WHITE)
                    self.run(c.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(c.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        set_grey_borders(tbl)
        return tbl

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            self.run(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, num, text, indent=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent
        self.run(p, f"{prefix}{num} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self, client_name=CLIENT_UPPER, contact_name=SIGNER):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)
        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", "Name: _________________________________",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)
        self.spacer(12)
        p = self.doc.add_paragraph()
        self.run(p, client_name, size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", f"Name: {contact_name}",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        self.doc.save(filename)
        size = os.path.getsize(filename)
        print(f"Created {filename}  ({size:,} bytes)")


# ================================================================
# MSA
# ================================================================
def build_msa():
    d = BrandedDoc()
    d.cover_page(
        "Master Service Agreement",
        CLIENT_FULL,
        client_name=CLIENT_FULL,
        date=EFFECTIVE_DATE,
    )
    d.page_break()

    d.part_header("MASTER SERVICE AGREEMENT")
    d.spacer()
    d.body(f"Agreement Number: {AGREEMENT_NO}")
    d.body(f"Effective Date: {EFFECTIVE_DATE}")
    d.spacer()
    d.body("This Master Service Agreement (\u201cAgreement\u201d) is entered into by and between:")
    d.spacer()
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", "")
    d.body("18 Technology Drive, Suite 141")
    d.body("Irvine, California 92618")
    d.spacer()
    d.body("and")
    d.spacer()
    d.body_bold(f"{CLIENT_FULL} (\u201c{CLIENT_SHORT}\u201d or \u201cClient\u201d)", "")
    d.body(CLIENT_ADDR_1)
    d.body(CLIENT_ADDR_2)
    d.body(f"Primary Contact: {SIGNER}")
    d.spacer()
    d.body("(collectively, the \u201cParties\u201d)")

    # SECTION 1
    d.spacer()
    d.section_header("SECTION 1 \u2014 SCOPE OF SERVICES")
    d.numbered("1.01.", "Services. Technijian shall provide the services described in the Schedules attached to this Agreement, which are incorporated herein by reference:")
    d.bullet("Schedule A \u2014 Monthly Managed Services (VDI, Endpoint Security, User Services, Domain/Site/IP, Shared Storage, and Virtual Staff Support)")
    d.bullet("Schedule B \u2014 Subscription and License Services (none in effect as of Effective Date)")
    d.bullet("Schedule C \u2014 Rate Card")
    d.body("Additional services may be provided through Statements of Work (\u201cSOWs\u201d) executed under this Agreement, including SOW-AFFG-003 (IT Compliance & VDI Implementation).")

    d.numbered("1.02.", "Standard of Care. Technijian shall perform all services in a professional and workmanlike manner, consistent with industry standards for managed IT service providers.")
    d.numbered("1.03.", "Service Level Agreement. The service levels applicable to the services are set forth in Schedule A. Technijian shall use commercially reasonable efforts to meet the service levels described therein.")
    d.numbered("1.04.", "Client Responsibilities. Client shall:")
    d.numbered("(a)", "Provide Technijian with reasonable access to Client\u2019s systems, facilities, and personnel as necessary for Technijian to perform the services;", indent=1)
    d.numbered("(b)", "Designate a primary point of contact for communications with Technijian;", indent=1)
    d.numbered("(c)", "Maintain current and accurate information regarding Client\u2019s systems and infrastructure;", indent=1)
    d.numbered("(d)", "Comply with all applicable laws and regulations in connection with its use of the services, including FINRA, SEC, and other securities-industry requirements; and", indent=1)
    d.numbered("(e)", "Be solely responsible for the security and management of Client\u2019s account credentials and passwords.", indent=1)
    d.numbered("1.05.", "Independent Contractor. Technijian is an independent contractor. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the Parties.")

    # SECTION 2
    d.spacer()
    d.section_header("SECTION 2 \u2014 TERM AND RENEWAL")
    d.numbered("2.01.", "Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the \u201cInitial Term\u201d).")
    d.numbered("2.02.", "Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a \u201cRenewal Term\u201d), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term. Technijian shall send Client a written renewal reminder at least thirty (30) days prior to each renewal date, which shall restate the auto-renewal terms and cancellation method.")
    d.numbered("2.03.", "Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party.")
    d.numbered("2.04.", "Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party:")
    d.numbered("(a)", "Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days after receiving written notice of the breach; or", indent=1)
    d.numbered("(b)", "Becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.", indent=1)

    d.numbered("2.05.", "Effect of Termination.")
    d.numbered("(a)", "Upon termination, Client shall pay all fees and charges for services rendered through the date of termination, including any remaining obligations for annual licenses and subscriptions procured on Client\u2019s behalf, and any unpaid balance of contracted Virtual Staff Support hours actually worked in excess of contracted levels.", indent=1)
    d.numbered("(b)", "Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, subject to payment of applicable fees.", indent=1)
    d.numbered("(c)", "Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format, provided Client is not in breach of this Agreement.", indent=1)
    d.numbered("(d)", "The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 9.09 (Personnel Transition Fee), and Section 10 (Data Protection).", indent=1)

    d.page_break()

    # SECTION 3
    d.section_header("SECTION 3 \u2014 PAYMENT")
    d.numbered("3.01.", "Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.")
    d.numbered("3.02.", "Invoice Types. Client may receive the following types of invoices from Technijian during the term of this Agreement. Each invoice will clearly identify its type, the applicable Schedule or SOW, and the billing period or delivery event.")
    d.numbered("(a)", "Monthly Service Invoice. Issued on the first business day of each month for recurring managed services under Schedule A. Billed in advance for the upcoming month.", indent=1)
    d.numbered("(b)", "Weekly In-Contract Invoice. Issued every Friday for Virtual Staff (contracted support) services performed under Schedule A during the preceding week (Monday through Friday). Each invoice includes: (i) a listing of each support ticket addressed; (ii) the assigned resource, role, and hours spent per ticket; (iii) a description of the work performed per ticket; (iv) whether work was performed during normal or after-hours; and (v) the current running balance for each contracted role.", indent=1)
    d.numbered("(c)", "Weekly Out-of-Contract Invoice. Issued every Friday for labor services performed outside the scope of any active Schedule or SOW \u2014 including ad-hoc support requests, emergency work, and services performed under a SOW with hourly billing. Each invoice includes: (i) a listing of each support ticket or task performed; (ii) the assigned resource, role, and applicable hourly rate from the Rate Card (Schedule C); (iii) time entries with hours billed per activity (in 15-minute increments); (iv) whether work was performed during normal or after-hours; and (v) the total hours and total amount for the week.", indent=1)
    d.numbered("(d)", "Equipment and Materials Invoice. Issued upon delivery or procurement of hardware, software licenses (perpetual), or other tangible goods on Client\u2019s behalf. Each invoice includes: (i) item description, manufacturer, and model/part number; (ii) quantity and unit price; (iii) applicable sales tax; (iv) shipping and handling charges, if any; and (v) total amount due. Title to equipment shall not pass to Client until payment is received in full.", indent=1)
    d.numbered("(e)", "Project Milestone Invoice. Issued upon completion of a project milestone as defined in an applicable SOW. The milestone, deliverables, and invoiced amount are as specified in the payment schedule of the SOW. Milestone invoices are billed in arrears upon acceptance of the deliverables or deemed acceptance under the SOW\u2019s acceptance provisions.", indent=1)
    d.numbered("3.03.", "Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified in the applicable Schedule or SOW.")
    d.numbered("3.04.", "Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated as simple interest from the date payment was due until the date payment is received in full. Late fees are payable in addition to the outstanding principal amount. The Parties acknowledge that the late fee represents a reasonable estimate of Technijian\u2019s administrative costs and damages resulting from late payment, including cash-flow disruption, collection overhead, and the cost of carrying accounts receivable, and is not intended as a penalty.")
    d.numbered("3.05.", "Disputed Invoices.")
    d.numbered("(a)", "Weekly Invoices (In-Contract and Out-of-Contract). Because weekly invoices include detailed ticket descriptions and time entries, Client shall have thirty (30) days from the invoice date to review and dispute any portion of a weekly invoice. Client shall notify Technijian in writing, specifying the ticket number(s) and nature of the dispute. Undisputed tickets and time entries on the same invoice shall remain payable by the due date. Failure to provide a timely written dispute notice within the thirty (30) day period shall constitute acceptance of all tickets and time entries on the invoice.", indent=1)
    d.numbered("(b)", "All Other Invoices. For monthly service invoices, equipment invoices, and project milestone invoices, Client shall notify Technijian in writing within fifteen (15) days of the invoice date if Client disputes any portion, specifying the nature and basis of the dispute. Client shall pay all undisputed amounts by the due date. Failure to provide a timely written dispute notice shall constitute acceptance of the invoice.", indent=1)
    d.numbered("(c)", "Resolution. The Parties shall work in good faith to resolve any invoice dispute within thirty (30) days of the dispute notice. If the dispute results in an adjustment, Technijian shall issue a credit memo or revised invoice within ten (10) business days of resolution.", indent=1)
    d.numbered("3.06.", "Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend services under the Schedule or SOW associated with the unpaid invoice until payment is received in full, including all accrued late fees. If Client fails to pay any undisputed invoice within sixty (60) days of the due date, Technijian may suspend all services under this Agreement and any related Schedules or SOWs. Recurring fees for suspended services shall continue to accrue for a period not to exceed thirty (30) days following the date of suspension, after which Technijian may terminate the affected Schedule, SOW, or this Agreement upon written notice. Suspension of services shall not relieve Client of its payment obligations.")
    d.numbered("3.07.", "Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian\u2019s income.")

    # SECTION 4
    d.spacer()
    d.section_header("SECTION 4 \u2014 CONFIDENTIALITY")
    d.numbered("4.01.", "Definition. \u201cConfidential Information\u201d means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information.")
    d.numbered("4.02.", "Obligations. Each Party shall:")
    d.numbered("(a)", "Hold the other Party\u2019s Confidential Information in confidence using at least the same degree of care it uses for its own confidential information, but not less than reasonable care;", indent=1)
    d.numbered("(b)", "Not disclose Confidential Information to third parties without prior written consent, except to employees, agents, and subcontractors who have a need to know and are bound by equivalent obligations; and", indent=1)
    d.numbered("(c)", "Not use Confidential Information for any purpose other than performing obligations under this Agreement.", indent=1)
    d.numbered("4.03.", "Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known to the receiving Party prior to disclosure, is independently developed, or is received from a third party without restriction.")
    d.numbered("4.04.", "Compelled Disclosure. If required by law or court order to disclose Confidential Information, the receiving Party shall provide prompt written notice to the disclosing Party and cooperate in seeking a protective order.")
    d.numbered("4.05.", "Duration. Confidentiality obligations shall survive termination for a period of three (3) years.")

    # SECTION 5
    d.spacer()
    d.section_header("SECTION 5 \u2014 LIMITATION OF LIABILITY")
    d.numbered("5.01.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, NEITHER PARTY\u2019S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM (THE \u201cSTANDARD CAP\u201d).")
    d.numbered("5.02.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL, REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE OR WHETHER EITHER PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.")
    d.numbered("5.03.", "ENHANCED CAP FOR CERTAIN CLAIMS. FOR CLAIMS ARISING FROM BREACHES OF SECTION 4 (CONFIDENTIALITY), SECTION 10 (DATA PROTECTION), INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, THE TOTAL AGGREGATE LIABILITY OF THE RESPONSIBLE PARTY SHALL NOT EXCEED THREE (3) TIMES THE STANDARD CAP DEFINED IN SECTION 5.01 (THE \u201cENHANCED CAP\u201d). THE ENHANCED CAP SHALL NOT APPLY TO LIABILITY ARISING FROM A PARTY\u2019S WILLFUL AND INTENTIONAL MISAPPROPRIATION OF THE OTHER PARTY\u2019S CONFIDENTIAL INFORMATION OR CLIENT DATA, FOR WHICH LIABILITY SHALL BE UNCAPPED.")
    d.numbered("5.04.", "Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data, Client acknowledges that Client is solely responsible for maintaining backup copies of its data.")

    d.page_break()

    # SECTION 6
    d.section_header("SECTION 6 \u2014 INDEMNIFICATION")
    d.numbered("6.01.", "By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from Technijian\u2019s gross negligence or willful misconduct in performing the services.")
    d.numbered("6.02.", "By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from Client\u2019s use of the services in violation of applicable law, Client\u2019s breach of this Agreement, or any data, content, or materials provided by Client.")
    d.numbered("6.03.", "Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party\u2019s prior written consent.")

    # SECTION 7
    d.spacer()
    d.section_header("SECTION 7 \u2014 INTELLECTUAL PROPERTY")
    d.numbered("7.01.", "Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes.")
    d.numbered("7.02.", "Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property.")
    d.numbered("7.03.", "Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW.")

    # SECTION 8
    d.spacer()
    d.section_header("SECTION 8 \u2014 DISPUTE RESOLUTION")
    d.numbered("8.01.", "Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations for a period of thirty (30) days.")
    d.numbered("8.02.", "Mediation. If not resolved, the Parties shall submit the dispute to mediation in Orange County, California, for a period not to exceed sixty (60) days.")
    d.numbered("8.03.", "Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association in Orange County, California.")
    d.numbered("8.04.", "Injunctive Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief to prevent irreparable harm.")

    # SECTION 9
    d.spacer()
    d.section_header("SECTION 9 \u2014 GENERAL PROVISIONS")
    d.numbered("9.01.", "Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties. " + SUPERSEDE_CLAUSE)
    d.numbered("9.02.", "Amendment. This Agreement may only be amended by a written instrument signed by both Parties.")
    d.numbered("9.03.", "Severability. If any provision is found invalid or unenforceable, the remaining provisions shall continue in full force.")
    d.numbered("9.04.", "Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party.")
    d.numbered("9.05.", "Assignment. Neither Party may assign this Agreement without prior written consent, except in connection with a merger, acquisition, or sale of substantially all assets.")
    d.numbered("9.06.", "Force Majeure.")
    d.numbered("(a)", "Neither Party shall be liable for delays or failures in performance caused by events beyond its reasonable control, including natural disasters, acts of government, labor disputes, pandemics, cyberattacks on critical infrastructure, or failures of major third-party infrastructure services that are not attributable to the affected Party\u2019s failure to implement reasonable redundancy (\u201cForce Majeure Event\u201d).", indent=1)
    d.numbered("(b)", "The affected Party shall notify the other Party in writing within five (5) business days of becoming aware of a Force Majeure Event and shall use commercially reasonable efforts to mitigate the impact and resume performance.", indent=1)
    d.numbered("(c)", "Payment obligations are not excused by a Force Majeure Event.", indent=1)
    d.numbered("(d)", "If a Force Majeure Event prevents performance of a material portion of the services for more than ninety (90) consecutive days, either Party may terminate the affected Schedule, SOW, or this Agreement upon fifteen (15) days written notice without liability for early termination fees, and Client shall pay only for services actually rendered through the date of termination.", indent=1)
    d.numbered("9.07.", "Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier.")
    d.numbered("9.08.", "Governing Law. This Agreement shall be governed by the laws of the State of California.")
    d.numbered("9.09.", "Personnel Transition Fee. The Parties acknowledge that each invests significant resources in recruiting, training, and retaining skilled personnel. If either Party hires (whether as an employee or independent contractor) any individual who was an employee of the other Party and who was directly involved in performing or receiving services under this Agreement, and such hiring occurs during the term of this Agreement or within twelve (12) months following termination, the hiring Party shall pay the other Party a personnel transition fee equal to 25% of the hired individual\u2019s first-year annual compensation (base salary or annualized contractor fees). This fee represents a reasonable estimate of the non-hiring Party\u2019s recruiting and training costs and is not intended as a restraint on trade or employment. This Section does not restrict any individual\u2019s right to seek or obtain employment, and shall not apply to individuals who: (a) respond to general public job postings or advertisements not specifically targeted at the other Party\u2019s employees; or (b) are referred by a third-party recruiting firm without the hiring Party\u2019s direction to target the other Party\u2019s employees.")
    d.numbered("9.10.", "Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.")

    # SECTION 10
    d.spacer()
    d.section_header("SECTION 10 \u2014 DATA PROTECTION")
    d.numbered("10.01.", "CCPA/CPRA Compliance. To the extent Technijian processes, stores, or has access to personal information (as defined under the California Consumer Privacy Act, as amended by the California Privacy Rights Act, Cal. Civ. Code \u00a7 1798.100 et seq., collectively \u201cCCPA\u201d) on behalf of Client, Technijian acts as a \u201cservice provider\u201d as defined under Cal. Civ. Code \u00a7 1798.140(ag) and shall:")
    d.numbered("(a)", "Process such personal information only as necessary to perform the services and in accordance with Client\u2019s documented instructions and this Agreement;", indent=1)
    d.numbered("(b)", "Not sell, share, retain, use, or disclose personal information for any purpose other than performing the services, including not using personal information for targeted advertising or cross-context behavioral advertising;", indent=1)
    d.numbered("(c)", "Not combine personal information received from Client with personal information received from other sources or collected from Technijian\u2019s own interactions with individuals, except as expressly permitted by the CCPA to perform the services;", indent=1)
    d.numbered("(d)", "Implement reasonable security measures appropriate to the nature of the personal information, consistent with the requirements of Section 10.03;", indent=1)
    d.numbered("(e)", "Cooperate with Client in responding to verifiable consumer rights requests under the CCPA (including access, deletion, correction, and opt-out requests) within ten (10) business days of Client\u2019s request;", indent=1)
    d.numbered("(f)", "Notify Client within five (5) business days if Technijian determines that it can no longer meet its obligations as a service provider under the CCPA;", indent=1)
    d.numbered("(g)", "Ensure that all subcontractors who process personal information on behalf of Client are bound by written agreements containing data protection obligations at least as protective as those in this Section 10, and Technijian shall remain liable for the acts and omissions of its subcontractors with respect to personal information;", indent=1)
    d.numbered("(h)", "Permit Client, upon thirty (30) days written notice and no more than once per twelve (12) month period, to audit or inspect Technijian\u2019s data processing practices to verify compliance with this Section 10, or, at Technijian\u2019s option, provide Client with a summary of a recent independent third-party audit (such as SOC 2 Type II) covering the relevant controls; and", indent=1)
    d.numbered("(i)", "Certify that Technijian understands the restrictions in this Section 10 and will comply with them.", indent=1)

    d.numbered("10.02.", "Security Incident Notification. If Technijian becomes aware of a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Client Data (\u201cSecurity Incident\u201d), Technijian shall: (a) notify Client in writing without unreasonable delay and in no event later than forty-eight (48) hours after becoming aware of the Security Incident; (b) provide Client with sufficient information to enable Client to comply with its obligations under California Civil Code \u00a7 1798.82 (data breach notification), applicable SEC Regulation S-P customer notification requirements, and FINRA Rule 4530 reporting obligations, including the categories and approximate number of records affected, the nature of the incident, and the measures taken or proposed to address it; (c) cooperate with Client\u2019s investigation of the Security Incident; and (d) take reasonable steps to contain and remediate the Security Incident.")
    d.numbered("10.03.", "Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards designed to protect Client Data from unauthorized access, use, or disclosure, consistent with industry standards for managed IT service providers. Such safeguards shall include, at a minimum: (a) encryption of Client Data in transit and at rest; (b) access controls limiting access to authorized personnel; (c) regular security assessments and vulnerability testing; and (d) employee security awareness training.")
    d.numbered("10.04.", "Regulatory Compliance. Client is a dually registered Investment Adviser and Broker-Dealer subject to SEC Regulation S-P (including the 2024 Amendments), FINRA Rule 4370, FINRA Rule 3110, and SEC Rule 17a-4. The Parties shall cooperate in good faith to implement the controls described in the applicable Schedules and SOWs to support Client\u2019s compliance with these requirements. If additional industry-specific data protection requirements apply, the Parties shall execute a separate addendum addressing the additional obligations applicable to the regulated data.")
    d.numbered("10.05.", "Data Return and Deletion. Upon termination of this Agreement or upon Client\u2019s written request, Technijian shall securely delete or return all Client Data in its possession within thirty (30) days, using methods consistent with NIST SP 800-88 or equivalent standards, and shall certify such deletion in writing upon request.")

    d.signatures(client_name=CLIENT_UPPER, contact_name=SIGNER)

    d.spacer()
    d.body("Schedules (delivered as separate documents):")
    d.bullet("Schedule A \u2014 Monthly Managed Services")
    d.bullet("Schedule B \u2014 Subscription and License Services (none in effect as of Effective Date)")
    d.bullet("Schedule C \u2014 Rate Card")
    d.spacer()
    d.body("Statements of Work delivered under this Agreement:")
    d.bullet("SOW-AFFG-001 \u2014 Sophos Firewall Deployment")
    d.bullet("SOW-AFFG-002 \u2014 Wireless Access Point Deployment")
    d.bullet("SOW-AFFG-003 \u2014 IT Compliance & VDI Implementation")

    d.footer_bar()
    d.save(os.path.join(OUT_DIR, "MSA-AFFG.docx"))


# ================================================================
# SCHEDULE A — Monthly Managed Services ($6,922.35/mo)
# ================================================================
def build_schedule_a():
    d = BrandedDoc()
    d.cover_page(
        "Schedule A",
        "Monthly Managed Services",
        client_name=CLIENT_FULL,
        date=EFFECTIVE_DATE,
    )
    d.page_break()

    d.part_header("SCHEDULE A \u2014 MONTHLY MANAGED SERVICES")
    d.spacer()
    d.body(f"Agreement Number: {AGREEMENT_NO}")
    d.body(f"Effective Date: {EFFECTIVE_DATE}")
    d.body(f"Parent Agreement: Master Service Agreement between Technijian, Inc. and {CLIENT_FULL}")
    d.spacer()
    d.body(
        f"This Schedule A sets forth the recurring managed services to be provided by Technijian "
        f"to {CLIENT_FULL} (\u201c{CLIENT_SHORT}\u201d) under the Master Service Agreement "
        f"({AGREEMENT_NO}). The total monthly investment is comprised of five parts: "
        "(1) VDI Workstations; (2) All-User Services; (3) Domain / Site / IP Services; "
        "(4) Shared Storage; and (5) Virtual Staff Support. Capitalized terms used and "
        "not defined herein have the meanings given to them in the MSA."
    )
    d.spacer()
    d.body_bold("Superseding Prior Agreement. ", SUPERSEDE_CLAUSE)
    d.spacer()

    # -- Part 1 — VDI Workstations --
    d.section_header("Part 1 \u2014 VDI Workstations (11 agents)")
    d.spacer()
    d.body(
        f"Eleven (11) dedicated Windows 11 personal-desktop virtual machines are provisioned \u2014 "
        f"one per {CLIENT_SHORT} agent. Each VDI instance is sized at 4 vCore / 16 GB RAM / 100 GB "
        "with a shared-bandwidth allocation and is protected by the endpoint security, monitoring, "
        "and auditing services listed below."
    )
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("VDI VM: MYVDI ($10) + 4x CL-VC ($14=$56) + 16x CL-GB ($6.25=$100) + CL-SBW ($30)", "MYVDI+VM", "$196.00", "11", "$2,156.00"),
            ("CrowdStrike Falcon EDR (on VDI)", "AVD", "$8.50", "11", "$93.50"),
            ("Huntress Managed Detection (on VDI)", "AVMH", "$6.00", "11", "$66.00"),
            ("My Secure Internet / DNS filtering (on VDI)", "SI", "$6.00", "11", "$66.00"),
            ("Patch Management (on VDI)", "PMW", "$4.00", "11", "$44.00"),
            ("My Remote / ScreenConnect (on VDI)", "MR", "$2.00", "11", "$22.00"),
            ("My Audit UAM+DLP / 1-Year retention", "AMDLP1Y", "$108.10", "11", "$1,189.10"),
            ("SSO / Multi-Factor", "SSO-2FA", "$12.00", "11", "$132.00"),
            ("Credential Manager", "CRM", "$5.00", "11", "$55.00"),
            ("Subtotal \u2014 VDI Workstations", "", "", "11 agents", "$3,823.60"),
        ],
        total_indices=[9],
    )

    # -- Part 2 — All-User Services --
    d.spacer()
    d.section_header("Part 2 \u2014 All-User Services (31 M365 users)")
    d.spacer()
    d.body(
        f"Services below apply to all thirty-one (31) {CLIENT_SHORT} Microsoft 365 users, "
        "regardless of whether the user operates via the VDI environment or a named-exception "
        "device approved by Client."
    )
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("Veeam 365 Backup (Exchange/SharePoint/OneDrive/Teams)", "V365", "$4.00", "31", "$124.00"),
            ("Phishing Training (Huntress SAT)", "PHT", "$6.00", "31", "$186.00"),
            ("Inky Anti-Spam Basic (full Inky package incl. archiving)", "ASB", "$4.25", "31", "$131.75"),
            ("Subtotal \u2014 All-User Services", "", "", "31 users", "$441.75"),
        ],
        total_indices=[3],
    )

    # -- Part 3 — Domain / Site / IP Services --
    d.spacer()
    d.section_header("Part 3 \u2014 Domain / Site / IP Services")
    d.spacer()
    d.body(
        "Domain security monitoring, site-level assessments, and IP-level penetration testing "
        "required to maintain the compliance posture described in SOW-AFFG-003."
    )
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("DMARC/DKIM Monitoring", "DKIM", "$20.00", "1 domain", "$20.00"),
            ("Site Assessment (Physical Office \u2014 1 Park Plaza, Irvine)", "SA", "$50.00", "1 site", "$50.00"),
            ("Site Assessment (VDI Environment)", "SA", "$50.00", "1 site", "$50.00"),
            ("Real-Time Pen Testing", "RTPT", "$7.00", "6 IPs", "$42.00"),
            ("Subtotal \u2014 Domain / Site / IP", "", "", "", "$162.00"),
        ],
        total_indices=[4],
    )

    # -- Part 4 — Shared Storage --
    d.spacer()
    d.section_header("Part 4 \u2014 Shared Storage")
    d.spacer()
    d.body(
        "Production storage supports FSLogix VDI profiles and cache; Backup storage supports "
        "Veeam 365 retention with 7-year immutability per SEC Rule 17a-4."
    )
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per TB", "Qty", "Monthly"],
        [
            ("Production Storage (VDI profiles, cache)", "TB-PSTR", "$200", "2 TB", "$400.00"),
            ("Backup Storage (V365 retention)", "TB-BSTR", "$50", "2 TB", "$100.00"),
            ("Subtotal \u2014 Shared Storage", "", "", "4 TB", "$500.00"),
        ],
        total_indices=[2],
    )
    d.spacer()
    d.body(
        "Note: Storage is sold in 1 TB increments only; sizes have been rounded up to the next "
        "whole TB."
    )

    # -- Part 5 — Virtual Staff Support --
    d.spacer()
    d.section_header("Part 5 \u2014 Virtual Staff Support (UNCHANGED from prior MSA)")
    d.spacer()
    d.body(
        "Virtual Staff Support hours and rates in this Part 5 are carried forward UNCHANGED "
        "from the prior Client Monthly Service Agreement."
    )
    d.spacer()
    d.styled_table(
        ["Role", "Office-POD", "Hours/mo", "Rate", "Monthly"],
        [
            ("CTO Advisory (Normal Hours)", "IRV-AD1", "3.00", "$225/hr", "$675.00"),
            ("US Tech Support (Normal Hours)", "IRV-TS1", "6.00", "$125/hr", "$750.00"),
            ("India Tech Support (Normal Hours)", "CHD-TS1", "22.00", "$15/hr", "$330.00"),
            ("India Tech Support (After-Hours)", "CHD-TS1", "8.00", "$30/hr", "$240.00"),
            ("Subtotal \u2014 Virtual Staff Support", "", "39.00 hrs", "", "$1,995.00"),
        ],
        total_indices=[4],
    )
    d.spacer()
    d.body(
        "Hours and rates in Part 5 are unchanged from the prior Monthly Service Agreement. "
        "These contracted hours are governed by the Cycle-Based Billing Model; any hours worked "
        "in excess of contracted hours accrue to the Unpaid Balance and are payable upon "
        "termination at the applicable rates."
    )

    # -- Monthly Investment Summary --
    d.page_break()
    d.section_header("Monthly Investment Summary")
    d.spacer()
    d.styled_table(
        ["Part", "Monthly"],
        [
            ("Part 1 \u2014 VDI Workstations (11 agents)", "$3,823.60"),
            ("Part 2 \u2014 All-User Services (31 users)", "$441.75"),
            ("Part 3 \u2014 Domain / Site / IP Services", "$162.00"),
            ("Part 4 \u2014 Shared Storage", "$500.00"),
            ("Part 5 \u2014 Virtual Staff Support (unchanged)", "$1,995.00"),
            ("TOTAL MONTHLY", "$6,922.35"),
        ],
        total_indices=[5],
    )
    d.spacer()
    d.body(
        "The TOTAL MONTHLY amount of $6,922.35 represents the contracted recurring fee. Any "
        "ad-hoc hourly labor, project work (including SOW-AFFG-003), equipment procurement, or "
        "other out-of-scope services are invoiced separately per Section 3.02 of the MSA and "
        "Schedule C."
    )

    # -- Changes Callout --
    d.spacer()
    d.section_header("CHANGES FROM PRIOR AGREEMENT")
    d.spacer()
    d.body(
        "The table below summarizes how charges under this Schedule A compare to the prior "
        "Client Monthly Service Agreement (DocuSign Envelope " + PRIOR_ENVELOPE + ", effective "
        "3/11/2026). The new monthly charges REPLACE the prior charges. This is not additive. "
        "Net increase to AFFG is +$4,127.85/month."
    )
    d.spacer()
    d.styled_table(
        ["Category", "Prior MSA", "New MSA", "Change"],
        [
            ("Workstations (16 physical)", "$424.00", "Replaced by VDI (11 agents)", "\u2014"),
            ("User services (V365 $77.50 + PHT $186)", "$263.50", "$441.75 (includes ASB $131.75, V365 at $4)", "+$178.25"),
            ("Domain + Site + IP services", "$112.00", "$162.00 (added VDI site assessment)", "+$50.00"),
            ("Shared storage", "\u2014", "$500.00 (new)", "+$500.00"),
            ("VDI + compliance (NEW)", "\u2014", "$4,323.60 (VDI VMs + My Audit UAM+DLP + SSO + CRM)", "+$4,323.60"),
            ("Virtual Staff Support", "$1,995.00", "$1,995.00 (UNCHANGED)", "$0.00"),
            ("TOTAL", "$2,794.50", "$6,922.35", "+$4,127.85"),
        ],
        total_indices=[6],
    )
    d.spacer()
    d.body_bold(
        "Important: ",
        "The new monthly charges REPLACE the prior charges. This is not additive. "
        "Net increase to AFFG is +$4,127.85/month."
    )

    d.footer_bar()
    d.save(os.path.join(OUT_DIR, "Schedule-A-AFFG.docx"))


# ================================================================
# SCHEDULE B — Subscription and License Services (placeholder)
# ================================================================
def build_schedule_b():
    d = BrandedDoc()
    d.cover_page(
        "Schedule B",
        "Subscription and License Services",
        client_name=CLIENT_FULL,
        date=EFFECTIVE_DATE,
    )
    d.page_break()

    d.part_header("SCHEDULE B \u2014 SUBSCRIPTION AND LICENSE SERVICES")
    d.spacer()
    d.body(f"Agreement Number: {AGREEMENT_NO}")
    d.body(f"Effective Date: {EFFECTIVE_DATE}")
    d.body(f"Parent Agreement: Master Service Agreement between Technijian, Inc. and {CLIENT_FULL}")
    d.spacer()
    d.body(
        f"This Schedule B governs any subscription services, software licenses, SaaS subscriptions, "
        f"SIP trunk services, and similar recurring third-party services procured by Technijian on "
        f"behalf of {CLIENT_FULL} under the Master Service Agreement ({AGREEMENT_NO}). "
        "Capitalized terms used and not defined herein have the meanings given to them in the MSA."
    )
    d.spacer()

    d.section_header("Current Subscription Services")
    d.spacer()
    d.body_bold("Status as of Effective Date: ", "No subscription or license services are currently in effect under this Schedule.")
    d.spacer()
    d.body(
        "Endpoint security, user services, and VDI-resident products (CrowdStrike, Huntress, "
        "My Secure Internet, Patch Management, My Remote, My Audit UAM+DLP, Veeam 365 Backup, "
        "Huntress SAT, and Inky Anti-Spam Basic) are provided under Schedule A as managed "
        "services and are not subject to this Schedule B."
    )

    d.spacer()
    d.section_header("Adding Subscription Services")
    d.spacer()
    d.body(
        "Subscription and license services may be added to this Schedule B at any time during the "
        "term of the Agreement by executing a Service Order (or written amendment) signed by both "
        "Parties. Each Service Order shall specify:"
    )
    d.bullet("The subscription product or license being procured")
    d.bullet("The quantity, unit price, and total monthly or annual fee")
    d.bullet("The subscription term and renewal terms (if different from the MSA)")
    d.bullet("Any vendor-specific terms that apply (such as EULA acknowledgment)")
    d.spacer()
    d.body(
        "Subscription fees are billed on the Monthly Recurring Subscription Invoice described in "
        "Section 3.02(a) of the MSA, in advance, on the first business day of each month."
    )

    d.spacer()
    d.section_header("Pass-Through and Client-Procured Licenses")
    d.spacer()
    d.body(
        f"If {CLIENT_SHORT} procures software or subscription licenses directly from a third-party "
        "vendor (rather than through Technijian), those licenses are not governed by this Schedule B. "
        "Client remains solely responsible for the terms, renewal, and payment of any directly-"
        "procured subscriptions, including Microsoft 365 licensing and any existing email-archive "
        "platform already operated by Client. Technijian will install, configure, and support such "
        "Client-procured licenses as part of the managed services under Schedule A, subject to "
        "vendor and license compatibility."
    )

    d.spacer()
    d.section_header("Termination of Subscriptions")
    d.spacer()
    d.body(
        "Upon termination of the MSA or this Schedule B, Client shall pay all outstanding fees "
        "for any subscription services through the end of the then-current vendor billing period "
        "(including any annual commitments procured on Client\u2019s behalf that are not cancellable). "
        "Technijian will cooperate in good faith to transfer any transferable subscriptions to "
        "Client or a successor provider."
    )

    d.footer_bar()
    d.save(os.path.join(OUT_DIR, "Schedule-B-AFFG.docx"))


# ================================================================
# SCHEDULE C — Rate Card
# ================================================================
def build_schedule_c():
    d = BrandedDoc()
    d.cover_page(
        "Schedule C",
        "Rate Card",
        client_name=CLIENT_FULL,
        date=EFFECTIVE_DATE,
    )
    d.page_break()

    d.part_header("SCHEDULE C \u2014 RATE CARD")
    d.spacer()
    d.body(f"Agreement Number: {AGREEMENT_NO}")
    d.body(f"Effective Date: {EFFECTIVE_DATE}")
    d.body(f"Parent Agreement: Master Service Agreement between Technijian, Inc. and {CLIENT_FULL}")
    d.body("Normal Business Hours: Monday \u2013 Friday, 8:00 AM \u2013 6:00 PM Pacific Time")
    d.spacer()
    d.body(
        f"This Schedule C sets forth the standard hourly rates applicable to labor services "
        f"performed under the Master Service Agreement ({AGREEMENT_NO}) that are not covered by "
        "an active contracted Schedule or Statement of Work. Capitalized terms used and not "
        "defined herein have the meanings given to them in the MSA."
    )
    d.spacer()

    d.section_header("US-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr"),
            ("Developer", "$150/hr", "N/A", "$125/hr"),
            ("Tech Support", "$150/hr", "$250/hr", "$125/hr"),
        ],
    )

    d.spacer()
    d.section_header("Offshore-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr"),
            ("Tech Support", "$15/hr", "$30/hr", "$10/hr"),
        ],
    )

    d.spacer()
    d.section_header("Project & Ad Hoc")
    d.spacer()
    d.styled_table(
        ["Service", "Rate", "Minimum"],
        [
            ("On-Site Support (US)", "$150/hr", "2-hr minimum"),
            ("Remote Support (ad hoc)", "$150/hr", "15-min increments"),
            ("Emergency / Critical", "$250/hr", "1-hr minimum"),
            ("Project Management", "$150/hr", "SOW-based"),
        ],
    )

    d.spacer()
    d.section_header("Rate Definitions")
    d.spacer()
    d.body_bold("Normal Business Hours: ", "Monday \u2013 Friday, 8:00 AM \u2013 6:00 PM Pacific Time, excluding US federal holidays.")
    d.body_bold("After-Hours: ", "All hours outside Normal Business Hours, including weekends and US federal holidays.")
    d.body_bold("Contracted Rate: ", "Discounted rate available when Client commits to the Cycle-Based Billing Model under an active Schedule.")
    d.body_bold("Hourly Rate: ", "Standard ad-hoc rate applied to work performed outside an active contracted Schedule.")

    d.footer_bar()
    d.save(os.path.join(OUT_DIR, "Schedule-C-AFFG.docx"))


# ================================================================
# SOW-AFFG-003 — IT Compliance & VDI Implementation
# ================================================================
def _load_tickets_by_phase():
    """Read the ticket CSV and bucket each row by phase."""
    by_phase = {}
    with open(TICKETS_CSV, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            phase = row["Phase"].strip()
            by_phase.setdefault(phase, []).append({
                "id": row["Ticket"].strip(),
                "title": row["Title"].strip(),
                "hours": row["Est Hours"].strip(),
                "assignee": row["Assigned To"].strip(),
            })
    return by_phase


PHASE_NARRATIVES = [
    (
        "Phase 1: VDI Foundation",
        "Phase 1 \u2014 VDI Foundation (Weeks 1\u20134)",
        (
            "Establish the Technijian VDI tenant (Technijian\u2019s privately-operated Virtual Desktop "
            "service) that will host all eleven (11) AFFG agents. Phase 1 deliverables include the "
            "Technijian VDI workspace and personal host pool, eleven dedicated Windows 11 Enterprise "
            "personal-desktop VMs, hardened VDI session policies "
            "(no clipboard, no local drives, no USB redirection, watermarking, 15-minute inactivity "
            "lock), FSLogix profile containers on the Technijian Production Storage pool, MFA and "
            "Entra Conditional Access, hands-on enrollment and training, My Audit UAM+DLP agent "
            "integration, and documented VDI egress IPs prepared for Schwab and Interactive Brokers "
            "allowlisting. This phase enforces the \u2018no client data on personal device\u2019 core "
            "design principle that underpins the AFFG compliance strategy."
        ),
    ),
    (
        "Phase 2: Access Lockdown",
        "Phase 2 \u2014 Access Lockdown (Weeks 3\u20136)",
        (
            "Lock down identity and network perimeters so that production access to AFFG data can "
            "only originate from the VDI environment. Phase 2 refines Entra Conditional Access to "
            "restrict M365 sign-ins to the VDI egress IPs across all 31 users, submits and activates "
            "Schwab Advisor Services and Interactive Brokers IP allowlists, blocks legacy "
            "authentication protocols, and deploys CrowdStrike Falcon EDR and Huntress Managed "
            "Detection inside each VDI desktop with Technijian SOC forwarding."
        ),
    ),
    (
        "Phase 3: Data Protection",
        "Phase 3 \u2014 Data Protection / DLP (Weeks 5\u20138)",
        (
            "Deploy Microsoft 365 Data Loss Prevention across Exchange Online, SharePoint, OneDrive, "
            "and Teams. Phase 3 defines custom Sensitive Information Types for Schwab and IBKR "
            "account-number patterns, locks down OneDrive external sharing, configures email DLP "
            "rules with TLS enforcement for custodian domains, and executes end-to-end exfiltration "
            "testing with sign-off by AFFG\u2019s registered principal."
        ),
    ),
    (
        "Phase 4: V365 Backup",
        "Phase 4 \u2014 V365 Backup (Weeks 7\u201310)",
        (
            "Stand up Veeam 365 Backup against the AFFG Microsoft 365 tenant \u2014 Exchange Online "
            "mailboxes, SharePoint sites, OneDrive accounts, and Teams data for all 31 users \u2014 "
            "with a 7-year immutable retention policy on the Backup Storage tier to satisfy SEC "
            "Rule 17a-4. Phase 4 also configures the M365 Compliance Center eDiscovery (Standard) "
            "workspace and coordinates with AFFG\u2019s existing email-archive platform so long-term "
            "mail retention remains authoritative there."
        ),
    ),
    (
        "Phase 5: Monitoring & Audit",
        "Phase 5 \u2014 Monitoring & Audit (Weeks 9\u201312)",
        (
            "Bring the supervisory and monitoring surface to full regulator-ready coverage. Phase 5 "
            "confirms My Audit UAM+DLP is operational across all 11 VDI desktops, enables detailed "
            "Technijian VDI session logging with central forwarding, builds insider-threat and anomaly alerting "
            "(off-hours sign-ins, impossible travel, mass-deletion, after-hours DLP), delivers the "
            "monthly compliance dashboard PDF, and documents the FINRA Rule 3110 supervisory audit "
            "trail architecture anchored on the dedicated-VM per-agent model."
        ),
    ),
    (
        "Phase 6: Validation & Testing",
        "Phase 6 \u2014 Validation & Testing (Weeks 11\u201314)",
        (
            "Independently validate the control environment. Phase 6 executes a formal penetration "
            "test against the VDI environment (external surface, egress/ingress controls, clipboard/"
            "drive/USB bypass attempts, cross-VM isolation, session takeover, privilege escalation), "
            "deploys Huntress SAT with financial-services-specific phishing content, baselines the "
            "first phishing simulation, and closes with a post-implementation gap assessment against "
            "the AFFG IT Compliance Strategy."
        ),
    ),
    (
        "Phase 7: Policy & Compliance Docs",
        "Phase 7 \u2014 Policy & Compliance Docs (Weeks 13\u201318)",
        (
            "Produce the regulator-ready documentation set. Phase 7 drafts the Written Incident "
            "Response Plan, breach notification procedures and templates (SEC Reg S-P 30-day customer, "
            "California Civil Code 1798.82 state AG, FINRA, and custodian), the Service Provider "
            "Oversight Policy (incorporating the SEC Reg S-P 72-hour vendor-incident notification "
            "clause), the BCP Customer Disclosure (FINRA Rule 4370) plus the Customer Information "
            "Disposal Policy, the annual sign-off process, and final packaging of all deliverables "
            "into the AFFG regulator-ready compliance binder."
        ),
    ),
]


def build_sow_compliance():
    d = BrandedDoc()
    # Cover
    d.cover_page(
        "Statement of Work",
        "SOW-AFFG-003: IT Compliance & VDI Implementation",
        client_name=CLIENT_FULL,
        date=EFFECTIVE_DATE,
    )
    d.page_break()

    # Reference section
    d.part_header("SOW-AFFG-003 \u2014 IT COMPLIANCE & VDI IMPLEMENTATION")
    d.spacer()
    d.body("SOW Number: SOW-AFFG-003")
    d.body(f"Effective Date: {EFFECTIVE_DATE}")
    d.body(f"Parent Agreement: {AGREEMENT_NO}")
    d.body(f"Client: {CLIENT_FULL} (\u201c{CLIENT_SHORT}\u201d)")
    d.body("Source: AFFG IT Compliance Strategy (REF: AFFG-CS-2026-04, April 2026)")
    d.body("Regulatory Scope: SEC Reg S-P (2024 Amendments), FINRA Rule 4370, FINRA Rule 3110, SEC Rule 17a-4")
    d.spacer()
    d.body_bold("Supersedes: ", SUPERSEDE_CLAUSE)
    d.spacer()

    # Section 1 — Overview
    d.section_header("1. Project Overview")
    d.body(
        f"Technijian will deliver a 7-phase implementation that migrates {CLIENT_SHORT}\u2019s eleven "
        "(11) registered agents onto the Technijian VDI (Virtual Desktop Infrastructure) environment "
        "\u2014 Technijian\u2019s privately-operated, managed Windows 11 desktop service \u2014 locks down "
        "access to the Microsoft 365 tenant and the Schwab / Interactive Brokers custodial APIs to "
        "VDI egress IPs only, deploys data-loss-prevention controls across Exchange Online, "
        "SharePoint, OneDrive, and Teams, and produces the full regulator-ready compliance "
        "documentation set required under SEC Regulation S-P (including the 2024 Amendments), "
        "FINRA Rule 4370, FINRA Rule 3110, and SEC Rule 17a-4."
    )
    d.spacer()
    d.body_bold("Objectives:", "")
    d.bullet("Deploy VDI for 11 agents with dedicated Windows 11 personal desktops (one per agent)")
    d.bullet("Restrict M365 and custodian-API access to VDI egress IPs only")
    d.bullet("Deploy M365 DLP, OneDrive external-sharing lockdown, and email DLP with custodian TLS enforcement")
    d.bullet("Stand up Veeam 365 Backup with 7-year immutable retention (SEC 17a-4)")
    d.bullet("Achieve 100% regulatory compliance posture (up from the current ~15% baseline) per the AFFG IT Compliance Strategy")
    d.bullet("Produce the complete regulator-ready compliance binder (IRP, breach notification, vendor oversight, BCP disclosure, info disposal, annual sign-off)")
    d.spacer()
    d.body_bold("Exclusions: ", "Technijian My Archive is NOT included under this SOW. AFFG operates its own email-archiving platform, which will remain authoritative for long-term mail retention. Coordination with that platform is included in Phase 4.")

    # Section 2 — 7-phase scope with ticket tables
    d.spacer()
    d.section_header("2. 7-Phase Implementation Scope")
    d.body(
        "Project work is organized into thirty-seven (37) numbered tickets distributed across "
        "seven (7) sequential phases, 135 total estimated hours. Each phase below includes the "
        "narrative description and the full ticket summary with estimated hours and assigned POD."
    )

    tickets_by_phase = _load_tickets_by_phase()
    for phase_key, phase_title, narrative in PHASE_NARRATIVES:
        d.spacer()
        d.section_header(phase_title)
        d.body(narrative)
        d.spacer()
        rows = []
        total_hrs = 0.0
        phase_rows = tickets_by_phase.get(phase_key, [])
        for t in phase_rows:
            rows.append((t["id"], t["title"], t["hours"], t["assignee"]))
            try:
                total_hrs += float(t["hours"])
            except ValueError:
                pass
        rows.append((
            f"Phase Subtotal",
            f"{len(phase_rows)} tickets",
            f"{total_hrs:g} hrs",
            "",
        ))
        d.styled_table(
            ["Ticket ID", "Title", "Est Hours", "Assignee"],
            rows,
            total_indices=[len(rows) - 1],
        )

    # Section 3 — Regulatory Control Mapping
    d.page_break()
    d.section_header("3. Regulatory Control Mapping")
    d.body(
        "The table below maps the deliverables of this SOW against the AFFG IT Compliance Strategy "
        "regulatory coverage (REF: AFFG-CS-2026-04)."
    )
    d.spacer()
    d.styled_table(
        ["Regulation", "Requirements Covered", "Primary Control(s)"],
        [
            ("SEC Regulation S-P (2024 Amendments)", "10 of 10",
             "VDI isolation, DLP, MFA, CCPA-aligned incident response, 30-day customer breach notification, disposal rule, service-provider oversight"),
            ("FINRA Rule 4370", "2 of 2",
             "BCP customer disclosure, documented business continuity + recovery in Technijian environment"),
            ("FINRA Rule 3110", "3 of 3",
             "Per-agent dedicated VMs, individual accountability, supervisory audit trail with immutable retention"),
            ("SEC Rule 17a-4", "3 of 3",
             "7-year immutable retention on V365 Backup storage, coordinated archival with AFFG\u2019s existing email archive, eDiscovery workspace"),
        ],
    )

    # Section 4 — Deliverables
    d.spacer()
    d.section_header("4. Deliverables")
    for item in [
        "Live Technijian VDI environment with 11 dedicated Windows 11 personal desktops",
        "VDI egress IP allowlists approved and active at Schwab Advisor Services and Interactive Brokers",
        "M365 DLP policies deployed across Exchange Online, SharePoint, OneDrive, and Teams, with custom Sensitive Information Types for Schwab and IBKR",
        "Veeam 365 Backup operational with 7-year immutable retention (SEC 17a-4)",
        "My Audit UAM+DLP agents deployed across all 11 VDI desktops with 1-year retention",
        "Monthly compliance dashboard PDF (MFA rate, DLP events, VDI uptime, backup success, pen-test status, phishing completion)",
        "Penetration test report covering VDI external surface, egress controls, isolation, and privilege escalation",
        "Written Incident Response Plan (IRP)",
        "Breach Notification procedures and templates (SEC Reg S-P 30-day, CA Civ. Code 1798.82, FINRA, custodians)",
        "Service Provider Oversight Policy including the SEC Reg S-P 72-hour vendor incident notification clause",
        "BCP Customer Disclosure (FINRA Rule 4370) + Customer Information Disposal Policy",
        "Annual compliance sign-off process (principal attestation, audit evidence package)",
        "Final regulator-ready compliance binder (encrypted PDF bundle)",
    ]:
        d.bullet(item)

    # Section 5 — Pricing
    d.spacer()
    d.section_header("5. Pricing and Payment")
    d.body(
        "Labor for this SOW is billed at the hourly (not contracted) rates shown below. Total "
        "estimated labor is 135 hours at $12,795.00."
    )
    d.spacer()
    d.styled_table(
        ["Role", "Rate", "Hours", "Labor"],
        [
            ("US Tech Support (IRV-TS1)", "$150/hr (hourly, not contracted)", "64", "$9,600.00"),
            ("India Tech Support (CHD-TS1)", "$45/hr (hourly, not contracted)", "71", "$3,195.00"),
            ("TOTAL", "", "135 hrs", "$12,795.00"),
        ],
        total_indices=[2],
    )
    d.spacer()
    d.body_bold("Payment Schedule:", "")
    d.bullet("50% ($6,397.50) due upon SOW execution")
    d.bullet("25% ($3,198.75) upon Phase 4 completion (V365 Backup operational)")
    d.bullet("25% ($3,198.75) upon Phase 7 delivery (compliance docs accepted)")

    # Section 6 — Assumptions
    d.spacer()
    d.section_header("6. Assumptions")
    for item in [
        f"{CLIENT_SHORT} provides administrative credentials to the Microsoft 365 tenant and Entra ID with sufficient privilege to deploy Conditional Access, DLP, and eDiscovery.",
        f"{CLIENT_SHORT} provides timely approvals for the Schwab Advisor Services and Interactive Brokers IP allowlist requests, including any required authorized-signer signatures.",
        f"{CLIENT_SHORT}\u2019s existing M365 tenant is licensed at Microsoft 365 E3 or E5 (or equivalent add-ons) such that Data Loss Prevention and Conditional Access are available.",
        f"{CLIENT_SHORT}\u2019s existing email archive platform meets SEC Rule 17a-4 retention and WORM requirements and can be referenced authoritatively in the Service Provider Oversight documentation.",
        "The 11 named AFFG agents are available for VDI onboarding and training sessions during Weeks 3\u20134.",
        "Coordination with third-party vendors (Schwab, Interactive Brokers, AFFG\u2019s email archiver) proceeds on the standard timelines of those vendors; Technijian\u2019s performance dates assume reasonable vendor responsiveness.",
    ]:
        d.bullet(item)

    # Section 7 — Exclusions
    d.spacer()
    d.section_header("7. Exclusions")
    for item in [
        "Technijian My Archive (AFFG operates its own email-archiving platform, which remains authoritative for long-term mail retention).",
        "Microsoft 365 license costs (AFFG procures M365 licenses directly).",
        "Third-party public cloud consumption costs (all Technijian VDI resources are provisioned within Technijian\u2019s private cloud and billed via Schedule A, Part 1).",
        "Hardware procurement (AFFG continues to use its existing end-user devices as thin-client access points into VDI; this SOW does not include new workstation purchases).",
        "Any remediation beyond the scope described in the phase narratives; significant remediation work arising from pen-test findings will be scoped under a separate change order.",
    ]:
        d.bullet(item)

    # Section 8 — Ongoing Services
    d.spacer()
    d.section_header("8. Ongoing Services")
    d.body(
        f"Upon completion of this SOW, {CLIENT_SHORT} is billed the ongoing recurring charges set "
        f"forth in Schedule A of {AGREEMENT_NO} ($6,922.35/month). The $12,795.00 price of this SOW "
        "is a one-time implementation fee and is separate from the Schedule A monthly charges. The "
        "recurring Schedule A charges cover the VDI environment, endpoint security on VDI, all-user "
        "services, domain/site/IP services, shared storage, and Virtual Staff Support (unchanged "
        "from the prior Monthly Service Agreement)."
    )

    # Signatures
    d.signatures(client_name=CLIENT_UPPER, contact_name=SIGNER)
    d.footer_bar()
    d.save(os.path.join(OUT_DIR, "SOW-AFFG-003-Compliance-Strategy.docx"))


# ================================================================
# EXECUTIVE SUMMARY
# ================================================================
def build_executive_summary():
    d = BrandedDoc()
    d.cover_page(
        "Executive Summary",
        "Compliance Strategy Implementation \u2014 VDI + Regulatory Controls",
        client_name="American Fundstars Financial Group LLC",
        date="May 1, 2026",
    )
    d.page_break()

    # Overview
    d.section_header("Overview")
    d.body(
        "Technijian is pleased to present this Executive Summary for American Fundstars Financial "
        "Group LLC (\u201cAFFG\u201d). This document describes the recommended transition from AFFG\u2019s "
        "current IT posture \u2014 eleven (11) registered agents operating on unmanaged personal "
        "laptops \u2014 to a dedicated Technijian VDI (Virtual Desktop Infrastructure) environment "
        "that achieves full alignment with SEC Regulation S-P (2024 Amendments), FINRA Rule 4370, "
        "FINRA Rule 3110, and SEC Rule 17a-4."
    )
    d.body(
        "As a dually-registered Investment Adviser and Broker-Dealer handling customer Non-Public "
        "Information (NPI), PII, and financial account data, AFFG\u2019s current compliance posture "
        "carries material regulatory exposure. The proposed solution brings that posture from an "
        "estimated ~15% compliance coverage to 100% across the four governing frameworks."
    )
    d.body_bold("Document Package: ", "")
    d.bullet("MSA-AFFG \u2014 Master Service Agreement (legal body)")
    d.bullet("Schedule A \u2014 Monthly Managed Services pricing ($6,922.35/mo)")
    d.bullet("Schedule B \u2014 Subscription and License Services placeholder")
    d.bullet("Schedule C \u2014 Rate Card")
    d.bullet("SOW-AFFG-003 \u2014 IT Compliance & VDI Implementation ($12,795 one-time)")

    # Current vs. Future State
    d.page_break()
    d.section_header("Current vs. Future State")
    d.spacer()
    d.styled_table(
        ["Dimension", "Current State", "Future State (Post-Implementation)"],
        [
            ("Compliance Posture", "~15% coverage", "100% coverage"),
            ("Agent Workstations", "11 unmanaged personal laptops", "11 dedicated Technijian VDI desktops"),
            ("Data Location", "Client data on personal devices", "Zero client data on personal device \u2014 VDI is the boundary"),
            ("Schwab / IBKR Access", "From any personal device", "VDI egress IPs only (allowlisted)"),
            ("M365 Access", "From any personal device", "VDI-sourced sessions only (Entra Conditional Access)"),
            ("DLP Controls", "None", "M365 DLP across Exchange / SharePoint / OneDrive / Teams"),
            ("Email Archive", "AFFG existing platform", "AFFG existing platform (unchanged \u2014 remains authoritative)"),
            ("M365 Data Backup", "None", "Veeam 365 Backup with 7-year immutable retention"),
            ("User Activity Monitoring", "None", "My Audit UAM+DLP on all 11 VDI desktops"),
            ("Security Awareness Training", "None", "Huntress SAT with monthly phishing simulations"),
            ("MFA", "Inconsistent", "Unified MFA on VDI + M365 + Schwab + IBKR"),
            ("Incident Response Plan", "None documented", "Written IRP + breach notification templates + 72-hour vendor clause"),
        ],
    )

    # What We Are Adding and Why
    d.page_break()
    d.section_header("New Services \u2014 What and Why")
    d.body(
        "Every service added is directly tied to a regulatory requirement. Below is the mapping of "
        "each new monthly service to the specific rule(s) it satisfies."
    )
    d.spacer()
    d.styled_table(
        ["Service", "Monthly", "Why Added (Regulatory Driver)"],
        [
            ("Technijian VDI Workstations (11)", "$2,156.00",
             "Core compliance boundary. SEC Reg S-P Safeguards (prevent unauthorized access to NPI); FINRA Rule 3110 individual accountability (each agent = one dedicated VM = clean audit trail); FINRA 4370 alternate work location."),
            ("Endpoint Security on VDI (11)", "$291.50",
             "CrowdStrike + Huntress + DNS filtering + Patch Mgmt + Remote. SEC Reg S-P Safeguards Rule requires endpoint protection. 24/7 SOC via Huntress is the defensible baseline for a regulated firm."),
            ("My Audit UAM+DLP / 1-Year (11)", "$1,189.10",
             "FINRA Rule 3110 supervisory audit trail; SEC Reg S-P access monitoring; insider threat detection. 1-year online retention satisfies immediate supervisor review; older data archived via V365 backup."),
            ("SSO / 2FA (11)", "$132.00",
             "SEC Reg S-P access control; FINRA 3110 named accounts; unified MFA across VDI, M365, Schwab, IBKR."),
            ("Credential Manager (11)", "$55.00",
             "SEC Reg S-P access control; eliminates passwords-in-email, passwords-on-sticky-notes; satisfies password-complexity and rotation requirements."),
            ("Veeam 365 Backup (31 users)", "$124.00",
             "SEC Rule 17a-4 recordkeeping (7-year immutable retention for M365 data); ransomware protection; SEC Reg S-P data recovery."),
            ("Phishing Training / Huntress SAT (31 users)", "$186.00",
             "SEC Reg S-P employee training requirement; FINRA 3110 ongoing training obligation; reduces social-engineering attack surface on a high-value financial target."),
            ("Inky Anti-Spam Basic (31 users)", "$131.75",
             "Email-threat prevention for a firm where email is the #1 compromise vector; includes archiving capability."),
            ("VDI Site Assessment", "$50.00",
             "Monthly vulnerability scan of the VDI environment (subnets, firewall rules, egress IPs). FINRA 4370 network integrity."),
            ("Physical Office Site Assessment", "$50.00",
             "Continuing monthly scan of the 1 Park Plaza office network (Sophos firewall + Wireless APs from SOW-001 / SOW-002)."),
            ("Production Storage (2 TB)", "$400.00",
             "FSLogix user profile containers + VDI cache storage. Required infrastructure for VDI."),
            ("Backup Storage (2 TB)", "$100.00",
             "Destination for Veeam 365 Backup with 7-year WORM immutability \u2014 direct SEC 17a-4 satisfaction."),
        ],
        total_indices=[],
    )
    d.spacer()
    d.body_bold("Not Included: ", "")
    d.bullet("Technijian My Archive \u2014 AFFG\u2019s existing email archiving platform remains authoritative for 7-year mail retention; we coordinate around it rather than duplicating it.")

    # Investment & Savings Analysis
    d.page_break()
    d.section_header("Monthly Investment Change")
    d.spacer()
    d.styled_table(
        ["Line Item", "Prior MSA", "New MSA", "Change"],
        [
            ("16 Physical Desktops (endpoint services)", "$424.00", "Replaced by VDI (11 agents)", "\u2014"),
            ("User Services (V365 + PHT)", "$263.50", "$441.75 (adds ASB, V365 at updated rate)", "+$178.25"),
            ("Domain / Site / IP (DKIM + SA + RTPT)", "$112.00", "$162.00 (adds VDI Site Assessment)", "+$50.00"),
            ("Shared Storage (Production + Backup)", "$0.00", "$500.00 (new)", "+$500.00"),
            ("Technijian VDI + My Audit + SSO + CRM", "$0.00", "$3,823.60 (new)", "+$3,823.60"),
            ("Virtual Staff Support (UNCHANGED)", "$1,995.00", "$1,995.00", "$0.00"),
            ("TOTAL MONTHLY", "$2,794.50", "$6,922.35", "+$4,127.85"),
        ],
        total_indices=[6],
    )
    d.spacer()
    d.body_bold("Key Point: ", "")
    d.body(
        "The new monthly charges REPLACE the prior workstation and user services \u2014 they are "
        "NOT additive. Virtual Staff Support hours and rates are unchanged. The +$4,127.85/month "
        "increase funds the compliance infrastructure required to avoid SEC enforcement actions, "
        "FINRA examination deficiency letters, and customer-data-breach liability."
    )

    # One-Time Implementation
    d.spacer()
    d.section_header("One-Time Implementation")
    d.body(
        "The 7-phase implementation (18 weeks, see SOW-AFFG-003) totals $12,795 in one-time "
        "professional services fees, blended across US and India technical resources:"
    )
    d.spacer()
    d.styled_table(
        ["Resource", "Rate", "Hours", "Labor"],
        [
            ("US Tech Support (IRV-TS1)", "$150/hr", "64", "$9,600.00"),
            ("India Tech Support (CHD-TS1)", "$45/hr", "71", "$3,195.00"),
            ("TOTAL", "", "135", "$12,795.00"),
        ],
        total_indices=[2],
    )
    d.spacer()
    d.body(
        "Payment schedule: 50% upon SOW execution / 25% upon Phase 4 completion / 25% upon Phase 7 "
        "delivery and compliance-documentation acceptance."
    )

    # Risk of Inaction
    d.page_break()
    d.section_header("Risk of Inaction")
    d.spacer()
    d.styled_table(
        ["Risk", "Potential Impact"],
        [
            ("SEC enforcement action under Reg S-P", "Monetary fines; cease-and-desist order; ongoing monitoring obligation"),
            ("FINRA examination deficiency letter", "Remediation timeline imposed by FINRA; potential suspension of activities"),
            ("Customer data breach (NPI or PII exposure)", "Customer notification obligations; litigation exposure; reputational damage"),
            ("SEC 17a-4 recordkeeping finding", "Retention-period penalties; rule 17a-4 certification revocation risk"),
            ("FINRA 3110 supervisory finding", "Supervisory-system deficiency letter; required remediation plan"),
        ],
    )

    # Year 1 Summary
    d.spacer()
    d.section_header("Year 1 Investment Summary")
    d.spacer()
    d.styled_table(
        ["Component", "Amount"],
        [
            ("One-Time Implementation (SOW-AFFG-003)", "$12,795.00"),
            ("Ongoing Monthly (12 x $6,922.35)", "$83,068.20"),
            ("TOTAL YEAR 1", "$95,863.20"),
        ],
        total_indices=[2],
    )
    d.spacer()
    d.body_bold("Note: ", "")
    d.body(
        "All figures are presented monthly; the 12-month total above is shown strictly for planning "
        "purposes and is not the contracted commitment amount. The Agreement term is twelve (12) "
        "months with automatic renewal unless notice of non-renewal is provided sixty (60) days "
        "before term end."
    )

    # Next Steps
    d.spacer()
    d.section_header("Next Steps")
    d.bullet("Review MSA-AFFG, Schedules A / B / C, and SOW-AFFG-003.")
    d.bullet("Confirm named agent list for the 11 VDI desktops.")
    d.bullet("Confirm AFFG Microsoft 365 license tier (E3 / E5) for DLP capability alignment.")
    d.bullet("Sign MSA + SOW-AFFG-003 to initiate Phase 1 (VDI Foundation).")
    d.bullet("Technijian kicks off within five (5) business days of execution.")

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "AFFG-Executive-Summary.docx"))


# ================================================================
# Main
# ================================================================
if __name__ == "__main__":
    build_msa()
    build_schedule_a()
    build_schedule_b()
    build_schedule_c()
    build_sow_compliance()
    build_executive_summary()
    print("\nAll six AFFG documents generated successfully.")
