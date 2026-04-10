"""
AMP Complete Document Suite Generator
Generates MSA and Proposal of Services for AmPac Business Capital
All using consistent Technijian Brand Guide 2026 formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Brand Colors (Technijian Brand Guide 2026) --
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "technijian-logo.png")


def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


def set_grey_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        return r

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_borders(tbl)

    def cover_page(self, title, subtitle, client_name="AmPac Business Capital", date="April 10, 2026"):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(LOGO_PATH, width=Inches(3.0))
        div = self.doc.add_table(rows=1, cols=3)
        div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in div.rows[0].cells:
            c.text = ""
            c.paragraphs[0].paragraph_format.space_before = Pt(0)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div.rows[0].cells[1], HEX_ORANGE)
        div.rows[0].cells[0].width = Inches(2.25)
        div.rows[0].cells[1].width = Inches(2.5)
        div.rows[0].cells[2].width = Inches(2.25)
        remove_borders(div)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Prepared for ", size=12, color=BRAND_GREY)
        self.run(p, client_name, size=12, bold=True, color=CORE_BLUE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.accent_bar(HEX_ORANGE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar = tbl.rows[0].cells[0]
        set_cell_shading(bar, HEX_BLUE)
        bar.width = Inches(0.08)
        bar.text = ""
        bar.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar.paragraphs[0].paragraph_format.space_after = Pt(0)
        tc = tbl.rows[0].cells[1]
        tc.text = ""
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_borders(tbl)

    def part_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_borders(tbl)

    def styled_table(self, headers, rows, total_indices=None):
        total_indices = total_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            set_cell_shading(c, HEX_BLUE)
            c.text = ""
            self.run(c.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for ri, rd in enumerate(rows):
            row = tbl.add_row()
            is_total = ri in total_indices
            for i, val in enumerate(rd):
                c = row.cells[i]
                c.text = ""
                if is_total:
                    set_cell_shading(c, HEX_OFF_WHITE)
                    self.run(c.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(c.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        set_grey_borders(tbl)
        return tbl

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            self.run(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, num, text, indent=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent
        self.run(p, f"{prefix}{num} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self, client_name="AMPAC BUSINESS CAPITAL", contact_name="[AUTHORIZED SIGNER]"):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)
        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", "Name: _________________________________",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)
        self.spacer(12)
        p = self.doc.add_paragraph()
        self.run(p, client_name, size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", f"Name: {contact_name}",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        self.doc.save(filename)
        print(f"Created {filename}")


# ================================================================
#  MSA
# ================================================================
def build_msa():
    d = BrandedDoc()
    d.cover_page("Master Service Agreement", "AmPac Business Capital")
    d.page_break()

    d.part_header("MASTER SERVICE AGREEMENT")
    d.spacer()
    d.body("Agreement Number: MSA-AMP-2026")
    d.body("Effective Date: April 10, 2026")
    d.spacer()
    d.body("This Master Service Agreement (\u201cAgreement\u201d) is entered into by and between:")
    d.spacer()
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", "")
    d.body("18 Technology Drive, Suite 141")
    d.body("Irvine, California 92618")
    d.spacer()
    d.body("and")
    d.spacer()
    d.body_bold("AmPac Business Capital (\u201cClient\u201d)", "")
    d.body("[CLIENT ADDRESS]")
    d.body("[CITY, STATE ZIP]")
    d.body("Primary Contact: [AUTHORIZED SIGNER]")
    d.spacer()
    d.body("(collectively, the \u201cParties\u201d)")

    # SECTION 1
    d.spacer()
    d.section_header("SECTION 1 \u2014 SCOPE OF SERVICES")

    d.numbered("1.01.", "Services. Technijian shall provide the services described in the Schedules attached to this Agreement, which are incorporated herein by reference:")
    d.bullet("Schedule A \u2014 Monthly Managed Services (Cybersecurity, Email Protection, Compliance, Support)")
    d.bullet("Schedule C \u2014 Rate Card")
    d.body("Additional services may be provided through Statements of Work (\u201cSOWs\u201d) executed under this Agreement.")

    d.numbered("1.02.", "Standard of Care. Technijian shall perform all services in a professional and workmanlike manner, consistent with industry standards for managed IT service providers.")

    d.numbered("1.03.", "Service Level Agreement. The service levels applicable to the services are set forth in Schedule A. Technijian shall use commercially reasonable efforts to meet the service levels described therein.")

    d.numbered("1.04.", "Client Responsibilities. Client shall:")
    d.numbered("(a)", "Provide Technijian with reasonable access to Client\u2019s systems, facilities, and personnel as necessary for Technijian to perform the services;", indent=1)
    d.numbered("(b)", "Designate a primary point of contact for communications with Technijian;", indent=1)
    d.numbered("(c)", "Maintain current and accurate information regarding Client\u2019s systems and infrastructure;", indent=1)
    d.numbered("(d)", "Comply with all applicable laws and regulations in connection with its use of the services; and", indent=1)
    d.numbered("(e)", "Be solely responsible for the security and management of Client\u2019s account credentials and passwords.", indent=1)

    d.numbered("1.05.", "Independent Contractor. Technijian is an independent contractor. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the Parties.")

    # SECTION 2
    d.spacer()
    d.section_header("SECTION 2 \u2014 TERM AND RENEWAL")

    d.numbered("2.01.", "Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the \u201cInitial Term\u201d).")
    d.numbered("2.02.", "Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a \u201cRenewal Term\u201d), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term. Technijian shall send Client a written renewal reminder at least thirty (30) days prior to each renewal date, which shall restate the auto-renewal terms and cancellation method.")
    d.numbered("2.03.", "Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party.")
    d.numbered("2.04.", "Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party:")
    d.numbered("(a)", "Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days after receiving written notice of the breach; or", indent=1)
    d.numbered("(b)", "Becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.", indent=1)

    d.numbered("2.05.", "Effect of Termination.")
    d.numbered("(a)", "Upon termination, Client shall pay all fees and charges for services rendered through the date of termination, including any remaining obligations for annual licenses and subscriptions procured on Client\u2019s behalf.", indent=1)
    d.numbered("(b)", "Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, subject to payment of applicable fees.", indent=1)
    d.numbered("(c)", "Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format, provided Client is not in breach of this Agreement.", indent=1)
    d.numbered("(d)", "The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 9.09 (Personnel Transition Fee), and Section 10 (Data Protection).", indent=1)

    d.page_break()

    # SECTION 3
    d.section_header("SECTION 3 \u2014 PAYMENT")

    d.numbered("3.01.", "Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.")
    d.numbered("3.02.", "Invoice Types. Client may receive the following types of invoices from Technijian during the term of this Agreement. Each invoice will clearly identify its type, the applicable Schedule or SOW, and the billing period or delivery event.")
    d.numbered("(a)", "Monthly Service Invoice. Issued on the first business day of each month for recurring managed services under Schedule A (cybersecurity, email protection, compliance, support). Billed in advance for the upcoming month.", indent=1)
    d.numbered("(b)", "Weekly In-Contract Invoice. Issued every Friday for Virtual Staff (contracted support) services performed under Schedule A during the preceding week (Monday through Friday). Each invoice includes: (i) a listing of each support ticket addressed; (ii) the assigned resource, role, and hours spent per ticket; (iii) a description of the work performed per ticket; (iv) whether work was performed during normal or after-hours; and (v) the current running balance for each contracted role.", indent=1)
    d.numbered("(c)", "Weekly Out-of-Contract Invoice. Issued every Friday for labor services performed outside the scope of any active Schedule or SOW \u2014 including ad-hoc support requests, emergency work, and services performed under a SOW with hourly billing. Each invoice includes: (i) a listing of each support ticket or task performed; (ii) the assigned resource, role, and applicable hourly rate from the Rate Card (Schedule C); (iii) time entries with hours billed per activity (in 15-minute increments); (iv) whether work was performed during normal or after-hours; and (v) the total hours and total amount for the week.", indent=1)
    d.numbered("(d)", "Equipment and Materials Invoice. Issued upon delivery or procurement of hardware, software licenses (perpetual), or other tangible goods on Client\u2019s behalf. Each invoice includes: (i) item description, manufacturer, and model/part number; (ii) quantity and unit price; (iii) applicable sales tax; (iv) shipping and handling charges, if any; and (v) total amount due. Title to equipment shall not pass to Client until payment is received in full.", indent=1)
    d.numbered("(e)", "Project Milestone Invoice. Issued upon completion of a project milestone as defined in an applicable SOW. The milestone, deliverables, and invoiced amount are as specified in the payment schedule of the SOW. Milestone invoices are billed in arrears upon acceptance of the deliverables or deemed acceptance under the SOW\u2019s acceptance provisions.", indent=1)
    d.numbered("3.03.", "Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified in the applicable Schedule or SOW.")
    d.numbered("3.04.", "Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated as simple interest from the date payment was due until the date payment is received in full. Late fees are payable in addition to the outstanding principal amount. The Parties acknowledge that the late fee represents a reasonable estimate of Technijian\u2019s administrative costs and damages resulting from late payment, including cash-flow disruption, collection overhead, and the cost of carrying accounts receivable, and is not intended as a penalty.")
    d.numbered("3.05.", "Disputed Invoices.")
    d.numbered("(a)", "Weekly Invoices (In-Contract and Out-of-Contract). Because weekly invoices include detailed ticket descriptions and time entries, Client shall have thirty (30) days from the invoice date to review and dispute any portion of a weekly invoice. Client shall notify Technijian in writing, specifying the ticket number(s) and nature of the dispute. Undisputed tickets and time entries on the same invoice shall remain payable by the due date. Failure to provide a timely written dispute notice within the thirty (30) day period shall constitute acceptance of all tickets and time entries on the invoice.", indent=1)
    d.numbered("(b)", "All Other Invoices. For monthly service invoices, equipment invoices, and project milestone invoices, Client shall notify Technijian in writing within fifteen (15) days of the invoice date if Client disputes any portion, specifying the nature and basis of the dispute. Client shall pay all undisputed amounts by the due date. Failure to provide a timely written dispute notice shall constitute acceptance of the invoice.", indent=1)
    d.numbered("(c)", "Resolution. The Parties shall work in good faith to resolve any invoice dispute within thirty (30) days of the dispute notice. If the dispute results in an adjustment, Technijian shall issue a credit memo or revised invoice within ten (10) business days of resolution.", indent=1)
    d.numbered("3.06.", "Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend services under the Schedule or SOW associated with the unpaid invoice until payment is received in full, including all accrued late fees. If Client fails to pay any undisputed invoice within sixty (60) days of the due date, Technijian may suspend all services under this Agreement and any related Schedules or SOWs. Recurring fees for suspended services shall continue to accrue for a period not to exceed thirty (30) days following the date of suspension, after which Technijian may terminate the affected Schedule, SOW, or this Agreement upon written notice. Suspension of services shall not relieve Client of its payment obligations.")
    d.numbered("3.07.", "Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian\u2019s income.")

    # SECTION 4
    d.spacer()
    d.section_header("SECTION 4 \u2014 CONFIDENTIALITY")

    d.numbered("4.01.", "Definition. \u201cConfidential Information\u201d means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information.")
    d.numbered("4.02.", "Obligations. Each Party shall:")
    d.numbered("(a)", "Hold the other Party\u2019s Confidential Information in confidence using at least the same degree of care it uses for its own confidential information, but not less than reasonable care;", indent=1)
    d.numbered("(b)", "Not disclose Confidential Information to third parties without prior written consent, except to employees, agents, and subcontractors who have a need to know and are bound by equivalent obligations; and", indent=1)
    d.numbered("(c)", "Not use Confidential Information for any purpose other than performing obligations under this Agreement.", indent=1)
    d.numbered("4.03.", "Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known to the receiving Party prior to disclosure, is independently developed, or is received from a third party without restriction.")
    d.numbered("4.04.", "Compelled Disclosure. If required by law or court order to disclose Confidential Information, the receiving Party shall provide prompt written notice to the disclosing Party and cooperate in seeking a protective order.")
    d.numbered("4.05.", "Duration. Confidentiality obligations shall survive termination for a period of three (3) years.")

    # SECTION 5
    d.spacer()
    d.section_header("SECTION 5 \u2014 LIMITATION OF LIABILITY")

    d.numbered("5.01.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, NEITHER PARTY\u2019S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM (THE \u201cSTANDARD CAP\u201d).")
    d.numbered("5.02.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL, REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE OR WHETHER EITHER PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.")
    d.numbered("5.03.", "ENHANCED CAP FOR CERTAIN CLAIMS. FOR CLAIMS ARISING FROM BREACHES OF SECTION 4 (CONFIDENTIALITY), SECTION 10 (DATA PROTECTION), INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, THE TOTAL AGGREGATE LIABILITY OF THE RESPONSIBLE PARTY SHALL NOT EXCEED THREE (3) TIMES THE STANDARD CAP DEFINED IN SECTION 5.01 (THE \u201cENHANCED CAP\u201d). THE ENHANCED CAP SHALL NOT APPLY TO LIABILITY ARISING FROM A PARTY\u2019S WILLFUL AND INTENTIONAL MISAPPROPRIATION OF THE OTHER PARTY\u2019S CONFIDENTIAL INFORMATION OR CLIENT DATA, FOR WHICH LIABILITY SHALL BE UNCAPPED.")
    d.numbered("5.04.", "Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data, Client acknowledges that Client is solely responsible for maintaining backup copies of its data.")

    d.page_break()

    # SECTION 6
    d.section_header("SECTION 6 \u2014 INDEMNIFICATION")

    d.numbered("6.01.", "By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from Technijian\u2019s gross negligence or willful misconduct in performing the services.")
    d.numbered("6.02.", "By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from Client\u2019s use of the services in violation of applicable law, Client\u2019s breach of this Agreement, or any data, content, or materials provided by Client.")
    d.numbered("6.03.", "Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party\u2019s prior written consent.")

    # SECTION 7
    d.spacer()
    d.section_header("SECTION 7 \u2014 INTELLECTUAL PROPERTY")

    d.numbered("7.01.", "Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes.")
    d.numbered("7.02.", "Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property.")
    d.numbered("7.03.", "Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW.")

    # SECTION 8
    d.spacer()
    d.section_header("SECTION 8 \u2014 DISPUTE RESOLUTION")

    d.numbered("8.01.", "Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations for a period of thirty (30) days.")
    d.numbered("8.02.", "Mediation. If not resolved, the Parties shall submit the dispute to mediation in Orange County, California, for a period not to exceed sixty (60) days.")
    d.numbered("8.03.", "Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association in Orange County, California.")
    d.numbered("8.04.", "Injunctive Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief to prevent irreparable harm.")

    # SECTION 9
    d.spacer()
    d.section_header("SECTION 9 \u2014 GENERAL PROVISIONS")

    d.numbered("9.01.", "Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties.")
    d.numbered("9.02.", "Amendment. This Agreement may only be amended by a written instrument signed by both Parties.")
    d.numbered("9.03.", "Severability. If any provision is found invalid or unenforceable, the remaining provisions shall continue in full force.")
    d.numbered("9.04.", "Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party.")
    d.numbered("9.05.", "Assignment. Neither Party may assign this Agreement without prior written consent, except in connection with a merger, acquisition, or sale of substantially all assets.")
    d.numbered("9.06.", "Force Majeure.")
    d.numbered("(a)", "Neither Party shall be liable for delays or failures in performance caused by events beyond its reasonable control, including natural disasters, acts of government, labor disputes, pandemics, cyberattacks on critical infrastructure, or failures of major third-party infrastructure services that are not attributable to the affected Party\u2019s failure to implement reasonable redundancy (\u201cForce Majeure Event\u201d).", indent=1)
    d.numbered("(b)", "The affected Party shall notify the other Party in writing within five (5) business days of becoming aware of a Force Majeure Event and shall use commercially reasonable efforts to mitigate the impact and resume performance.", indent=1)
    d.numbered("(c)", "Payment obligations are not excused by a Force Majeure Event.", indent=1)
    d.numbered("(d)", "If a Force Majeure Event prevents performance of a material portion of the services for more than ninety (90) consecutive days, either Party may terminate the affected Schedule, SOW, or this Agreement upon fifteen (15) days written notice without liability for early termination fees, and Client shall pay only for services actually rendered through the date of termination.", indent=1)
    d.numbered("9.07.", "Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier.")
    d.numbered("9.08.", "Governing Law. This Agreement shall be governed by the laws of the State of California.")
    d.numbered("9.09.", "Personnel Transition Fee. The Parties acknowledge that each invests significant resources in recruiting, training, and retaining skilled personnel. If either Party hires (whether as an employee or independent contractor) any individual who was an employee of the other Party and who was directly involved in performing or receiving services under this Agreement, and such hiring occurs during the term of this Agreement or within twelve (12) months following termination, the hiring Party shall pay the other Party a personnel transition fee equal to 25% of the hired individual\u2019s first-year annual compensation (base salary or annualized contractor fees). This fee represents a reasonable estimate of the non-hiring Party\u2019s recruiting and training costs and is not intended as a restraint on trade or employment. This Section does not restrict any individual\u2019s right to seek or obtain employment, and shall not apply to individuals who: (a) respond to general public job postings or advertisements not specifically targeted at the other Party\u2019s employees; or (b) are referred by a third-party recruiting firm without the hiring Party\u2019s direction to target the other Party\u2019s employees.")
    d.numbered("9.10.", "Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.")

    # SECTION 10
    d.spacer()
    d.section_header("SECTION 10 \u2014 DATA PROTECTION")

    d.numbered("10.01.", "CCPA/CPRA Compliance. To the extent Technijian processes, stores, or has access to personal information (as defined under the California Consumer Privacy Act, as amended by the California Privacy Rights Act, Cal. Civ. Code \u00a7 1798.100 et seq., collectively \u201cCCPA\u201d) on behalf of Client, Technijian acts as a \u201cservice provider\u201d as defined under Cal. Civ. Code \u00a7 1798.140(ag) and shall:")
    d.numbered("(a)", "Process such personal information only as necessary to perform the services and in accordance with Client\u2019s documented instructions and this Agreement;", indent=1)
    d.numbered("(b)", "Not sell, share, retain, use, or disclose personal information for any purpose other than performing the services, including not using personal information for targeted advertising or cross-context behavioral advertising;", indent=1)
    d.numbered("(c)", "Not combine personal information received from Client with personal information received from other sources or collected from Technijian\u2019s own interactions with individuals, except as expressly permitted by the CCPA to perform the services;", indent=1)
    d.numbered("(d)", "Implement reasonable security measures appropriate to the nature of the personal information, consistent with the requirements of Section 10.03;", indent=1)
    d.numbered("(e)", "Cooperate with Client in responding to verifiable consumer rights requests under the CCPA (including access, deletion, correction, and opt-out requests) within ten (10) business days of Client\u2019s request;", indent=1)
    d.numbered("(f)", "Notify Client within five (5) business days if Technijian determines that it can no longer meet its obligations as a service provider under the CCPA;", indent=1)
    d.numbered("(g)", "Ensure that all subcontractors who process personal information on behalf of Client are bound by written agreements containing data protection obligations at least as protective as those in this Section 10, and Technijian shall remain liable for the acts and omissions of its subcontractors with respect to personal information;", indent=1)
    d.numbered("(h)", "Permit Client, upon thirty (30) days written notice and no more than once per twelve (12) month period, to audit or inspect Technijian\u2019s data processing practices to verify compliance with this Section 10, or, at Technijian\u2019s option, provide Client with a summary of a recent independent third-party audit (such as SOC 2 Type II) covering the relevant controls; and", indent=1)
    d.numbered("(i)", "Certify that Technijian understands the restrictions in this Section 10 and will comply with them.", indent=1)

    d.numbered("10.02.", "Security Incident Notification. If Technijian becomes aware of a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Client Data (\u201cSecurity Incident\u201d), Technijian shall: (a) notify Client in writing without unreasonable delay and in no event later than forty-eight (48) hours after becoming aware of the Security Incident; (b) provide Client with sufficient information to enable Client to comply with its obligations under California Civil Code \u00a7 1798.82 (data breach notification), including the categories and approximate number of records affected, the nature of the incident, and the measures taken or proposed to address it; (c) cooperate with Client\u2019s investigation of the Security Incident; and (d) take reasonable steps to contain and remediate the Security Incident.")

    d.numbered("10.03.", "Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards designed to protect Client Data from unauthorized access, use, or disclosure, consistent with industry standards for managed IT service providers. Such safeguards shall include, at a minimum: (a) encryption of Client Data in transit and at rest; (b) access controls limiting access to authorized personnel; (c) regular security assessments and vulnerability testing; and (d) employee security awareness training.")

    d.numbered("10.04.", "Regulatory Compliance. If Client is subject to HIPAA, PCI DSS, GDPR, or other industry-specific data protection requirements, the Parties shall execute a separate addendum addressing the additional obligations applicable to the regulated data.")

    d.numbered("10.05.", "Data Return and Deletion. Upon termination of this Agreement or upon Client\u2019s written request, Technijian shall securely delete or return all Client Data in its possession within thirty (30) days, using methods consistent with NIST SP 800-88 or equivalent standards, and shall certify such deletion in writing upon request.")

    d.signatures()

    d.spacer()
    d.body("Schedules:")
    d.bullet("Schedule A \u2014 Monthly Managed Services")
    d.bullet("Schedule C \u2014 Rate Card")

    # ── SCHEDULE A ──
    d.page_break()
    d.part_header("SCHEDULE A \u2014 MONTHLY MANAGED SERVICES")
    d.spacer()
    d.body("Agreement Number: MSA-AMP-2026")
    d.body("Effective Date: April 10, 2026")
    d.spacer()

    # Desktop/Endpoint Security
    d.section_header("Part 1 \u2014 Desktop / Endpoint Security")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("CrowdStrike Falcon EDR", "AVD", "$8.50/endpoint", "40", "$340.00"),
            ("Huntress Managed Detection", "AVMH", "$6.00/endpoint", "40", "$240.00"),
            ("Cisco Umbrella DNS Filtering", "SI", "$6.00/endpoint", "40", "$240.00"),
            ("Patch Management (ManageEngine)", "PMW", "$4.00/endpoint", "40", "$160.00"),
            ("My Remote (ScreenConnect)", "MR", "$2.00/endpoint", "40", "$80.00"),
            ("Subtotal \u2014 Endpoint Security", "", "", "40 endpoints", "$1,060.00"),
        ],
        total_indices=[5],
    )

    # Email Protection
    d.spacer()
    d.section_header("Part 2 \u2014 Email Protection")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("Inky Anti-Phishing + Archiving", "ASP", "$4.25/user", "34", "$144.50"),
            ("DMARC/DKIM Monitoring", "DKIM", "$20.00/domain", "1", "$20.00"),
            ("Site Assessment (Network Detective)", "SA", "$50.00/site", "1", "$50.00"),
            ("Email Backup Storage", "EBS", "$50.00/TB", "2 TB", "$100.00"),
            ("Subtotal \u2014 Email Protection", "", "", "", "$314.50"),
        ],
        total_indices=[4],
    )

    # Compliance Add-Ons
    d.spacer()
    d.section_header("Part 3 \u2014 Compliance Add-Ons")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("Veeam 365 Backup", "V365", "$2.50/user", "34", "$85.00"),
            ("Phishing Training (Security Awareness)", "PHT", "$6.00/user", "34", "$204.00"),
            ("Subtotal \u2014 Compliance", "", "", "", "$289.00"),
        ],
        total_indices=[2],
    )

    # Virtual Staff Support
    d.spacer()
    d.section_header("Part 4 \u2014 Virtual Staff Support")
    d.spacer()
    d.body("40 hours per month (1 hour per device). 50/50 day/night split; day hours split 50/50 US/India.")
    d.spacer()
    d.styled_table(
        ["Role", "Hours", "Rate", "Monthly"],
        [
            ("US Tech Support (Day)", "10", "$125/hr (contracted)", "$1,250.00"),
            ("India Night Support (US biz hrs)", "20", "$30/hr (contracted)", "$600.00"),
            ("India Day Support (US after-hrs)", "10", "$15/hr (contracted)", "$150.00"),
            ("Subtotal \u2014 Virtual Staff Support", "40", "", "$2,000.00"),
        ],
        total_indices=[3],
    )

    # Grand Total
    d.spacer()
    d.section_header("Monthly Investment Summary")
    d.spacer()
    d.styled_table(
        ["Category", "Monthly"],
        [
            ("Desktop / Endpoint Security (40 endpoints)", "$1,060.00"),
            ("Email Protection (34 users, 1 domain, 1 site)", "$314.50"),
            ("Compliance Add-Ons (34 users)", "$289.00"),
            ("Virtual Staff Support (40 hrs/mo)", "$2,000.00"),
            ("TOTAL MONTHLY INVESTMENT", "$3,663.50"),
        ],
        total_indices=[4],
    )

    # ── SCHEDULE C ──
    d.page_break()
    d.part_header("SCHEDULE C \u2014 RATE CARD")
    d.spacer()
    d.body("Agreement Number: MSA-AMP-2026")
    d.body("Effective Date: April 10, 2026")
    d.body("Normal Business Hours: Monday \u2013 Friday, 8:00 AM \u2013 6:00 PM Pacific Time")
    d.spacer()

    d.section_header("US-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr"),
            ("Developer", "$150/hr", "N/A", "$125/hr"),
            ("Tech Support", "$150/hr", "$250/hr", "$125/hr"),
        ],
    )

    d.spacer()
    d.section_header("Offshore-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr"),
            ("Tech Support", "$15/hr", "$30/hr", "$10/hr"),
        ],
    )

    d.spacer()
    d.section_header("Project & Ad Hoc")
    d.spacer()
    d.styled_table(
        ["Service", "Rate", "Minimum"],
        [
            ("On-Site Support (US)", "$150/hr", "2-hr minimum"),
            ("Remote Support (ad hoc)", "$150/hr", "15-min increments"),
            ("Emergency / Critical", "$250/hr", "1-hr minimum"),
            ("Project Management", "$150/hr", "SOW-based"),
        ],
    )

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "MSA-AMP.docx"))


# ================================================================
#  PROPOSAL OF SERVICES
# ================================================================
def build_proposal():
    d = BrandedDoc()

    # ── Cover Page ──
    d.cover_page(
        "Proposal of Services",
        "Managed IT Services, Cybersecurity & Compliance Solutions",
        client_name="AmPac Business Capital",
        date="April 2026",
    )
    d.page_break()

    # ── Page 2: Overview ──
    d.part_header("OVERVIEW")
    d.spacer()
    d.body(
        "Technijian, Inc. is pleased to present this Proposal of Services to AmPac Business Capital "
        "(\u201cAMPAC\u201d). As an SBA 504 lender and Community Development Financial Institution (CDFI) "
        "operating across California, Arizona, and Nevada, AMPAC handles sensitive borrower financial "
        "data subject to stringent federal and state regulatory requirements."
    )
    d.spacer()
    d.body(
        "This proposal outlines a comprehensive managed IT services, cybersecurity, and compliance "
        "program designed specifically for AMPAC\u2019s regulatory environment, operational footprint, "
        "and risk profile. Every service was selected to address a specific compliance obligation, "
        "security gap, or operational need."
    )
    d.spacer()
    d.section_header("Environment Summary")
    d.spacer()
    d.styled_table(
        ["Component", "Count"],
        [
            ("Desktops / Endpoints", "40"),
            ("Users", "34"),
            ("Locations", "1"),
            ("Domains", "1"),
            ("States of Operation", "CA, AZ, NV"),
        ],
    )
    d.footer_bar()

    # ── Page 3: Regulatory Landscape ──
    d.page_break()
    d.part_header("REGULATORY LANDSCAPE")
    d.spacer()
    d.body(
        "As an SBA 504 lender and CDFI, AMPAC is subject to a complex web of overlapping federal and "
        "state regulations governing data security, consumer protection, and financial integrity. "
        "The following frameworks directly impact AMPAC\u2019s IT and cybersecurity obligations:"
    )
    d.spacer()
    d.styled_table(
        ["Framework", "Requirement"],
        [
            ("SBA Lending Standards", "Secure handling of borrower financial data, audit trails, records retention"),
            ("CDFI Fund Requirements", "Data integrity, reporting accuracy, program compliance documentation"),
            ("ECOA (Reg B)", "Fair lending data protection, anti-discrimination controls on borrower records"),
            ("AML/BSA", "Suspicious activity monitoring, secure transaction records, SAR filing controls"),
            ("Safeguards Rule (GLBA)", "Written infosec program, CISO designation, encryption, access controls, employee training, incident response, annual risk assessments"),
            ("FCRA", "Consumer credit data protection, secure disposal of credit reports"),
            ("CRA", "Community reinvestment data integrity, geographic lending data security"),
            ("CCPA / GLBA Privacy", "Consumer data protection, breach notification, opt-out rights, data minimization"),
        ],
    )
    d.spacer()
    d.section_header("Safeguards Rule \u2014 Key IT Requirements")
    d.spacer()
    d.body(
        "The FTC\u2019s Safeguards Rule (16 CFR Part 314), updated in 2023, imposes specific "
        "cybersecurity requirements on financial institutions including SBA lenders. AMPAC must maintain:"
    )
    d.bullet("A written information security program with a designated qualified individual (CISO)")
    d.bullet("Risk assessments conducted at least annually")
    d.bullet("Encryption of customer information in transit and at rest")
    d.bullet("Multi-factor authentication for all systems accessing customer data")
    d.bullet("Access controls limiting information access to authorized personnel")
    d.bullet("Continuous monitoring and detection of unauthorized activity")
    d.bullet("Employee security awareness training program")
    d.bullet("Incident response plan with documented procedures")
    d.bullet("Oversight of service providers with access to customer information")
    d.spacer()
    d.body(
        "Technijian\u2019s proposed service stack directly addresses each of these requirements, "
        "as detailed in the following sections."
    )
    d.footer_bar()

    # ── Pages 4-5: Service Categories ──
    d.page_break()
    d.part_header("SERVICE CATEGORIES & JUSTIFICATIONS")

    # Endpoint Security
    d.spacer()
    d.section_header("Endpoint Security \u2014 $1,060.00/mo (40 endpoints)")
    d.spacer()

    d.body_bold("CrowdStrike Falcon EDR ", "($8.50/endpoint \u2014 $340.00/mo)")
    d.body(
        "Next-generation endpoint detection and response (EDR) required by the Safeguards Rule for "
        "continuous threat detection. CrowdStrike uses AI-driven behavioral analysis to detect malware, "
        "ransomware, and fileless attacks on every workstation handling borrower financial data. "
        "Provides the audit trail and incident forensics required for regulatory examinations."
    )
    d.spacer()

    d.body_bold("Huntress Managed Detection & Response ", "($6.00/endpoint \u2014 $240.00/mo)")
    d.body(
        "24/7 Security Operations Center (SOC) monitoring with human threat hunters who investigate "
        "and remediate persistent threats. Addresses the Safeguards Rule requirement for continuous "
        "monitoring of information systems. Critical for detecting targeted attacks against financial "
        "institutions, including persistent footholds that automated tools miss."
    )
    d.spacer()

    d.body_bold("Cisco Umbrella DNS Filtering ", "($6.00/endpoint \u2014 $240.00/mo)")
    d.body(
        "DNS-layer security that blocks connections to malicious domains before they are established. "
        "Prevents data exfiltration of borrower PII and financial records by blocking command-and-control "
        "communications. Addresses Safeguards Rule network security controls and provides policy-based "
        "web filtering to enforce acceptable use."
    )
    d.spacer()

    d.body_bold("Patch Management \u2014 ManageEngine ", "($4.00/endpoint \u2014 $160.00/mo)")
    d.body(
        "Automated patch deployment for operating systems and third-party applications. Addresses the "
        "Safeguards Rule requirement for known vulnerability remediation. Prevents exploitation of "
        "unpatched financial software and maintains compliance with security baseline standards."
    )
    d.spacer()

    d.body_bold("My Remote \u2014 ScreenConnect ", "($2.00/endpoint \u2014 $80.00/mo)")
    d.body(
        "Encrypted remote management tool enabling Technijian to provide secure support without "
        "exposing AMPAC\u2019s network. All sessions are logged and auditable, compliant with "
        "Safeguards Rule access controls and service provider oversight requirements."
    )

    # Email Protection
    d.spacer()
    d.section_header("Email Protection \u2014 $314.50/mo (34 users, 1 domain, 1 site)")
    d.spacer()

    d.body_bold("Inky Anti-Phishing + Archiving ", "($4.25/user \u2014 $144.50/mo)")
    d.body(
        "AI-powered email threat detection that identifies and neutralizes phishing, spear-phishing, "
        "and Business Email Compromise (BEC) attacks. Financial institutions are the #1 target for "
        "BEC, which accounts for over $2.7 billion in annual losses (FBI IC3). Built-in email "
        "archiving satisfies SBA/CDFI record retention requirements and provides searchable audit "
        "access for regulatory examinations."
    )
    d.spacer()

    d.body_bold("DMARC/DKIM Monitoring ", "($20.00/domain \u2014 $20.00/mo)")
    d.body(
        "Email authentication protocols (SPF, DKIM, DMARC) prevent bad actors from spoofing "
        "AMPAC\u2019s domain to send fraudulent emails to borrowers, partners, or regulators. "
        "Protects AMPAC\u2019s brand reputation and prevents domain-based phishing attacks."
    )
    d.spacer()

    d.body_bold("Site Assessment \u2014 Network Detective ", "($50.00/site \u2014 $50.00/mo)")
    d.body(
        "Quarterly network vulnerability scans and security assessments. Directly satisfies the "
        "Safeguards Rule requirement for regular security assessments and generates the documentation "
        "needed for annual risk assessment reporting."
    )
    d.spacer()

    d.body_bold("Email Backup Storage ", "($50.00/TB \u00d7 2 TB \u2014 $100.00/mo)")
    d.body(
        "Redundant backup of all email data independent of Microsoft\u2019s infrastructure. Satisfies "
        "Safeguards Rule data protection requirements and ensures CDFI audit access to historical "
        "communications. Protects against ransomware and accidental deletion."
    )

    d.page_break()

    # Compliance Add-Ons
    d.section_header("Compliance Add-Ons \u2014 $289.00/mo (34 users)")
    d.spacer()

    d.body_bold("Veeam 365 Backup ", "($2.50/user \u2014 $85.00/mo)")
    d.body(
        "Complete Microsoft 365 data backup covering Exchange, SharePoint, OneDrive, and Teams. "
        "The Safeguards Rule requires data recovery capabilities and business continuity planning. "
        "Veeam provides granular recovery of individual emails, files, and SharePoint items, "
        "protecting against ransomware targeting financial data stored in the cloud."
    )
    d.spacer()

    d.body_bold("Phishing Training \u2014 Security Awareness ", "($6.00/user \u2014 $204.00/mo)")
    d.body(
        "Monthly simulated phishing campaigns and security awareness training for all 34 users. "
        "The Safeguards Rule mandates employee security awareness training for all personnel with "
        "access to customer information. Financial institutions face elevated social engineering "
        "risk due to the high value of borrower data. Program includes compliance tracking and "
        "reporting for regulatory examinations."
    )

    # Virtual Staff Support
    d.spacer()
    d.section_header("Virtual Staff Support \u2014 $2,000.00/mo (40 hrs)")
    d.spacer()

    d.body(
        "AMPAC receives 40 hours of dedicated technical support per month (1 hour per device), "
        "delivered by a blended US and India team providing 24/7 coverage:"
    )
    d.spacer()
    d.styled_table(
        ["Role", "Hours", "Rate", "Monthly", "Coverage"],
        [
            ("US Tech Support (Day)", "10", "$125/hr", "$1,250.00", "Complex issues, escalations, on-site"),
            ("India Night Support", "20", "$30/hr", "$600.00", "US business hours monitoring & tickets"),
            ("India Day Support", "10", "$15/hr", "$150.00", "US after-hours & overnight response"),
            ("Total", "40", "", "$2,000.00", "24/7 coverage"),
        ],
        total_indices=[3],
    )
    d.spacer()
    d.body_bold("Why this model? ", "")
    d.bullet("US technicians handle complex issues, security incidents, and vendor escalations during AMPAC\u2019s business hours")
    d.bullet("India Night team (operating during US daytime) provides first-response triage, routine ticket resolution, and monitoring")
    d.bullet("India Day team (operating during US nighttime) ensures after-hours and overnight coverage for critical alerts")
    d.bullet("Contracted rates save AMPAC approximately 35% compared to ad-hoc hourly rates from Schedule C")
    d.footer_bar()

    # ── Page 6: Monthly Investment Summary ──
    d.page_break()
    d.part_header("MONTHLY INVESTMENT SUMMARY")
    d.spacer()

    d.styled_table(
        ["Category", "Details", "Monthly"],
        [
            ("Desktop / Endpoint Security", "40 endpoints \u00d7 $26.50", "$1,060.00"),
            ("Email Protection", "34 users, 1 domain, 1 site, 2 TB backup", "$314.50"),
            ("Compliance Add-Ons", "34 users (backup + training)", "$289.00"),
            ("Virtual Staff Support", "40 hrs/mo blended US/India", "$2,000.00"),
            ("TOTAL MONTHLY INVESTMENT", "", "$3,663.50"),
        ],
        total_indices=[4],
    )

    d.spacer()
    d.section_header("Ad-Hoc Rate Comparison")
    d.spacer()
    d.body(
        "The contracted support model provides significant savings compared to purchasing the same "
        "hours at Technijian\u2019s standard ad-hoc rates (Schedule C):"
    )
    d.spacer()
    d.styled_table(
        ["Component", "Contracted", "Ad-Hoc Equivalent", "Savings"],
        [
            ("10 hrs US Tech Support", "$1,250.00", "$1,500.00 ($150/hr)", "$250.00"),
            ("20 hrs India Night", "$600.00", "$900.00 ($45/hr)", "$300.00"),
            ("10 hrs India Day", "$150.00", "$450.00 ($45/hr)", "$300.00"),
            ("Support Total", "$2,000.00", "$2,850.00", "$850.00 (30%)"),
        ],
        total_indices=[3],
    )
    d.spacer()
    d.body(
        "AMPAC saves $850.00 per month ($10,200 annually) on support labor alone by committing to "
        "the contracted model versus ad-hoc billing."
    )
    d.footer_bar()

    # ── Page 7: SLA Commitments ──
    d.page_break()
    d.part_header("SERVICE LEVEL COMMITMENTS")
    d.spacer()

    d.section_header("Response Times by Priority")
    d.spacer()
    d.styled_table(
        ["Priority", "Definition", "Response Time", "Update Frequency"],
        [
            ("Critical (P1)", "System-wide outage, security breach, or data loss event", "15 minutes", "Every 30 minutes"),
            ("High (P2)", "Major system degraded, multiple users impacted", "30 minutes", "Every 1 hour"),
            ("Medium (P3)", "Single user impacted, workaround available", "2 hours", "Every 4 hours"),
            ("Low (P4)", "Informational, enhancement request, or scheduled task", "4 hours", "Next business day"),
        ],
    )

    d.spacer()
    d.section_header("Infrastructure & Availability")
    d.spacer()
    d.bullet("99.9% uptime", bold_prefix="Infrastructure Uptime SLA: ")
    d.bullet("All AMPAC-facing services monitored 24/7/365 with automated alerting", bold_prefix="Monitoring: ")
    d.bullet("Named point of contact for all service delivery, quarterly business reviews", bold_prefix="Dedicated Account Manager: ")
    d.bullet("Monthly reports covering ticket volume, resolution times, security events, patch compliance, and phishing training scores", bold_prefix="Monthly Service Reports: ")
    d.bullet("Quarterly vulnerability assessment reports aligned with Safeguards Rule annual risk assessment", bold_prefix="Quarterly Security Reviews: ")
    d.footer_bar()

    # ── Page 8: Rate Card ──
    d.page_break()
    d.part_header("RATE CARD")
    d.spacer()
    d.body("Normal Business Hours: Monday \u2013 Friday, 8:00 AM \u2013 6:00 PM Pacific Time")
    d.spacer()

    d.section_header("US-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr"),
            ("Developer", "$150/hr", "N/A", "$125/hr"),
            ("Tech Support", "$150/hr", "$250/hr", "$125/hr"),
        ],
    )

    d.spacer()
    d.section_header("Offshore-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr"),
            ("Tech Support", "$15/hr", "$30/hr", "$10/hr"),
        ],
    )

    d.spacer()
    d.section_header("Project & Ad Hoc")
    d.spacer()
    d.styled_table(
        ["Service", "Rate", "Minimum"],
        [
            ("On-Site Support (US)", "$150/hr", "2-hr minimum"),
            ("Remote Support (ad hoc)", "$150/hr", "15-min increments"),
            ("Emergency / Critical", "$250/hr", "1-hr minimum"),
            ("Project Management", "$150/hr", "SOW-based"),
        ],
    )
    d.footer_bar()

    # ── Page 9: Agreement Structure & Next Steps ──
    d.page_break()
    d.part_header("AGREEMENT STRUCTURE & NEXT STEPS")
    d.spacer()

    d.section_header("Agreement Structure")
    d.spacer()
    d.body("The engagement between Technijian and AMPAC is governed by the following documents:")
    d.spacer()
    d.bullet("Governs the overall relationship, payment terms, confidentiality, liability, data protection, and general provisions. 12-month initial term with automatic annual renewal.", bold_prefix="Master Service Agreement (MSA): ")
    d.bullet("Attached to the MSA. Details all monthly managed services, pricing, and SLA commitments.", bold_prefix="Schedule A \u2014 Monthly Managed Services: ")
    d.bullet("Attached to the MSA. Full rate card for all labor categories (US, offshore, project, emergency).", bold_prefix="Schedule C \u2014 Rate Card: ")
    d.bullet("Executed separately for any project-based work (migrations, deployments, consulting engagements) with defined scope, timeline, and budget.", bold_prefix="Statements of Work (SOWs): ")

    d.spacer()
    d.section_header("Next Steps")
    d.spacer()

    d.numbered("1.", "Review this Proposal of Services and the accompanying Master Service Agreement (MSA-AMP-2026).")
    d.numbered("2.", "Schedule a call with Technijian to discuss any questions, customizations, or scope adjustments.")
    d.numbered("3.", "Execute the MSA and Schedules (electronic signature via DocuSign).")
    d.numbered("4.", "Technijian begins onboarding (typically 5\u20137 business days):")
    d.bullet("Agent deployment to all 40 endpoints (CrowdStrike, Huntress, Umbrella, Patch Management, ScreenConnect)")
    d.bullet("Email protection configuration (Inky, DMARC/DKIM)")
    d.bullet("Veeam 365 backup enrollment for all 34 users")
    d.bullet("Phishing training program launch")
    d.bullet("Support team onboarding and escalation path configuration")
    d.numbered("5.", "First monthly service report delivered within 30 days of go-live.")

    d.spacer()
    d.section_header("Contact")
    d.spacer()
    d.body_bold("Ravi Jain", " \u2014 CEO, Technijian, Inc.")
    d.body("rjain@technijian.com  |  949.379.8499")
    d.body("18 Technology Drive, Suite 141, Irvine, CA 92618")
    d.body("technijian.com")

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "AMP-Proposal-of-Services.docx"))


if __name__ == "__main__":
    build_msa()
    build_proposal()
    print("\nAll documents generated.")
