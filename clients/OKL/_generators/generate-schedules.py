"""
OKL Schedule A & Schedule C Generator
Uses Technijian Brand Guide 2026 formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")


def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_table_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


def set_light_grey_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        run = paragraph.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return run

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_table_borders(tbl)

    def cover_page(self, title, subtitle, msa_ref="MSA-OKL-2026", date="March 23, 2026"):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(LOGO_PATH, width=Inches(3.0))

        # Orange divider
        div_tbl = self.doc.add_table(rows=1, cols=3)
        div_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in div_tbl.rows[0].cells:
            cell.text = ""
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div_tbl.rows[0].cells[1], HEX_ORANGE)
        div_tbl.rows[0].cells[0].width = Inches(2.25)
        div_tbl.rows[0].cells[1].width = Inches(2.5)
        div_tbl.rows[0].cells[2].width = Inches(2.25)
        remove_table_borders(div_tbl)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Oaktree Law", size=12, bold=True, color=CORE_BLUE)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)

        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

        self.accent_bar(HEX_ORANGE)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar_cell = tbl.rows[0].cells[0]
        set_cell_shading(bar_cell, HEX_BLUE)
        bar_cell.width = Inches(0.08)
        bar_cell.text = ""
        bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        title_cell = tbl.rows[0].cells[1]
        title_cell.text = ""
        p = title_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_table_borders(tbl)

    def part_header(self, title):
        """Full-width blue banner for Part headers."""
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_table_borders(tbl)

    def styled_table(self, headers, rows, total_row_indices=None):
        total_row_indices = total_row_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            set_cell_shading(cell, HEX_BLUE)
            cell.text = ""
            self.run(cell.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for row_idx, row_data in enumerate(rows):
            row = tbl.add_row()
            is_total = row_idx in total_row_indices
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                cell.text = ""
                if is_total:
                    set_cell_shading(cell, HEX_OFF_WHITE)
                    self.run(cell.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(cell.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        set_light_grey_borders(tbl)
        return tbl

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold_prefix(self, bold_text, text, indent=False):
        if indent:
            p = self.doc.add_paragraph(style="List Bullet")
        else:
            p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def numbered_item(self, number, text, indent_level=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent_level
        self.run(p, f"{prefix}{number} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)

        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________",
                      "Name: _________________________________",
                      "Title: _________________________________",
                      "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

        self.spacer(12)
        p = self.doc.add_paragraph()
        self.run(p, "OAKTREE LAW", size=11, bold=True, color=DARK_CHARCOAL)
        sigs = [
            "By: ___________________________________",
            "Name: Ed Pits",
            "Title: _________________________________",
            "Date: _________________________________",
        ]
        for label in sigs:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)


# ════════════════════════════════════════════════════════════════
#  SCHEDULE A -- MONTHLY MANAGED SERVICES
# ════════════════════════════════════════════════════════════════

def build_schedule_a():
    d = BrandedDoc()

    # Cover page
    d.cover_page("Schedule A", "Monthly Managed Services")
    d.doc.add_page_break()

    # Header block
    d.section_header("SCHEDULE A \u2014 MONTHLY MANAGED SERVICES")
    d.body("Attached to Master Service Agreement MSA-OKL-2026")
    d.body("Effective Date: March 23, 2026")
    d.body(
        "This Schedule describes the Monthly Managed Services provided by Technijian, Inc. "
        "(\u201cTechnijian\u201d) to Oaktree Law (\u201cClient\u201d) under the Master Service Agreement."
    )
    d.spacer()

    # ── PART 1: ONLINE SERVICES ──
    d.part_header("PART 1 \u2014 ONLINE SERVICES")
    d.spacer()

    d.section_header("1.1 Description")
    d.body(
        "Online Services include managed infrastructure, security, monitoring, and related IT services "
        "delivered on a recurring monthly basis. Services are selected from the Technijian Price List "
        "and itemized in the attached Service Order."
    )

    d.section_header("1.2 Service Categories")
    d.body(
        "Online Services are organized into the following categories. Specific services, quantities, "
        "and pricing are detailed in the Service Order attached to this Schedule."
    )
    d.styled_table(
        ["Category", "Description"],
        [
            ("Cloud Infrastructure", "Virtual machines (vCores, memory, bandwidth), production/backup storage"),
            ("Server Management", "Patch management, image backup, antivirus (CrowdStrike, Huntress), remote access"),
            ("Microsoft Licensing", "Windows Server licensing (Std 2-Core packs)"),
            ("Backup & Recovery", "Veeam image backup, backup storage"),
        ],
    )

    d.spacer()
    d.section_header("1.3 Service Order")
    d.body(
        "The specific services, quantities, and monthly pricing for Client are detailed below. "
        "The Service Order may be updated by mutual written agreement."
    )
    d.styled_table(
        ["Service", "Code", "Qty", "Unit Price", "Monthly"],
        [
            ("VM1 \u2014 vCores (Primary Server)", "CL-VC", "4", "$6.25", "$25.00"),
            ("VM1 \u2014 Memory (GB)", "CL-GB", "16", "$0.63", "$10.08"),
            ("VM1 \u2014 Shared Bandwidth", "CL-SBW", "1", "$15.00", "$15.00"),
            ("VM2 \u2014 vCores (Cloudbrink)", "CL-VC", "2", "$6.25", "$12.50"),
            ("VM2 \u2014 Memory (GB)", "CL-GB", "4", "$0.63", "$2.52"),
            ("VM2 \u2014 Shared Bandwidth", "CL-SBW", "1", "$15.00", "$15.00"),
            ("Production Storage", "TB-PSTR", "1 TB", "$200.00", "$200.00"),
            ("Backup Storage", "TB-BSTR", "1 TB", "$50.00", "$50.00"),
            ("Infrastructure Subtotal", "", "", "", "$330.10"),
            ("Windows Server Std \u2014 2-Core Pack", "MS-STD", "4", "$5.25", "$21.00"),
            ("Licensing Subtotal", "", "", "", "$21.00"),
            ("CrowdStrike \u2014 Server", "AVS", "2", "$10.50", "$21.00"),
            ("Huntress \u2014 Server", "AVHS", "2", "$6.00", "$12.00"),
            ("Patch Management", "PMW", "2", "$4.00", "$8.00"),
            ("My Remote", "MR", "2", "$2.00", "$4.00"),
            ("Security Subtotal", "", "", "", "$45.00"),
            ("Image Backup (Veeam)", "IB", "2", "$15.00", "$30.00"),
            ("Backup Subtotal", "", "", "", "$30.00"),
            ("TOTAL MONTHLY", "", "", "", "$426.10"),
            ("TOTAL ANNUAL", "", "", "", "$5,113.20"),
        ],
        total_row_indices=[8, 10, 15, 17, 18, 19],
    )

    d.spacer()
    d.section_header("1.4 Service Levels")
    d.styled_table(
        ["Service Level", "Target"],
        [
            ("Infrastructure Uptime", "99.9% monthly (excluding scheduled maintenance)"),
            ("Scheduled Maintenance", "Tuesday evenings and Saturdays (with advance notice)"),
            ("Critical Incident Response", "Within 1 hour of notification"),
            ("Standard Support Response", "Within 4 business hours"),
            ("Emergency Maintenance", "As needed with reasonable notice"),
        ],
    )

    d.spacer()
    d.section_header("1.4a Service Level Remedies")
    d.numbered_item("(a)", "Service Credits. If Technijian fails to meet the Infrastructure Uptime target of 99.9% in any calendar month, Client shall be entitled to a service credit equal to 5% of the monthly recurring charges for the affected service for each full 0.1% below the target, up to a maximum credit of 25% of the monthly recurring charges for that service.")
    d.numbered_item("(b)", "Chronic Failure. If Technijian fails to meet the Infrastructure Uptime target for three (3) or more consecutive months, Client shall have the right to terminate the affected service without penalty upon thirty (30) days written notice.")
    d.numbered_item("(c)", "Credit Requests. To receive a service credit, Client must submit a written request within thirty (30) days of the end of the affected month. Service credits shall be applied against future invoices and shall not be paid as cash refunds.")
    d.numbered_item("(d)", "Exclusions. Service level targets and remedies do not apply during scheduled maintenance windows, force majeure events, or outages caused by Client\u2019s actions, third-party services not managed by Technijian, or factors outside Technijian\u2019s reasonable control.")
    d.numbered_item("(e)", "Sole Remedy. Service credits under this Section 1.4a are Client\u2019s sole and exclusive remedy for Technijian\u2019s failure to meet the applicable service levels.")

    d.spacer()
    d.section_header("1.5 Monitoring and Reporting")
    d.body("Technijian shall provide:")
    d.numbered_item("(a)", "24/7 monitoring of Client\u2019s infrastructure included in the Service Order;")
    d.numbered_item("(b)", "Monthly service reports summarizing uptime, incidents, and support activity; and")
    d.numbered_item("(c)", "Quarterly service reviews with Client\u2019s designated representative.")
    d.numbered_item("(d)", "Escalation Path. Support requests shall be escalated as follows: Tier 1 (initial response) \u2014 assigned technician; Tier 2 (within 4 hours if unresolved) \u2014 senior engineer or team lead; Tier 3 (within 8 hours if unresolved) \u2014 department manager or CTO Advisory.")

    d.doc.add_page_break()

    # ── PART 2: SIP TRUNK SERVICES ──
    d.part_header("PART 2 \u2014 SIP TRUNK SERVICES")
    d.spacer()

    d.section_header("2.1 Description")
    d.body("SIP Trunk Services include voice-over-IP telephony, SIP trunking, and related telecommunications services.")

    d.section_header("2.2 Service Components")
    d.styled_table(
        ["Component", "Description"],
        [
            ("SIP Trunk", "Primary voice connectivity with failover"),
            ("Voice Package", "Bundled calling plans (domestic, long distance)"),
            ("DID Numbers", "Direct inward dialing numbers"),
            ("E911", "Emergency services routing"),
        ],
    )

    d.spacer()
    d.section_header("2.3 Service Levels")
    d.styled_table(
        ["Service Level", "Target"],
        [
            ("Voice Uptime", "99.9% monthly"),
            ("Call Quality (MOS)", "4.0 or higher"),
            ("Number Porting", "Completed within 10 business days of request"),
        ],
    )

    d.spacer()
    d.section_header("2.4 SIP Trunk Service Level Remedies")
    d.body(
        "The service level remedies set forth in Section 1.4a (Service Credits, Chronic Failure, "
        "Credit Requests, Exclusions, and Sole Remedy) shall apply equally to SIP Trunk Service Levels, "
        'with "Voice Uptime" substituted for "Infrastructure Uptime" where applicable.'
    )

    d.body(
        "Note: SIP Trunk Services are not included in the current Service Order. This section is included "
        "for reference should Client elect to add SIP Trunk Services in the future.",
        space_after=4,
    )

    d.doc.add_page_break()

    # ── PART 3: VIRTUAL STAFF ──
    d.part_header("PART 3 \u2014 VIRTUAL STAFF (CONTRACTED SUPPORT)")
    d.spacer()

    d.section_header("3.1 Description")
    d.body(
        "Virtual Staff services provide Client with dedicated technology support personnel on a contracted "
        "basis. This service operates on a cycle-based billing model as described below."
    )

    d.section_header("3.2 Support Roles")
    d.styled_table(
        ["Role", "Location", "Hours", "Contracted Rate", "Hourly Rate", "After-Hours Rate"],
        [
            ("CTO Advisory", "United States", "Normal Business Hours", "$225/hr", "$250/hr", "$350/hr"),
            ("Developer", "United States", "Normal Business Hours", "$125/hr", "$150/hr", "N/A"),
            ("Tech Support", "United States", "Normal Business Hours", "$125/hr", "$150/hr", "$250/hr"),
            ("Developer", "Offshore", "US Business Hours", "$30/hr", "$45/hr", "N/A"),
            ("SEO Specialist", "Offshore", "US Business Hours", "$30/hr", "$45/hr", "N/A"),
            ("Tech Support", "Offshore", "US Business / After-Hours", "$10/hr", "$15/hr", "$30/hr"),
        ],
    )

    d.spacer()
    d.section_header("3.3 Cycle-Based Billing Model")

    d.numbered_item("(a)", "Billing Cycle. Client selects a billing cycle of 3, 6, or 12 months (the \u201cCycle\u201d). The purpose of the Cycle is to provide Client with a structured path to eliminate any unpaid hour balance by the end of each Cycle, thereby avoiding cancellation fees.")

    d.numbered_item("(b)", "Monthly Billed Amount Calculation. The fixed monthly billing rate for each role is calculated as follows:")
    d.numbered_item("1.", "At the start of each new Cycle, Technijian calculates the average monthly hours consumed per role during the previous Cycle, excluding the final month of that Cycle.", indent_level=1)
    d.numbered_item("2.", "This average is then adjusted to account for any unpaid hour balance carried forward from the previous Cycle, so that the new monthly billed amount is set at a level designed to bring the unpaid balance to zero by the end of the current Cycle.", indent_level=1)
    d.numbered_item("3.", "The monthly billed amount for each role equals the adjusted monthly billed hours multiplied by the applicable Contracted Rate from the Rate Card (Schedule C).", indent_level=1)

    d.numbered_item("(c)", "Running Balance. Technijian maintains a running balance for each role:")
    d.numbered_item("1.", "At the start of each month, the running balance is adjusted by adding the actual hours used during the previous month and subtracting the monthly billed hours for that month.", indent_level=1)
    d.numbered_item("2.", "A positive running balance indicates hours consumed in excess of billed amounts (an unpaid hour balance that the Cycle is designed to resolve).", indent_level=1)
    d.numbered_item("3.", "A negative running balance indicates hours billed in excess of consumption. A negative balance does not entitle Client to a credit or refund; it simply means Client has no unpaid hour balance (the unpaid balance is zero). The underlying usage data from months producing a negative balance is used in calculating the average for the next Cycle per Section 3.3(b).", indent_level=1)

    d.numbered_item("(d)", "Cycle Reconciliation. At the end of each Cycle:")
    d.numbered_item("1.", "The running balance for each role is reconciled.", indent_level=1)
    d.numbered_item("2.", "Any net positive balance (hours consumed but not yet billed) is carried forward into the next Cycle. The next Cycle\u2019s monthly billed amount will be recalculated per Section 3.3(b) to absorb this balance and target a zero unpaid balance by the end of the next Cycle.", indent_level=1)
    d.numbered_item("3.", "Any net negative balance (hours billed but not consumed) resets to zero. No credit or refund is issued. The actual usage data from the completed Cycle is used to calculate the monthly billed amount for the next Cycle per Section 3.3(b).", indent_level=1)

    d.numbered_item("(e)", "Cancellation. If Client terminates Virtual Staff services or this Agreement:")
    d.numbered_item("1.", "Any positive running balance (actual hours exceeding billed hours) shall become immediately due and payable. The unpaid hours shall be invoiced at the applicable Hourly Rate from the Rate Card (Schedule C), not the Contracted Rate. This reflects the rate that would have applied had Client engaged Technijian on an ad hoc hourly basis without a Cycle commitment.", indent_level=1)
    d.numbered_item("2.", "Any negative running balance (billed hours exceeding actual hours) does not entitle Client to a credit, refund, or offset against the final invoice. The unpaid balance is simply zero and no further amounts are due from either Party with respect to that role.", indent_level=1)
    d.numbered_item("3.", "The Cycle-Based Billing Model provides Client with Contracted Rates that are lower than the standard Hourly Rates as consideration for Client\u2019s commitment to the Cycle. The application of Hourly Rates upon cancellation reflects the removal of that commitment and the corresponding rate benefit.", indent_level=1)

    d.spacer()
    d.section_header("3.4 Weekly Service Reports")
    d.body("Technijian shall provide Client with weekly service reports that detail:")
    d.numbered_item("(a)", "Each support ticket addressed during the period;")
    d.numbered_item("(b)", "The role, resource name, and hours spent per ticket;")
    d.numbered_item("(c)", "Whether the work was performed during normal or after-hours;")
    d.numbered_item("(d)", "A description of the work performed; and")
    d.numbered_item("(e)", "The current running balance for each role.")

    d.spacer()
    d.section_header("3.5 New Client Onboarding")
    d.body("For new Clients without a previous Cycle history, the initial Cycle billing shall be based on:")
    d.numbered_item("(a)", "A mutually agreed-upon estimated monthly hours per role, documented in the Service Order; and")
    d.numbered_item("(b)", "Actual usage tracking beginning immediately, with the first Cycle reconciliation occurring at the end of the initial Cycle period.")

    d.spacer()
    d.section_header("3.6 Acceptable Use")
    d.body(
        "Client shall direct Virtual Staff only to perform work within the scope of the applicable role "
        "description and the services contemplated by this Agreement. Client shall not direct or request Virtual Staff to:"
    )
    d.numbered_item("(a)", "Perform any activity that violates applicable law, regulation, or third-party rights;")
    d.numbered_item("(b)", "Access, process, or transmit data in violation of data protection laws or Client\u2019s own data handling policies;")
    d.numbered_item("(c)", "Perform work for any entity other than Client without Technijian\u2019s prior written consent; or")
    d.numbered_item("(d)", "Perform work outside the scope of the role for which they are contracted without a corresponding change to the Service Order.")
    d.body(
        "Technijian reserves the right to reassign or remove Virtual Staff if Client directs work that falls "
        "outside these guidelines, without affecting Client\u2019s billing obligations under Section 3.3."
    )

    d.body(
        "Note: Virtual Staff services are not included in the current Service Order. This section is included "
        "for reference should Client elect to add Virtual Staff services in the future.",
        space_after=4,
    )

    d.doc.add_page_break()

    # ── GENERAL TERMS ──
    d.part_header("GENERAL TERMS FOR THIS SCHEDULE")
    d.spacer()

    d.section_header("Changes to Services")
    d.body(
        "Either Party may request changes to the services described in this Schedule by providing thirty (30) "
        "days written notice. Changes to quantities, roles, or service levels shall be documented in an updated "
        "Service Order signed by both Parties."
    )

    d.section_header("Pricing Adjustments")
    d.body(
        "Pricing for services under this Schedule is subject to the Rate Card (Schedule C). Technijian may "
        "adjust rates upon sixty (60) days written notice, effective at the start of the next Renewal Term of the Agreement."
    )

    # Signatures
    d.signatures()
    d.footer_bar()

    d.doc.save(os.path.join(os.path.dirname(__file__), "..", "02_MSA", "Schedule-A-OKL.docx"))
    print("Created Schedule-A-OKL.docx")


# ════════════════════════════════════════════════════════════════
#  SCHEDULE C -- RATE CARD
# ════════════════════════════════════════════════════════════════

def build_schedule_c():
    d = BrandedDoc()

    # Cover page
    d.cover_page("Schedule C", "Rate Card")
    d.doc.add_page_break()

    # Header block
    d.section_header("SCHEDULE C \u2014 RATE CARD")
    d.body("Attached to Master Service Agreement MSA-OKL-2026")
    d.body("Effective Date: March 23, 2026")
    d.spacer()

    # ── SECTION 1: VIRTUAL STAFF RATES ──
    d.part_header("1. VIRTUAL STAFF RATES")
    d.spacer()

    d.section_header("1.1 United States \u2014 Based Staff")
    d.styled_table(
        ["Role", "Hourly Rate", "After-Hours Rate", "Contracted Rate", "Description"],
        [
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr", "Strategic technology leadership and advisory"),
            ("Developer", "$150/hr", "N/A", "$125/hr", "Software development and engineering"),
            ("Tech Support", "$150/hr", "$250/hr", "$125/hr", "Technical support and systems administration"),
        ],
    )

    d.spacer()
    d.section_header("1.2 Offshore \u2014 Based Staff")
    d.styled_table(
        ["Role", "Hourly Rate", "After-Hours Rate", "Contracted Rate", "Description"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr", "Software development and engineering"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr", "Search engine optimization and digital marketing"),
            ("Tech Support", "$15/hr", "$30/hr", "$10/hr", "Technical support and systems administration"),
        ],
    )

    d.spacer()
    d.body_bold_prefix("Normal Business Hours: ", "Monday through Friday, 8:00 AM to 6:00 PM Pacific Time, excluding US federal holidays.")
    d.body_bold_prefix("After-Hours: ", "All hours outside of Normal Business Hours, including weekends and US federal holidays.")
    d.body_bold_prefix("Contracted Rate: ", "The discounted hourly rate applied when Client commits to a Cycle under the Cycle-Based Billing Model described in Schedule A, Part 3. Contracted Rates are available only for Virtual Staff services billed through the Cycle model.")
    d.body_bold_prefix("Hourly Rate: ", "The standard rate for ad hoc (non-contracted) services. This rate is also applied to calculate cancellation fees on any unpaid hour balance upon termination of Virtual Staff services, as described in Schedule A, Section 3.3(e).")

    d.doc.add_page_break()

    # ── SECTION 3: ONLINE SERVICES ──
    d.part_header("3. ONLINE SERVICES \u2014 INFRASTRUCTURE")
    d.spacer()

    d.section_header("3.1 Cloud Infrastructure")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Cloud VM \u2014 vCore", "CL-VC", "Per vCore", "M", "$6.25"),
            ("Cloud VM \u2014 Memory", "CL-GB", "Per GB", "M", "$0.63"),
            ("Cloud VM \u2014 Shared Bandwidth", "CL-SBW", "Per connection", "M", "$15.00"),
            ("Hosting \u2014 CDN", "HC", "Per domain", "M", "$30.00"),
            ("Production Storage", "TB-PSTR", "Per TB", "M", "$200.00"),
            ("Replicated Storage", "TB-RSTR", "Per TB", "M", "$100.00"),
            ("Backup Storage", "TB-BSTR", "Per TB", "M", "$50.00"),
        ],
    )

    d.spacer()
    d.section_header("3.2 Server Management")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Patch Management", "PMW", "Per server", "M", "$4.00"),
            ("Image Backup (Veeam)", "IB", "Per server", "M", "$15.00"),
            ("AV Protection \u2014 Server (CrowdStrike)", "AVS", "Per server", "M", "$10.50"),
            ("AVM Protection \u2014 Server (MalwareBytes)", "AVMS", "Per server", "M", "$16.00"),
            ("AVH Protection \u2014 Server (Huntress)", "AVHS", "Per server", "M", "$6.00"),
            ("My Secure Internet (DNS Filtering)", "SI", "Per server", "M", "$6.00"),
            ("My Remote", "MR", "Per server", "M", "$2.00"),
            ("Health Monitoring (SNMP)", "SHM", "Per device", "M", "$2.00"),
            ("Syslog Monitoring (SNMP)", "SSM", "Per device", "M", "$2.00"),
            ("My Ops \u2014 Net", "OPS-NET", "Per server", "M", "$3.25"),
            ("My Ops \u2014 Config", "OPS-BKP", "Per switch", "M", "$6.00"),
            ("My Ops \u2014 Traffic", "OPS-TR", "Per firewall", "M", "$14.00"),
            ("My Ops \u2014 Port", "OPS-PRT", "Per port", "M", "$0.25"),
            ("My Ops \u2014 Storage", "OPS-ST", "Per disk", "M", "$4.75"),
            ("My Ops \u2014 Wifi", "OPS-WF", "Per AP", "M", "$1.00"),
            ("MyAudit Server", "AMS", "Per server", "M", "$252.00"),
        ],
    )

    d.spacer()
    d.section_header("3.3 Desktop Management")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Patch Management", "PMW", "Per desktop", "M", "$4.00"),
            ("Patch Management (Mac OS)", "PMMAC", "Per desktop", "M", "$11.00"),
            ("AV Protection \u2014 Desktop (CrowdStrike)", "AVD", "Per desktop", "M", "$8.50"),
            ("AVM Protection \u2014 Desktop (MalwareBytes)", "AVMD", "Per desktop", "M", "$5.00"),
            ("AVH Protection \u2014 Desktop (Huntress)", "AVMH", "Per desktop", "M", "$6.00"),
            ("My Secure Internet (DNS Filtering)", "SI", "Per desktop", "M", "$6.00"),
            ("My Remote", "MR", "Per desktop", "M", "$2.00"),
            ("My Ops \u2014 Net", "OPS-NET", "Per desktop", "M", "$3.25"),
            ("Audit Monitoring (Base)", "AM", "Per desktop", "M", "$9.50"),
            ("Audit Monitoring (30 Day)", "AM30", "Per desktop", "M", "$12.50"),
            ("Audit Monitoring (90 Day)", "AM90", "Per desktop", "M", "$15.00"),
            ("Audit Monitoring (6 Months)", "AM6M", "Per desktop", "M", "$25.00"),
            ("Audit Monitoring (1 Year)", "AM1Y", "Per desktop", "M", "$45.00"),
        ],
    )

    d.spacer()
    d.section_header("3.3a Advanced Audit Monitoring")
    d.body("Advanced audit monitoring tiers include User Activity Monitoring (UAM) and Data Loss Prevention (DLP).")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Audit Monitoring UAM (Base + UAM)", "AMUAM", "Per desktop", "M", "$42.00"),
            ("Audit Monitoring UAM (30 Day)", "AMUAM30", "Per desktop", "M", "$44.90"),
            ("Audit Monitoring UAM (90 Day)", "AMUAM90", "Per desktop", "M", "$47.40"),
            ("Audit Monitoring UAM (6 Months)", "AMUAM6M", "Per desktop", "M", "$57.40"),
            ("Audit Monitoring UAM (1 Year)", "AMUAM1Y", "Per desktop", "M", "$77.40"),
            ("Audit Monitoring DLP (UAM + DLP)", "AMDLP", "Per desktop", "M", "$49.00"),
            ("Audit Monitoring DLP (30 Day)", "AMDLP30", "Per desktop", "M", "$51.90"),
            ("Audit Monitoring DLP (90 Day)", "AMDLP90", "Per desktop", "M", "$57.30"),
            ("Audit Monitoring DLP (6 Months)", "AMDLP6M", "Per desktop", "M", "$72.70"),
            ("Audit Monitoring DLP (1 Year)", "AMDLP1Y", "Per desktop", "M", "$108.10"),
        ],
    )

    d.doc.add_page_break()

    d.spacer()
    d.section_header("3.4 Network and Security")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Sophos 1C-4G (Up to 100 Mbps)", "SO-1C4G", "Per appliance", "Y", "$145.00"),
            ("Sophos 2C-4G (Up to 300 Mbps)", "SO-2C4G", "Per appliance", "Y", "$350.00"),
            ("Sophos 4C-6G (Up to 500 Mbps)", "SO-4C6G", "Per appliance", "Y", "$510.00"),
            ("Sophos 6C-8G (Up to 750 Mbps)", "SO-6C8G", "Per appliance", "Y", "$855.00"),
            ("Sophos 8C-16G (Up to 1 Gbps)", "SO-8C16G", "Per appliance", "Y", "$1,330.00"),
            ("Sophos 16C-24G (Up to 1.5 Gbps)", "SO-16C24G", "Per appliance", "Y", "$2,435.00"),
            ("Sophos UC-UG (Over 1.5 Gbps)", "SO-UCG", "Per appliance", "Y", "$3,320.00"),
            ("VeloCloud SD-WAN 30M", "VC-30M", "Per appliance", "Y", "$200.00"),
            ("VeloCloud SD-WAN 50M", "VC-50M", "Per appliance", "Y", "$250.00"),
            ("VeloCloud SD-WAN 100M", "VC-100M", "Per appliance", "Y", "$310.00"),
            ("VeloCloud SD-WAN 200M", "VC-200M", "Per appliance", "Y", "$385.00"),
            ("Edge Appliance 16GB", "Edge-16M", "Per appliance", "Y", "$100.00"),
            ("Real-Time Penetration Testing", "RTPT", "Per IP", "M", "$7.00"),
            ("Vulnerability Assessment", "PT", "Per device", "M", "$1.00"),
            ("Network Assessment", "NA", "Per device", "M", "$1.00"),
        ],
    )

    d.spacer()
    d.section_header("3.5 Email, Compliance, and Identity")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("E-Mail Archiving", "EA", "Per user", "M", "$4.00"),
            ("DMARC/DKIM Management", "DKIM", "Per domain", "M", "$20.00"),
            ("Site Assessment", "SA", "Per domain", "M", "$50.00"),
            ("Veeam 365 Backup", "V365", "Per user", "M", "$2.50"),
            ("Anti-Spam Basic", "ASB", "Per user", "M", "$4.25"),
            ("Anti-Spam Standard (+ Encryption)", "ASA", "Per user", "M", "$6.25"),
            ("Anti-Spam Professional", "ASP", "Per user", "M", "$10.25"),
            ("Anti-Spam GSuite", "ASG", "Per user", "M", "$6.25"),
            ("Phishing Training", "PHT", "Per user", "M", "$6.00"),
            ("SSO / 2FA Login", "SSO-2FA", "Per user", "M", "$12.00"),
            ("SSO / Desktop", "SSO-DSK", "Per user", "M", "$6.00"),
            ("Credential Manager", "CRM", "Per user", "M", "$5.00"),
        ],
    )

    d.doc.add_page_break()

    # ── SECTION 4: SIP TRUNK, HOSTING, LICENSING ──
    d.part_header("4. SIP TRUNK, HOSTING, AND LICENSING SERVICES")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("SIP Trunk", "\u2014", "Per trunk", "M", "Per Service Order"),
            ("Voice Package", "\u2014", "Per package", "M", "Per Service Order"),
            ("DID Number", "\u2014", "Per number", "M", "Per Service Order"),
            ("Hosting \u2014 Web IIS", "HIIS", "Per domain", "M", "$10.00"),
            ("Hosting \u2014 Web WordPress", "HWP", "Per domain", "M", "$10.00"),
            ("Hosting \u2014 Web Magento", "HMG", "Per domain", "M", "$10.00"),
            ("My VDI", "MYVDI", "Per user", "M", "$10.00"),
            ("MDM (iOS)", "MDMIOS", "Per device", "M", "$4.00"),
            ("My Disk (File Sharing & Backup)", "MDU", "Per user", "M", "$16.00"),
            ("Veeam VBR (Backup & Replication)", "VBR", "Per VM", "M", "$13.00"),
            ("Veeam One (Backup Management)", "VONE", "Per VM", "M", "$3.00"),
            ("365 Azure Directory P2", "MS-ADP2", "Per user", "M", "$9.00"),
            ("MSFT \u2014 RDP SAL", "MS-WRDS", "Per license", "M", "$9.20"),
            ("MSFT \u2014 SQL Std SAL", "MS-SSSE", "Per license", "M", "$18.50"),
            ("MSFT \u2014 SQL Std Core", "MS-SSSC", "Per license", "M", "$175.00"),
            ("MSFT \u2014 Sharepoint SAL", "MS-SPSE", "Per license", "M", "$4.00"),
            ("MSFT \u2014 Exchange Std SAL", "MS-ESS", "Per license", "M", "$3.50"),
            ("MSFT \u2014 Exchange Ent SAL", "MS-EESS", "Per license", "M", "$4.00"),
            ("MSFT \u2014 Std Core 2 LIC", "MS-STD", "Per license", "M", "$5.25"),
            ("Foxit PDF Business", "FX-BSPDF", "Per license", "Y", "$14.00"),
        ],
    )

    d.doc.add_page_break()

    # ── SECTION 5: TERMS ──
    d.part_header("5. TERMS")
    d.spacer()

    d.section_header("5.01 Rate Adjustments")
    d.body(
        "Technijian may adjust the rates in this Schedule upon sixty (60) days written notice to Client. "
        "Adjusted rates shall take effect at the start of the next Renewal Term of the Agreement."
    )

    d.section_header("5.02 Volume Discounts")
    d.body(
        "Volume-based pricing may be negotiated and documented in the applicable Service Order or Subscription Order."
    )

    d.section_header("5.03 Minimum Billing")
    d.body(
        "Contracted Virtual Staff hours are billed per the Cycle-Based Billing Model described in Schedule A, "
        "Part 3, at the Contracted Rate. Ad hoc (non-contracted) support is billed at the Hourly Rate in "
        "15-minute increments. Upon cancellation or termination, any unpaid hour balance shall be invoiced at "
        "the applicable Hourly Rate as set forth in Schedule A, Section 3.3(e)."
    )

    d.section_header("5.04 Relationship to Price List")
    d.body(
        "Technijian maintains a separate Services Price List that provides the current catalog of available "
        "services and list prices. The Price List may be updated by Technijian from time to time and shall be "
        "made available to Client upon request. In the event of a conflict between the Price List and this Rate "
        "Card, the rates documented in this Rate Card (Schedule C) shall govern for the duration of the "
        "then-current term. Any rate change shall be subject to the sixty (60) day written notice and Renewal "
        "Term effective date requirements set forth in Section 5.01 above and MSA Section 9.02."
    )

    # Signatures
    d.signatures()
    d.footer_bar()

    d.doc.save(os.path.join(os.path.dirname(__file__), "..", "02_MSA", "Schedule-C-OKL.docx"))
    print("Created Schedule-C-OKL.docx")


if __name__ == "__main__":
    build_schedule_a()
    build_schedule_c()
