"""
OKL — Server Replacement Options (CLIENT-FACING)
Branded proposal with 3 tiers. Shows ONLY client investment.
No cost-plus breakdown, no markup labels, no outlet references.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")

# Markup is applied silently; NOT shown in the client doc
MARKUP_PCT = 0.20


def shade(cell, hex_color):
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    )


def fix_table_layout(tbl, col_widths_inches):
    """
    Force Word to honor exact column widths by:
      1. Setting tblLayout to 'fixed' (not autofit).
      2. Emitting a tblGrid with explicit gridCol widths in dxa (twentieths of a point).
      3. Setting tblW to the total width in dxa.
      4. Writing tcW on every cell (belt + suspenders).
    Without this, Word silently stretches tables past the page margin.
    """
    dxas = [int(w * 1440) for w in col_widths_inches]  # 1 inch = 1440 dxa
    total_dxa = sum(dxas)
    tbl_pr = tbl._element.tblPr

    # Fixed layout
    existing = tbl_pr.find(qn("w:tblLayout"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(parse_xml(
        f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'
    ))

    # Table width
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_pr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>'
    ))

    # Grid
    existing_grid = tbl._element.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl._element.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>' + \
        "".join(f'<w:gridCol w:w="{d}"/>' for d in dxas) + \
        "</w:tblGrid>"
    # tblGrid must come right after tblPr
    tbl._element.insert(list(tbl._element).index(tbl_pr) + 1,
                        parse_xml(grid_xml))

    # Per-cell widths
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._element.get_or_add_tcPr()
            existing_tcw = tc_pr.find(qn("w:tcW"))
            if existing_tcw is not None:
                tc_pr.remove(existing_tcw)
            tc_pr.append(parse_xml(
                f'<w:tcW {nsdecls("w")} w:w="{dxas[i]}" w:type="dxa"/>'
            ))


def keep_rows_together(tbl):
    """Prevent rows from breaking across pages."""
    for row in tbl.rows:
        tr_pr = row._element.get_or_add_trPr()
        tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def keep_with_next_paragraph(doc):
    """Mark the last paragraph as keep-with-next so a heading stays with its table."""
    p = doc.paragraphs[-1]
    p_pr = p._element.get_or_add_pPr()
    p_pr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))


def light_borders(tbl):
    tbl_pr = tbl._element.tblPr
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    ))


def no_borders(tbl):
    tbl_pr = tbl._element.tblPr
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    ))


def set_font(run, size=10, bold=False, color=None, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>'
        )
        rPr.append(rFonts)


def add_heading(doc, text, size=16, color=CORE_BLUE, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=color)
    return p


def add_body(doc, text, size=10, bold=False, color=DARK_CHARCOAL, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_hero(doc, title, subtitle):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    shade(cell, HEX_DARK)
    cell.paragraphs[0].paragraph_format.space_before = Pt(20)
    cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    r = cell.paragraphs[0].add_run(title)
    set_font(r, size=24, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(20)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11, bold=False, color=RGBColor(0xCC, 0xCC, 0xCC))
    no_borders(tbl)
    fix_table_layout(tbl, [7.1])
    keep_rows_together(tbl)
    doc.add_paragraph()


def add_kv_table(doc, rows, col_widths=(2.2, 4.9)):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    light_borders(tbl)
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[i]
        shade(row.cells[0], HEX_OFF_WHITE)
        for j, txt in enumerate((k, v)):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            r = c.paragraphs[0].add_run(str(txt))
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)
    fix_table_layout(tbl, list(col_widths))
    keep_rows_together(tbl)
    doc.add_paragraph()
    return tbl


def add_option_card(doc, tier_tag, model, spec_rows, cost_basis, recommended=False):
    """
    Client-facing option card rendered as a SINGLE table so Word keeps the
    entire card on one page. Only the client investment is shown.
    """
    client_price = round(cost_basis * (1 + MARKUP_PCT), 2)

    # Single table: 1 header + N spec + 1 footer, all 2 cols
    total_rows = 1 + len(spec_rows) + 1
    tbl = doc.add_table(rows=total_rows, cols=2)
    tbl.autofit = False
    light_borders(tbl)

    # Column widths: label col narrower, detail col wider
    LEFT_W = 2.1
    RIGHT_W = 5.0
    fix_table_layout(tbl, [LEFT_W, RIGHT_W])

    # --- Header row (blue + orange) ---
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    shade(left, HEX_BLUE)
    shade(right, HEX_ORANGE)

    left_title = f"{tier_tag}  —  {model}"
    if recommended:
        left_title = f"{tier_tag}  ★ RECOMMENDED  —  {model}"

    left.paragraphs[0].paragraph_format.space_before = Pt(6)
    left.paragraphs[0].paragraph_format.space_after = Pt(6)
    r1 = left.paragraphs[0].add_run(left_title)
    set_font(r1, size=12, bold=True, color=WHITE)

    right.paragraphs[0].paragraph_format.space_before = Pt(6)
    right.paragraphs[0].paragraph_format.space_after = Pt(6)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = right.paragraphs[0].add_run(f"${client_price:,.2f}")
    set_font(r2, size=14, bold=True, color=WHITE)

    # --- Spec rows ---
    for i, (k, v) in enumerate(spec_rows):
        row = tbl.rows[1 + i]
        shade(row.cells[0], HEX_OFF_WHITE)
        for j, txt in enumerate((k, v)):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            r = c.paragraphs[0].add_run(str(txt))
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)

    # --- Investment footer row (dark charcoal across both cells) ---
    foot_left = tbl.rows[-1].cells[0]
    foot_right = tbl.rows[-1].cells[1]
    shade(foot_left, HEX_DARK)
    shade(foot_right, HEX_DARK)

    foot_left.paragraphs[0].paragraph_format.space_before = Pt(4)
    foot_left.paragraphs[0].paragraph_format.space_after = Pt(4)
    rfl = foot_left.paragraphs[0].add_run("Your Investment")
    set_font(rfl, size=11, bold=True, color=WHITE)

    foot_right.paragraphs[0].paragraph_format.space_before = Pt(4)
    foot_right.paragraphs[0].paragraph_format.space_after = Pt(4)
    foot_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rfr = foot_right.paragraphs[0].add_run(f"${client_price:,.2f}")
    set_font(rfr, size=12, bold=True, color=WHITE)

    # Keep the entire card on one page
    keep_rows_together(tbl)
    doc.add_paragraph()
    return client_price


def add_footer_strip(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.cell(0, 0)
    shade(c, HEX_OFF_WHITE)
    c.paragraphs[0].paragraph_format.space_before = Pt(8)
    c.paragraphs[0].paragraph_format.space_after = Pt(8)
    r = c.paragraphs[0].add_run(text)
    set_font(r, size=9, bold=False, color=BRAND_GREY)
    fix_table_layout(tbl, [7.1])
    keep_rows_together(tbl)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(LOGO_PATH, width=Inches(1.8))

    add_hero(
        doc,
        "Server Replacement Proposal",
        "Prepared for Oaktree Law  |  April 15, 2026",
    )

    # Opening letter
    add_heading(doc, "Dear Ed,", size=13, color=DARK_CHARCOAL)
    add_body(doc,
        "Thank you for the opportunity to present replacement options for Oaktree Law's "
        "on-premises server. The current system has served the firm well, but after nearly "
        "a decade of service it is out of manufacturer warranty, low on memory, and running "
        "near storage capacity. Replacing it now — on your timeline, not under emergency "
        "conditions — protects the firm's productivity and data."
    )
    add_body(doc,
        "Below are three Dell PowerEdge rack server options, each sized to give Oaktree Law "
        "meaningful headroom over the current system. Every option is backed by Dell factory "
        "warranty and Technijian's white-glove configuration, deployment, and support."
    )

    # Current environment
    add_heading(doc, "Your Current Server", size=14)
    add_kv_table(doc, [
        ("Platform", "Dell PowerEdge 1U Rack Server (circa 2016)"),
        ("Processor", "Intel Xeon E3-1220 v5  (4 cores, 3.0 GHz)"),
        ("Memory", "14 GB"),
        ("Storage", "119 GB system drive + 779 GB data drive (~900 GB total)"),
        ("Warranty Status", "Out of Dell manufacturer support — end-of-life platform"),
        ("Data Drive Utilization", "94% full (approximately 60 GB free)"),
    ])

    # Options
    add_heading(
        doc, "Replacement Options", size=18, color=CORE_ORANGE, space_before=14
    )

    add_option_card(doc,
        tier_tag="OPTION A  —  ESSENTIAL",
        model="Dell PowerEdge R350",
        spec_rows=[
            ("Processor", "Intel Xeon E-2334  (4 cores / 8 threads, 3.4 GHz)"),
            ("Memory", "32 GB ECC — 2.3x your current capacity"),
            ("Storage", "2 x 2 TB enterprise drives in RAID 1 (mirrored)"),
            ("Usable Storage", "~2 TB of protected storage — 2.2x your current capacity"),
            ("Power", "Single enterprise-grade power supply"),
            ("Remote Management", "Dell iDRAC9 Basic"),
            ("Form Factor", "1U Rack — fits your existing rack footprint"),
            ("Warranty", "1 Year Dell ProSupport Next-Business-Day included"),
        ],
        cost_basis=1850.00,
    )

    add_option_card(doc,
        tier_tag="OPTION B  —  PROFESSIONAL",
        model="Dell PowerEdge R350",
        spec_rows=[
            ("Processor", "Intel Xeon E-2378  (8 cores / 16 threads, 2.6 GHz)"),
            ("Memory", "64 GB ECC — 4.5x your current capacity"),
            ("Storage", "4 x 2 TB enterprise drives in RAID 10"),
            ("Usable Storage", "~4 TB of protected storage, high performance"),
            ("Power", "Dual hot-plug redundant power supplies"),
            ("Remote Management", "Dell iDRAC9 Enterprise"),
            ("Form Factor", "1U Rack with hot-plug drive bays"),
            ("Warranty", "1 Year Dell ProSupport Next-Business-Day included"),
        ],
        cost_basis=3200.00,
        recommended=True,
    )

    add_option_card(doc,
        tier_tag="OPTION C  —  PREMIUM",
        model="Dell PowerEdge R360  (Current Generation, 2024)",
        spec_rows=[
            ("Processor", "Intel Xeon E-2488  (8 cores / 16 threads, 3.2 GHz)"),
            ("Memory", "64 GB DDR5 ECC — fastest available memory"),
            ("Storage", "4 x 2 TB enterprise drives in RAID 10"),
            ("Usable Storage", "~4 TB of protected storage, highest throughput"),
            ("Power", "Dual hot-plug Titanium-rated redundant power supplies"),
            ("Remote Management", "Dell iDRAC9 Enterprise"),
            ("Form Factor", "1U Rack with hot-plug drive bays"),
            ("Warranty", "1 Year Dell ProSupport Next-Business-Day included"),
        ],
        cost_basis=4250.00,
    )

    # Comparison
    add_heading(doc, "At-a-Glance Comparison", size=14)
    comp = doc.add_table(rows=6, cols=5)
    comp.autofit = False
    light_borders(comp)
    headers = ["Feature", "Current", "Option A", "Option B", "Option C"]
    data = [
        ("Processor cores",           "4",        "4",        "8",        "8"),
        ("Memory",                    "14 GB",    "32 GB",    "64 GB",    "64 GB DDR5"),
        ("Usable storage",            "~900 GB",  "~2 TB",    "~4 TB",    "~4 TB"),
        ("Redundant power",           "No",       "No",       "Yes",      "Yes"),
        ("Under manufacturer warranty", "No",     "Yes",      "Yes",      "Yes"),
    ]
    widths = [1.9, 1.1, 1.3, 1.3, 1.5]
    for j, h in enumerate(headers):
        c = comp.rows[0].cells[j]
        shade(c, HEX_BLUE)
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True, color=WHITE)
    for i, row_data in enumerate(data, start=1):
        for j, txt in enumerate(row_data):
            c = comp.rows[i].cells[j]
            if j == 0:
                shade(c, HEX_OFF_WHITE)
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            r = c.paragraphs[0].add_run(txt)
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)
    fix_table_layout(comp, widths)
    keep_rows_together(comp)
    doc.add_paragraph()

    # Recommendation
    add_heading(doc, "Our Recommendation", size=14, color=CORE_ORANGE)
    add_body(doc,
        "For Oaktree Law's workload profile, we recommend Option B (PowerEdge R350 Professional). "
        "It delivers double the processor cores, four-and-a-half times the memory, four times the "
        "protected storage, and — importantly — dual redundant power supplies so that a single "
        "power-supply failure never takes the firm offline. It is the right balance of performance, "
        "reliability, and investment."
    )

    # What's included
    add_heading(doc, "What's Included With Every Option", size=14)
    bullets = [
        "Dell factory warranty with 1-year Next-Business-Day ProSupport coverage",
        "Technijian procurement, pre-delivery inspection, and burn-in testing",
        "Full BIOS, firmware, and iDRAC configuration to Technijian standards",
        "Asset tagging, documentation, and addition to your managed-services inventory",
        "Ground freight and delivery to Oaktree Law's server room",
        "Coordination of warranty claims and advance-replacement parts for the full warranty term",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(b)
        set_font(r, size=10, color=DARK_CHARCOAL)

    # Not included
    add_heading(doc, "Services Quoted Separately", size=14)
    bullets2 = [
        "On-site installation, rack-and-stack, cabling, and physical cutover",
        "Data migration from the current server to the new hardware",
        "Operating system and hypervisor licensing (Windows Server, VMware, or equivalent)",
        "Any application re-installation or re-licensing required by the migration",
    ]
    for b in bullets2:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(b)
        set_font(r, size=10, color=DARK_CHARCOAL)

    # Terms
    add_heading(doc, "Proposal Terms", size=14)
    add_kv_table(doc, [
        ("Delivery Lead Time", "5 to 10 business days from purchase order"),
        ("Freight", "Standard ground freight included; expedite available on request"),
        ("Payment Terms", "50% at purchase order, 50% on delivery"),
        ("Proposal Validity", "14 days from the date of this proposal"),
        ("Pricing Basis", "All pricing in US Dollars; taxes (if applicable) additional"),
    ])

    # Next steps
    add_heading(doc, "Next Steps", size=14, color=CORE_ORANGE)
    add_body(doc,
        "To proceed, simply reply to this proposal with your selected option and Technijian "
        "will issue a purchase order, place the hardware on reserve, and schedule the "
        "installation window with your team. If you would like to discuss the options or "
        "customize a configuration, please use the booking link in Ravi Jain's email "
        "signature to find a time that works for you."
    )

    # Signature block
    doc.add_paragraph()
    add_body(doc, "Respectfully submitted,", size=10, color=DARK_CHARCOAL)
    add_body(doc, "Ravi Jain", size=11, bold=True, color=CORE_BLUE)
    add_body(doc, "Chief Executive Officer", size=10, color=BRAND_GREY)
    add_body(doc, "Technijian, Inc.  |  rjain@technijian.com", size=10, color=BRAND_GREY)

    # Footer
    doc.add_paragraph()
    add_footer_strip(doc,
        "Technijian, Inc.  |  18 Technology Drive, Suite 141, Irvine, CA 92618  |  technijian.com  |  "
        "This proposal is confidential and prepared exclusively for Oaktree Law."
    )

    out = os.path.join(
        os.path.dirname(__file__), "..", "04_Quotes", "OKL-Server-Replacement-Proposal.docx"
    )
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
