"""
OKL — 4 Branded Technijian Quotes from Dell Quote #3000201473092.1
One quote per server + one for additional items.
Client-facing: shows only "Your Investment" — no cost basis or markup.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")
BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "04_Quotes")

MARKUP_PCT = 0.20


def shade(cell, hex_color):
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    )


def fix_table_layout(tbl, col_widths_inches):
    dxas = [int(w * 1440) for w in col_widths_inches]
    total_dxa = sum(dxas)
    tbl_pr = tbl._element.tblPr
    existing = tbl_pr.find(qn("w:tblLayout"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_pr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>'))
    existing_grid = tbl._element.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl._element.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>' + \
        "".join(f'<w:gridCol w:w="{d}"/>' for d in dxas) + "</w:tblGrid>"
    tbl._element.insert(list(tbl._element).index(tbl_pr) + 1, parse_xml(grid_xml))
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._element.get_or_add_tcPr()
            existing_tcw = tc_pr.find(qn("w:tcW"))
            if existing_tcw is not None:
                tc_pr.remove(existing_tcw)
            tc_pr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxas[i]}" w:type="dxa"/>'))


def keep_rows_together(tbl):
    for row in tbl.rows:
        tr_pr = row._element.get_or_add_trPr()
        tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def light_borders(tbl):
    tbl_pr = tbl._element.tblPr
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    ))


def no_borders(tbl):
    tbl_pr = tbl._element.tblPr
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    ))


def set_font(run, size=10, bold=False, color=None, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>'
        )
        rPr.append(rFonts)


def add_heading(doc, text, size=16, color=CORE_BLUE, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=color)
    return p


def add_body(doc, text, size=10, bold=False, color=DARK_CHARCOAL, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_hero(doc, title, subtitle):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    shade(cell, HEX_DARK)
    cell.paragraphs[0].paragraph_format.space_before = Pt(20)
    cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    r = cell.paragraphs[0].add_run(title)
    set_font(r, size=24, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(20)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11, bold=False, color=RGBColor(0xCC, 0xCC, 0xCC))
    no_borders(tbl)
    fix_table_layout(tbl, [7.1])
    keep_rows_together(tbl)
    doc.add_paragraph()


def add_kv_table(doc, rows, col_widths=(2.2, 4.9)):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    light_borders(tbl)
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[i]
        shade(row.cells[0], HEX_OFF_WHITE)
        for j, txt in enumerate((k, v)):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            r = c.paragraphs[0].add_run(str(txt))
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)
    fix_table_layout(tbl, list(col_widths))
    keep_rows_together(tbl)
    doc.add_paragraph()
    return tbl


def add_line_item_table(doc, items, total_label, total_amount):
    """Table with Item | Description | Investment columns, plus a total row."""
    num_rows = 1 + len(items) + 1  # header + items + total
    tbl = doc.add_table(rows=num_rows, cols=3)
    tbl.autofit = False
    light_borders(tbl)
    widths = [2.0, 3.5, 1.6]
    fix_table_layout(tbl, widths)

    # Header
    headers = ["Item", "Description", "Investment"]
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        shade(c, HEX_BLUE)
        c.paragraphs[0].paragraph_format.space_before = Pt(4)
        c.paragraphs[0].paragraph_format.space_after = Pt(4)
        if j == 2:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True, color=WHITE)

    # Item rows
    for i, (item, desc, price) in enumerate(items):
        row = tbl.rows[1 + i]
        shade(row.cells[0], HEX_OFF_WHITE)
        for j, txt in enumerate((item, desc, f"${price:,.2f}")):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            if j == 2:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = c.paragraphs[0].add_run(str(txt))
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)

    # Total row
    total_row = tbl.rows[-1]
    shade(total_row.cells[0], HEX_DARK)
    shade(total_row.cells[1], HEX_DARK)
    shade(total_row.cells[2], HEX_DARK)
    total_row.cells[0].paragraphs[0].paragraph_format.space_before = Pt(4)
    total_row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(4)
    r_label = total_row.cells[0].paragraphs[0].add_run(total_label)
    set_font(r_label, size=11, bold=True, color=WHITE)
    # Leave cell 1 empty
    total_row.cells[1].paragraphs[0].paragraph_format.space_before = Pt(4)
    total_row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(4)
    # Total amount in cell 2
    total_row.cells[2].paragraphs[0].paragraph_format.space_before = Pt(4)
    total_row.cells[2].paragraphs[0].paragraph_format.space_after = Pt(4)
    total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_amt = total_row.cells[2].paragraphs[0].add_run(f"${total_amount:,.2f}")
    set_font(r_amt, size=12, bold=True, color=WHITE)

    keep_rows_together(tbl)
    doc.add_paragraph()
    return tbl


def add_footer_strip(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.cell(0, 0)
    shade(c, HEX_OFF_WHITE)
    c.paragraphs[0].paragraph_format.space_before = Pt(8)
    c.paragraphs[0].paragraph_format.space_after = Pt(8)
    r = c.paragraphs[0].add_run(text)
    set_font(r, size=9, bold=False, color=BRAND_GREY)
    fix_table_layout(tbl, [7.1])
    keep_rows_together(tbl)


def markup(cost):
    return round(cost * (1 + MARKUP_PCT), 2)


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    return doc


def add_quote_details(doc, quote_number):
    add_kv_table(doc, [
        ("Quote Number", quote_number),
        ("Prepared For", "Oaktree Law"),
        ("Attention", "Ed Pits"),
        ("Prepared By", "Ravi Jain, CEO — Technijian, Inc."),
        ("Date", "April 16, 2026"),
        ("Valid Through", "April 30, 2026"),
        ("Dell Reference", "Dell Quote #3000201473092.1"),
    ])


def add_terms(doc):
    add_heading(doc, "Terms & Conditions", size=14)
    add_kv_table(doc, [
        ("Equipment Condition", "Dell Outlet — Certified Refurbished with Dell factory warranty"),
        ("Warranty", "5-Year Dell ProSupport with Next-Business-Day On-Site Service"),
        ("Delivery", "Estimated delivery within 7-10 business days of purchase order"),
        ("Freight", "Standard ground shipping included"),
        ("Payment Terms", "50% at purchase order, 50% on delivery"),
        ("Validity", "This quote is valid for 14 days from the date above"),
    ])


def add_whats_included(doc):
    add_heading(doc, "Included With Your Equipment", size=14)
    bullets = [
        "Technijian procurement and order management",
        "Pre-delivery configuration validation and burn-in testing",
        "BIOS, firmware, and iDRAC configuration to Technijian standards",
        "Asset tagging, documentation, and inventory registration",
        "Standard ground freight and delivery coordination",
        "Warranty claims coordination for the full 5-year ProSupport term",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(b)
        set_font(r, size=10, color=DARK_CHARCOAL)


def add_not_included(doc):
    add_heading(doc, "Services Quoted Separately", size=14)
    bullets = [
        "On-site installation, rack-and-stack, cabling, and physical cutover",
        "Operating system and hypervisor licensing",
        "Data migration from existing servers to new hardware",
        "Application re-installation or re-licensing",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(b)
        set_font(r, size=10, color=DARK_CHARCOAL)


def add_signature(doc):
    doc.add_paragraph()
    add_body(doc, "To proceed, reply to this quote with your approval and Technijian will "
             "place the order with Dell immediately. If you have questions or would like "
             "to discuss, please use the booking link in my email signature.")
    doc.add_paragraph()
    add_body(doc, "Respectfully submitted,", size=10, color=DARK_CHARCOAL)
    add_body(doc, "Ravi Jain", size=11, bold=True, color=CORE_BLUE)
    add_body(doc, "Chief Executive Officer", size=10, color=BRAND_GREY)
    add_body(doc, "Technijian, Inc.  |  rjain@technijian.com", size=10, color=BRAND_GREY)


def add_confidentiality_footer(doc):
    doc.add_paragraph()
    add_footer_strip(doc,
        "Technijian, Inc.  |  18 Technology Drive, Suite 141, Irvine, CA 92618  |  technijian.com  |  "
        "This quote is confidential and prepared exclusively for Oaktree Law."
    )


# ─────────────────────────────────────────────────────────────────
# QUOTE 1: PowerEdge R650
# ─────────────────────────────────────────────────────────────────
def build_quote_r650():
    doc = new_doc()
    server_cost = 18245.44
    client_price = markup(server_cost)

    add_hero(doc,
        "Equipment Quote — Dell PowerEdge R650",
        "Prepared for Oaktree Law  |  April 16, 2026",
    )

    add_quote_details(doc, "TQ-OKL-2026-R650")

    add_heading(doc, "Equipment Summary", size=14)
    add_body(doc,
        "The Dell PowerEdge R650 is a high-performance 1U rack server built for demanding "
        "workloads. This configuration features dual Intel Xeon Gold processors with 32 cores "
        "total, 512 GB of ECC memory, enterprise SAS storage with hardware RAID, and extensive "
        "10/25GbE networking — all backed by 5-year Dell ProSupport."
    )

    add_heading(doc, "Configuration Details", size=14)
    add_kv_table(doc, [
        ("Server", "Dell PowerEdge R650 (Gen15, Certified Refurbished)"),
        ("Processors", "2x Intel Xeon Gold 6346 (16 cores / 32 threads each, up to 3.60 GHz, 36MB Cache)"),
        ("Total Cores", "32 cores / 64 threads"),
        ("Memory", "512 GB DDR4 ECC RDIMM (16x 32GB, 3200 MT/s)"),
        ("Storage", "2x 1.2 TB 10K RPM SAS 12Gbps 2.5-inch Hot-Plug Drives"),
        ("RAID Controller", "PERC H755 SAS (Hardware RAID 1)"),
        ("Networking — OCP", "Broadcom 57504 Quad Port 10/25GbE SFP28 OCP NIC 3.0"),
        ("Networking — PCIe", "2x Broadcom 57414 Dual Port 10/25GbE SFP28 Adapter"),
        ("Onboard LOM", "Broadcom 5720 Dual Port 1GbE"),
        ("Power", "2x 1100W Dual Hot-Plug Redundant Titanium PSU (1+1)"),
        ("Remote Management", "iDRAC9 Basic (Enterprise upgrade available separately)"),
        ("Security", "TPM 2.0"),
        ("Form Factor", "1U Rack — 2.5-inch chassis, up to 8 hot-plug drive bays, 3 PCIe slots"),
        ("Rails", "1U Rack Rails with Cable Management Arm"),
        ("Warranty", "5-Year Dell ProSupport Next-Business-Day On-Site Service"),
    ])

    add_heading(doc, "Your Investment", size=16, color=CORE_ORANGE)
    add_line_item_table(doc,
        items=[
            ("PowerEdge R650", "Dual Xeon Gold 6346, 512GB RAM, 2x1.2TB SAS RAID 1, 5yr ProSupport", client_price),
        ],
        total_label="Total Investment",
        total_amount=client_price,
    )

    add_whats_included(doc)
    add_not_included(doc)
    add_terms(doc)
    add_signature(doc)
    add_confidentiality_footer(doc)

    out = os.path.join(BASE_DIR, "OKL-Quote-R650.docx")
    doc.save(out)
    print(f"Saved: {out}  |  Client price: ${client_price:,.2f}")
    return client_price


# ─────────────────────────────────────────────────────────────────
# QUOTE 2: PowerEdge XR7620
# ─────────────────────────────────────────────────────────────────
def build_quote_xr7620():
    doc = new_doc()
    server_cost = 23539.12
    client_price = markup(server_cost)

    add_hero(doc,
        "Equipment Quote — Dell PowerEdge XR7620",
        "Prepared for Oaktree Law  |  April 16, 2026",
    )

    add_quote_details(doc, "TQ-OKL-2026-XR7620")

    add_heading(doc, "Equipment Summary", size=14)
    add_body(doc,
        "The Dell PowerEdge XR7620 is a rugged, GPU-enabled server designed for AI inference, "
        "edge computing, and high-performance workloads. This configuration includes dual Intel "
        "Xeon Gold processors, 256 GB of DDR5 memory, an NVIDIA L4 GPU with 24 GB of dedicated "
        "video memory, NVMe flash storage, and 5-year Dell ProSupport."
    )

    add_heading(doc, "Configuration Details", size=14)
    add_kv_table(doc, [
        ("Server", "Dell PowerEdge XR7620 (Gen16, Certified Refurbished)"),
        ("Processors", "2x Intel Xeon Gold 6426Y (16 cores / 32 threads each, up to 4.10 GHz, 37.5MB Cache)"),
        ("Total Cores", "32 cores / 64 threads"),
        ("Memory", "256 GB DDR5 ECC RDIMM (16x 16GB, 5600 MT/s)"),
        ("GPU", "NVIDIA L4 — 24 GB GDDR6, 72W TDP, PCIe Gen4 (AI inference / video encoding)"),
        ("Boot Storage", "2x 480 GB M.2 PCIe NVMe Gen4 SSD (BOSS-N1 RAID 1)"),
        ("Data Storage", "2x 3.84 TB PCIe U.2 NVMe Gen4 SSD (direct-attached)"),
        ("Total Raw Storage", "8.64 TB NVMe flash"),
        ("Networking — OCP", "Intel X710-T4L Quad Port 10GbE BASE-T OCP NIC 3.0"),
        ("Power", "2x 1800W Dual Hot-Plug Redundant Titanium PSU (1+1), 200-240V"),
        ("Remote Management", "iDRAC9 Basic (Enterprise upgrade available separately)"),
        ("Security", "TPM 2.0"),
        ("Form Factor", "Rugged short-depth chassis, rear-port access, up to 4 NVMe direct, 5 PCIe slots"),
        ("Rails", "1U Rack Rails with Cable Management Arm"),
        ("Warranty", "5-Year Dell ProSupport Next-Business-Day On-Site Service"),
    ])

    add_heading(doc, "Your Investment", size=16, color=CORE_ORANGE)
    add_line_item_table(doc,
        items=[
            ("PowerEdge XR7620", "Dual Xeon Gold 6426Y, 256GB DDR5, NVIDIA L4 GPU, 8.64TB NVMe, 5yr ProSupport", client_price),
        ],
        total_label="Total Investment",
        total_amount=client_price,
    )

    add_whats_included(doc)
    add_not_included(doc)
    add_terms(doc)
    add_signature(doc)
    add_confidentiality_footer(doc)

    out = os.path.join(BASE_DIR, "OKL-Quote-XR7620.docx")
    doc.save(out)
    print(f"Saved: {out}  |  Client price: ${client_price:,.2f}")
    return client_price


# ─────────────────────────────────────────────────────────────────
# QUOTE 3: PowerEdge R760
# ─────────────────────────────────────────────────────────────────
def build_quote_r760():
    doc = new_doc()
    server_cost = 15445.88
    client_price = markup(server_cost)

    add_hero(doc,
        "Equipment Quote — Dell PowerEdge R760",
        "Prepared for Oaktree Law  |  April 16, 2026",
    )

    add_quote_details(doc, "TQ-OKL-2026-R760")

    add_heading(doc, "Equipment Summary", size=14)
    add_body(doc,
        "The Dell PowerEdge R760 is a versatile 2U rack server with flexible storage options, "
        "ideal for virtualization, database, and mixed workloads. This configuration features "
        "dual Intel Xeon Silver processors, 256 GB of DDR5 memory, a mix of SAS and SSD storage "
        "across multiple drive bays, and 5-year Dell ProSupport."
    )

    add_heading(doc, "Configuration Details", size=14)
    add_kv_table(doc, [
        ("Server", "Dell PowerEdge R760 (Gen16, Certified Refurbished)"),
        ("Processors", "2x Intel Xeon Silver 4509Y (8 cores / 16 threads each, up to 4.10 GHz, 22.5MB Cache)"),
        ("Total Cores", "16 cores / 32 threads"),
        ("Memory", "256 GB DDR5 ECC RDIMM (4x 64GB, 5600 MT/s)"),
        ("Boot Storage", "2x 480 GB M.2 PCIe NVMe Gen4 SSD (BOSS-N1 RAID 1)"),
        ("SAS Storage", "5x 1.2 TB 10K RPM SAS 12Gbps 2.5-inch Hot-Plug Drives"),
        ("SSD Storage", "1x 960 GB SATA SSD 6Gbps 2.5-inch Hot-Plug (Read Intensive)"),
        ("Total Raw Storage", "7.92 TB (boot) + 6.96 TB (data)"),
        ("HBA Controllers", "Dell HBA355i (front) + Dell HBA355e (external)"),
        ("Networking — OCP", "Broadcom 5720 Quad Port 1GbE BASE-T OCP NIC 3.0"),
        ("Networking — PCIe", "2x Broadcom 57416 Dual Port 10GbE BASE-T Adapter"),
        ("Power", "2x 1400W Dual Hot-Plug Redundant Titanium PSU (1+1)"),
        ("Remote Management", "iDRAC9 Basic (Enterprise upgrade available separately)"),
        ("Security", "TPM 2.0"),
        ("Form Factor", "2U Rack — up to 24 x 2.5-inch hot-plug drive bays (SAS/SATA/NVMe), 8 PCIe slots"),
        ("Rails", "2U ReadyRails Sliding Rails with Cable Management Arm"),
        ("Warranty", "5-Year Dell ProSupport Next-Business-Day On-Site Service"),
    ])

    add_heading(doc, "Your Investment", size=16, color=CORE_ORANGE)
    add_line_item_table(doc,
        items=[
            ("PowerEdge R760", "Dual Xeon Silver 4509Y, 256GB DDR5, 5x1.2TB SAS + 960GB SSD, 5yr ProSupport", client_price),
        ],
        total_label="Total Investment",
        total_amount=client_price,
    )

    add_whats_included(doc)
    add_not_included(doc)
    add_terms(doc)
    add_signature(doc)
    add_confidentiality_footer(doc)

    out = os.path.join(BASE_DIR, "OKL-Quote-R760.docx")
    doc.save(out)
    print(f"Saved: {out}  |  Client price: ${client_price:,.2f}")
    return client_price


# ─────────────────────────────────────────────────────────────────
# QUOTE 4: Additional Items (iDRAC Kits + Shipping)
# ─────────────────────────────────────────────────────────────────
def build_quote_additional():
    doc = new_doc()

    idrac_15g_cost = 199.00
    idrac_16g_cost = 199.00  # x2
    shipping_cost = 118.47

    idrac_15g_client = markup(idrac_15g_cost)
    idrac_16g_client = markup(idrac_16g_cost)
    shipping_client = markup(shipping_cost)

    total_client = idrac_15g_client + (2 * idrac_16g_client) + shipping_client

    add_hero(doc,
        "Equipment Quote — Additional Items",
        "Prepared for Oaktree Law  |  April 16, 2026",
    )

    add_quote_details(doc, "TQ-OKL-2026-ADDL")

    add_heading(doc, "Additional Items", size=14)
    add_body(doc,
        "The following items complement the server equipment quoted separately. The iDRAC9 "
        "Enterprise upgrade kits enable full remote management capabilities including virtual "
        "console, virtual media, and advanced monitoring — essential for remote server "
        "administration and reducing on-site support visits."
    )

    add_heading(doc, "Item Details", size=14)
    add_kv_table(doc, [
        ("iDRAC9 Enterprise 15G Kit", "Customer-installable upgrade for PowerEdge R650; "
         "enables virtual console, virtual media, remote firmware updates, "
         "and advanced server monitoring via web interface"),
        ("iDRAC9 Enterprise 16G Kit", "Customer-installable upgrade for PowerEdge XR7620 and R760; "
         "enables virtual console, virtual media, remote firmware updates, "
         "and advanced server monitoring via web interface"),
    ])

    add_heading(doc, "Your Investment", size=16, color=CORE_ORANGE)
    add_line_item_table(doc,
        items=[
            ("iDRAC9 Enterprise 15G", "Enterprise remote management kit for PowerEdge R650", idrac_15g_client),
            ("iDRAC9 Enterprise 16G", "Enterprise remote management kit for PowerEdge XR7620", idrac_16g_client),
            ("iDRAC9 Enterprise 16G", "Enterprise remote management kit for PowerEdge R760", idrac_16g_client),
            ("Freight & Handling", "Standard ground delivery for all equipment", shipping_client),
        ],
        total_label="Total Investment",
        total_amount=total_client,
    )

    add_whats_included(doc)
    add_terms(doc)
    add_signature(doc)
    add_confidentiality_footer(doc)

    out = os.path.join(BASE_DIR, "OKL-Quote-Additional.docx")
    doc.save(out)
    print(f"Saved: {out}  |  Client price: ${total_client:,.2f}")
    return total_client


# ─────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    p1 = build_quote_r650()
    p2 = build_quote_xr7620()
    p3 = build_quote_r760()
    p4 = build_quote_additional()

    grand_total = p1 + p2 + p3 + p4
    print()
    print("=" * 60)
    print(f"  Quote 1 — R650:        ${p1:>12,.2f}")
    print(f"  Quote 2 — XR7620:      ${p2:>12,.2f}")
    print(f"  Quote 3 — R760:        ${p3:>12,.2f}")
    print(f"  Quote 4 — Additional:  ${p4:>12,.2f}")
    print(f"  {'-' * 40}")
    print(f"  GRAND TOTAL:           ${grand_total:>12,.2f}")
    print(f"  (Dell cost: $57,945.91  |  Markup: ${grand_total - 57945.91:,.2f})")
    print("=" * 60)
