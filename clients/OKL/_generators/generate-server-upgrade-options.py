"""
OKL — Dell Outlet Server Upgrade Options
Branded proposal doc with 3 config tiers + 20% Technijian markup.

Prices are ESTIMATES based on typical Dell Outlet scratch-and-dent pricing
for refurbished PowerEdge 1U rack servers. Must be validated against live
outlet.us.dell.com inventory before quoting.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")


def shade(cell, hex_color):
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    )


def light_borders(tbl):
    tbl_pr = tbl._element.tblPr
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    ))


def set_font(run, size=10, bold=False, color=None, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>')
        rPr.append(rFonts)


def add_heading(doc, text, size=16, color=CORE_BLUE, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=color)
    return p


def add_body(doc, text, size=10, bold=False, color=DARK_CHARCOAL, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_hero(doc, title, subtitle):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(7.0)
    shade(cell, HEX_DARK)
    cell.paragraphs[0].paragraph_format.space_before = Pt(18)
    cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    r = cell.paragraphs[0].add_run(title)
    set_font(r, size=22, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(18)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11, bold=False, color=RGBColor(0xCC, 0xCC, 0xCC))
    # remove default borders
    tbl_pr = tbl._element.tblPr
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    ))
    doc.add_paragraph()


def add_kv_table(doc, rows, header=None, col_widths=(2.4, 4.6)):
    tbl = doc.add_table(rows=(1 if header else 0) + len(rows), cols=2)
    tbl.autofit = False
    light_borders(tbl)
    ri = 0
    if header:
        hdr = tbl.rows[0]
        for i, txt in enumerate(header):
            c = hdr.cells[i]
            shade(c, HEX_BLUE)
            c.width = Inches(col_widths[i])
            c.paragraphs[0].paragraph_format.space_before = Pt(4)
            c.paragraphs[0].paragraph_format.space_after = Pt(4)
            r = c.paragraphs[0].add_run(txt)
            set_font(r, size=10, bold=True, color=WHITE)
        ri = 1
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[ri + i]
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
        shade(row.cells[0], HEX_OFF_WHITE)
        for j, txt in enumerate((k, v)):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            r = c.paragraphs[0].add_run(str(txt))
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)
    doc.add_paragraph()
    return tbl


def add_option_card(doc, tier_tag, model, spec_rows, outlet_price, markup_pct=0.20):
    """Render one option tier as a card-style table."""
    markup = round(outlet_price * markup_pct, 2)
    client_price = round(outlet_price + markup, 2)

    # Header strip
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    light_borders(tbl)
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    left.width = Inches(4.5)
    right.width = Inches(2.5)
    shade(left, HEX_BLUE)
    shade(right, HEX_ORANGE)
    left.paragraphs[0].paragraph_format.space_before = Pt(6)
    left.paragraphs[0].paragraph_format.space_after = Pt(6)
    r1 = left.paragraphs[0].add_run(f"{tier_tag}  —  {model}")
    set_font(r1, size=13, bold=True, color=WHITE)
    right.paragraphs[0].paragraph_format.space_before = Pt(6)
    right.paragraphs[0].paragraph_format.space_after = Pt(6)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = right.paragraphs[0].add_run(f"Client: ${client_price:,.2f}")
    set_font(r2, size=13, bold=True, color=WHITE)

    # Spec body
    spec_tbl = doc.add_table(rows=len(spec_rows), cols=2)
    spec_tbl.autofit = False
    light_borders(spec_tbl)
    for i, (k, v) in enumerate(spec_rows):
        row = spec_tbl.rows[i]
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(5.0)
        shade(row.cells[0], HEX_OFF_WHITE)
        for j, txt in enumerate((k, v)):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            r = c.paragraphs[0].add_run(str(txt))
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)

    # Price breakdown
    price_tbl = doc.add_table(rows=3, cols=2)
    price_tbl.autofit = False
    light_borders(price_tbl)
    breakdown = [
        ("Dell Outlet Estimated Price", f"${outlet_price:,.2f}"),
        (f"Technijian Markup ({int(markup_pct*100)}%)", f"${markup:,.2f}"),
        ("Client Investment", f"${client_price:,.2f}"),
    ]
    for i, (k, v) in enumerate(breakdown):
        row = price_tbl.rows[i]
        row.cells[0].width = Inches(4.5)
        row.cells[1].width = Inches(2.5)
        is_total = (i == 2)
        if is_total:
            shade(row.cells[0], HEX_DARK)
            shade(row.cells[1], HEX_DARK)
        else:
            shade(row.cells[0], HEX_OFF_WHITE)
            shade(row.cells[1], HEX_OFF_WHITE)
        for j, txt in enumerate((k, v)):
            c = row.cells[j]
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            if j == 1:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = c.paragraphs[0].add_run(str(txt))
            set_font(
                r, size=11 if is_total else 10,
                bold=(is_total or j == 0),
                color=WHITE if is_total else DARK_CHARCOAL,
            )
    doc.add_paragraph()
    return client_price


def add_footer_strip(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.cell(0, 0)
    c.width = Inches(7.0)
    shade(c, HEX_OFF_WHITE)
    c.paragraphs[0].paragraph_format.space_before = Pt(8)
    c.paragraphs[0].paragraph_format.space_after = Pt(8)
    r = c.paragraphs[0].add_run(text)
    set_font(r, size=9, bold=False, color=BRAND_GREY)


def build():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(LOGO_PATH, width=Inches(1.8))

    add_hero(
        doc,
        "Server Replacement Options",
        "Oaktree Law  |  Dell PowerEdge 1U Rack  |  Prepared April 15, 2026",
    )

    # Current environment
    add_heading(doc, "Current Environment", size=14)
    add_kv_table(doc, [
        ("Service Tag", "2KPZZV2"),
        ("Platform", "Dell PowerEdge (Gen13, ~2016 era)"),
        ("CPU", "Intel Xeon E3-1220 v5 (4 cores / 4 threads, 3.00 GHz)"),
        ("Memory", "14 GB DDR4 ECC"),
        ("Storage", "119 GB C: + 779 GB D:  (~900 GB usable)"),
        ("Form Factor", "1U Rack"),
        ("Warranty Status", "Out of Dell ProSupport; end-of-life platform"),
    ])

    # Why replace
    add_heading(doc, "Why Replace Now", size=14)
    add_body(doc,
        "The E3-1220 v5 platform shipped in 2015 and is past Dell's ProSupport lifecycle. "
        "Memory is at 14 GB (bordering capacity for modern workloads), and the D: drive is "
        "94% full. Any of the options below delivers at minimum 2x the compute, 2.3x the RAM, "
        "and headroom to 4 TB of usable RAID-protected storage on current or near-current "
        "generation hardware with factory warranty coverage."
    )

    # Options
    add_heading(doc, "Replacement Options", size=16, color=CORE_ORANGE, space_before=14)

    totals = {}

    totals["A"] = add_option_card(doc,
        tier_tag="OPTION A  —  GOOD",
        model="Dell PowerEdge R250  (Gen15, 2022)",
        spec_rows=[
            ("CPU", "Intel Xeon E-2334 (4 cores / 8 threads, 3.4 GHz, Rocket Lake)"),
            ("Memory", "32 GB DDR4 ECC (2x16 GB)"),
            ("Storage", "2 x 2 TB SATA 7.2K, RAID 1 (PERC H355)"),
            ("Usable Storage", "~2 TB protected"),
            ("Networking", "2 x 1 GbE onboard"),
            ("Power", "Single 450W PSU (non-redundant)"),
            ("Management", "iDRAC9 Basic"),
            ("Form Factor", "1U Rack (cabled)"),
            ("Warranty", "Dell Outlet 1-Year ProSupport included"),
        ],
        outlet_price=1850.00,
    )

    totals["B"] = add_option_card(doc,
        tier_tag="OPTION B  —  BETTER  (Recommended)",
        model="Dell PowerEdge R350  (Gen15, 2022)",
        spec_rows=[
            ("CPU", "Intel Xeon E-2378 (8 cores / 16 threads, 2.6 GHz, Rocket Lake)"),
            ("Memory", "64 GB DDR4 ECC (2x32 GB)"),
            ("Storage", "4 x 2 TB SATA 7.2K, RAID 10 (PERC H755)"),
            ("Usable Storage", "~4 TB protected, high IOPS"),
            ("Networking", "2 x 1 GbE onboard + optional 10 GbE"),
            ("Power", "Dual 600W hot-plug PSU (redundant)"),
            ("Management", "iDRAC9 Enterprise"),
            ("Form Factor", "1U Rack (hot-plug bays)"),
            ("Warranty", "Dell Outlet 1-Year ProSupport included"),
        ],
        outlet_price=3200.00,
    )

    totals["C"] = add_option_card(doc,
        tier_tag="OPTION C  —  BEST",
        model="Dell PowerEdge R360  (Gen16, 2024)",
        spec_rows=[
            ("CPU", "Intel Xeon E-2488 (8 cores / 16 threads, 3.2 GHz, Raptor Lake)"),
            ("Memory", "64 GB DDR5 ECC (2x32 GB)"),
            ("Storage", "4 x 2 TB SATA 7.2K, RAID 10 (PERC H965i)"),
            ("Usable Storage", "~4 TB protected, DDR5 throughput"),
            ("Networking", "2 x 1 GbE onboard + optional 10/25 GbE"),
            ("Power", "Dual 700W hot-plug PSU (redundant, Titanium)"),
            ("Management", "iDRAC9 Enterprise"),
            ("Form Factor", "1U Rack (hot-plug bays)"),
            ("Warranty", "Dell Outlet 1-Year ProSupport included"),
        ],
        outlet_price=4250.00,
    )

    # Comparison summary
    add_heading(doc, "Side-by-Side Summary", size=14)
    comp = doc.add_table(rows=6, cols=4)
    comp.autofit = False
    light_borders(comp)
    headers = ["Spec", "Option A — R250", "Option B — R350", "Option C — R360"]
    data = [
        ("CPU cores / threads", "4 / 8", "8 / 16", "8 / 16"),
        ("Memory", "32 GB DDR4", "64 GB DDR4", "64 GB DDR5"),
        ("Usable storage", "~2 TB (RAID 1)", "~4 TB (RAID 10)", "~4 TB (RAID 10)"),
        ("Redundant PSU", "No", "Yes", "Yes"),
        ("Generation", "Gen15 (2022)", "Gen15 (2022)", "Gen16 (2024)"),
    ]
    for j, h in enumerate(headers):
        c = comp.rows[0].cells[j]
        shade(c, HEX_BLUE)
        c.width = Inches([1.6, 1.8, 1.8, 1.8][j])
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True, color=WHITE)
    for i, row_data in enumerate(data, start=1):
        for j, txt in enumerate(row_data):
            c = comp.rows[i].cells[j]
            c.width = Inches([1.6, 1.8, 1.8, 1.8][j])
            if j == 0:
                shade(c, HEX_OFF_WHITE)
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            r = c.paragraphs[0].add_run(txt)
            set_font(r, size=10, bold=(j == 0), color=DARK_CHARCOAL)
    doc.add_paragraph()

    # Investment summary
    add_heading(doc, "Investment Summary (includes 20% Technijian markup)", size=14)
    inv = doc.add_table(rows=4, cols=4)
    inv.autofit = False
    light_borders(inv)
    inv_headers = ["Option", "Dell Outlet Est.", "Markup (20%)", "Client Investment"]
    inv_rows = [
        ("A — R250 (Good)",     1850.00, 370.00,  totals["A"]),
        ("B — R350 (Better)",   3200.00, 640.00,  totals["B"]),
        ("C — R360 (Best)",     4250.00, 850.00,  totals["C"]),
    ]
    for j, h in enumerate(inv_headers):
        c = inv.rows[0].cells[j]
        shade(c, HEX_BLUE)
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)
        if j >= 1:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True, color=WHITE)
    for i, (label, outlet, markup, client) in enumerate(inv_rows, start=1):
        vals = [label, f"${outlet:,.2f}", f"${markup:,.2f}", f"${client:,.2f}"]
        for j, txt in enumerate(vals):
            c = inv.rows[i].cells[j]
            if j == 0:
                shade(c, HEX_OFF_WHITE)
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            if j >= 1:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = c.paragraphs[0].add_run(txt)
            set_font(r, size=10, bold=(j == 0 or j == 3), color=DARK_CHARCOAL)
    doc.add_paragraph()

    # Recommendation
    add_heading(doc, "Technijian Recommendation", size=14, color=CORE_ORANGE)
    add_body(doc,
        "Option B (PowerEdge R350) is the best value for Oaktree Law's workload profile. "
        "It doubles CPU core count, quadruples memory, and provides redundant power and "
        "RAID 10 storage at roughly $1.7K less than the Gen16 R360 while remaining "
        "well within Dell's current support lifecycle. For a legal-practice workload "
        "where uptime matters more than bleeding-edge throughput, the R350 is the "
        "right balance of headroom, reliability, and cost."
    )

    # Delivery & terms
    add_heading(doc, "Delivery & Terms", size=14)
    add_kv_table(doc, [
        ("Procurement Lead Time", "5-10 business days (Dell Outlet ships from Round Rock, TX)"),
        ("Installation & Rack Services", "Quoted separately (typical 4-8 labor hours at contracted rates)"),
        ("OS & Hypervisor Licensing", "Not included; quoted separately"),
        ("Freight", "Included (standard ground) - expedite available"),
        ("Payment Terms", "50% at PO, 50% on delivery"),
        ("Validity", "14 days from issue date - Dell Outlet inventory is first-come, first-served"),
    ])

    # Disclaimer
    add_heading(doc, "Important Notes", size=12, color=BRAND_GREY)
    add_footer_strip(doc,
        "Pricing shown for Dell Outlet hardware is an ESTIMATE based on typical scratch-and-dent "
        "and refurbished inventory pricing for comparable PowerEdge configurations. Actual outlet "
        "inventory and pricing changes daily. Technijian will confirm live Dell Outlet SKU, exact "
        "configuration, and final pricing at the time of PO issuance. The 20% markup covers "
        "Technijian procurement, configuration validation, pre-ship burn-in testing, asset tagging, "
        "and one-year advance-replacement coordination with Dell."
    )

    # Save
    out = os.path.join(os.path.dirname(__file__), "..", "04_Quotes", "OKL-Server-Upgrade-Options.docx")
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
