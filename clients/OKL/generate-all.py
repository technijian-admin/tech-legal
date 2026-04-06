"""
OKL Complete Document Suite Generator
Generates MSA, SOW, Schedule A, Schedule C, Executive Summary
All using consistent Technijian Brand Guide 2026 formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors (Technijian Brand Guide 2026) ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "technijian-logo.png")


def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


def set_grey_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        return r

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_borders(tbl)

    def cover_page(self, title, subtitle, date="March 23, 2026"):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(LOGO_PATH, width=Inches(3.0))
        # Orange divider
        div = self.doc.add_table(rows=1, cols=3)
        div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in div.rows[0].cells:
            c.text = ""
            c.paragraphs[0].paragraph_format.space_before = Pt(0)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div.rows[0].cells[1], HEX_ORANGE)
        div.rows[0].cells[0].width = Inches(2.25)
        div.rows[0].cells[1].width = Inches(2.5)
        div.rows[0].cells[2].width = Inches(2.25)
        remove_borders(div)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Prepared for ", size=12, color=BRAND_GREY)
        self.run(p, "Oaktree Law", size=12, bold=True, color=CORE_BLUE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.accent_bar(HEX_ORANGE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar = tbl.rows[0].cells[0]
        set_cell_shading(bar, HEX_BLUE)
        bar.width = Inches(0.08)
        bar.text = ""
        bar.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar.paragraphs[0].paragraph_format.space_after = Pt(0)
        tc = tbl.rows[0].cells[1]
        tc.text = ""
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_borders(tbl)

    def part_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_borders(tbl)

    def styled_table(self, headers, rows, total_indices=None):
        total_indices = total_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            set_cell_shading(c, HEX_BLUE)
            c.text = ""
            self.run(c.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for ri, rd in enumerate(rows):
            row = tbl.add_row()
            is_total = ri in total_indices
            for i, val in enumerate(rd):
                c = row.cells[i]
                c.text = ""
                if is_total:
                    set_cell_shading(c, HEX_OFF_WHITE)
                    self.run(c.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(c.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        set_grey_borders(tbl)
        return tbl

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            self.run(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, num, text, indent=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent
        self.run(p, f"{prefix}{num} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self, client_name="OAKTREE LAW", contact_name="Ed Pits"):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)
        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", "Name: _________________________________",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)
        self.spacer(12)
        p = self.doc.add_paragraph()
        self.run(p, client_name, size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", f"Name: {contact_name}",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        self.doc.save(filename)
        print(f"Created {filename}")


# ════════════════════════════════════════════════════════════════
#  MSA
# ════════════════════════════════════════════════════════════════
def build_msa():
    d = BrandedDoc()
    d.cover_page("Master Service Agreement", "Oaktree Law")
    d.page_break()

    d.part_header("MASTER SERVICE AGREEMENT")
    d.spacer()
    d.body("Agreement Number: MSA-OKL-2026")
    d.body("Effective Date: March 23, 2026")
    d.spacer()
    d.body("This Master Service Agreement (\u201cAgreement\u201d) is entered into by and between:")
    d.spacer()
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", "")
    d.body("18 Technology Drive, Suite 141")
    d.body("Irvine, California 92618")
    d.spacer()
    d.body("and")
    d.spacer()
    d.body_bold("Oaktree Law (\u201cClient\u201d)", "")
    d.body("[CLIENT ADDRESS]")
    d.body("[CITY, STATE ZIP]")
    d.body("Primary Contact: Ed Pits")
    d.spacer()
    d.body("(collectively, the \u201cParties\u201d)")

    # SECTION 1
    d.spacer()
    d.section_header("SECTION 1 \u2014 SCOPE OF SERVICES")

    d.numbered("1.01.", "Services. Technijian shall provide the services described in the Schedules attached to this Agreement, which are incorporated herein by reference:")
    d.bullet("Schedule A \u2014 Monthly Managed Services (Cloud Hosting, Security, Monitoring)")
    d.bullet("Schedule B \u2014 Subscription and License Services")
    d.bullet("Schedule C \u2014 Rate Card")
    d.body("Additional services may be provided through Statements of Work (\u201cSOWs\u201d) executed under this Agreement.")

    d.numbered("1.02.", "Standard of Care. Technijian shall perform all services in a professional and workmanlike manner, consistent with industry standards for managed IT service providers.")

    d.numbered("1.03.", "Service Level Agreement. The service levels applicable to the services are set forth in Schedule A. Technijian shall use commercially reasonable efforts to meet the service levels described therein.")

    d.numbered("1.04.", "Client Responsibilities. Client shall:")
    d.numbered("(a)", "Provide Technijian with reasonable access to Client\u2019s systems, facilities, and personnel as necessary for Technijian to perform the services;", indent=1)
    d.numbered("(b)", "Designate a primary point of contact for communications with Technijian;", indent=1)
    d.numbered("(c)", "Maintain current and accurate information regarding Client\u2019s systems and infrastructure;", indent=1)
    d.numbered("(d)", "Comply with all applicable laws and regulations in connection with its use of the services; and", indent=1)
    d.numbered("(e)", "Be solely responsible for the security and management of Client\u2019s account credentials and passwords.", indent=1)

    d.numbered("1.05.", "Independent Contractor. Technijian is an independent contractor. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the Parties.")

    # SECTION 2
    d.spacer()
    d.section_header("SECTION 2 \u2014 TERM AND RENEWAL")

    d.numbered("2.01.", "Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the \u201cInitial Term\u201d).")
    d.numbered("2.02.", "Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a \u201cRenewal Term\u201d), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term. Technijian shall send Client a written renewal reminder at least thirty (30) days prior to each renewal date, which shall restate the auto-renewal terms and cancellation method.")
    d.numbered("2.03.", "Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party.")
    d.numbered("2.04.", "Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party:")
    d.numbered("(a)", "Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days after receiving written notice of the breach; or", indent=1)
    d.numbered("(b)", "Becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.", indent=1)

    d.numbered("2.05.", "Effect of Termination.")
    d.numbered("(a)", "Upon termination, Client shall pay all fees and charges for services rendered through the date of termination, including any remaining obligations for annual licenses and subscriptions procured on Client\u2019s behalf.", indent=1)
    d.numbered("(b)", "Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, subject to payment of applicable fees.", indent=1)
    d.numbered("(c)", "Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format, provided Client is not in breach of this Agreement.", indent=1)
    d.numbered("(d)", "The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 9.09 (Personnel Transition Fee), Section 10 (Data Protection), and Section 11 (Insurance).", indent=1)

    d.page_break()

    # SECTION 3
    d.section_header("SECTION 3 \u2014 PAYMENT")

    d.numbered("3.01.", "Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.")
    d.numbered("3.02.", "Invoice Types. Client may receive the following types of invoices from Technijian during the term of this Agreement. Each invoice will clearly identify its type, the applicable Schedule or SOW, and the billing period or delivery event.")
    d.numbered("(a)", "Monthly Service Invoice. Issued on the first business day of each month for recurring managed services under Schedule A (Online Services, infrastructure, monitoring, desktop/server management). Billed in advance for the upcoming month.", indent=1)
    d.numbered("(b)", "Monthly Recurring Subscription Invoice. Issued on the first business day of each month for subscription and license services under Schedule B (software licenses, SaaS subscriptions, SIP trunk services). Billed in advance for the upcoming month.", indent=1)
    d.numbered("(c)", "Weekly In-Contract Invoice. Issued every Friday for Virtual Staff (contracted support) services performed under Schedule A, Part 3, during the preceding week (Monday through Friday). Each invoice includes: (i) a listing of each support ticket addressed; (ii) the assigned resource, role, and hours spent per ticket; (iii) a description of the work performed per ticket; (iv) whether work was performed during normal or after-hours; and (v) the current running balance for each contracted role. The weekly in-contract invoice is issued for transparency and tracking purposes; the actual billed amount is governed by the cycle-based billing model described in Schedule A, Section 3.3.", indent=1)
    d.numbered("(d)", "Weekly Out-of-Contract Invoice. Issued every Friday for labor services performed outside the scope of any active Schedule or SOW \u2014 including ad-hoc support requests, emergency work, and services performed under a SOW with hourly billing (such as CTO Advisory engagements). Each invoice includes: (i) a listing of each support ticket or task performed; (ii) the assigned resource, role, and applicable hourly rate from the Rate Card (Schedule C); (iii) time entries with hours billed per activity (in 15-minute increments); (iv) whether work was performed during normal or after-hours; and (v) the total hours and total amount for the week.", indent=1)
    d.numbered("(e)", "Equipment and Materials Invoice. Issued upon delivery or procurement of hardware, software licenses (perpetual), or other tangible goods on Client\u2019s behalf. Each invoice includes: (i) item description, manufacturer, and model/part number; (ii) quantity and unit price; (iii) applicable sales tax; (iv) shipping and handling charges, if any; and (v) total amount due. Title to equipment shall not pass to Client until payment is received in full.", indent=1)
    d.numbered("(f)", "Project Milestone Invoice. Issued upon completion of a project milestone as defined in an applicable SOW. The milestone, deliverables, and invoiced amount are as specified in the payment schedule of the SOW. Milestone invoices are billed in arrears upon acceptance of the deliverables or deemed acceptance under the SOW\u2019s acceptance provisions.", indent=1)
    d.numbered("3.03.", "Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified in the applicable Schedule or SOW.")
    d.numbered("3.04.", "Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated as simple interest from the date payment was due until the date payment is received in full. Late fees are payable in addition to the outstanding principal amount. The Parties acknowledge that the late fee represents a reasonable estimate of Technijian\u2019s administrative costs and damages resulting from late payment, including cash-flow disruption, collection overhead, and the cost of carrying accounts receivable, and is not intended as a penalty.")
    d.numbered("3.05.", "Disputed Invoices.")
    d.numbered("(a)", "Weekly Invoices (In-Contract and Out-of-Contract). Because weekly invoices include detailed ticket descriptions and time entries, Client shall have thirty (30) days from the invoice date to review and dispute any portion of a weekly invoice. Client shall notify Technijian in writing, specifying the ticket number(s) and nature of the dispute. Undisputed tickets and time entries on the same invoice shall remain payable by the due date. Failure to provide a timely written dispute notice within the thirty (30) day period shall constitute acceptance of all tickets and time entries on the invoice.", indent=1)
    d.numbered("(b)", "All Other Invoices. For monthly service invoices, monthly recurring subscription invoices, equipment invoices, and project milestone invoices, Client shall notify Technijian in writing within fifteen (15) days of the invoice date if Client disputes any portion, specifying the nature and basis of the dispute. Client shall pay all undisputed amounts by the due date. Failure to provide a timely written dispute notice shall constitute acceptance of the invoice.", indent=1)
    d.numbered("(c)", "Resolution. The Parties shall work in good faith to resolve any invoice dispute within thirty (30) days of the dispute notice. If the dispute results in an adjustment, Technijian shall issue a credit memo or revised invoice within ten (10) business days of resolution.", indent=1)
    d.numbered("3.06.", "Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend services under the Schedule or SOW associated with the unpaid invoice until payment is received in full, including all accrued late fees. If Client fails to pay any undisputed invoice within sixty (60) days of the due date, Technijian may suspend all services under this Agreement and any related Schedules or SOWs. Recurring fees for suspended services shall continue to accrue for a period not to exceed thirty (30) days following the date of suspension, after which Technijian may terminate the affected Schedule, SOW, or this Agreement upon written notice. Suspension of services shall not relieve Client of its payment obligations.")
    d.numbered("3.07.", "Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian\u2019s income.")

    # SECTION 4
    d.spacer()
    d.section_header("SECTION 4 \u2014 CONFIDENTIALITY")

    d.numbered("4.01.", "Definition. \u201cConfidential Information\u201d means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information.")
    d.numbered("4.02.", "Obligations. Each Party shall:")
    d.numbered("(a)", "Hold the other Party\u2019s Confidential Information in confidence using at least the same degree of care it uses for its own confidential information, but not less than reasonable care;", indent=1)
    d.numbered("(b)", "Not disclose Confidential Information to third parties without prior written consent, except to employees, agents, and subcontractors who have a need to know and are bound by equivalent obligations; and", indent=1)
    d.numbered("(c)", "Not use Confidential Information for any purpose other than performing obligations under this Agreement.", indent=1)
    d.numbered("4.03.", "Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known to the receiving Party prior to disclosure, is independently developed, or is received from a third party without restriction.")
    d.numbered("4.04.", "Compelled Disclosure. If required by law or court order to disclose Confidential Information, the receiving Party shall provide prompt written notice to the disclosing Party and cooperate in seeking a protective order.")
    d.numbered("4.05.", "Duration. Confidentiality obligations shall survive termination for a period of three (3) years.")

    # SECTION 5
    d.spacer()
    d.section_header("SECTION 5 \u2014 LIMITATION OF LIABILITY")

    d.numbered("5.01.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, NEITHER PARTY\u2019S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM (THE \u201cSTANDARD CAP\u201d).")
    d.numbered("5.02.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL, REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE OR WHETHER EITHER PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.")
    d.numbered("5.03.", "ENHANCED CAP FOR CERTAIN CLAIMS. FOR CLAIMS ARISING FROM BREACHES OF SECTION 4 (CONFIDENTIALITY), SECTION 10 (DATA PROTECTION), INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, THE TOTAL AGGREGATE LIABILITY OF THE RESPONSIBLE PARTY SHALL NOT EXCEED THREE (3) TIMES THE STANDARD CAP DEFINED IN SECTION 5.01 (THE \u201cENHANCED CAP\u201d). THE ENHANCED CAP SHALL NOT APPLY TO LIABILITY ARISING FROM A PARTY\u2019S WILLFUL AND INTENTIONAL MISAPPROPRIATION OF THE OTHER PARTY\u2019S CONFIDENTIAL INFORMATION OR CLIENT DATA, FOR WHICH LIABILITY SHALL BE UNCAPPED.")
    d.numbered("5.04.", "Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data, Client acknowledges that Client is solely responsible for maintaining backup copies of its data.")

    d.page_break()

    # SECTION 6
    d.section_header("SECTION 6 \u2014 INDEMNIFICATION")

    d.numbered("6.01.", "By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from Technijian\u2019s gross negligence or willful misconduct in performing the services.")
    d.numbered("6.02.", "By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from Client\u2019s use of the services in violation of applicable law, Client\u2019s breach of this Agreement, or any data, content, or materials provided by Client.")
    d.numbered("6.03.", "Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party\u2019s prior written consent.")

    # SECTION 7
    d.spacer()
    d.section_header("SECTION 7 \u2014 INTELLECTUAL PROPERTY")

    d.numbered("7.01.", "Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes.")
    d.numbered("7.02.", "Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property.")
    d.numbered("7.03.", "Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW.")

    # SECTION 8
    d.spacer()
    d.section_header("SECTION 8 \u2014 DISPUTE RESOLUTION")

    d.numbered("8.01.", "Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations for a period of thirty (30) days.")
    d.numbered("8.02.", "Mediation. If not resolved, the Parties shall submit the dispute to mediation in Orange County, California, for a period not to exceed sixty (60) days.")
    d.numbered("8.03.", "Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association in Orange County, California.")
    d.numbered("8.04.", "Injunctive Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief to prevent irreparable harm.")

    # SECTION 9
    d.spacer()
    d.section_header("SECTION 9 \u2014 GENERAL PROVISIONS")

    d.numbered("9.01.", "Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties.")
    d.numbered("9.02.", "Amendment. This Agreement may only be amended by a written instrument signed by both Parties.")
    d.numbered("9.03.", "Severability. If any provision is found invalid or unenforceable, the remaining provisions shall continue in full force.")
    d.numbered("9.04.", "Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party.")
    d.numbered("9.05.", "Assignment. Neither Party may assign this Agreement without prior written consent, except in connection with a merger, acquisition, or sale of substantially all assets.")
    d.numbered("9.06.", "Force Majeure.")
    d.numbered("(a)", "Neither Party shall be liable for delays or failures in performance caused by events beyond its reasonable control, including natural disasters, acts of government, labor disputes, pandemics, cyberattacks on critical infrastructure, or failures of major third-party infrastructure services that are not attributable to the affected Party\u2019s failure to implement reasonable redundancy (\u201cForce Majeure Event\u201d).", indent=1)
    d.numbered("(b)", "The affected Party shall notify the other Party in writing within five (5) business days of becoming aware of a Force Majeure Event and shall use commercially reasonable efforts to mitigate the impact and resume performance.", indent=1)
    d.numbered("(c)", "Payment obligations are not excused by a Force Majeure Event.", indent=1)
    d.numbered("(d)", "If a Force Majeure Event prevents performance of a material portion of the services for more than ninety (90) consecutive days, either Party may terminate the affected Schedule, SOW, or this Agreement upon fifteen (15) days written notice without liability for early termination fees, and Client shall pay only for services actually rendered through the date of termination.", indent=1)
    d.numbered("9.07.", "Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier.")
    d.numbered("9.08.", "Governing Law. This Agreement shall be governed by the laws of the State of California.")
    d.numbered("9.09.", "Personnel Transition Fee. The Parties acknowledge that each invests significant resources in recruiting, training, and retaining skilled personnel. If either Party hires (whether as an employee or independent contractor) any individual who was an employee of the other Party and who was directly involved in performing or receiving services under this Agreement, and such hiring occurs during the term of this Agreement or within twelve (12) months following termination, the hiring Party shall pay the other Party a personnel transition fee equal to 25% of the hired individual\u2019s first-year annual compensation (base salary or annualized contractor fees). This fee represents a reasonable estimate of the non-hiring Party\u2019s recruiting and training costs and is not intended as a restraint on trade or employment. This Section does not restrict any individual\u2019s right to seek or obtain employment, and shall not apply to individuals who: (a) respond to general public job postings or advertisements not specifically targeted at the other Party\u2019s employees; or (b) are referred by a third-party recruiting firm without the hiring Party\u2019s direction to target the other Party\u2019s employees.")
    d.numbered("9.10.", "Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.")

    # SECTION 10
    d.spacer()
    d.section_header("SECTION 10 \u2014 DATA PROTECTION")

    d.numbered("10.01.", "CCPA/CPRA Compliance. To the extent Technijian processes, stores, or has access to personal information (as defined under the California Consumer Privacy Act, as amended by the California Privacy Rights Act, Cal. Civ. Code \u00a7 1798.100 et seq., collectively \u201cCCPA\u201d) on behalf of Client, Technijian acts as a \u201cservice provider\u201d as defined under Cal. Civ. Code \u00a7 1798.140(ag) and shall:")
    d.numbered("(a)", "Process such personal information only as necessary to perform the services and in accordance with Client\u2019s documented instructions and this Agreement;", indent=1)
    d.numbered("(b)", "Not sell, share, retain, use, or disclose personal information for any purpose other than performing the services, including not using personal information for targeted advertising or cross-context behavioral advertising;", indent=1)
    d.numbered("(c)", "Not combine personal information received from Client with personal information received from other sources or collected from Technijian\u2019s own interactions with individuals, except as expressly permitted by the CCPA to perform the services;", indent=1)
    d.numbered("(d)", "Implement reasonable security measures appropriate to the nature of the personal information, consistent with the requirements of Section 10.03;", indent=1)
    d.numbered("(e)", "Cooperate with Client in responding to verifiable consumer rights requests under the CCPA (including access, deletion, correction, and opt-out requests) within ten (10) business days of Client\u2019s request;", indent=1)
    d.numbered("(f)", "Notify Client within five (5) business days if Technijian determines that it can no longer meet its obligations as a service provider under the CCPA;", indent=1)
    d.numbered("(g)", "Ensure that all subcontractors who process personal information on behalf of Client are bound by written agreements containing data protection obligations at least as protective as those in this Section 10, and Technijian shall remain liable for the acts and omissions of its subcontractors with respect to personal information;", indent=1)
    d.numbered("(h)", "Permit Client, upon thirty (30) days written notice and no more than once per twelve (12) month period, to audit or inspect Technijian\u2019s data processing practices to verify compliance with this Section 10, or, at Technijian\u2019s option, provide Client with a summary of a recent independent third-party audit (such as SOC 2 Type II) covering the relevant controls; and", indent=1)
    d.numbered("(i)", "Certify that Technijian understands the restrictions in this Section 10 and will comply with them.", indent=1)

    d.numbered("10.02.", "Security Incident Notification. If Technijian becomes aware of a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Client Data (\u201cSecurity Incident\u201d), Technijian shall: (a) notify Client in writing without unreasonable delay and in no event later than forty-eight (48) hours after becoming aware of the Security Incident; (b) provide Client with sufficient information to enable Client to comply with its obligations under California Civil Code \u00a7 1798.82 (data breach notification), including the categories and approximate number of records affected, the nature of the incident, and the measures taken or proposed to address it; (c) cooperate with Client\u2019s investigation of the Security Incident; and (d) take reasonable steps to contain and remediate the Security Incident.")

    d.numbered("10.03.", "Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards designed to protect Client Data from unauthorized access, use, or disclosure, consistent with industry standards for managed IT service providers. Such safeguards shall include, at a minimum: (a) encryption of Client Data in transit and at rest; (b) access controls limiting access to authorized personnel; (c) regular security assessments and vulnerability testing; and (d) employee security awareness training.")

    d.numbered("10.04.", "Regulatory Compliance. If Client is subject to HIPAA, PCI DSS, GDPR, or other industry-specific data protection requirements, the Parties shall execute a separate addendum addressing the additional obligations applicable to the regulated data.")

    d.numbered("10.05.", "Data Return and Deletion. Upon termination of this Agreement or upon Client\u2019s written request, Technijian shall securely delete or return all Client Data in its possession within thirty (30) days, using methods consistent with NIST SP 800-88 or equivalent standards, and shall certify such deletion in writing upon request.")

    d.signatures()
    d.spacer()
    d.body("Schedules:")
    d.bullet("Schedule A \u2014 Monthly Managed Services")
    d.bullet("Schedule B \u2014 Subscription and License Services")
    d.bullet("Schedule C \u2014 Rate Card")
    d.footer_bar()
    d.save("MSA-OKL.docx")


# ════════════════════════════════════════════════════════════════
#  SOW-001 SERVER MIGRATION
# ════════════════════════════════════════════════════════════════
def build_sow():
    d = BrandedDoc()
    d.cover_page("Statement of Work", "Server Migration & Cloud Hosting\nSOW-OKL-001")
    d.page_break()

    d.part_header("STATEMENT OF WORK")
    d.spacer()
    d.body("SOW Number: SOW-OKL-001")
    d.body("Effective Date: March 23, 2026")
    d.body("Master Service Agreement: MSA-OKL-2026")
    d.spacer()
    d.body("This Statement of Work (\u201cSOW\u201d) is entered into by and between:")
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", " \u2014 18 Technology Drive, Suite 141, Irvine, California 92618")
    d.body("and")
    d.body_bold("Oaktree Law (\u201cClient\u201d)", " \u2014 Primary Contact: Ed Pits")

    # 1. PROJECT OVERVIEW
    d.spacer()
    d.section_header("1. PROJECT OVERVIEW")
    d.body_bold("1.1 Project Title: ", "Physical Server Migration to Technijian Cloud with Cloudbrink ZTNA and OneDrive Migration")
    d.body("1.2 Project Description: Oaktree Law currently operates a physical on-premises server that requires migration to a cloud-hosted environment. This SOW covers:")
    d.bullet("Full migration of the existing physical server to Technijian\u2019s private cloud datacenter")
    d.bullet("Deployment of a Cloudbrink Zero Trust Network Access (ZTNA) virtual appliance for secure user access")
    d.bullet("Migration of approximately 779 GB of shared folders to Microsoft OneDrive / SharePoint Online")
    d.bullet("Deployment of security and management agents (CrowdStrike, Huntress, Patch Management, My Remote)")
    d.bullet("Ongoing managed hosting, security, and support services")

    d.spacer()
    d.body_bold("1.3 Current Environment", "")
    d.styled_table(
        ["Component", "Specification"],
        [
            ("Processor", "Intel Xeon E3-1220 v5 @ 3.00 GHz (2 processors)"),
            ("Memory", "14.0 GB RAM"),
            ("Local Drive (C:)", "119 GB total / 3.33 GB free"),
            ("Shared Storage (D:)", "779 GB total / 60.7 GB free"),
            ("Internet", "500 / 500 Mbps symmetric"),
        ],
    )

    # 2. SCOPE OF WORK
    d.spacer()
    d.section_header("2. SCOPE OF WORK")
    d.body_bold("2.1 In Scope", "")
    d.bullet("Full discovery and assessment of the existing physical server environment")
    d.bullet("Provisioning of two (2) virtual machines in the Technijian datacenter")
    d.bullet("Physical-to-Virtual (P2V) migration of the existing server to VM1")
    d.bullet("Deployment and configuration of Cloudbrink ZTNA virtual appliance on VM2")
    d.bullet("Migration of shared folders (~779 GB) to Microsoft OneDrive / SharePoint Online")
    d.bullet("Installation of security agents: CrowdStrike, Huntress, Patch Management, My Remote")
    d.bullet("Deployment of Veeam backup for both virtual machines")
    d.bullet("End-to-end testing, validation, and go-live cutover")
    d.bullet("30-day post-migration support period")

    d.spacer()
    d.body_bold("2.2 Out of Scope", "")
    d.bullet("Desktop/workstation support, upgrades, or reimaging")
    d.bullet("Email migration or Microsoft 365 tenant configuration")
    d.bullet("Line-of-business application reconfiguration beyond basic validation")
    d.bullet("Procurement of Microsoft 365 / OneDrive licenses (Client responsibility)")
    d.bullet("Physical decommissioning or disposal of the old server hardware")
    d.bullet("Cloudbrink per-user subscription licensing (billed separately by Cloudbrink)")
    d.bullet("Network equipment upgrades at Client\u2019s office")

    d.spacer()
    d.body_bold("2.3 Assumptions", "")
    d.bullet("Client will provide administrative credentials and access to the existing server")
    d.bullet("Client will ensure all critical data is backed up prior to migration start")
    d.bullet("Client\u2019s 500/500 Mbps internet is sufficient for Cloudbrink performance")
    d.bullet("Client has or will procure Microsoft 365 licenses with OneDrive/SharePoint storage")
    d.bullet("Migration work will be performed during off-hours to minimize disruption")

    d.page_break()

    # 3. PROJECT PHASES
    d.section_header("3. PROJECT PHASES")

    phases = [
        ("Phase 1: Discovery & Assessment", "Comprehensive audit of the existing physical server including installed roles, services, applications, shared folder structure, permissions, and network configuration."),
        ("Phase 2: Cloud Environment Provisioning", "Provision and configure the target environment in Technijian datacenter: two virtual machines, storage, networking, and firewall rules."),
        ("Phase 3: Physical-to-Virtual Server Migration", "Migrate the server OS and applications (~119 GB local volume) to VM1 via P2V. Includes pre-migration backup, driver remediation, Windows activation, service validation, and network cutover."),
        ("Phase 4: Shared Folder Migration to OneDrive / SharePoint", "Migrate ~718 GB of shared folder data to SharePoint Online / OneDrive for Business. Includes site architecture design, pre-migration file remediation, NTFS-to-SharePoint permission mapping, incremental migration runs, and per-workstation OneDrive sync client setup."),
        ("Phase 5: Cloudbrink ZTNA Deployment", "Deploy Cloudbrink ZTNA virtual appliance on VM2. Includes tenant configuration, connector deployment, Microsoft Entra ID integration, application resource definitions, per-user ZTNA policies, split tunneling, and agent deployment."),
        ("Phase 6: Security & Management Agent Deployment", "Install and configure all security and management agents on both virtual machines."),
        ("Phase 7: Testing, Validation & Go-Live", "End-to-end testing of migrated server, OneDrive/SharePoint access, Cloudbrink ZTNA connectivity, and all security agents. Coordinated production cutover during a planned maintenance window."),
        ("Phase 8: Post-Migration Support (30 Days)", "30-day post-migration support period for issue resolution, performance tuning, and user assistance."),
    ]
    for title, desc in phases:
        d.body_bold(title + ". ", desc)

    # Phase detail tables
    d.spacer()
    d.section_header("3.1 Phase Labor Breakdown")

    phase_tables = [
        ("Phase 1: Discovery & Assessment", [
            ("CTO Advisory", "Migration strategy, architecture planning, risk assessment", "$250/hr", "2", "Week 1"),
            ("Tech Support (US)", "Server audit, inventory, and documentation", "$150/hr", "2", "Week 1"),
        ]),
        ("Phase 2: Cloud Environment Provisioning", [
            ("Tech Support (US)", "VM provisioning and OS install", "$150/hr", "4", "Week 1\u20132"),
            ("Tech Support (US)", "Network and storage configuration", "$150/hr", "2", "Week 1\u20132"),
            ("Tech Support (Offshore)", "Backup setup (Veeam)", "$45/hr", "2", "Week 2"),
        ]),
        ("Phase 3: Server Migration (P2V)", [
            ("Tech Support (US)", "Pre-migration backup verification and snapshot", "$150/hr", "2", "Week 2"),
            ("Tech Support (US)", "P2V conversion setup (Disk2VHD / Veeam)", "$150/hr", "2", "Week 2\u20133"),
            ("Tech Support (Offshore)", "Data transfer monitoring (~119 GB OS volume)", "$45/hr", "2", "Week 2\u20133"),
            ("Tech Support (US)", "Driver cleanup, HAL remediation, VM tools install", "$150/hr", "2", "Week 3"),
            ("Tech Support (US)", "Windows activation, licensing validation", "$150/hr", "1", "Week 3"),
            ("Tech Support (US)", "Server role/service validation (AD, DNS, DHCP, file shares, print)", "$150/hr", "3", "Week 3"),
            ("Tech Support (US)", "DNS/IP reconfiguration and network cutover", "$150/hr", "2", "Week 3"),
            ("Tech Support (Offshore)", "Performance validation and baseline comparison", "$45/hr", "2", "Week 3"),
        ]),
        ("Phase 4: OneDrive / SharePoint Migration", [
            ("CTO Advisory", "SharePoint site collection architecture and library design", "$250/hr", "2", "Week 3"),
            ("Tech Support (US)", "SharePoint Migration Tool setup and configuration", "$150/hr", "2", "Week 3"),
            ("Tech Support (Offshore)", "Pre-migration file scan and remediation", "$45/hr", "3", "Week 3"),
            ("Tech Support (US)", "Initial bulk migration kickoff (~718 GB)", "$150/hr", "2", "Week 3\u20134"),
            ("Tech Support (Offshore)", "Migration monitoring, failed item resolution, delta sync", "$45/hr", "3", "Week 4"),
            ("Tech Support (US)", "NTFS-to-SharePoint permission mapping and validation", "$150/hr", "4", "Week 4"),
            ("Tech Support (Offshore)", "OneDrive sync client deployment on user workstations", "$45/hr", "2", "Week 4\u20135"),
            ("Tech Support (US)", "User training session and quick-reference guide", "$150/hr", "2", "Week 5"),
            ("Tech Support (Offshore)", "Post-migration validation (file counts, integrity, access)", "$45/hr", "2", "Week 5"),
        ]),
        ("Phase 5: Cloudbrink ZTNA Deployment", [
            ("CTO Advisory", "ZTNA architecture, policy design, application resource planning", "$250/hr", "2", "Week 5"),
            ("Tech Support (US)", "Cloudbrink tenant setup and connector deployment on VM2", "$150/hr", "2", "Week 5"),
            ("Tech Support (US)", "Microsoft Entra ID identity provider integration", "$150/hr", "2", "Week 5"),
            ("Tech Support (US)", "Split tunneling, routing, and performance tuning", "$150/hr", "1", "Week 5"),
            ("Tech Support (Offshore)", "Cloudbrink agent deployment to end-user devices", "$45/hr", "3", "Week 5\u20136"),
            ("Tech Support (US)", "Multi-location access testing and troubleshooting", "$150/hr", "2", "Week 6"),
        ]),
        ("Phase 6: Security Agent Deployment", [
            ("Tech Support (US)", "Agent installation and configuration", "$150/hr", "2", "Week 6"),
            ("Tech Support (Offshore)", "Validation and reporting verification", "$45/hr", "2", "Week 6"),
        ]),
        ("Phase 7: Testing & Go-Live", [
            ("CTO Advisory", "Go-live coordination, cutover plan review, client sign-off", "$250/hr", "2", "Week 6"),
            ("Tech Support (US)", "End-to-end functional testing (server, apps, shares)", "$150/hr", "3", "Week 6"),
            ("Tech Support (US)", "Coordinated go-live cutover (DNS, firewall, routing)", "$150/hr", "2", "Week 6\u20137"),
            ("Tech Support (Offshore)", "Cloudbrink and OneDrive access validation with end users", "$45/hr", "1", "Week 7"),
            ("Tech Support (Offshore)", "Post-cutover monitoring (first 48 hours)", "$45/hr", "2", "Week 7"),
        ]),
        ("Phase 8: Post-Migration Support", [
            ("Tech Support (US)", "Escalated issue resolution and performance tuning", "$150/hr", "2", "Weeks 7\u201311"),
            ("Tech Support (Offshore)", "Routine support tickets, monitoring, user assistance", "$45/hr", "6", "Weeks 7\u201311"),
        ]),
    ]

    for phase_name, rows in phase_tables:
        d.body_bold(phase_name, "")
        d.styled_table(
            ["Role", "Description", "Rate", "Hours", "Timeline"],
            rows,
        )
        d.spacer()

    d.page_break()

    # 4. PRICING AND PAYMENT
    d.section_header("4. PRICING AND PAYMENT")
    d.spacer()

    d.body_bold("4.1 Rate Card", "")
    d.styled_table(
        ["Role", "Location", "Rate"],
        [
            ("CTO Advisory", "United States", "$250.00/hr"),
            ("Tech Support", "United States", "$150.00/hr"),
            ("Tech Support", "Offshore (India)", "$45.00/hr"),
        ],
    )

    d.spacer()
    d.body_bold("4.2 One-Time Migration Labor Summary", "")
    d.styled_table(
        ["Phase", "Description", "CTO ($250)", "US Tech ($150)", "Offshore ($45)", "Total Hrs", "Cost"],
        [
            ("1", "Discovery & Assessment", "2", "2", "\u2014", "4", "$800"),
            ("2", "Cloud Environment Provisioning", "\u2014", "6", "2", "8", "$990"),
            ("3", "Server Migration (P2V)", "\u2014", "12", "4", "16", "$1,980"),
            ("4", "OneDrive / SharePoint Migration", "2", "10", "10", "22", "$2,450"),
            ("5", "Cloudbrink ZTNA Deployment", "2", "7", "3", "12", "$1,685"),
            ("6", "Security Agent Deployment", "\u2014", "2", "2", "4", "$390"),
            ("7", "Testing & Go-Live", "2", "5", "3", "10", "$1,385"),
            ("8", "Post-Migration Support", "\u2014", "2", "6", "8", "$570"),
            ("TOTAL", "", "8", "46", "30", "84", "$10,250"),
        ],
        total_indices=[8],
    )
    d.body("Pricing Type: Estimate Cost \u2014 Actual time billed at the applicable rate. If hours are projected to exceed the estimate by more than 10%, Technijian will notify Client before proceeding.")

    d.spacer()
    d.body_bold("4.3 Ongoing Monthly Services \u2014 Technijian Datacenter", "")
    d.styled_table(
        ["Service", "Code", "Qty", "Unit Price", "Monthly"],
        [
            ("VM1 \u2014 vCores (Primary Server)", "CL-VC", "4", "$6.25", "$25.00"),
            ("VM1 \u2014 Memory (GB)", "CL-GB", "16", "$0.63", "$10.08"),
            ("VM1 \u2014 Shared Bandwidth", "CL-SBW", "1", "$15.00", "$15.00"),
            ("VM2 \u2014 vCores (Cloudbrink)", "CL-VC", "2", "$6.25", "$12.50"),
            ("VM2 \u2014 Memory (GB)", "CL-GB", "4", "$0.63", "$2.52"),
            ("VM2 \u2014 Shared Bandwidth", "CL-SBW", "1", "$15.00", "$15.00"),
            ("Production Storage", "TB-PSTR", "1 TB", "$200.00", "$200.00"),
            ("Backup Storage", "TB-BSTR", "1 TB", "$50.00", "$50.00"),
            ("Infrastructure Subtotal", "", "", "", "$330.10"),
            ("Windows Server Std \u2014 2-Core Pack", "MS-STD", "4", "$5.25", "$21.00"),
            ("Licensing Subtotal", "", "", "", "$21.00"),
            ("CrowdStrike \u2014 Server", "AVS", "2", "$10.50", "$21.00"),
            ("Huntress \u2014 Server", "AVHS", "2", "$6.00", "$12.00"),
            ("Patch Management", "PMW", "2", "$4.00", "$8.00"),
            ("My Remote", "MR", "2", "$2.00", "$4.00"),
            ("Security Subtotal", "", "", "", "$45.00"),
            ("Image Backup (Veeam)", "IB", "2", "$15.00", "$30.00"),
            ("Backup Subtotal", "", "", "", "$30.00"),
            ("TECHNIJIAN TOTAL MONTHLY", "", "", "", "$426.10"),
            ("TECHNIJIAN TOTAL ANNUAL", "", "", "", "$5,113.20"),
        ],
        total_indices=[8, 10, 15, 17, 18, 19],
    )

    d.spacer()
    d.body_bold("4.4 Azure Equivalent Monthly Cost (For Comparison)", "")
    d.styled_table(
        ["Service", "Azure SKU", "Monthly"],
        [
            ("VM1 \u2014 Primary Server", "D4s v5 (4 vCPU, 16 GB) Windows", "$281.00"),
            ("VM2 \u2014 Cloudbrink Appliance", "B2s (2 vCPU, 4 GB) Linux", "$31.00"),
            ("Managed Disk", "1 TB Premium SSD (P30)", "$123.00"),
            ("Azure Backup", "Recovery Services Vault (1 TB)", "$55.00"),
            ("Outbound Data Transfer", "~500 GB egress / month", "$44.00"),
            ("Static Public IP", "1 Standard Static IP", "$4.00"),
            ("Windows Server License", "Included in VM pricing", "$0.00"),
            ("CrowdStrike \u2014 Server", "2 servers", "$21.00"),
            ("Huntress \u2014 Server", "2 servers", "$12.00"),
            ("Patch Management", "2 servers", "$8.00"),
            ("My Remote", "2 servers", "$4.00"),
            ("Azure Monitor", "Basic monitoring & alerts", "$15.00"),
            ("AZURE TOTAL MONTHLY", "", "$598.00"),
            ("AZURE TOTAL ANNUAL", "", "$7,176.00"),
        ],
        total_indices=[12, 13],
    )

    d.spacer()
    d.body_bold("4.5 Cost Comparison Summary", "")
    d.styled_table(
        ["", "Technijian Datacenter", "Microsoft Azure", "Savings"],
        [
            ("Monthly Cost", "$426.10", "$598.00", "$171.90 (28.7%)"),
            ("Annual Cost", "$5,113.20", "$7,176.00", "$2,062.80 (28.7%)"),
            ("3-Year Cost", "$15,339.60", "$21,528.00", "$6,188.40 (28.7%)"),
        ],
    )

    d.spacer()
    d.body_bold("Technijian Datacenter Advantages:", "")
    d.bullet("28.7% lower monthly cost compared to Azure pay-as-you-go pricing")
    d.bullet("$2,062.80 annual savings on hosting alone")
    d.bullet("No egress / bandwidth charges \u2014 shared bandwidth included")
    d.bullet("No surprise Azure consumption charges \u2014 predictable, fixed monthly billing")
    d.bullet("Local Technijian support team manages infrastructure end-to-end")
    d.bullet("Single vendor for hosting, security, backup, and support")
    d.bullet("Included Veeam backup \u2014 enterprise-grade with faster recovery")
    d.bullet("No Azure expertise required from Client")

    d.spacer()
    d.body_bold("4.6 Payment Schedule", "")
    d.body("All invoices are due and payable within thirty (30) days of the invoice date.")
    d.styled_table(
        ["Milestone", "Invoiced", "Amount"],
        [
            ("Project kickoff (before Phase 1)", "Upon SOW execution", "$5,125.00 (50%)"),
            ("Go-live completion (after Phase 7)", "Upon go-live", "$5,125.00 (50%)"),
            ("Monthly managed services", "1st of each month", "$426.10/month"),
            ("ONE-TIME PROJECT TOTAL", "", "$10,250.00"),
        ],
        total_indices=[3],
    )

    d.page_break()

    # 5. CLIENT RESPONSIBILITIES
    d.section_header("5. CLIENT RESPONSIBILITIES")
    d.body("Client shall:")
    d.numbered("(a)", "Provide administrative credentials and physical/remote access to the existing server;")
    d.numbered("(b)", "Designate Ed Pits (or alternate) as the authorized point of contact;")
    d.numbered("(c)", "Review and approve deliverables within five (5) business days of submission;")
    d.numbered("(d)", "Ensure all critical data is backed up prior to migration start;")
    d.numbered("(e)", "Inform users of planned service changes, maintenance windows, and downtime;")
    d.numbered("(f)", "Procure and maintain Microsoft 365 licenses with adequate OneDrive/SharePoint storage (~779 GB); and")
    d.numbered("(g)", "Coordinate with Cloudbrink for per-user ZTNA subscription licensing.")

    # 6. CHANGE MANAGEMENT
    d.spacer()
    d.section_header("6. CHANGE MANAGEMENT")
    d.numbered("6.01.", "Any changes to the scope, schedule, or pricing of this SOW must be documented in a written Change Order signed by both Parties.")
    d.numbered("6.02.", "If Client requests out-of-scope work, Technijian shall provide a Change Order detailing the additional work, estimated hours, and cost impact.")
    d.numbered("6.03.", "Technijian shall not proceed with out-of-scope work without an approved Change Order, except where delay would harm Client\u2019s systems.")

    # 7. ACCEPTANCE
    d.spacer()
    d.section_header("7. ACCEPTANCE")
    d.numbered("7.01.", "Upon completion of each phase, Technijian shall notify Client that deliverables are ready for review.")
    d.numbered("7.02.", "Client shall review and provide written acceptance or a description of deficiencies within five (5) business days.")
    d.numbered("7.03.", "If Client does not respond within the review period, deliverables shall be deemed accepted.")

    # 8. GOVERNING TERMS
    d.spacer()
    d.section_header("8. GOVERNING TERMS")
    d.body("The terms of the Master Service Agreement (MSA-OKL-2026) shall govern this SOW. In the event of a conflict between this SOW and the MSA, the MSA shall prevail unless this SOW expressly states otherwise.")

    d.signatures()
    d.footer_bar()
    d.save("SOW-001-ServerMigration.docx")


if __name__ == "__main__":
    build_msa()
    build_sow()
    print("\nAll documents generated.")
