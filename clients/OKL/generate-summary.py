"""
OKL Executive Summary Generator
Uses Technijian Brand Guide 2026 formatting (tech-branding repo)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors (from Technijian_Brand_Guide_2026) ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex versions for shading
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"

FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = BRAND_GREY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# ── Helper Functions ──

def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def accent_bar(color_hex, height_pt=4):
    """Full-width colored accent bar using a 1-row table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pf = p._element.get_or_add_pPr()
    spacing = pf.makeelement(qn("w:spacing"), {
        qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
    })
    pf.append(spacing)
    # Remove table borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def run_styled(paragraph, text, size=11, bold=False, color=None, italic=False, font=FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def section_header(title):
    """Section header with blue left-bar accent (matches build-proposal.js pattern)."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Blue accent bar cell (narrow)
    bar_cell = tbl.rows[0].cells[0]
    set_cell_shading(bar_cell, HEX_BLUE)
    bar_cell.width = Inches(0.08)
    bar_cell.text = ""
    bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Title cell
    title_cell = tbl.rows[0].cells[1]
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run_styled(p, title, size=14, bold=True, color=CORE_BLUE)

    # Remove table borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def styled_table(headers, rows, col_widths=None, total_row_indices=None):
    """Create a branded table with light grey borders and blue header row."""
    total_row_indices = total_row_indices or []
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        run_styled(p, h, size=9, bold=True, color=WHITE)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = tbl.add_row()
        is_total = row_idx in total_row_indices
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if is_total:
                set_cell_shading(cell, HEX_OFF_WHITE)
                run_styled(p, str(val), size=9, bold=True, color=DARK_CHARCOAL)
            else:
                run_styled(p, str(val), size=9, color=BRAND_GREY)

    # Apply light grey borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return tbl


def bullet_item(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_styled(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)
    return p


def body_text(text, space_after=8):
    p = doc.add_paragraph()
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

# Top blue accent bar
accent_bar(HEX_BLUE, height_pt=6)

# Spacer
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Centered logo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO_PATH):
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(3.0))

# Orange divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
# Create a centered orange line using a table
div_tbl = doc.add_table(rows=1, cols=3)
div_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in div_tbl.rows[0].cells:
    cell.text = ""
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
mid_cell = div_tbl.rows[0].cells[1]
mid_cell.width = Inches(2.5)
set_cell_shading(mid_cell, HEX_ORANGE)
div_tbl.rows[0].cells[0].width = Inches(2.25)
div_tbl.rows[0].cells[2].width = Inches(2.25)
# Remove borders
div_pr = div_tbl._element.tblPr
div_borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    "</w:tblBorders>"
)
div_pr.append(div_borders)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(8)
run_styled(p, "Executive Summary", size=26, bold=True, color=DARK_CHARCOAL)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run_styled(p, "Server Migration & Managed Cloud Hosting", size=14, color=BRAND_GREY)

# Prepared for
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run_styled(p, "Prepared for ", size=12, color=BRAND_GREY)
run_styled(p, "Oaktree Law", size=12, bold=True, color=CORE_BLUE)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run_styled(p, "March 2026", size=11, color=BRAND_GREY)

# Spacer
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Bottom orange accent bar
accent_bar(HEX_ORANGE, height_pt=6)

# Confidential
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)

# Contact
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(
    p,
    "Technijian  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)


# ════════════════════════════════════════════════════════════════
# PAGE 2 - OVERVIEW + TIMELINE + ONE-TIME COST
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Overview")
body_text(
    "Technijian proposes migrating Oaktree Law\u2019s aging physical server to a fully managed "
    "cloud environment hosted in Technijian\u2019s private datacenter. The project includes server "
    "virtualization, migration of 779 GB of shared data to Microsoft OneDrive/SharePoint, deployment "
    "of Cloudbrink Zero Trust Network Access (ZTNA) for secure remote connectivity, and installation "
    "of enterprise security agents across all systems."
)
body_text(
    "Upon completion, Oaktree Law will have a modern, secure, and fully managed IT infrastructure with "
    "predictable monthly costs, enterprise-grade backup, and 24/7 monitoring \u2014 eliminating the risk "
    "and maintenance burden of on-premises hardware."
)

# What's Included
section_header("What\u2019s Included")
items = [
    ("Cloud-Hosted Servers: ", "2 virtual machines in Technijian\u2019s private datacenter replacing your physical server"),
    ("Data Migration: ", "779 GB of shared folders migrated to OneDrive/SharePoint with full permission mapping"),
    ("Zero Trust Access: ", "Cloudbrink ZTNA for secure access from anywhere \u2014 no VPN needed"),
    ("Enterprise Security: ", "CrowdStrike endpoint protection, Huntress threat hunting, automated patch management"),
    ("Managed Backup: ", "Veeam image-level backup of both VMs with daily snapshots and rapid recovery"),
    ("30-Day Post-Migration Support: ", "Dedicated support for issue resolution and user assistance"),
]
for bold_text, desc in items:
    bullet_item(desc, bold_prefix=bold_text)

# Project Timeline
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Project Timeline")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Phase", "Description", "Timeline"],
    [
        ("1", "Discovery & Assessment", "Week 1"),
        ("2", "Cloud Environment Provisioning", "Weeks 1\u20132"),
        ("3", "Physical-to-Virtual Server Migration", "Weeks 2\u20133"),
        ("4", "OneDrive / SharePoint Data Migration", "Weeks 3\u20135"),
        ("5", "Cloudbrink ZTNA Deployment", "Weeks 5\u20136"),
        ("6", "Security Agent Deployment", "Week 6"),
        ("7", "Testing, Validation & Go-Live", "Weeks 6\u20137"),
        ("8", "Post-Migration Support (30 Days)", "Weeks 7\u201311"),
    ],
    col_widths=[0.6, 4.0, 1.4],
)

# One-Time Migration Cost
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("One-Time Migration Cost")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Role", "Rate", "Hours", "Cost"],
    [
        ("CTO Advisory (US)", "$250/hr", "8", "$2,000"),
        ("Tech Support (US)", "$150/hr", "46", "$6,900"),
        ("Tech Support (Offshore)", "$45/hr", "30", "$1,350"),
        ("TOTAL", "", "84 hours", "$10,250"),
    ],
    col_widths=[2.5, 1.2, 1.2, 1.1],
    total_row_indices=[3],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run_styled(p, "Payment Schedule: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(p, "50% at project kickoff ($5,125)  +  50% at go-live ($5,125)", size=10, color=BRAND_GREY)

p = doc.add_paragraph()
run_styled(p, "Pricing Type: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "Estimate \u2014 actual time billed at applicable rate. If hours exceed estimate by >10%, "
    "Technijian will notify before proceeding.",
    size=10, color=BRAND_GREY,
)


# ════════════════════════════════════════════════════════════════
# PAGE 3 - MONTHLY SERVICES + COMPARISON + SUMMARY
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Ongoing Monthly Services")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Services Included", "Monthly"],
    [
        ("Cloud Infrastructure", "2 VMs (vCores, Memory, Shared BW) + 1 TB Prod + 1 TB Backup Storage", "$330.10"),
        ("Microsoft Licensing", "Windows Server Std \u2014 4x 2-Core Packs", "$21.00"),
        ("Security & Monitoring", "CrowdStrike (2), Huntress (2), Patch Mgmt (2), My Remote (2)", "$45.00"),
        ("Backup & Recovery", "Veeam Image Backup (2 VMs)", "$30.00"),
        ("TOTAL MONTHLY", "", "$426.10"),
        ("TOTAL ANNUAL", "", "$5,113.20"),
    ],
    col_widths=[1.8, 3.2, 1.0],
    total_row_indices=[4, 5],
)

# Cost Comparison
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Cost Comparison: Technijian vs. Azure")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Period", "Technijian", "Azure", "Savings"],
    [
        ("Monthly", "$426.10", "$598.00", "$171.90 (28.7%)"),
        ("Annual", "$5,113.20", "$7,176.00", "$2,062.80 (28.7%)"),
        ("3-Year", "$15,339.60", "$21,528.00", "$6,188.40 (28.7%)"),
    ],
    col_widths=[1.2, 1.5, 1.5, 1.8],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run_styled(p, "Technijian Datacenter Advantages:", size=10, bold=True, color=DARK_CHARCOAL)

advantages = [
    "Predictable, fixed monthly billing \u2014 no surprise Azure consumption charges",
    "No egress or bandwidth charges \u2014 shared bandwidth included",
    "Single vendor for hosting, security, backup, and support",
    "Enterprise-grade Veeam backup with faster recovery than Azure Backup",
    "Local Technijian support team manages infrastructure end-to-end",
]
for adv in advantages:
    bullet_item(adv)

# Total Investment Summary
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Total Investment Summary")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Item", "Amount"],
    [
        ("One-Time Migration (84 hours)", "$10,250.00"),
        ("Monthly Managed Services", "$426.10/month"),
        ("Year 1 Total (migration + 12 months hosting)", "$15,363.20"),
        ("Year 2+ Annual Cost (hosting only)", "$5,113.20/year"),
    ],
    col_widths=[4.0, 2.0],
    total_row_indices=[2],
)

# Agreement Structure
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Agreement Structure")
agreements = [
    ("Master Service Agreement (MSA): ", "Governs the overall relationship, payment terms, confidentiality, liability, and dispute resolution for a 12-month initial term with automatic annual renewal."),
    ("SOW-001 \u2014 Server Migration: ", "Defines the 8-phase migration project scope, deliverables, timeline, and one-time pricing. Governed by the MSA."),
    ("Schedule A \u2014 Monthly Services: ", "Details the ongoing managed hosting, security, and monitoring services with SLA commitments."),
    ("Schedule C \u2014 Rate Card: ", "Establishes hourly rates for any additional work or change orders."),
]
for bold_text, desc in agreements:
    bullet_item(desc, bold_prefix=bold_text)

# Next Steps
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Next Steps")
steps = [
    "Review this summary and the attached agreements",
    "Confirm client address and any questions with Ed Pits",
    "Sign the MSA, Schedule A, Schedule C, and SOW-001",
    "Technijian begins Phase 1 Discovery within 5 business days of execution",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run_styled(p, f"{i}. ", size=10, bold=True, color=CORE_BLUE)
    run_styled(p, step, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)

# Footer divider
doc.add_paragraph().paragraph_format.space_after = Pt(8)
accent_bar(HEX_BLUE, height_pt=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run_styled(
    p,
    "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(
    p,
    "CONFIDENTIAL \u2014 Prepared exclusively for Oaktree Law",
    size=8, italic=True, color=BRAND_GREY,
)

doc.save("OKL-Executive-Summary.docx")
print("Created OKL-Executive-Summary.docx")
