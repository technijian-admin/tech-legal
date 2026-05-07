#!/usr/bin/env python3
"""Convert the four TALY MSA Markdown sources to branded DOCX for Robert's review.

Outputs alongside the .md files:
  MSA-TALY-2026.docx
  Schedule-A-TALY.docx
  Schedule-B-TALY.docx
  Schedule-C-TALY.docx

Branding: Technijian Brand Guide 2026 (Open Sans, blue/orange/charcoal).
Not a from-scratch styled MSA — simple, readable, branded conversion of the
authoritative .md sources for client review.
"""
from __future__ import annotations

import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent

# Brand colors (hex w/o #)
BLUE = "006DB6"
ORANGE = "F67D4B"
CHARCOAL = "1A1A2E"
GREY = "59595B"
OFF_WHITE = "F8F9FA"
WHITE = "FFFFFF"
LIGHT_GREY = "DEE2E6"
FONT = "Open Sans"

LOGO_URL = "https://technijian.com/wp-content/uploads/2026/03/technijian-logo-full-color-600x125-1.png"
LOGO_PATH = HERE / "_technijian-logo.png"


def _hex_to_rgb(hexstr: str) -> RGBColor:
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


def _ensure_logo() -> Path | None:
    if LOGO_PATH.exists() and LOGO_PATH.stat().st_size > 1000:
        return LOGO_PATH
    try:
        urllib.request.urlretrieve(LOGO_URL, LOGO_PATH)
        return LOGO_PATH
    except Exception as e:
        print(f"  warning: could not fetch logo ({e})")
        return None


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = LIGHT_GREY, sz: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _add_run(paragraph, text: str, *, size: int = 10, bold: bool = False,
             italic: bool = False, color: str = GREY, font: str = FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = _hex_to_rgb(color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    return run


# ── Inline Markdown parser: bold + italic across a single line ──
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|_.+?_)")


def _add_inline_runs(paragraph, text: str, *, size: int = 10, color: str = GREY):
    """Add a paragraph with inline **bold** / *italic* support."""
    text = text.replace("**:", "**:").replace("&mdash;", "—").replace("&ndash;", "–")
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) >= 4:
            _add_run(paragraph, piece[2:-2], size=size, bold=True, color=CHARCOAL)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) >= 2:
            _add_run(paragraph, piece[1:-1], size=size, italic=True, color=color)
        elif piece.startswith("_") and piece.endswith("_") and len(piece) >= 2:
            _add_run(paragraph, piece[1:-1], size=size, italic=True, color=color)
        else:
            _add_run(paragraph, piece, size=size, color=color)


def _new_paragraph(doc, *, alignment=None, space_after: int = 4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.2
    if alignment is not None:
        p.alignment = alignment
    return p


def _add_heading(doc, text: str, level: int) -> None:
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    colors = {1: BLUE, 2: CHARCOAL, 3: CHARCOAL, 4: CHARCOAL}
    p = _new_paragraph(doc, space_after=6)
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    _add_run(p, text, size=sizes.get(level, 11), bold=True, color=colors.get(level, CHARCOAL))


def _is_table_separator(row: str) -> bool:
    cells = [c.strip() for c in row.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-+:?", c or "") for c in cells)


def _split_table_row(row: str) -> list[str]:
    cells = row.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.autofit = True
    # Header row
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, BLUE)
        _set_cell_borders(cell, color=BLUE, sz="4")
        cell.text = ""
        para = cell.paragraphs[0]
        _add_run(para, txt, size=10, bold=True, color=WHITE)
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = OFF_WHITE if r_idx % 2 == 0 else WHITE
        for c_idx, txt in enumerate(row):
            if c_idx >= n_cols:
                break
            cell = cells[c_idx]
            _set_cell_shading(cell, bg)
            _set_cell_borders(cell, color=LIGHT_GREY, sz="4")
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline_runs(para, txt, size=10)


def _add_blockquote(doc, lines: list[str]) -> None:
    p = _new_paragraph(doc, space_after=6)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3)
    pf.space_before = Pt(4)
    text = " ".join(line.lstrip("> ").strip() for line in lines)
    _add_inline_runs(p, text, size=10, color=CHARCOAL)


def _cover_block(doc: Document, title: str, subtitle: str, agreement_line: str) -> None:
    logo = _ensure_logo()
    if logo:
        p = _new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        run = p.add_run()
        run.add_picture(str(logo), width=Inches(2.6))
    p = _new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_run(p, title, size=22, bold=True, color=BLUE)
    if subtitle:
        p = _new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        _add_run(p, subtitle, size=14, bold=True, color=CHARCOAL)
    if agreement_line:
        p = _new_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        _add_run(p, agreement_line, size=11, italic=True, color=GREY)
    # Orange accent rule
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, ORANGE)
    cell.text = ""
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    # 4pt bar
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), "9000")
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    _new_paragraph(doc, space_after=6)


def render_md_to_docx(md_path: Path, docx_path: Path,
                      cover_title: str, cover_subtitle: str,
                      cover_agreement: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    base = doc.styles["Normal"]
    base.font.name = FONT
    base.font.size = Pt(10)
    base.font.color.rgb = _hex_to_rgb(GREY)

    _cover_block(doc, cover_title, cover_subtitle, cover_agreement)

    text = md_path.read_text(encoding="utf-8")
    # Strip YAML front-matter if present
    if text.startswith("---\n"):
        end = text.find("\n---", 4)
        if end != -1:
            text = text[end + 4:].lstrip("\n")

    # Skip top-level # title (we already rendered it on the cover)
    lines = text.splitlines()
    skipped_top_h1 = False

    i = 0
    block_quote: list[str] = []
    while i < len(lines):
        line = lines[i]

        # Blockquote consume
        if line.startswith(">"):
            block_quote.append(line)
            i += 1
            while i < len(lines) and lines[i].startswith(">"):
                block_quote.append(lines[i])
                i += 1
            _add_blockquote(doc, block_quote)
            block_quote = []
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line.strip()):
            p = _new_paragraph(doc, space_after=4)
            _add_run(p, "", size=8)
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            _set_cell_shading(cell, BLUE)
            cell.text = ""
            _new_paragraph(doc, space_after=6)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            depth = len(m.group(1))
            heading_text = m.group(2).strip()
            if depth == 1 and not skipped_top_h1:
                skipped_top_h1 = True
                i += 1
                continue
            _add_heading(doc, heading_text, max(1, depth - 1) if depth > 1 else 1)
            i += 1
            continue

        # Tables
        if line.lstrip().startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _split_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, header, rows)
            _new_paragraph(doc, space_after=4)
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                txt = re.sub(r"^\s*[-*]\s+", "", lines[i])
                p = _new_paragraph(doc, space_after=2)
                p.paragraph_format.left_indent = Inches(0.25)
                _add_run(p, "•  ", bold=True, color=BLUE, size=10)
                _add_inline_runs(p, txt, size=10)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                m_num = re.match(r"^\s*(\d+)\.\s+(.*)$", lines[i])
                num = m_num.group(1)
                txt = m_num.group(2)
                p = _new_paragraph(doc, space_after=2)
                p.paragraph_format.left_indent = Inches(0.25)
                _add_run(p, f"{num}.  ", bold=True, color=BLUE, size=10)
                _add_inline_runs(p, txt, size=10)
                i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Plain paragraph (consume continuation lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].lstrip().startswith("|")
            or lines[i].lstrip().startswith("#")
            or lines[i].startswith(">")
            or re.match(r"^\s*[-*]\s+", lines[i])
            or re.match(r"^\s*\d+\.\s+", lines[i])
            or re.fullmatch(r"-{3,}|\*{3,}|_{3,}", lines[i].strip())
        ):
            para_lines.append(lines[i])
            i += 1
        para = _new_paragraph(doc, space_after=4)
        _add_inline_runs(para, " ".join(s.strip() for s in para_lines), size=10)

    doc.save(docx_path)


def main() -> None:
    docs = [
        ("MSA-TALY-2026.md",   "MSA-TALY-2026.docx",
         "MASTER SERVICE AGREEMENT", "Talley & Associates",
         "Agreement Number: MSA-TALY-2026"),
        ("Schedule-A-TALY.md", "Schedule-A-TALY.docx",
         "SCHEDULE A", "Monthly Managed Services",
         "Attached to Master Service Agreement MSA-TALY-2026"),
        ("Schedule-B-TALY.md", "Schedule-B-TALY.docx",
         "SCHEDULE B", "Subscription and License Services",
         "Attached to Master Service Agreement MSA-TALY-2026"),
        ("Schedule-C-TALY.md", "Schedule-C-TALY.docx",
         "SCHEDULE C", "Rate Card",
         "Attached to Master Service Agreement MSA-TALY-2026"),
    ]
    for src, dst, title, subtitle, agreement in docs:
        src_path = HERE / src
        dst_path = HERE / dst
        print(f"Generating {dst_path.name}...")
        render_md_to_docx(src_path, dst_path, title, subtitle, agreement)
        print(f"  wrote {dst_path.stat().st_size:,} bytes")
    print("Done.")


if __name__ == "__main__":
    main()
