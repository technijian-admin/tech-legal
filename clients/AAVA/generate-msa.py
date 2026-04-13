"""
AAVA MSA Generator
Generates Master Service Agreement for Aventine at Aliso Viejo Apartments
This is a renewal MSA replacing an older Monthly Service Agreement.

Key changes vs. prior agreement:
  1. Site Assessment ($50/mo) removed
  2. USA Tech Support moved from contracted to ad-hoc hourly (per Rate Card)
  3. India support hours held at prior contracted levels (4.55 + 1.77 hrs)
     Unpaid balance is disclosed but NOT amortized — it will continue to grow
     at approximately $98.70/mo if usage stays at current actuals.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Brand Colors (Technijian Brand Guide 2026) --
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "technijian-logo.png")


def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


def set_grey_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        return r

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_borders(tbl)

    def cover_page(self, title, subtitle, client_name="AmPac Business Capital", date="April 10, 2026"):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(LOGO_PATH, width=Inches(3.0))
        div = self.doc.add_table(rows=1, cols=3)
        div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in div.rows[0].cells:
            c.text = ""
            c.paragraphs[0].paragraph_format.space_before = Pt(0)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div.rows[0].cells[1], HEX_ORANGE)
        div.rows[0].cells[0].width = Inches(2.25)
        div.rows[0].cells[1].width = Inches(2.5)
        div.rows[0].cells[2].width = Inches(2.25)
        remove_borders(div)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Prepared for ", size=12, color=BRAND_GREY)
        self.run(p, client_name, size=12, bold=True, color=CORE_BLUE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.accent_bar(HEX_ORANGE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar = tbl.rows[0].cells[0]
        set_cell_shading(bar, HEX_BLUE)
        bar.width = Inches(0.08)
        bar.text = ""
        bar.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar.paragraphs[0].paragraph_format.space_after = Pt(0)
        tc = tbl.rows[0].cells[1]
        tc.text = ""
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_borders(tbl)

    def part_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_borders(tbl)

    def styled_table(self, headers, rows, total_indices=None):
        total_indices = total_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            set_cell_shading(c, HEX_BLUE)
            c.text = ""
            self.run(c.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for ri, rd in enumerate(rows):
            row = tbl.add_row()
            is_total = ri in total_indices
            for i, val in enumerate(rd):
                c = row.cells[i]
                c.text = ""
                if is_total:
                    set_cell_shading(c, HEX_OFF_WHITE)
                    self.run(c.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(c.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        set_grey_borders(tbl)
        return tbl

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            self.run(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, num, text, indent=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent
        self.run(p, f"{prefix}{num} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self, client_name="AMPAC BUSINESS CAPITAL", contact_name="[AUTHORIZED SIGNER]"):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)
        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", "Name: _________________________________",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)
        self.spacer(12)
        p = self.doc.add_paragraph()
        self.run(p, client_name, size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", f"Name: {contact_name}",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        self.doc.save(filename)
        print(f"Created {filename}")


# ================================================================
#  MSA
# ================================================================
def build_msa():
    d = BrandedDoc()
    d.cover_page(
        "Master Service Agreement",
        "Aventine at Aliso Viejo Apartments",
        client_name="Aventine at Aliso Viejo Apartments",
        date="May 1, 2026",
    )
    d.page_break()

    d.part_header("MASTER SERVICE AGREEMENT")
    d.spacer()
    d.body("Agreement Number: MSA-AAVA-2026")
    d.body("Effective Date: May 1, 2026")
    d.spacer()
    d.body("This Master Service Agreement (\u201cAgreement\u201d) is entered into by and between:")
    d.spacer()
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", "")
    d.body("18 Technology Drive, Suite 141")
    d.body("Irvine, California 92618")
    d.spacer()
    d.body("and")
    d.spacer()
    d.body_bold("Aventine at Aliso Viejo Apartments (\u201cAAVA\u201d or \u201cClient\u201d)", "")
    d.body("22501 Chase")
    d.body("Aliso Viejo, California 92656")
    d.body("Primary Contact: [AUTHORIZED SIGNER]")
    d.spacer()
    d.body("(collectively, the \u201cParties\u201d)")

    # SECTION 1
    d.spacer()
    d.section_header("SECTION 1 \u2014 SCOPE OF SERVICES")

    d.numbered("1.01.", "Services. Technijian shall provide the services described in the Schedules attached to this Agreement, which are incorporated herein by reference:")
    d.bullet("Schedule A \u2014 Monthly Managed Services (Endpoint Security and Virtual Staff Support)")
    d.bullet("Schedule B \u2014 Subscription and License Services (none in effect as of Effective Date)")
    d.bullet("Schedule C \u2014 Rate Card")
    d.body("Additional services may be provided through Statements of Work (\u201cSOWs\u201d) executed under this Agreement.")

    d.numbered("1.02.", "Standard of Care. Technijian shall perform all services in a professional and workmanlike manner, consistent with industry standards for managed IT service providers.")

    d.numbered("1.03.", "Service Level Agreement. The service levels applicable to the services are set forth in Schedule A. Technijian shall use commercially reasonable efforts to meet the service levels described therein.")

    d.numbered("1.04.", "Client Responsibilities. Client shall:")
    d.numbered("(a)", "Provide Technijian with reasonable access to Client\u2019s systems, facilities, and personnel as necessary for Technijian to perform the services;", indent=1)
    d.numbered("(b)", "Designate a primary point of contact for communications with Technijian;", indent=1)
    d.numbered("(c)", "Maintain current and accurate information regarding Client\u2019s systems and infrastructure;", indent=1)
    d.numbered("(d)", "Comply with all applicable laws and regulations in connection with its use of the services; and", indent=1)
    d.numbered("(e)", "Be solely responsible for the security and management of Client\u2019s account credentials and passwords.", indent=1)

    d.numbered("1.05.", "Independent Contractor. Technijian is an independent contractor. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the Parties.")

    # SECTION 2
    d.spacer()
    d.section_header("SECTION 2 \u2014 TERM AND RENEWAL")

    d.numbered("2.01.", "Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the \u201cInitial Term\u201d).")
    d.numbered("2.02.", "Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a \u201cRenewal Term\u201d), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term. Technijian shall send Client a written renewal reminder at least thirty (30) days prior to each renewal date, which shall restate the auto-renewal terms and cancellation method.")
    d.numbered("2.03.", "Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party.")
    d.numbered("2.04.", "Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party:")
    d.numbered("(a)", "Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days after receiving written notice of the breach; or", indent=1)
    d.numbered("(b)", "Becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.", indent=1)

    d.numbered("2.05.", "Effect of Termination.")
    d.numbered("(a)", "Upon termination, Client shall pay all fees and charges for services rendered through the date of termination, including any remaining obligations for annual licenses and subscriptions procured on Client\u2019s behalf, and any unamortized portion of the prior unpaid balance described in the Unpaid Balance Acknowledgment of Schedule A.", indent=1)
    d.numbered("(b)", "Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, subject to payment of applicable fees.", indent=1)
    d.numbered("(c)", "Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format, provided Client is not in breach of this Agreement.", indent=1)
    d.numbered("(d)", "The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 9.09 (Personnel Transition Fee), and Section 10 (Data Protection).", indent=1)

    d.page_break()

    # SECTION 3
    d.section_header("SECTION 3 \u2014 PAYMENT")

    d.numbered("3.01.", "Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.")
    d.numbered("3.02.", "Invoice Types. Client may receive the following types of invoices from Technijian during the term of this Agreement. Each invoice will clearly identify its type, the applicable Schedule or SOW, and the billing period or delivery event.")
    d.numbered("(a)", "Monthly Service Invoice. Issued on the first business day of each month for recurring managed services under Schedule A (endpoint security and contracted virtual staff support). Billed in advance for the upcoming month.", indent=1)
    d.numbered("(b)", "Weekly In-Contract Invoice. Issued every Friday for Virtual Staff (contracted support) services performed under Schedule A during the preceding week (Monday through Friday). Each invoice includes: (i) a listing of each support ticket addressed; (ii) the assigned resource, role, and hours spent per ticket; (iii) a description of the work performed per ticket; (iv) whether work was performed during normal or after-hours; and (v) the current running balance for each contracted role.", indent=1)
    d.numbered("(c)", "Weekly Out-of-Contract Invoice. Issued every Friday for labor services performed outside the scope of any active Schedule or SOW \u2014 including ad-hoc support requests (such as USA Tech Support, which under this Agreement is billed hourly per Schedule C), emergency work, and services performed under a SOW with hourly billing. Each invoice includes: (i) a listing of each support ticket or task performed; (ii) the assigned resource, role, and applicable hourly rate from the Rate Card (Schedule C); (iii) time entries with hours billed per activity (in 15-minute increments); (iv) whether work was performed during normal or after-hours; and (v) the total hours and total amount for the week.", indent=1)
    d.numbered("(d)", "Equipment and Materials Invoice. Issued upon delivery or procurement of hardware, software licenses (perpetual), or other tangible goods on Client\u2019s behalf. Each invoice includes: (i) item description, manufacturer, and model/part number; (ii) quantity and unit price; (iii) applicable sales tax; (iv) shipping and handling charges, if any; and (v) total amount due. Title to equipment shall not pass to Client until payment is received in full.", indent=1)
    d.numbered("(e)", "Project Milestone Invoice. Issued upon completion of a project milestone as defined in an applicable SOW. The milestone, deliverables, and invoiced amount are as specified in the payment schedule of the SOW. Milestone invoices are billed in arrears upon acceptance of the deliverables or deemed acceptance under the SOW\u2019s acceptance provisions.", indent=1)
    d.numbered("3.03.", "Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified in the applicable Schedule or SOW.")
    d.numbered("3.04.", "Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated as simple interest from the date payment was due until the date payment is received in full. Late fees are payable in addition to the outstanding principal amount. The Parties acknowledge that the late fee represents a reasonable estimate of Technijian\u2019s administrative costs and damages resulting from late payment, including cash-flow disruption, collection overhead, and the cost of carrying accounts receivable, and is not intended as a penalty.")
    d.numbered("3.05.", "Disputed Invoices.")
    d.numbered("(a)", "Weekly Invoices (In-Contract and Out-of-Contract). Because weekly invoices include detailed ticket descriptions and time entries, Client shall have thirty (30) days from the invoice date to review and dispute any portion of a weekly invoice. Client shall notify Technijian in writing, specifying the ticket number(s) and nature of the dispute. Undisputed tickets and time entries on the same invoice shall remain payable by the due date. Failure to provide a timely written dispute notice within the thirty (30) day period shall constitute acceptance of all tickets and time entries on the invoice.", indent=1)
    d.numbered("(b)", "All Other Invoices. For monthly service invoices, equipment invoices, and project milestone invoices, Client shall notify Technijian in writing within fifteen (15) days of the invoice date if Client disputes any portion, specifying the nature and basis of the dispute. Client shall pay all undisputed amounts by the due date. Failure to provide a timely written dispute notice shall constitute acceptance of the invoice.", indent=1)
    d.numbered("(c)", "Resolution. The Parties shall work in good faith to resolve any invoice dispute within thirty (30) days of the dispute notice. If the dispute results in an adjustment, Technijian shall issue a credit memo or revised invoice within ten (10) business days of resolution.", indent=1)
    d.numbered("3.06.", "Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend services under the Schedule or SOW associated with the unpaid invoice until payment is received in full, including all accrued late fees. If Client fails to pay any undisputed invoice within sixty (60) days of the due date, Technijian may suspend all services under this Agreement and any related Schedules or SOWs. Recurring fees for suspended services shall continue to accrue for a period not to exceed thirty (30) days following the date of suspension, after which Technijian may terminate the affected Schedule, SOW, or this Agreement upon written notice. Suspension of services shall not relieve Client of its payment obligations.")
    d.numbered("3.07.", "Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian\u2019s income.")

    # SECTION 4
    d.spacer()
    d.section_header("SECTION 4 \u2014 CONFIDENTIALITY")

    d.numbered("4.01.", "Definition. \u201cConfidential Information\u201d means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information.")
    d.numbered("4.02.", "Obligations. Each Party shall:")
    d.numbered("(a)", "Hold the other Party\u2019s Confidential Information in confidence using at least the same degree of care it uses for its own confidential information, but not less than reasonable care;", indent=1)
    d.numbered("(b)", "Not disclose Confidential Information to third parties without prior written consent, except to employees, agents, and subcontractors who have a need to know and are bound by equivalent obligations; and", indent=1)
    d.numbered("(c)", "Not use Confidential Information for any purpose other than performing obligations under this Agreement.", indent=1)
    d.numbered("4.03.", "Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known to the receiving Party prior to disclosure, is independently developed, or is received from a third party without restriction.")
    d.numbered("4.04.", "Compelled Disclosure. If required by law or court order to disclose Confidential Information, the receiving Party shall provide prompt written notice to the disclosing Party and cooperate in seeking a protective order.")
    d.numbered("4.05.", "Duration. Confidentiality obligations shall survive termination for a period of three (3) years.")

    # SECTION 5
    d.spacer()
    d.section_header("SECTION 5 \u2014 LIMITATION OF LIABILITY")

    d.numbered("5.01.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, NEITHER PARTY\u2019S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM (THE \u201cSTANDARD CAP\u201d).")
    d.numbered("5.02.", "EXCEPT AS PROVIDED IN SECTION 5.03 BELOW, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL, REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE OR WHETHER EITHER PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.")
    d.numbered("5.03.", "ENHANCED CAP FOR CERTAIN CLAIMS. FOR CLAIMS ARISING FROM BREACHES OF SECTION 4 (CONFIDENTIALITY), SECTION 10 (DATA PROTECTION), INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, THE TOTAL AGGREGATE LIABILITY OF THE RESPONSIBLE PARTY SHALL NOT EXCEED THREE (3) TIMES THE STANDARD CAP DEFINED IN SECTION 5.01 (THE \u201cENHANCED CAP\u201d). THE ENHANCED CAP SHALL NOT APPLY TO LIABILITY ARISING FROM A PARTY\u2019S WILLFUL AND INTENTIONAL MISAPPROPRIATION OF THE OTHER PARTY\u2019S CONFIDENTIAL INFORMATION OR CLIENT DATA, FOR WHICH LIABILITY SHALL BE UNCAPPED.")
    d.numbered("5.04.", "Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data, Client acknowledges that Client is solely responsible for maintaining backup copies of its data.")

    d.page_break()

    # SECTION 6
    d.section_header("SECTION 6 \u2014 INDEMNIFICATION")

    d.numbered("6.01.", "By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from Technijian\u2019s gross negligence or willful misconduct in performing the services.")
    d.numbered("6.02.", "By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from Client\u2019s use of the services in violation of applicable law, Client\u2019s breach of this Agreement, or any data, content, or materials provided by Client.")
    d.numbered("6.03.", "Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party\u2019s prior written consent.")

    # SECTION 7
    d.spacer()
    d.section_header("SECTION 7 \u2014 INTELLECTUAL PROPERTY")

    d.numbered("7.01.", "Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes.")
    d.numbered("7.02.", "Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property.")
    d.numbered("7.03.", "Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW.")

    # SECTION 8
    d.spacer()
    d.section_header("SECTION 8 \u2014 DISPUTE RESOLUTION")

    d.numbered("8.01.", "Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations for a period of thirty (30) days.")
    d.numbered("8.02.", "Mediation. If not resolved, the Parties shall submit the dispute to mediation in Orange County, California, for a period not to exceed sixty (60) days.")
    d.numbered("8.03.", "Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association in Orange County, California.")
    d.numbered("8.04.", "Injunctive Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief to prevent irreparable harm.")

    # SECTION 9
    d.spacer()
    d.section_header("SECTION 9 \u2014 GENERAL PROVISIONS")

    d.numbered("9.01.", "Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties and supersedes the prior Client Monthly Service Agreement between the Parties, except with respect to amounts owed thereunder, which are addressed in the Unpaid Balance Acknowledgment in Schedule A.")
    d.numbered("9.02.", "Amendment. This Agreement may only be amended by a written instrument signed by both Parties.")
    d.numbered("9.03.", "Severability. If any provision is found invalid or unenforceable, the remaining provisions shall continue in full force.")
    d.numbered("9.04.", "Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party.")
    d.numbered("9.05.", "Assignment. Neither Party may assign this Agreement without prior written consent, except in connection with a merger, acquisition, or sale of substantially all assets.")
    d.numbered("9.06.", "Force Majeure.")
    d.numbered("(a)", "Neither Party shall be liable for delays or failures in performance caused by events beyond its reasonable control, including natural disasters, acts of government, labor disputes, pandemics, cyberattacks on critical infrastructure, or failures of major third-party infrastructure services that are not attributable to the affected Party\u2019s failure to implement reasonable redundancy (\u201cForce Majeure Event\u201d).", indent=1)
    d.numbered("(b)", "The affected Party shall notify the other Party in writing within five (5) business days of becoming aware of a Force Majeure Event and shall use commercially reasonable efforts to mitigate the impact and resume performance.", indent=1)
    d.numbered("(c)", "Payment obligations are not excused by a Force Majeure Event.", indent=1)
    d.numbered("(d)", "If a Force Majeure Event prevents performance of a material portion of the services for more than ninety (90) consecutive days, either Party may terminate the affected Schedule, SOW, or this Agreement upon fifteen (15) days written notice without liability for early termination fees, and Client shall pay only for services actually rendered through the date of termination.", indent=1)
    d.numbered("9.07.", "Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier.")
    d.numbered("9.08.", "Governing Law. This Agreement shall be governed by the laws of the State of California.")
    d.numbered("9.09.", "Personnel Transition Fee. The Parties acknowledge that each invests significant resources in recruiting, training, and retaining skilled personnel. If either Party hires (whether as an employee or independent contractor) any individual who was an employee of the other Party and who was directly involved in performing or receiving services under this Agreement, and such hiring occurs during the term of this Agreement or within twelve (12) months following termination, the hiring Party shall pay the other Party a personnel transition fee equal to 25% of the hired individual\u2019s first-year annual compensation (base salary or annualized contractor fees). This fee represents a reasonable estimate of the non-hiring Party\u2019s recruiting and training costs and is not intended as a restraint on trade or employment. This Section does not restrict any individual\u2019s right to seek or obtain employment, and shall not apply to individuals who: (a) respond to general public job postings or advertisements not specifically targeted at the other Party\u2019s employees; or (b) are referred by a third-party recruiting firm without the hiring Party\u2019s direction to target the other Party\u2019s employees.")
    d.numbered("9.10.", "Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.")

    # SECTION 10
    d.spacer()
    d.section_header("SECTION 10 \u2014 DATA PROTECTION")

    d.numbered("10.01.", "CCPA/CPRA Compliance. To the extent Technijian processes, stores, or has access to personal information (as defined under the California Consumer Privacy Act, as amended by the California Privacy Rights Act, Cal. Civ. Code \u00a7 1798.100 et seq., collectively \u201cCCPA\u201d) on behalf of Client, Technijian acts as a \u201cservice provider\u201d as defined under Cal. Civ. Code \u00a7 1798.140(ag) and shall:")
    d.numbered("(a)", "Process such personal information only as necessary to perform the services and in accordance with Client\u2019s documented instructions and this Agreement;", indent=1)
    d.numbered("(b)", "Not sell, share, retain, use, or disclose personal information for any purpose other than performing the services, including not using personal information for targeted advertising or cross-context behavioral advertising;", indent=1)
    d.numbered("(c)", "Not combine personal information received from Client with personal information received from other sources or collected from Technijian\u2019s own interactions with individuals, except as expressly permitted by the CCPA to perform the services;", indent=1)
    d.numbered("(d)", "Implement reasonable security measures appropriate to the nature of the personal information, consistent with the requirements of Section 10.03;", indent=1)
    d.numbered("(e)", "Cooperate with Client in responding to verifiable consumer rights requests under the CCPA (including access, deletion, correction, and opt-out requests) within ten (10) business days of Client\u2019s request;", indent=1)
    d.numbered("(f)", "Notify Client within five (5) business days if Technijian determines that it can no longer meet its obligations as a service provider under the CCPA;", indent=1)
    d.numbered("(g)", "Ensure that all subcontractors who process personal information on behalf of Client are bound by written agreements containing data protection obligations at least as protective as those in this Section 10, and Technijian shall remain liable for the acts and omissions of its subcontractors with respect to personal information;", indent=1)
    d.numbered("(h)", "Permit Client, upon thirty (30) days written notice and no more than once per twelve (12) month period, to audit or inspect Technijian\u2019s data processing practices to verify compliance with this Section 10, or, at Technijian\u2019s option, provide Client with a summary of a recent independent third-party audit (such as SOC 2 Type II) covering the relevant controls; and", indent=1)
    d.numbered("(i)", "Certify that Technijian understands the restrictions in this Section 10 and will comply with them.", indent=1)

    d.numbered("10.02.", "Security Incident Notification. If Technijian becomes aware of a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Client Data (\u201cSecurity Incident\u201d), Technijian shall: (a) notify Client in writing without unreasonable delay and in no event later than forty-eight (48) hours after becoming aware of the Security Incident; (b) provide Client with sufficient information to enable Client to comply with its obligations under California Civil Code \u00a7 1798.82 (data breach notification), including the categories and approximate number of records affected, the nature of the incident, and the measures taken or proposed to address it; (c) cooperate with Client\u2019s investigation of the Security Incident; and (d) take reasonable steps to contain and remediate the Security Incident.")

    d.numbered("10.03.", "Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards designed to protect Client Data from unauthorized access, use, or disclosure, consistent with industry standards for managed IT service providers. Such safeguards shall include, at a minimum: (a) encryption of Client Data in transit and at rest; (b) access controls limiting access to authorized personnel; (c) regular security assessments and vulnerability testing; and (d) employee security awareness training.")

    d.numbered("10.04.", "Regulatory Compliance. If Client is subject to HIPAA, PCI DSS, GDPR, or other industry-specific data protection requirements, the Parties shall execute a separate addendum addressing the additional obligations applicable to the regulated data.")

    d.numbered("10.05.", "Data Return and Deletion. Upon termination of this Agreement or upon Client\u2019s written request, Technijian shall securely delete or return all Client Data in its possession within thirty (30) days, using methods consistent with NIST SP 800-88 or equivalent standards, and shall certify such deletion in writing upon request.")

    d.signatures(client_name="AVENTINE AT ALISO VIEJO APARTMENTS", contact_name="[AUTHORIZED SIGNER]")

    d.spacer()
    d.body("Schedules (delivered as separate documents):")
    d.bullet("Schedule A \u2014 Monthly Managed Services")
    d.bullet("Schedule B \u2014 Subscription and License Services (none in effect as of Effective Date)")
    d.bullet("Schedule C \u2014 Rate Card")

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "MSA-AAVA.docx"))


# ================================================================
# SCHEDULE A - MONTHLY MANAGED SERVICES
# ================================================================
def build_schedule_a():
    d = BrandedDoc()
    d.cover_page(
        "Schedule A",
        "Monthly Managed Services",
        client_name="Aventine at Aliso Viejo Apartments",
        date="May 1, 2026",
    )
    d.page_break()

    d.part_header("SCHEDULE A \u2014 MONTHLY MANAGED SERVICES")
    d.spacer()
    d.body("Agreement Number: MSA-AAVA-2026")
    d.body("Effective Date: May 1, 2026")
    d.body("Parent Agreement: Master Service Agreement between Technijian, Inc. and Aventine at Aliso Viejo Apartments")
    d.spacer()
    d.body(
        "This Schedule A sets forth the recurring managed services to be provided by Technijian to "
        "Aventine at Aliso Viejo Apartments under the Master Service Agreement (MSA-AAVA-2026). "
        "The total monthly investment is comprised of two parts: (1) Endpoint Security and "
        "(2) Virtual Staff Support \u2014 Contracted (Offshore Only). USA Tech Support is provided on "
        "an ad-hoc hourly basis under Schedule C and is not part of the contracted monthly amount. "
        "Capitalized terms used and not defined herein have the meanings given to them in the MSA."
    )
    d.spacer()

    # Part 1 — Endpoint Security
    d.section_header("Part 1 \u2014 Endpoint Security")
    d.spacer()
    d.body("Eight (8) desktops are protected under the following per-endpoint services:")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Per Unit", "Qty", "Monthly"],
        [
            ("CrowdStrike Falcon EDR", "AVD", "$8.50/endpoint", "8", "$68.00"),
            ("Huntress Managed Detection", "AVMH", "$6.00/endpoint", "8", "$48.00"),
            ("My Secure Internet (DNS)", "SI", "$6.00/endpoint", "8", "$48.00"),
            ("Patch Management", "PMW", "$4.00/endpoint", "8", "$32.00"),
            ("My Remote (ScreenConnect)", "MR", "$2.00/endpoint", "8", "$16.00"),
            ("Subtotal \u2014 Endpoint Security", "", "", "8 endpoints", "$212.00"),
        ],
        total_indices=[5],
    )

    # Part 2 — Virtual Staff Support — Contracted (Offshore Only)
    d.spacer()
    d.section_header("Part 2 \u2014 Virtual Staff Support \u2014 Contracted (Offshore Only)")
    d.spacer()
    d.body(
        "Contracted hours are held at the same levels as the prior Client Monthly Service "
        "Agreement. Client\u2019s current actual usage exceeds these contracted hours, which means "
        "any excess hours worked each month accrue to the Unpaid Balance described below and "
        "remain payable per the terms of this Agreement."
    )
    d.spacer()
    d.styled_table(
        ["Role", "Contracted (hrs)", "Rate", "Monthly"],
        [
            ("India Tech \u2014 Normal Hours", "4.55", "$15/hr", "$68.25"),
            ("India Tech \u2014 After Hours", "1.77", "$30/hr", "$53.10"),
            ("Subtotal \u2014 Virtual Staff Support", "6.32", "", "$121.35"),
        ],
        total_indices=[2],
    )

    # Unpaid Balance Acknowledgment
    d.spacer()
    d.section_header("UNPAID BALANCE ACKNOWLEDGMENT")
    d.spacer()
    d.body(
        "The Parties acknowledge that as of the Effective Date of this Agreement, Client has an "
        "accumulated unpaid support hours balance of 56.70 hours (40.41 India Tech Normal + "
        "16.29 India Tech After-Hours) from the prior Client Monthly Service Agreement, "
        "representing $1,094.85 in deferred support fees. This Unpaid Balance carries forward "
        "under this Agreement and remains payable by Client."
    )
    d.spacer()
    d.body(
        "Based on the prior three (3) months\u2019 actual usage (India Tech Normal averaging 9.91 "
        "hrs/month and India Tech After-Hours averaging 2.38 hrs/month), Client\u2019s actual monthly "
        "support usage exceeds the contracted hours in Part 2 by approximately 5.36 India Normal "
        "hours and 0.61 India After-Hours hours per month. Accordingly, the Unpaid Balance is "
        "expected to continue growing at an estimated rate of approximately $98.70 per month "
        "(5.36 hrs \u00d7 $15/hr + 0.61 hrs \u00d7 $30/hr) unless actual usage decreases or Client elects "
        "to increase the contracted hours in a written amendment to this Agreement."
    )
    d.spacer()
    d.body(
        "The Unpaid Balance, including any growth during the Initial Term, shall be billed at the "
        "rates set forth in Schedule C (Rate Card) upon termination or non-renewal of this "
        "Agreement per Section 2.05(a) and shall be due before the Agreement is fully terminated. "
        "Client may elect at any time, by written notice to Technijian, to increase the contracted "
        "hours in Schedule A, Part 2 in order to reduce or eliminate growth of the Unpaid Balance."
    )
    d.spacer()
    d.styled_table(
        ["Component", "Hours", "Rate", "Amount"],
        [
            ("India Tech \u2014 Normal Hours (unpaid)", "40.41", "$15/hr", "$606.15"),
            ("India Tech \u2014 After-Hours (unpaid)", "16.29", "$30/hr", "$488.70"),
            ("Total Unpaid Balance (carry-forward)", "56.70", "", "$1,094.85"),
            ("Projected Monthly Growth", "5.97", "", "$98.70"),
        ],
        total_indices=[2],
    )

    # Part 3 — USA Tech Support (Hourly — Ad-Hoc)
    d.spacer()
    d.section_header("Part 3 \u2014 USA Tech Support (Hourly \u2014 Ad-Hoc)")
    d.spacer()
    d.body(
        "USA Tech Support is NOT part of the monthly contracted amount under this renewal "
        "Agreement. When USA-based support is required, it is billed on an ad-hoc hourly basis "
        "at the rates set forth in Schedule C (Rate Card):"
    )
    d.bullet("$150/hr \u2014 Normal Business Hours")
    d.bullet("$200/hr \u2014 After-Hours")
    d.bullet("Two (2) hour minimum per engagement; 15-minute increments thereafter")
    d.spacer()
    d.body(
        "Ad-hoc USA Tech Support hours are invoiced on the Weekly Out-of-Contract Invoice "
        "described in Section 3.02(c) and are reported with full ticket-level detail."
    )

    # Monthly Investment Summary
    d.spacer()
    d.section_header("Monthly Investment Summary")
    d.spacer()
    d.styled_table(
        ["Category", "Monthly"],
        [
            ("Endpoint Security (8 desktops)", "$212.00"),
            ("Virtual Staff Support \u2014 Contracted (6.32 hrs/mo)", "$121.35"),
            ("USA Tech Support (Ad-Hoc \u2014 per Schedule C)", "Billed when used"),
            ("TOTAL MONTHLY", "$333.35"),
        ],
        total_indices=[3],
    )
    d.spacer()
    d.body(
        "Note: The TOTAL MONTHLY amount of $333.35 represents the contracted recurring fee. "
        "Any ad-hoc USA Tech Support, project work, equipment procurement, or other out-of-scope "
        "services are invoiced separately per Section 3.02 and Schedule C. The Unpaid Balance of "
        "$1,094.85 (see Unpaid Balance Acknowledgment) is carried forward from the prior agreement "
        "and is expected to grow by approximately $98.70 per month at current usage levels."
    )

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "Schedule-A-AAVA.docx"))


# ================================================================
# SCHEDULE B - SUBSCRIPTION AND LICENSE SERVICES
# ================================================================
def build_schedule_b():
    d = BrandedDoc()
    d.cover_page(
        "Schedule B",
        "Subscription and License Services",
        client_name="Aventine at Aliso Viejo Apartments",
        date="May 1, 2026",
    )
    d.page_break()

    d.part_header("SCHEDULE B \u2014 SUBSCRIPTION AND LICENSE SERVICES")
    d.spacer()
    d.body("Agreement Number: MSA-AAVA-2026")
    d.body("Effective Date: May 1, 2026")
    d.body("Parent Agreement: Master Service Agreement between Technijian, Inc. and Aventine at Aliso Viejo Apartments")
    d.spacer()
    d.body(
        "This Schedule B governs any subscription services, software licenses, SaaS subscriptions, "
        "SIP trunk services, and similar recurring third-party services procured by Technijian on "
        "behalf of Aventine at Aliso Viejo Apartments under the Master Service Agreement "
        "(MSA-AAVA-2026). Capitalized terms used and not defined herein have the meanings given "
        "to them in the MSA."
    )
    d.spacer()

    d.section_header("Current Subscription Services")
    d.spacer()
    d.body_bold("Status as of Effective Date: ", "No subscription or license services are currently in effect under this Schedule.")
    d.spacer()
    d.body(
        "Endpoint security services (CrowdStrike, Huntress, My Secure Internet, Patch Management, "
        "and My Remote) are provided under Schedule A, Part 1 as managed services and are not "
        "subject to this Schedule B."
    )

    d.spacer()
    d.section_header("Adding Subscription Services")
    d.spacer()
    d.body(
        "Subscription and license services may be added to this Schedule B at any time during the "
        "term of the Agreement by executing a Service Order (or written amendment) signed by both "
        "Parties. Each Service Order shall specify:"
    )
    d.bullet("The subscription product or license being procured")
    d.bullet("The quantity, unit price, and total monthly or annual fee")
    d.bullet("The subscription term and renewal terms (if different from the MSA)")
    d.bullet("Any vendor-specific terms that apply (such as EULA acknowledgment)")
    d.spacer()
    d.body(
        "Subscription fees are billed on the Monthly Recurring Subscription Invoice described in "
        "Section 3.02(b) of the MSA, in advance, on the first business day of each month."
    )

    d.spacer()
    d.section_header("Pass-Through and Client-Procured Licenses")
    d.spacer()
    d.body(
        "If Client procures software or subscription licenses directly from a third-party vendor "
        "(rather than through Technijian), those licenses are not governed by this Schedule B. "
        "Client remains solely responsible for the terms, renewal, and payment of any directly-"
        "procured subscriptions. Technijian will install, configure, and support such Client-"
        "procured licenses as part of the managed services under Schedule A, subject to vendor "
        "and license compatibility."
    )

    d.spacer()
    d.section_header("Termination of Subscriptions")
    d.spacer()
    d.body(
        "Upon termination of the MSA or this Schedule B, Client shall pay all outstanding fees "
        "for any subscription services through the end of the then-current vendor billing period "
        "(including any annual commitments procured on Client\u2019s behalf that are not cancellable). "
        "Technijian will cooperate in good faith to transfer any transferable subscriptions to "
        "Client or a successor provider."
    )

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "Schedule-B-AAVA.docx"))


# ================================================================
# SCHEDULE C - RATE CARD
# ================================================================
def build_schedule_c():
    d = BrandedDoc()
    d.cover_page(
        "Schedule C",
        "Rate Card",
        client_name="Aventine at Aliso Viejo Apartments",
        date="May 1, 2026",
    )
    d.page_break()

    d.part_header("SCHEDULE C \u2014 RATE CARD")
    d.spacer()
    d.body("Agreement Number: MSA-AAVA-2026")
    d.body("Effective Date: May 1, 2026")
    d.body("Parent Agreement: Master Service Agreement between Technijian, Inc. and Aventine at Aliso Viejo Apartments")
    d.body("Normal Business Hours: Monday \u2013 Friday, 8:00 AM \u2013 6:00 PM Pacific Time")
    d.spacer()
    d.body(
        "This Schedule C sets forth the standard hourly rates applicable to labor services "
        "performed under the Master Service Agreement (MSA-AAVA-2026) that are not covered by an "
        "active contracted Schedule or Statement of Work. Capitalized terms used and not defined "
        "herein have the meanings given to them in the MSA."
    )
    d.spacer()

    d.section_header("US-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr"),
            ("Developer", "$150/hr", "N/A", "$125/hr"),
            ("Tech Support", "$150/hr", "$250/hr", "$125/hr"),
        ],
    )

    d.spacer()
    d.section_header("Offshore-Based Staff")
    d.spacer()
    d.styled_table(
        ["Role", "Hourly", "After-Hours", "Contracted"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr"),
            ("Tech Support", "$15/hr", "$30/hr", "$10/hr"),
        ],
    )

    d.spacer()
    d.section_header("Project & Ad Hoc")
    d.spacer()
    d.styled_table(
        ["Service", "Rate", "Minimum"],
        [
            ("On-Site Support (US)", "$150/hr", "2-hr minimum"),
            ("Remote Support (ad hoc)", "$150/hr", "15-min increments"),
            ("Emergency / Critical", "$250/hr", "1-hr minimum"),
            ("Project Management", "$150/hr", "SOW-based"),
        ],
    )

    d.spacer()
    d.section_header("Rate Definitions")
    d.spacer()
    d.body_bold("Normal Business Hours: ", "Monday \u2013 Friday, 8:00 AM \u2013 6:00 PM Pacific Time, excluding US federal holidays.")
    d.body_bold("After-Hours: ", "All hours outside Normal Business Hours, including weekends and US federal holidays.")
    d.body_bold("Contracted Rate: ", "Discounted rate available when Client commits to the Cycle-Based Billing Model under an active Schedule.")
    d.body_bold("Hourly Rate: ", "Standard ad-hoc rate applied to work performed outside an active contracted Schedule.")

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "Schedule-C-AAVA.docx"))


# ================================================================
# EXECUTIVE SUMMARY
# ================================================================
def build_executive_summary():
    d = BrandedDoc()
    d.cover_page(
        "Executive Summary",
        "Renewal Agreement \u2014 Cost Reduction & Unpaid Balance Review",
        client_name="Aventine at Aliso Viejo Apartments",
        date="May 1, 2026",
    )
    d.page_break()

    # -- Overview --
    d.section_header("Overview")
    d.body(
        "Technijian is pleased to present this renewal of managed IT services for Aventine at "
        "Aliso Viejo Apartments (\u201cAAVA\u201d), effective May 1, 2026. This Executive Summary explains "
        "the key changes in the renewal, the immediate monthly savings being delivered, and an "
        "important update regarding the accumulated unpaid support hours balance carried forward "
        "from the prior agreement."
    )
    d.body(
        "The renewal reduces AAVA\u2019s monthly recurring cost by $206.25 \u2014 a 38% reduction versus "
        "the April 1, 2026 invoice \u2014 while maintaining the same endpoint security coverage across "
        "all eight (8) desktops."
    )

    # -- Monthly Cost Comparison --
    d.spacer()
    d.section_header("Monthly Cost Comparison")
    d.spacer()
    d.styled_table(
        ["Line Item", "Prior (4/1 Invoice)", "New (5/1 MSA)", "Change"],
        [
            ("Endpoint Security (8 desktops)", "$212.00", "$212.00", "No change"),
            ("Site Assessment", "$50.00", "Removed", "\u2212 $50.00"),
            ("USA Tech Support \u2014 Contracted (1.25 hrs)", "$156.25", "Ad-hoc hourly only", "\u2212 $156.25"),
            ("India Tech \u2014 Normal (4.55 hrs @ $15)", "$68.25", "$68.25", "No change"),
            ("India Tech \u2014 After-Hours (1.77 hrs @ $30)", "$53.10", "$53.10", "No change"),
            ("TOTAL MONTHLY", "$539.60", "$333.35", "\u2212 $206.25"),
        ],
        total_indices=[5],
    )

    # -- What Changed --
    d.spacer()
    d.section_header("What Changed \u2014 And Why")
    d.spacer()
    d.body_bold("1. Site Assessment removed (\u2212 $50/mo).", "")
    d.body(
        "The monthly Site Assessment charge has been eliminated. Technijian will continue to "
        "monitor network health through the included endpoint security stack (CrowdStrike, "
        "Huntress, My Secure Internet, Patch Management), with no reduction in day-to-day "
        "visibility or protection."
    )
    d.spacer()
    d.body_bold("2. USA Tech Support moved to ad-hoc hourly (\u2212 $156.25/mo).", "")
    d.body(
        "USA-based technician hours are no longer included in the monthly contracted amount. "
        "Based on recent usage, AAVA consumed zero (0) USA Tech hours in the most recent reporting "
        "period, meaning the contracted $156.25/month was being charged without being used. Under "
        "the renewal, USA support is available on demand at the standard Rate Card rates "
        "($150/hr Normal, $200/hr After-Hours, 2-hour minimum). If AAVA does not call on USA Tech "
        "Support, there is no charge."
    )
    d.spacer()
    d.body_bold("3. Endpoint security and offshore support stay the same.", "")
    d.body(
        "All eight (8) desktops continue to receive CrowdStrike, Huntress, DNS filtering, patch "
        "management, and remote access under the same pricing. India Tech support hours remain "
        "at the prior contracted levels (4.55 Normal + 1.77 After-Hours)."
    )

    # -- Unpaid Balance Section --
    d.page_break()
    d.part_header("IMPORTANT \u2014 UNPAID SUPPORT HOURS BALANCE")
    d.spacer()
    d.section_header("What Is the Unpaid Balance?")
    d.body(
        "Under Technijian\u2019s cycle-based billing model, each contracted support role has a "
        "committed number of monthly hours. When AAVA\u2019s actual usage of a role exceeds the "
        "contracted hours, the extra hours are documented through ticketing but not billed on "
        "the regular monthly invoice \u2014 they accumulate as an \u201cUnpaid Balance\u201d that becomes due "
        "if the agreement is terminated or non-renewed."
    )
    d.spacer()
    d.body_bold("Current Unpaid Balance as of April 1, 2026:", "")
    d.styled_table(
        ["Role", "Unpaid Hours", "Rate", "Amount"],
        [
            ("India Tech \u2014 Normal Hours", "40.41", "$15/hr", "$606.15"),
            ("India Tech \u2014 After-Hours", "16.29", "$30/hr", "$488.70"),
            ("TOTAL UNPAID BALANCE", "56.70", "", "$1,094.85"),
        ],
        total_indices=[2],
    )

    d.spacer()
    d.section_header("Why the Balance Is Growing")
    d.body(
        "AAVA\u2019s actual monthly India Tech usage is currently higher than the contracted hours in "
        "the agreement. In the most recent reporting period (March 2026), actual usage compared "
        "against contracted hours was:"
    )
    d.spacer()
    d.styled_table(
        ["Role", "Contracted (hrs)", "Actual March (hrs)", "Over-Use (hrs)", "Over-Use ($)"],
        [
            ("India Tech \u2014 Normal", "4.55", "9.91", "5.36", "$80.40"),
            ("India Tech \u2014 After-Hours", "1.77", "2.38", "0.61", "$18.30"),
            ("TOTAL PER MONTH", "6.32", "12.29", "5.97", "$98.70"),
        ],
        total_indices=[2],
    )
    d.spacer()
    d.body(
        "At the current usage pattern, the Unpaid Balance will grow by approximately $98.70 per "
        "month. Over the 12-month renewal term, that would add roughly $1,184 to the Unpaid "
        "Balance, bringing the total at the end of the term to approximately $2,279 if no action "
        "is taken and usage stays at March 2026 levels."
    )

    # -- How to Bring It Down --
    d.spacer()
    d.section_header("How to Bring the Unpaid Balance Down")
    d.body(
        "The Unpaid Balance can only decrease if AAVA\u2019s actual monthly usage falls below the "
        "contracted hours. Every hour of India Tech support that AAVA does not need in a given "
        "month is applied against the Unpaid Balance. Examples:"
    )
    d.spacer()
    d.styled_table(
        ["Scenario", "India Normal (actual)", "India AH (actual)", "Monthly Impact on Unpaid"],
        [
            ("Usage stays at March levels", "9.91 hrs", "2.38 hrs", "Grows by $98.70"),
            ("Usage drops to contracted", "4.55 hrs", "1.77 hrs", "Unchanged ($0)"),
            ("Usage drops 25% below contracted", "3.41 hrs", "1.33 hrs", "Reduces by ~$30.50"),
            ("Usage drops 50% below contracted", "2.28 hrs", "0.89 hrs", "Reduces by ~$60.45"),
        ],
    )
    d.spacer()
    d.body(
        "Put simply: to pay down the Unpaid Balance, actual monthly support usage must drop "
        "below the contracted hours (4.55 India Normal + 1.77 India After-Hours). Any month "
        "where usage exceeds those levels will add to the Unpaid Balance; any month where usage "
        "is below will reduce it."
    )

    # -- Recommendations --
    d.page_break()
    d.section_header("Recommendations")
    d.spacer()
    d.body_bold("Option A \u2014 Maintain the renewal as-is and target lower usage.", "")
    d.body(
        "AAVA receives the full $206.25/month savings immediately. Technijian will continue to "
        "provide detailed weekly ticket reports showing time entries by resource, role, and "
        "issue. Client reviews those reports to identify patterns and opportunities to reduce "
        "support consumption (training, documentation, recurring issues worth permanent fixes). "
        "This is the most common path for clients whose support volume is tied to fixable root "
        "causes."
    )
    d.spacer()
    d.body_bold("Option B \u2014 Increase contracted hours to match actual usage.", "")
    d.body(
        "If AAVA expects support demand to stay at March 2026 levels, the contracted hours can "
        "be increased by written amendment so the monthly invoice matches actual usage and the "
        "Unpaid Balance stops growing. Increasing India Normal to approximately 10 hrs and "
        "India After-Hours to approximately 2.5 hrs would cost an additional ~$96/month but "
        "would freeze the Unpaid Balance at its current level."
    )
    d.spacer()
    d.body_bold("Option C \u2014 Pay down the existing Unpaid Balance directly.", "")
    d.body(
        "Client may elect, at any time and by written notice, to settle all or any portion of "
        "the Unpaid Balance via a one-time payment at the contracted rates ($15/hr India Normal, "
        "$30/hr India After-Hours). This eliminates the termination liability exposure without "
        "changing the monthly rate."
    )

    # -- Service Level & What You Get --
    d.spacer()
    d.section_header("Services Included at $333.35 / Month")
    d.spacer()
    d.styled_table(
        ["Category", "What\u2019s Included"],
        [
            ("Endpoint Security", "CrowdStrike Falcon EDR, Huntress Managed Detection, My Secure Internet (DNS), Patch Management, My Remote \u2014 on all 8 desktops"),
            ("Offshore Support", "4.55 hrs India Normal + 1.77 hrs India After-Hours per month, with full ticket documentation"),
            ("USA Support", "Available on demand at $150/hr (Normal) or $200/hr (After-Hours), 2-hr minimum"),
            ("Service Level", "4-hour response to remote support requests, 24-hour emergency onsite"),
            ("Reporting", "Detailed weekly ticket report showing resource, role, hours, and work performed per ticket"),
        ],
    )

    # -- Path Forward --
    d.spacer()
    d.section_header("Path Forward")
    d.spacer()
    d.body_bold("1. Review this summary and the renewal MSA.", "")
    d.body(
        "The attached Master Service Agreement (MSA-AAVA-2026) governs the renewal and contains "
        "the full Unpaid Balance Acknowledgment in Schedule A."
    )
    d.spacer()
    d.body_bold("2. Sign and return the MSA before April 30, 2026.", "")
    d.body(
        "This ensures uninterrupted service and the new pricing takes effect with the May 2026 "
        "invoice."
    )
    d.spacer()
    d.body_bold("3. Schedule a usage review after 60-90 days.", "")
    d.body(
        "Technijian will review actual support consumption versus contracted hours and work with "
        "AAVA to identify opportunities to reduce recurring support demand \u2014 the most reliable "
        "path to paying down the Unpaid Balance."
    )

    d.spacer()
    d.section_header("Questions")
    d.body(
        "Please direct any questions about this summary, the renewal MSA, or the Unpaid Balance "
        "to Ravi Jain at rjain@technijian.com or 949.379.8499."
    )

    d.footer_bar()
    d.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), "AAVA-Executive-Summary.docx"))


if __name__ == "__main__":
    build_msa()
    build_schedule_a()
    build_schedule_b()
    build_schedule_c()
    build_executive_summary()
