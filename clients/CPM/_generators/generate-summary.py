"""
California Mobile Home Park Management Co. (CPM) Executive Summary Generator
MSA Managed Services — Desktop Security + Email Protection + Virtual Staff
Uses Technijian Brand Guide 2026 formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors (from Technijian_Brand_Guide_2026) ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex versions for shading
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"

FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = BRAND_GREY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# ── Helper Functions ──

def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def accent_bar(color_hex, height_pt=4):
    """Full-width colored accent bar using a 1-row table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pf = p._element.get_or_add_pPr()
    spacing = pf.makeelement(qn("w:spacing"), {
        qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
    })
    pf.append(spacing)
    # Remove table borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def run_styled(paragraph, text, size=11, bold=False, color=None, italic=False, font=FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def section_header(title):
    """Section header with blue left-bar accent."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Blue accent bar cell (narrow)
    bar_cell = tbl.rows[0].cells[0]
    set_cell_shading(bar_cell, HEX_BLUE)
    bar_cell.width = Inches(0.08)
    bar_cell.text = ""
    bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Title cell
    title_cell = tbl.rows[0].cells[1]
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run_styled(p, title, size=14, bold=True, color=CORE_BLUE)

    # Remove table borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def styled_table(headers, rows, col_widths=None, total_row_indices=None):
    """Create a branded table with light grey borders and blue header row."""
    total_row_indices = total_row_indices or []
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        run_styled(p, h, size=9, bold=True, color=WHITE)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = tbl.add_row()
        is_total = row_idx in total_row_indices
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if is_total:
                set_cell_shading(cell, HEX_OFF_WHITE)
                run_styled(p, str(val), size=9, bold=True, color=DARK_CHARCOAL)
            else:
                run_styled(p, str(val), size=9, color=BRAND_GREY)

    # Apply light grey borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return tbl


def bullet_item(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_styled(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)
    return p


def body_text(text, space_after=8):
    p = doc.add_paragraph()
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ════════════════════════════════════════════════════════════════
# CLIENT DATA
# ════════════════════════════════════════════════════════════════

CLIENT_NAME = "California Mobile Home Park Management Co."
CLIENT_SHORT = "CalParkMgmt"
CLIENT_CODE = "CPM"
CLIENT_WEBSITE = "calparkmgmt.com"

USERS = 7
DESKTOPS = 7      # endpoints = desktops
LOCATIONS = 1     # single location

# ── Per-Unit Monthly Pricing (from PriceList_20260323) ──
PRICE_CROWDSTRIKE = 8.50       # AVD - CrowdStrike per desktop
PRICE_HUNTRESS = 6.00          # AVMH - Huntress per desktop
PRICE_UMBRELLA = 6.00          # SI - My Secure Internet (Cisco Umbrella DNS) per endpoint
PRICE_PATCH_MGMT = 4.00        # PMW - Patch Management per endpoint
PRICE_MY_REMOTE = 2.00         # MR - My Remote per endpoint
PRICE_INKY = 10.25             # ASP - Anti-Spam Professional (Inky) per user
PRICE_DMARC_DKIM = 20.00       # DKIM - DMARC/DKIM Monitoring per domain
PRICE_SITE_ASSESSMENT = 50.00  # SA - Site/365 Assessment per domain

# Support contracted rates (from Rate Card / Schedule C)
RATE_US_TECH = 125.00          # US Tech Support — contracted rate
RATE_INDIA_NIGHT = 30.00      # India night (after-hours in India) = US business hours
RATE_INDIA_DAY = 15.00        # India day (normal hours in India) = US after-hours

# Support hours breakdown (standard formula: 1 hr per device, 50/50 day/night, day 50/50 US/India)
TOTAL_SUPPORT_HOURS = 7       # 1 hr per desktop
DAY_HOURS = 3.5               # 50% day
NIGHT_HOURS = 3.5             # 50% night
US_DAY_HOURS = 1.75           # 50% of daytime = US
INDIA_NIGHT_HOURS = 1.75      # 50% of daytime = India (night shift, US biz hours)
INDIA_DAY_HOURS = 3.5         # all nighttime = India (day shift, US after-hours)

# Onsite support
RATE_ONSITE = 150.00
ONSITE_MIN_HOURS = 2

# ── Calculated Costs ──
# Desktop/Endpoint Security
cost_crowdstrike = DESKTOPS * PRICE_CROWDSTRIKE
cost_huntress = DESKTOPS * PRICE_HUNTRESS
cost_umbrella = DESKTOPS * PRICE_UMBRELLA
cost_patch = DESKTOPS * PRICE_PATCH_MGMT
cost_remote = DESKTOPS * PRICE_MY_REMOTE
subtotal_desktop = cost_crowdstrike + cost_huntress + cost_umbrella + cost_patch + cost_remote

# Email Security
cost_inky = USERS * PRICE_INKY
cost_dmarc = 1 * PRICE_DMARC_DKIM  # 1 domain
cost_assessment = LOCATIONS * PRICE_SITE_ASSESSMENT
subtotal_email = cost_inky + cost_dmarc + cost_assessment

# Virtual Staff
cost_us_support = US_DAY_HOURS * RATE_US_TECH
cost_india_night = INDIA_NIGHT_HOURS * RATE_INDIA_NIGHT
cost_india_day = INDIA_DAY_HOURS * RATE_INDIA_DAY
subtotal_support = cost_us_support + cost_india_night + cost_india_day

# Totals
total_monthly = subtotal_desktop + subtotal_email + subtotal_support


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

# Top blue accent bar
accent_bar(HEX_BLUE, height_pt=6)

# Spacer
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Centered logo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO_PATH):
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(3.0))

# Orange divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
div_tbl = doc.add_table(rows=1, cols=3)
div_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in div_tbl.rows[0].cells:
    cell.text = ""
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
mid_cell = div_tbl.rows[0].cells[1]
mid_cell.width = Inches(2.5)
set_cell_shading(mid_cell, HEX_ORANGE)
div_tbl.rows[0].cells[0].width = Inches(2.25)
div_tbl.rows[0].cells[2].width = Inches(2.25)
div_pr = div_tbl._element.tblPr
div_borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    "</w:tblBorders>"
)
div_pr.append(div_borders)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(8)
run_styled(p, "Executive Summary", size=26, bold=True, color=DARK_CHARCOAL)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run_styled(p, "Managed Desktop Security, Email Protection & Support Services", size=14, color=BRAND_GREY)

# Prepared for
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run_styled(p, "Prepared for ", size=12, color=BRAND_GREY)
run_styled(p, CLIENT_NAME, size=12, bold=True, color=CORE_BLUE)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run_styled(p, "March 2026", size=11, color=BRAND_GREY)

# Spacer
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Bottom orange accent bar
accent_bar(HEX_ORANGE, height_pt=6)

# Confidential
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)

# Contact
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(
    p,
    "Technijian  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)


# ════════════════════════════════════════════════════════════════
# PAGE 2 - ENVIRONMENT OVERVIEW + WHAT'S INCLUDED
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Overview")
body_text(
    f"Technijian proposes a comprehensive Managed Services Agreement for {CLIENT_NAME}, "
    "delivering enterprise-grade desktop security, endpoint protection, email security, and "
    f"dedicated virtual staff support. This solution is designed to secure and manage "
    f"{CLIENT_NAME}\u2019s {USERS} users across {DESKTOPS} desktops, providing layered "
    "protection with predictable monthly costs."
)
body_text(
    f"Upon engagement, {CLIENT_NAME} will have a fully managed security and support infrastructure "
    "including next-generation endpoint detection and response, DNS-layer protection, "
    "automated patch management, advanced email threat prevention, and a blended US/offshore "
    "support team providing coverage across business hours and after-hours."
)

# Current Environment
section_header("Current Environment")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Details"],
    [
        ("Active Users", f"{USERS}"),
        ("Managed Desktops", f"{DESKTOPS}"),
        ("Office Locations", f"{LOCATIONS}"),
        ("Cloud Platform", "Microsoft 365 (Exchange Online, SharePoint, OneDrive, Teams)"),
        ("Industry", "Mobile Home Park Management / Property Management"),
        ("Website", CLIENT_WEBSITE),
    ],
    col_widths=[1.8, 4.2],
)

# Why Each Service
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Why Each Service")
body_text(
    f"Every component in this proposal addresses a specific risk or operational need for "
    f"{CLIENT_NAME}\u2019s environment. Below is the justification for each service."
)

services_justification = [
    ("CrowdStrike Falcon EDR",
     "Next-generation endpoint detection and response deployed on all managed desktops.",
     f"With {DESKTOPS} desktops handling sensitive tenant information, lease agreements, and "
     f"financial records, {CLIENT_NAME} needs real-time threat detection that goes beyond "
     "traditional antivirus. CrowdStrike uses AI-driven behavioral analysis to detect ransomware, "
     "fileless attacks, and zero-day exploits before they execute. Property management companies "
     "are increasingly targeted by ransomware groups due to the volume of personal and financial "
     "data they process \u2014 Social Security numbers, bank account details, and lease records."),

    ("Huntress Managed Detection",
     "24/7 SOC-backed threat hunting and persistent foothold detection with human-led analysis.",
     "CrowdStrike stops known and behavioral threats in real time, but sophisticated attackers "
     "establish persistence mechanisms that evade automated tools. Huntress provides a dedicated "
     "Security Operations Center that actively hunts for footholds, backdoors, and living-off-the-land "
     "techniques. This layered approach closes the gap between automated detection and human-led "
     "threat intelligence \u2014 critical for an environment handling tenant PII and financial transactions."),

    ("Cisco Umbrella (My Secure Internet)",
     "DNS-layer security that blocks malicious domains, phishing sites, and command-and-control "
     "callbacks before a connection is ever established.",
     "Most cyberattacks begin with a DNS query \u2014 whether a phishing link in an email, a "
     "malicious ad, or a compromised website. Cisco Umbrella intercepts these requests at the "
     "DNS layer, blocking threats before they reach the desktop. This provides consistent protection "
     f"for all {DESKTOPS} of {CLIENT_NAME}\u2019s desktops, especially important when staff access "
     "cloud applications (M365, property management software) and process online rent payments."),

    ("Patch Management (ManageEngine Endpoint Central Plus)",
     "Automated OS and third-party application patching powered by ManageEngine Endpoint Central Plus, "
     "a unified endpoint management platform with software deployment, vulnerability scanning, and "
     "configuration management built in.",
     "Unpatched vulnerabilities are the #1 attack vector for ransomware and malware. With "
     f"{DESKTOPS} desktops running Windows plus business applications, keeping software current is "
     "critical. Endpoint Central Plus goes beyond basic patching \u2014 it provides automated "
     "vulnerability scanning to identify missing patches by severity, supports 850+ third-party "
     "application updates, and includes software deployment and configuration management from a "
     "single console. Scheduled deployment windows ensure security updates are applied consistently "
     "without disrupting daily operations."),

    ("My Remote (ConnectWise ScreenConnect)",
     "Secure remote access and support tool powered by ConnectWise ScreenConnect, providing "
     "on-demand and unattended access to all managed desktops.",
     f"Technijian\u2019s support team needs secure, on-demand access to {CLIENT_NAME}\u2019s "
     "desktops to troubleshoot issues, deploy software, and perform maintenance. ScreenConnect "
     "provides lightning-fast remote sessions with role-based access controls, full session audit "
     "logging, and unattended access for after-hours maintenance \u2014 eliminating the need for "
     "users to install ad hoc remote tools that introduce security risk."),

    ("Inky Anti-Phishing Protection",
     "AI-powered email security that detects and blocks phishing, business email compromise (BEC), "
     "and social engineering attacks targeting Microsoft 365 mailboxes.",
     f"Email is the primary attack vector for business email compromise, and with {USERS} users "
     "on Microsoft 365, every mailbox is a potential entry point. Inky uses computer vision and "
     "machine learning to analyze every inbound email, detecting impersonation attempts, brand "
     "forgery, and zero-day phishing that bypasses Microsoft\u2019s built-in filters. Property "
     "management companies frequently receive fraudulent wire transfer requests, fake vendor "
     "invoices, and tenant impersonation emails \u2014 Inky catches these before they reach the inbox."),

    ("DMARC/DKIM Monitoring",
     "Email authentication monitoring and reporting to prevent unauthorized parties from spoofing "
     f"{CLIENT_NAME}\u2019s domain.",
     "Without DMARC and DKIM enforcement, attackers can send emails that appear to come from "
     f"{CLIENT_NAME}\u2019s domain \u2014 targeting your tenants, vendors, and partners with "
     "fraudulent invoices or requests. Ongoing monitoring ensures your email authentication "
     "policies are correctly configured, alerts on spoofing attempts, and provides visibility "
     "into who is sending email on your behalf."),

    ("Monthly Site Assessment (Network Detective by RapidFire Tools)",
     f"Comprehensive security and compliance assessment for {CLIENT_NAME}\u2019s "
     "office and Microsoft 365 tenant, powered by Network Detective.",
     "Network Detective by RapidFire Tools performs automated, deep-scan assessments that go far "
     "beyond surface-level checks. Each monthly scan produces detailed reports covering network "
     "vulnerabilities, Active Directory health, M365 Secure Score analysis, external attack "
     "surface exposure, and compliance readiness. The tool generates executive-ready risk "
     f"scorecards that track {CLIENT_NAME}\u2019s security posture over time \u2014 giving "
     "leadership clear visibility into risk trends without requiring manual audit effort."),

    ("Virtual Staff Support",
     f"{TOTAL_SUPPORT_HOURS} hours/month of blended US and offshore technical support with "
     "50/50 day/night coverage.",
     f"With {USERS} users and {DESKTOPS} desktops running business-critical property management "
     "applications, the support demand is ongoing. The blended model provides US-based technicians "
     "for complex escalations during business hours, India-based night-shift staff for cost-effective "
     "daytime coverage, and India-based day-shift staff for after-hours support \u2014 ensuring "
     "issues don\u2019t wait until the next business day."),
]

for svc_name, svc_desc, svc_why in services_justification:
    p = doc.add_paragraph()
    run_styled(p, svc_name, size=11, bold=True, color=CORE_BLUE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run_styled(p, svc_desc, size=10, italic=True, color=DARK_CHARCOAL)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run_styled(p, "Why: ", size=10, bold=True, color=CORE_ORANGE)
    run_styled(p, svc_why, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(6)


# ════════════════════════════════════════════════════════════════
# PAGE 3 - DESKTOP SECURITY PRICING
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Desktop & Endpoint Security")
body_text(
    f"Enterprise cybersecurity stack deployed across all {DESKTOPS} managed desktops, "
    "providing layered protection from endpoint to network edge."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Unit", "Qty", "Unit Price", "Monthly"],
    [
        ("CrowdStrike Falcon EDR", "Per desktop", str(DESKTOPS), f"${PRICE_CROWDSTRIKE:.2f}", f"${cost_crowdstrike:,.2f}"),
        ("Huntress Managed Detection", "Per desktop", str(DESKTOPS), f"${PRICE_HUNTRESS:.2f}", f"${cost_huntress:,.2f}"),
        ("My Secure Internet (Cisco Umbrella)", "Per desktop", str(DESKTOPS), f"${PRICE_UMBRELLA:.2f}", f"${cost_umbrella:,.2f}"),
        ("Patch Mgmt (Endpoint Central Plus)", "Per desktop", str(DESKTOPS), f"${PRICE_PATCH_MGMT:.2f}", f"${cost_patch:,.2f}"),
        ("My Remote (ScreenConnect)", "Per desktop", str(DESKTOPS), f"${PRICE_MY_REMOTE:.2f}", f"${cost_remote:,.2f}"),
        ("SUBTOTAL \u2014 DESKTOP SECURITY", "", "", "", f"${subtotal_desktop:,.2f}"),
    ],
    col_widths=[2.2, 1.0, 0.5, 0.9, 1.0],
    total_row_indices=[5],
)

# Email Security
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Email Security & Compliance")
body_text(
    f"Advanced email protection for all {USERS} Microsoft 365 users with domain authentication "
    "monitoring and monthly security assessment."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Unit", "Qty", "Unit Price", "Monthly"],
    [
        ("Inky Anti-Phishing (Anti-Spam Pro)", "Per user", str(USERS), f"${PRICE_INKY:.2f}", f"${cost_inky:,.2f}"),
        ("DMARC/DKIM Monitoring", "Per domain", "1", f"${PRICE_DMARC_DKIM:.2f}", f"${cost_dmarc:,.2f}"),
        ("Site Assessment (Network Detective)", "Per site", str(LOCATIONS), f"${PRICE_SITE_ASSESSMENT:.2f}", f"${cost_assessment:,.2f}"),
        ("SUBTOTAL \u2014 EMAIL & COMPLIANCE", "", "", "", f"${subtotal_email:,.2f}"),
    ],
    col_widths=[2.2, 1.0, 0.5, 0.9, 1.0],
    total_row_indices=[3],
)


# ════════════════════════════════════════════════════════════════
# PAGE 4 - VIRTUAL STAFF + TOTAL INVESTMENT
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Virtual Staff \u2014 Contracted Support")
body_text(
    f"{TOTAL_SUPPORT_HOURS} hours per month of dedicated remote technical support, split 50/50 between "
    "daytime and nighttime coverage. Daytime hours are further split 50/50 between US-based "
    "and India-based (night shift) technicians for cost-optimized coverage during business hours."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Role", "Location", "Shift", "Hours/Mo", "Rate", "Monthly"],
    [
        ("Tech Support", "United States", "Day (US Biz Hours)", f"{US_DAY_HOURS:.2f}", f"${RATE_US_TECH:.2f}/hr", f"${cost_us_support:,.2f}"),
        ("Tech Support", "India \u2014 Night", "Day (US Biz Hours)", f"{INDIA_NIGHT_HOURS:.2f}", f"${RATE_INDIA_NIGHT:.2f}/hr", f"${cost_india_night:,.2f}"),
        ("Tech Support", "India \u2014 Day", "Night (US After-Hours)", f"{INDIA_DAY_HOURS:.2f}", f"${RATE_INDIA_DAY:.2f}/hr", f"${cost_india_day:,.2f}"),
        ("SUBTOTAL \u2014 VIRTUAL STAFF", "", "", f"{TOTAL_SUPPORT_HOURS}", "", f"${subtotal_support:,.2f}"),
    ],
    col_widths=[1.2, 1.1, 1.3, 0.7, 0.9, 0.9],
    total_row_indices=[3],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run_styled(p, "Naming Convention:", size=10, bold=True, color=DARK_CHARCOAL)
naming_items = [
    ("\u201cIndia \u2014 Night\u201d ", "refers to India-based staff working during US business hours (nighttime in India)"),
    ("\u201cIndia \u2014 Day\u201d ", "refers to India-based staff working during US after-hours (daytime in India)"),
]
for bold_text, desc in naming_items:
    bullet_item(desc, bold_prefix=bold_text)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run_styled(p, "Billing Model: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "Cycle-based billing with reconciliation. Initial cycle based on estimated "
    f"{TOTAL_SUPPORT_HOURS} hours/month per the Service Order. Actual usage tracked from day one "
    "with reconciliation at end of each cycle period.",
    size=10, color=BRAND_GREY,
)

# On-Site Support
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("On-Site Support")
body_text(
    "On-site technical support is available for issues that cannot be resolved remotely. "
    "On-site visits are billed at the rates below with no trip charges."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Rate", "Minimum", "Trip Charge"],
    [
        ("On-Site Support (US)", f"${RATE_ONSITE:.2f}/hr", f"{ONSITE_MIN_HOURS}-hour minimum", "No trip charges"),
    ],
    col_widths=[2.0, 1.2, 1.5, 1.3],
)

# Total Investment Summary
doc.add_paragraph().paragraph_format.space_after = Pt(8)
section_header("Total Monthly Investment Summary")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Monthly"],
    [
        ("Desktop & Endpoint Security", f"${subtotal_desktop:,.2f}"),
        ("Email Security & Compliance", f"${subtotal_email:,.2f}"),
        ("Virtual Staff Support", f"${subtotal_support:,.2f}"),
        ("TOTAL MONTHLY", f"${total_monthly:,.2f}"),
    ],
    col_widths=[3.5, 2.5],
    total_row_indices=[3],
)


# ════════════════════════════════════════════════════════════════
# PAGE 5 - LABOR RATE CARD + SLA + AGREEMENT STRUCTURE + NEXT STEPS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Labor Rate Card")
body_text(
    "The following rates apply to all services provided under this Agreement. "
    "Contracted rates are available when Client commits to a billing cycle under the "
    "Cycle-Based Billing Model described in Schedule A."
)

# US Staff Rates
doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph()
run_styled(p, "United States \u2014 Based Staff", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)
styled_table(
    ["Role", "Hourly Rate", "After-Hours Rate", "Contracted Rate", "Description"],
    [
        ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr", "Strategic technology leadership and advisory"),
        ("Developer", "$150/hr", "N/A", "$125/hr", "Software development and engineering"),
        ("Tech Support", "$150/hr", "$250/hr", "$125/hr", "Technical support and systems administration"),
    ],
    col_widths=[1.2, 0.9, 1.0, 1.0, 2.0],
)

# Offshore Staff Rates
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
run_styled(p, "Offshore \u2014 Based Staff", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)
styled_table(
    ["Role", "Hourly Rate", "After-Hours Rate", "Contracted Rate", "Description"],
    [
        ("Developer", "$45/hr", "N/A", "$30/hr", "Software development and engineering"),
        ("SEO Specialist", "$45/hr", "N/A", "$30/hr", "Search engine optimization and digital marketing"),
        ("Tech Support", "$15/hr", "$30/hr", "$10/hr", "Technical support and systems administration"),
    ],
    col_widths=[1.2, 0.9, 1.0, 1.0, 2.0],
)

# Project Rates
doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
run_styled(p, "Project & Ad Hoc Rates", size=11, bold=True, color=CORE_BLUE)
p.paragraph_format.space_after = Pt(4)
styled_table(
    ["Service", "Rate", "Minimum", "Notes"],
    [
        ("On-Site Support (US)", "$150/hr", "2-hour minimum", "No trip charges"),
        ("Remote Support (ad hoc, non-contracted)", "$150/hr", "15-min increments", "Billed in 15-minute increments"),
        ("Emergency / Critical Response", "$250/hr", "1-hour minimum", "After-hours and critical incidents"),
        ("Project Management", "$150/hr", "N/A", "For SOW-based engagements"),
    ],
    col_widths=[2.2, 0.8, 1.1, 2.0],
)

# Rate definitions
doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph()
run_styled(p, "Normal Business Hours: ", size=9, bold=True, color=DARK_CHARCOAL)
run_styled(p, "Monday through Friday, 8:00 AM to 6:00 PM Pacific Time, excluding US federal holidays.", size=9, color=BRAND_GREY)
p = doc.add_paragraph()
run_styled(p, "After-Hours: ", size=9, bold=True, color=DARK_CHARCOAL)
run_styled(p, "All hours outside of Normal Business Hours, including weekends and US federal holidays.", size=9, color=BRAND_GREY)
p = doc.add_paragraph()
run_styled(p, "Contracted Rate: ", size=9, bold=True, color=DARK_CHARCOAL)
run_styled(p, "Discounted hourly rate applied when Client commits to a Cycle under the Cycle-Based Billing Model (Schedule A, Part 3).", size=9, color=BRAND_GREY)

# SLA
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Service Level Commitments")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service Level", "Target"],
    [
        ("Critical Incident Response", "Within 1 hour of notification"),
        ("Standard Support Response", "Within 4 business hours"),
        ("Scheduled Maintenance", "Tuesday evenings and Saturdays (with advance notice)"),
        ("24/7 Monitoring", "All desktops and email security continuously monitored"),
        ("Monthly Reporting", "Service reports summarizing incidents, uptime, and support activity"),
        ("Quarterly Reviews", f"Scheduled service reviews with {CLIENT_NAME}\u2019s designated representative"),
    ],
    col_widths=[2.2, 3.8],
)

# Agreement Structure
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Agreement Structure")
agreements = [
    ("Master Service Agreement (MSA): ",
     "Governs the overall relationship, payment terms, confidentiality, liability, and dispute "
     "resolution for a 12-month initial term with automatic annual renewal."),
    ("Schedule A \u2014 Monthly Services: ",
     "Details the managed desktop security, endpoint protection, email security, and virtual staff "
     "services with SLA commitments."),
    ("Schedule B \u2014 Subscription Services: ",
     "Covers any third-party software subscriptions and licenses managed by Technijian on "
     f"{CLIENT_NAME}\u2019s behalf."),
    ("Schedule C \u2014 Rate Card: ",
     "Establishes hourly rates for all labor, including virtual staff, on-site support, "
     "emergency response, and ad hoc services."),
]
for bold_text, desc in agreements:
    bullet_item(desc, bold_prefix=bold_text)

# Next Steps
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Next Steps")
steps = [
    "Review this summary and the attached agreements",
    "Confirm any questions with your Technijian account representative",
    "Sign the MSA, Schedule A, Schedule B, Schedule C",
    "Technijian begins onboarding and agent deployment within 5 business days of execution",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run_styled(p, f"{i}. ", size=10, bold=True, color=CORE_BLUE)
    run_styled(p, step, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)

# Footer divider
doc.add_paragraph().paragraph_format.space_after = Pt(8)
accent_bar(HEX_BLUE, height_pt=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run_styled(
    p,
    "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(
    p,
    f"CONFIDENTIAL \u2014 Prepared exclusively for {CLIENT_NAME}",
    size=8, italic=True, color=BRAND_GREY,
)

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "02_MSA", "CPM-Executive-Summary.docx")
doc.save(OUTPUT)
print(f"Created {OUTPUT}")
