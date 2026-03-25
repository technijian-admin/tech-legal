"""
California Mobile Home Park Management Co. (CPM) MSA DOCX Generator
Uses Technijian Brand Guide 2026 formatting — full MSA template with all sections
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors (Technijian Brand Guide 2026) ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"
HEX_LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "technijian-logo.png")

CLIENT_NAME = "California Mobile Home Park Management Co."
CLIENT_SHORT = "CalParkMgmt"
CLIENT_CODE = "CPM"


def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def remove_borders(tbl):
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)


class BrandedDoc:
    def __init__(self):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10)
        style.font.color.rgb = BRAND_GREY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    def run(self, paragraph, text, size=10, bold=False, color=None, italic=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        return r

    def accent_bar(self, color_hex, height_pt=6):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, color_hex)
        cell.text = ""
        p = cell.paragraphs[0]
        pf = p._element.get_or_add_pPr()
        spacing = pf.makeelement(qn("w:spacing"), {
            qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
        })
        pf.append(spacing)
        remove_borders(tbl)

    def cover_page(self, title, subtitle, date="March 25, 2026"):
        self.accent_bar(HEX_BLUE)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(LOGO_PATH, width=Inches(3.0))
        # Orange divider
        div = self.doc.add_table(rows=1, cols=3)
        div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in div.rows[0].cells:
            c.text = ""
            c.paragraphs[0].paragraph_format.space_before = Pt(0)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(div.rows[0].cells[1], HEX_ORANGE)
        div.rows[0].cells[0].width = Inches(2.25)
        div.rows[0].cells[1].width = Inches(2.5)
        div.rows[0].cells[2].width = Inches(2.25)
        remove_borders(div)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        self.run(p, title, size=24, bold=True, color=DARK_CHARCOAL)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, subtitle, size=13, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Prepared for ", size=12, color=BRAND_GREY)
        self.run(p, CLIENT_NAME, size=12, bold=True, color=CORE_BLUE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, date, size=10, color=BRAND_GREY)
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.accent_bar(HEX_ORANGE)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        self.run(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def section_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar = tbl.rows[0].cells[0]
        set_cell_shading(bar, HEX_BLUE)
        bar.width = Inches(0.08)
        bar.text = ""
        bar.paragraphs[0].paragraph_format.space_before = Pt(0)
        bar.paragraphs[0].paragraph_format.space_after = Pt(0)
        tc = tbl.rows[0].cells[1]
        tc.text = ""
        p = tc.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=13, bold=True, color=CORE_BLUE)
        remove_borders(tbl)

    def part_header(self, title):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self.run(p, title, size=12, bold=True, color=WHITE)
        remove_borders(tbl)

    def body(self, text, space_after=6):
        p = self.doc.add_paragraph()
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def body_bold(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def body_bold_prefix(self, bold_text, text):
        p = self.doc.add_paragraph()
        self.run(p, bold_text, size=9, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=9, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def styled_table(self, headers, rows, total_indices=None):
        total_indices = total_indices or []
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            set_cell_shading(c, HEX_BLUE)
            c.text = ""
            self.run(c.paragraphs[0], h, size=9, bold=True, color=WHITE)
        for ri, rd in enumerate(rows):
            row = tbl.add_row()
            is_total = ri in total_indices
            for i, val in enumerate(rd):
                c = row.cells[i]
                c.text = ""
                if is_total:
                    set_cell_shading(c, HEX_OFF_WHITE)
                    self.run(c.paragraphs[0], str(val), size=9, bold=True, color=DARK_CHARCOAL)
                else:
                    self.run(c.paragraphs[0], str(val), size=9, color=BRAND_GREY)
        # Apply light grey borders
        tbl_pr = tbl._element.tblPr
        borders_xml = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_LIGHT_GREY}"/>'
            "</w:tblBorders>"
        )
        tbl_pr.append(borders_xml)
        return tbl

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            self.run(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, num, text, indent=0):
        p = self.doc.add_paragraph()
        prefix = "    " * indent
        self.run(p, f"{prefix}{num} ", size=10, bold=True, color=CORE_BLUE)
        self.run(p, text, size=10, color=BRAND_GREY)
        p.paragraph_format.space_after = Pt(3)
        return p

    def spacer(self, pts=4):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pts)

    def signatures(self, client_name=CLIENT_NAME.upper()):
        self.doc.add_paragraph()
        self.section_header("SIGNATURES")
        self.spacer(8)
        p = self.doc.add_paragraph()
        self.run(p, "TECHNIJIAN, INC.", size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", "Name: _________________________________",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)
        self.spacer(12)
        p = self.doc.add_paragraph()
        self.run(p, client_name, size=11, bold=True, color=DARK_CHARCOAL)
        for label in ["By: ___________________________________", "Name: _________________________________",
                      "Title: _________________________________", "Date: _________________________________"]:
            p = self.doc.add_paragraph()
            self.run(p, label, size=10, color=BRAND_GREY)
            p.paragraph_format.space_after = Pt(2)

    def footer_bar(self):
        self.spacer(8)
        self.accent_bar(HEX_BLUE, height_pt=2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        self.run(p, "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com", size=8, color=BRAND_GREY)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
        self.doc.save(path)
        print(f"Created {path}")


# ════════════════════════════════════════════════════════════════
#  MSA
# ════════════════════════════════════════════════════════════════
def build_msa():
    d = BrandedDoc()
    d.cover_page("Master Service Agreement", "Managed Desktop Security & Email Protection")
    d.page_break()

    d.part_header("MASTER SERVICE AGREEMENT")
    d.spacer()
    d.body("Agreement Number: MSA-CPM-2026")
    d.body("Effective Date: March 25, 2026")
    d.spacer()
    d.body("This Master Service Agreement (\u201cAgreement\u201d) is entered into by and between:")
    d.spacer()
    d.body_bold("Technijian, Inc. (\u201cTechnijian\u201d)", "")
    d.body("18 Technology Drive, Suite 141")
    d.body("Irvine, California 92618")
    d.spacer()
    d.body("and")
    d.spacer()
    d.body_bold(f"{CLIENT_NAME} (\u201cClient\u201d)", "")
    d.body("[CLIENT ADDRESS]")
    d.body("[CITY, STATE ZIP]")
    d.body("Website: calparkmgmt.com")
    d.spacer()
    d.body("(collectively, the \u201cParties\u201d)")

    # RECITALS
    d.spacer()
    d.section_header("RECITALS")
    d.body("WHEREAS, Technijian provides managed IT services, cybersecurity, cloud infrastructure, and related technology solutions; and")
    d.body(f"WHEREAS, Client desires to engage Technijian to provide certain services as described in the Schedules attached hereto, including but not limited to endpoint security, email protection, and ongoing managed desktop support;")
    d.body("NOW, THEREFORE, for good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:")

    # SECTION 1
    d.spacer()
    d.section_header("SECTION 1 \u2014 SCOPE OF SERVICES")

    d.numbered("1.01.", "Services. Technijian shall provide the services described in the Schedules attached to this Agreement, which are incorporated herein by reference:")
    d.bullet("Schedule A \u2014 Monthly Managed Services (Desktop Security, Email Protection, Monitoring)")
    d.bullet("Schedule B \u2014 Subscription and License Services")
    d.bullet("Schedule C \u2014 Rate Card")
    d.body("Additional services may be provided through Statements of Work (\u201cSOWs\u201d) executed under this Agreement. Each SOW shall be signed by authorized representatives of both Parties and shall reference this Agreement by number. Upon execution, each SOW is incorporated into and governed by this Agreement.")

    d.numbered("1.02.", "Standard of Care. Technijian shall perform all services in a professional and workmanlike manner, consistent with industry standards for managed IT service providers.")
    d.numbered("1.03.", "Service Level Agreement. The service levels applicable to the services are set forth in Schedule A. Technijian shall use commercially reasonable efforts to meet the service levels described therein.")

    d.numbered("1.04.", "Client Responsibilities. Client shall:")
    d.numbered("(a)", "Provide Technijian with reasonable access to Client\u2019s systems, facilities, and personnel as necessary for Technijian to perform the services;", indent=1)
    d.numbered("(b)", "Designate a primary point of contact for communications with Technijian;", indent=1)
    d.numbered("(c)", "Maintain current and accurate information regarding Client\u2019s systems and infrastructure;", indent=1)
    d.numbered("(d)", "Comply with all applicable laws and regulations in connection with its use of the services; and", indent=1)
    d.numbered("(e)", "Be solely responsible for the security and management of Client\u2019s account credentials and passwords.", indent=1)

    d.numbered("1.05.", "Independent Contractor. Technijian is an independent contractor. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the Parties.")
    d.numbered("1.06.", "Client Data. \u201cClient Data\u201d means any data, information, content, records, or files that belong to Client or Client\u2019s customers, employees, or agents, that are stored, processed, or transmitted using the services, including personal information as defined under the California Consumer Privacy Act.")
    d.numbered("1.07.", "Currency. All fees under this Agreement are stated and payable in United States Dollars (USD).")
    d.numbered("1.08.", "Order of Precedence. In the event of a conflict between this Agreement and any Schedule, SOW, or Service Order, the following order of precedence shall apply (highest to lowest): (a) the applicable SOW or Service Order; (b) the applicable Schedule; (c) this Master Service Agreement.")
    d.numbered("1.09.", "Representations and Warranties. Each Party represents and warrants that: (a) it has the legal power and authority to enter into this Agreement; (b) the execution of this Agreement has been duly authorized by all necessary corporate action; (c) this Agreement constitutes a valid and binding obligation enforceable against it in accordance with its terms; and (d) its performance under this Agreement will not violate any applicable law, regulation, or existing contractual obligation.")
    d.numbered("1.10.", "Subcontractors. Technijian may engage subcontractors, including offshore personnel, to perform services under this Agreement, provided that: (a) Technijian shall remain fully responsible for all services performed by its subcontractors; (b) all subcontractors shall be bound by confidentiality and data protection obligations at least as protective as those in this Agreement; and (c) Technijian shall be liable for the acts and omissions of its subcontractors as if they were Technijian\u2019s own.")

    # SECTION 2
    d.spacer()
    d.section_header("SECTION 2 \u2014 TERM AND RENEWAL")

    d.numbered("2.01.", "Initial Term. This Agreement shall commence on the Effective Date and continue for a period of twelve (12) months (the \u201cInitial Term\u201d).")
    d.numbered("2.02.", "Renewal. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve (12) month periods (each a \u201cRenewal Term\u201d), unless either Party provides written notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term.")
    d.numbered("2.03.", "Termination for Convenience. Either Party may terminate this Agreement for any reason upon sixty (60) days written notice to the other Party.")
    d.numbered("2.04.", "Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party:")
    d.numbered("(a)", "Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days after receiving written notice of the breach; or", indent=1)
    d.numbered("(b)", "Becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.", indent=1)

    d.numbered("2.05.", "Effect of Termination.")
    d.numbered("(a)", "Upon termination for any reason, all fees and charges for services rendered through the date of termination shall become immediately due and payable, including: (i) any remaining obligations for annual licenses and subscriptions procured on Client\u2019s behalf; (ii) all accrued and unpaid late fees under Section 3.04; and (iii) any early termination fees or minimum commitment obligations set forth in the applicable Schedule or SOW.", indent=1)
    d.numbered("(b)", "Technijian shall provide reasonable transition assistance for a period of up to thirty (30) days following termination, provided that Client has paid all amounts owed under this Agreement in full.", indent=1)
    d.numbered("(c)", "Technijian shall return all Client Data in its possession within thirty (30) days of termination, in a commercially standard format, provided Client is current on all payment obligations and is not otherwise in breach of this Agreement.", indent=1)
    d.numbered("(d)", "The following sections shall survive termination: Section 3 (Payment), Section 4 (Confidentiality), Section 5 (Limitation of Liability), Section 6 (Indemnification), Section 7 (Intellectual Property), Section 8 (Dispute Resolution), Section 10 (Data Protection), and Section 11 (Insurance).", indent=1)

    d.page_break()

    # SECTION 3
    d.section_header("SECTION 3 \u2014 PAYMENT")

    d.numbered("3.01.", "Fees. Client shall pay fees for the services as set forth in the applicable Schedule, SOW, or invoice. Fees are exclusive of applicable taxes.")
    d.numbered("3.02.", "Invoicing. Technijian shall invoice Client monthly in advance for recurring services and upon delivery for one-time services, unless otherwise specified in the applicable Schedule or SOW.")
    d.numbered("3.03.", "Payment Terms. All invoices are due and payable within thirty (30) days of the invoice date, unless otherwise specified.")
    d.numbered("3.04.", "Late Payment. Invoices not paid within terms shall accrue a late fee of 1.5% per month (or the maximum rate permitted by law, whichever is less) on the unpaid balance, calculated from the date payment was due until the date payment is received in full. Late fees shall compound monthly and are payable in addition to the outstanding principal amount.")
    d.numbered("3.05.", "Disputed Invoices. If Client disputes any portion of an invoice, Client shall notify Technijian in writing within fifteen (15) days of the invoice date, specifying the nature and basis of the dispute with reasonable particularity. Client shall pay all undisputed amounts by the due date. Failure to provide a timely written dispute notice shall constitute acceptance of the invoice.")
    d.numbered("3.06.", "Suspension of Services. If Client fails to pay any undisputed invoice within thirty (30) days of the due date, Technijian may, upon ten (10) days written notice, suspend all services under this Agreement until payment is received in full, including all accrued late fees.")
    d.numbered("3.07.", "Acceleration. Upon the occurrence of any of the following events, all fees, charges, and amounts owing under this Agreement shall become immediately due and payable: (a) Client fails to pay any undisputed invoice within forty-five (45) days of the due date; (b) Client terminates this Agreement while any invoices remain unpaid; (c) Client becomes insolvent or files for bankruptcy; or (d) Client is the subject of a material adverse change in its financial condition.")
    d.numbered("3.08.", "Collection Costs and Attorney\u2019s Fees. If Technijian engages any attorney, collection agency, or other third party to collect amounts owed under this Agreement, Client shall be liable for all costs of collection, including reasonable attorney\u2019s fees, collection agency fees, court costs, and all costs of appeal.")
    d.numbered("3.09.", "Right of Setoff and Lien. (a) Technijian shall have the right to set off any amounts owed by Client against any amounts Technijian may owe to Client. (b) Technijian shall retain a lien on all work product, deliverables, and materials in its possession until all amounts owed by Client are paid in full. (c) In the event of non-payment, Technijian may withhold transition assistance, data return, and credential transfers until all amounts owed are paid in full.")
    d.numbered("3.10.", "Grant of Security Interest (UCC). To secure payment of all amounts owed, Client hereby grants to Technijian a continuing security interest in: (i) all accounts receivable; (ii) all general intangibles; (iii) all equipment procured by Technijian on Client\u2019s behalf; (iv) all proceeds of the foregoing. Client authorizes Technijian to file a UCC-1 Financing Statement to perfect this security interest.")
    d.numbered("3.11.", "Credit Reporting and Collections. Technijian reserves the right to report delinquent accounts to commercial credit reporting agencies and to assign delinquent accounts to third-party collection agencies, in each case after sixty (60) days of non-payment and ten (10) days written notice to Client.")
    d.numbered("3.12.", "Taxes. Client shall be responsible for all applicable sales, use, and other taxes arising from the services, excluding taxes based on Technijian\u2019s income.")

    # SECTION 4
    d.spacer()
    d.section_header("SECTION 4 \u2014 CONFIDENTIALITY")

    d.numbered("4.01.", "Definition. \u201cConfidential Information\u201d means any non-public information disclosed by either Party to the other in connection with this Agreement, including business, technical, and financial information.")
    d.numbered("4.02.", "Obligations. Each Party shall:")
    d.numbered("(a)", "Hold the other Party\u2019s Confidential Information in confidence using at least the same degree of care it uses for its own confidential information, but not less than reasonable care;", indent=1)
    d.numbered("(b)", "Not disclose Confidential Information to third parties without prior written consent, except to employees, agents, and subcontractors who have a need to know and are bound by equivalent obligations; and", indent=1)
    d.numbered("(c)", "Not use Confidential Information for any purpose other than performing obligations under this Agreement.", indent=1)
    d.numbered("4.03.", "Exclusions. Confidential Information does not include information that is or becomes publicly available through no fault of the receiving Party, was known to the receiving Party prior to disclosure, is independently developed, or is received from a third party without restriction.")
    d.numbered("4.04.", "Compelled Disclosure. If required by law or court order to disclose Confidential Information, the receiving Party shall provide prompt written notice to the disclosing Party and cooperate in seeking a protective order.")
    d.numbered("4.05.", "Duration. Confidentiality obligations shall survive termination for a period of three (3) years.")

    # SECTION 5
    d.spacer()
    d.section_header("SECTION 5 \u2014 LIMITATION OF LIABILITY")

    d.numbered("5.01.", "EXCEPT FOR BREACHES OF SECTION 4 (CONFIDENTIALITY), WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, NEITHER PARTY\u2019S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM.")
    d.numbered("5.02.", "EXCEPT FOR BREACHES OF SECTION 4, SECTION 10, INDEMNIFICATION OBLIGATIONS UNDER SECTION 6, WILLFUL MISCONDUCT, OR GROSS NEGLIGENCE, IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, BUSINESS OPPORTUNITY, OR GOODWILL.")
    d.numbered("5.03.", "Data Liability. While Technijian shall use commercially reasonable efforts to protect Client Data, Client acknowledges that: (a) Client is solely responsible for maintaining backup copies of its data; (b) Technijian\u2019s liability for data loss shall be limited to commercially reasonable efforts to restore data from available backups; and (c) Technijian shall not be liable for data loss caused by Client\u2019s actions, third-party attacks, or events beyond Technijian\u2019s reasonable control.")

    d.page_break()

    # SECTION 6
    d.section_header("SECTION 6 \u2014 INDEMNIFICATION")

    d.numbered("6.01.", "By Technijian. Technijian shall indemnify, defend, and hold harmless Client from and against any third-party claims arising from Technijian\u2019s gross negligence or willful misconduct in performing the services.")
    d.numbered("6.02.", "By Client. Client shall indemnify, defend, and hold harmless Technijian from and against any third-party claims arising from: (a) Client\u2019s use of the services in violation of applicable law; (b) Client\u2019s breach of this Agreement; or (c) any data, content, or materials provided by Client.")
    d.numbered("6.03.", "Procedure. The indemnified Party shall provide prompt written notice of any claim, cooperate with the indemnifying Party in the defense, and not settle any claim without the indemnifying Party\u2019s prior written consent.")

    # SECTION 7
    d.spacer()
    d.section_header("SECTION 7 \u2014 INTELLECTUAL PROPERTY")

    d.numbered("7.01.", "Technijian IP. Technijian retains all right, title, and interest in its proprietary tools, methodologies, software, and processes.")
    d.numbered("7.02.", "Client IP. Client retains all right, title, and interest in its data, content, and pre-existing intellectual property.")
    d.numbered("7.03.", "Custom Development. Ownership of any custom software or materials developed under a SOW shall be governed by the terms of that SOW.")

    # SECTION 8
    d.spacer()
    d.section_header("SECTION 8 \u2014 DISPUTE RESOLUTION")

    d.numbered("8.01.", "Escalation. The Parties shall first attempt to resolve any dispute through good faith negotiations for a period of thirty (30) days.")
    d.numbered("8.02.", "Mediation. If not resolved, the Parties shall submit the dispute to mediation in Orange County, California, for a period not to exceed sixty (60) days.")
    d.numbered("8.03.", "Arbitration. If mediation fails, any remaining dispute shall be resolved by binding arbitration administered by the American Arbitration Association in Orange County, California, before a single arbitrator.")
    d.numbered("8.04.", "Fees \u2014 Payment Collection Actions. In any Collection Action, the prevailing Party shall be entitled to recover all reasonable costs and expenses, including attorney\u2019s fees, court costs, and all costs of appeal.")
    d.numbered("8.05.", "Fees \u2014 All Other Disputes. Except as provided in Section 8.04 and Section 3.08, in any non-collection dispute, each Party shall bear its own attorney\u2019s fees and costs.")
    d.numbered("8.06.", "Injunctive Relief. Nothing in this Section shall prevent either Party from seeking injunctive or other equitable relief to prevent irreparable harm.")

    # SECTION 9
    d.spacer()
    d.section_header("SECTION 9 \u2014 GENERAL PROVISIONS")

    d.numbered("9.01.", "Entire Agreement. This Agreement, together with its Schedules and any SOWs, constitutes the entire agreement between the Parties.")
    d.numbered("9.02.", "Amendment. This Agreement may only be amended by a written instrument signed by both Parties. Technijian may update its Rate Card (Schedule C) upon sixty (60) days written notice to Client, effective at the start of the next Renewal Term.")
    d.numbered("9.03.", "Severability. If any provision is found invalid or unenforceable, the remaining provisions shall continue in full force.")
    d.numbered("9.04.", "Waiver. No waiver of any provision shall be effective unless in writing and signed by the waiving Party.")
    d.numbered("9.05.", "Assignment. Neither Party may assign this Agreement without prior written consent, except in connection with a merger, acquisition, or sale of substantially all assets.")
    d.numbered("9.06.", "Force Majeure. Neither Party shall be liable for delays or failures caused by events beyond its reasonable control.")
    d.numbered("9.07.", "Notices. All notices shall be in writing and delivered by email with confirmation, certified mail, or nationally recognized overnight courier.")
    d.numbered("9.08.", "Governing Law. This Agreement shall be governed by the laws of the State of California without regard to conflict of law principles.")
    d.numbered("9.09.", "Non-Solicitation. During the term and for one (1) year following termination, neither Party shall directly solicit for employment any employee of the other Party involved in this Agreement, without prior written consent.")
    d.numbered("9.10.", "Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original.")

    # SECTION 10
    d.spacer()
    d.section_header("SECTION 10 \u2014 DATA PROTECTION")

    d.numbered("10.01.", "CCPA Compliance. To the extent Technijian processes personal information on behalf of Client, Technijian shall: (a) process such information only as necessary to perform the services; (b) not sell, share, or use personal information for any purpose other than performing the services; (c) implement reasonable security measures; and (d) cooperate with Client in responding to consumer rights requests under the CCPA.")
    d.numbered("10.02.", "Security Incident Notification. If Technijian becomes aware of a Security Incident, Technijian shall: (a) notify Client within seventy-two (72) hours; (b) provide sufficient information to enable Client to comply with data breach notification laws; (c) cooperate with Client\u2019s investigation; and (d) take reasonable steps to contain and remediate the incident.")
    d.numbered("10.03.", "Data Security. Technijian shall implement and maintain administrative, technical, and physical safeguards including: (a) encryption of Client Data in transit and at rest; (b) access controls; (c) regular security assessments; and (d) employee security awareness training.")
    d.numbered("10.04.", "Regulatory Compliance. If Client is subject to HIPAA, PCI DSS, GDPR, or other data protection requirements, the Parties shall execute a separate addendum addressing the additional obligations.")
    d.numbered("10.05.", "Data Return and Deletion. Upon termination or upon Client\u2019s written request, Technijian shall securely delete or return all Client Data within thirty (30) days, using methods consistent with NIST SP 800-88, and shall certify such deletion in writing upon request.")

    # SECTION 11
    d.spacer()
    d.section_header("SECTION 11 \u2014 INSURANCE")

    d.numbered("11.01.", "Required Coverage. During the term, Technijian shall maintain: (a) Commercial General Liability insurance with limits of not less than $1,000,000 per occurrence and $2,000,000 aggregate; (b) Professional Liability (E&O) insurance with limits of not less than $1,000,000 per claim and $2,000,000 aggregate; (c) Cyber Liability insurance with limits of not less than $1,000,000 per claim; and (d) Workers\u2019 Compensation insurance as required by the laws of the State of California.")
    d.numbered("11.02.", "Certificates of Insurance. Upon Client\u2019s written request, Technijian shall provide certificates of insurance and at least thirty (30) days\u2019 prior written notice of any material change or cancellation.")

    d.signatures()
    d.spacer()
    d.body("Schedules:")
    d.bullet("Schedule A \u2014 Monthly Managed Services")
    d.bullet("Schedule B \u2014 Subscription and License Services")
    d.bullet("Schedule C \u2014 Rate Card")
    d.footer_bar()

    # ════════════════════════════════════════════════════════════════
    #  SCHEDULE C — RATE CARD (Appendix)
    # ════════════════════════════════════════════════════════════════
    d.page_break()
    d.part_header("SCHEDULE C \u2014 RATE CARD")
    d.spacer()
    d.body("Attached to Master Service Agreement MSA-CPM-2026")
    d.body("Effective Date: March 25, 2026")
    d.spacer()

    # ── SECTION 1: VIRTUAL STAFF RATES ──
    d.part_header("1. VIRTUAL STAFF RATES")
    d.spacer()

    d.section_header("1.1 United States \u2014 Based Staff")
    d.styled_table(
        ["Role", "Hourly Rate", "After-Hours Rate", "Contracted Rate", "Description"],
        [
            ("CTO Advisory", "$250/hr", "$350/hr", "$225/hr", "Strategic technology leadership and advisory"),
            ("Developer", "$150/hr", "N/A", "$125/hr", "Software development and engineering"),
            ("Tech Support", "$150/hr", "$250/hr", "$125/hr", "Technical support and systems administration"),
        ],
    )

    d.spacer()
    d.section_header("1.2 Offshore \u2014 Based Staff")
    d.styled_table(
        ["Role", "Hourly Rate", "After-Hours Rate", "Contracted Rate", "Description"],
        [
            ("Developer", "$45/hr", "N/A", "$30/hr", "Software development and engineering"),
            ("SEO Specialist", "$45/hr", "N/A", "$30/hr", "Search engine optimization and digital marketing"),
            ("Tech Support", "$15/hr", "$30/hr", "$10/hr", "Technical support and systems administration"),
        ],
    )

    d.spacer()
    d.body_bold_prefix("Normal Business Hours: ", "Monday through Friday, 8:00 AM to 6:00 PM Pacific Time, excluding US federal holidays.")
    d.body_bold_prefix("After-Hours: ", "All hours outside of Normal Business Hours, including weekends and US federal holidays.")
    d.body_bold_prefix("Contracted Rate: ", "The discounted hourly rate applied when Client commits to a Cycle under the Cycle-Based Billing Model described in Schedule A, Part 3. Contracted Rates are available only for Virtual Staff services billed through the Cycle model.")
    d.body_bold_prefix("Hourly Rate: ", "The standard rate for ad hoc (non-contracted) services. This rate is also applied to calculate cancellation fees on any unpaid hour balance upon termination of Virtual Staff services, as described in Schedule A, Section 3.3(e).")

    # ── SECTION 2: PROJECT RATES ──
    d.spacer()
    d.part_header("2. PROJECT RATES")
    d.spacer()
    d.styled_table(
        ["Service", "Rate", "Minimum", "Notes"],
        [
            ("On-Site Support (US)", "$150/hr", "2-hour minimum", "No trip charges"),
            ("Remote Support (ad hoc, non-contracted)", "$150/hr", "15-min increments", "Billed in 15-minute increments"),
            ("Emergency / Critical Response", "$250/hr", "1-hour minimum", "After-hours and critical incidents"),
            ("Project Management", "$150/hr", "N/A", "For SOW-based engagements"),
        ],
    )

    d.page_break()

    # ── SECTION 3: ONLINE SERVICES ──
    d.part_header("3. ONLINE SERVICES \u2014 INFRASTRUCTURE")
    d.spacer()

    d.section_header("3.1 Cloud Infrastructure")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Cloud VM \u2014 vCore", "CL-VC", "Per vCore", "M", "$6.25"),
            ("Cloud VM \u2014 Memory", "CL-GB", "Per GB", "M", "$0.63"),
            ("Cloud VM \u2014 Shared Bandwidth", "CL-SBW", "Per connection", "M", "$15.00"),
            ("Hosting \u2014 CDN", "HC", "Per domain", "M", "$30.00"),
            ("Production Storage", "TB-PSTR", "Per TB", "M", "$200.00"),
            ("Replicated Storage", "TB-RSTR", "Per TB", "M", "$100.00"),
            ("Backup Storage", "TB-BSTR", "Per TB", "M", "$50.00"),
        ],
    )

    d.spacer()
    d.section_header("3.2 Server Management")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Patch Management", "PMW", "Per server", "M", "$4.00"),
            ("Image Backup (Veeam)", "IB", "Per server", "M", "$15.00"),
            ("AV Protection \u2014 Server (CrowdStrike)", "AVS", "Per server", "M", "$10.50"),
            ("AVM Protection \u2014 Server (MalwareBytes)", "AVMS", "Per server", "M", "$16.00"),
            ("AVH Protection \u2014 Server (Huntress)", "AVHS", "Per server", "M", "$6.00"),
            ("My Secure Internet (DNS Filtering)", "SI", "Per server", "M", "$6.00"),
            ("My Remote", "MR", "Per server", "M", "$2.00"),
            ("Health Monitoring (SNMP)", "SHM", "Per device", "M", "$2.00"),
            ("Syslog Monitoring (SNMP)", "SSM", "Per device", "M", "$2.00"),
            ("My Ops \u2014 Net", "OPS-NET", "Per server", "M", "$3.25"),
            ("My Ops \u2014 Config", "OPS-BKP", "Per switch", "M", "$6.00"),
            ("My Ops \u2014 Traffic", "OPS-TR", "Per firewall", "M", "$14.00"),
            ("My Ops \u2014 Port", "OPS-PRT", "Per port", "M", "$0.25"),
            ("My Ops \u2014 Storage", "OPS-ST", "Per disk", "M", "$4.75"),
            ("My Ops \u2014 Wifi", "OPS-WF", "Per AP", "M", "$1.00"),
            ("MyAudit Server", "AMS", "Per server", "M", "$252.00"),
        ],
    )

    d.spacer()
    d.section_header("3.3 Desktop Management")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Patch Management", "PMW", "Per desktop", "M", "$4.00"),
            ("Patch Management (Mac OS)", "PMMAC", "Per desktop", "M", "$11.00"),
            ("AV Protection \u2014 Desktop (CrowdStrike)", "AVD", "Per desktop", "M", "$8.50"),
            ("AVM Protection \u2014 Desktop (MalwareBytes)", "AVMD", "Per desktop", "M", "$5.00"),
            ("AVH Protection \u2014 Desktop (Huntress)", "AVMH", "Per desktop", "M", "$6.00"),
            ("My Secure Internet (DNS Filtering)", "SI", "Per desktop", "M", "$6.00"),
            ("My Remote", "MR", "Per desktop", "M", "$2.00"),
            ("My Ops \u2014 Net", "OPS-NET", "Per desktop", "M", "$3.25"),
            ("Audit Monitoring (Base)", "AM", "Per desktop", "M", "$9.50"),
            ("Audit Monitoring (30 Day)", "AM30", "Per desktop", "M", "$12.50"),
            ("Audit Monitoring (90 Day)", "AM90", "Per desktop", "M", "$15.00"),
            ("Audit Monitoring (6 Months)", "AM6M", "Per desktop", "M", "$25.00"),
            ("Audit Monitoring (1 Year)", "AM1Y", "Per desktop", "M", "$45.00"),
        ],
    )

    d.spacer()
    d.section_header("3.4 Network and Security")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("Sophos 1C-4G (Up to 100 Mbps)", "SO-1C4G", "Per appliance", "Y", "$145.00"),
            ("Sophos 2C-4G (Up to 300 Mbps)", "SO-2C4G", "Per appliance", "Y", "$350.00"),
            ("Sophos 4C-6G (Up to 500 Mbps)", "SO-4C6G", "Per appliance", "Y", "$510.00"),
            ("Sophos 6C-8G (Up to 750 Mbps)", "SO-6C8G", "Per appliance", "Y", "$855.00"),
            ("Sophos 8C-16G (Up to 1 Gbps)", "SO-8C16G", "Per appliance", "Y", "$1,330.00"),
            ("Sophos 16C-24G (Up to 1.5 Gbps)", "SO-16C24G", "Per appliance", "Y", "$2,435.00"),
            ("Sophos UC-UG (Over 1.5 Gbps)", "SO-UCG", "Per appliance", "Y", "$3,320.00"),
            ("VeloCloud SD-WAN 30M", "VC-30M", "Per appliance", "Y", "$200.00"),
            ("VeloCloud SD-WAN 50M", "VC-50M", "Per appliance", "Y", "$250.00"),
            ("VeloCloud SD-WAN 100M", "VC-100M", "Per appliance", "Y", "$310.00"),
            ("VeloCloud SD-WAN 200M", "VC-200M", "Per appliance", "Y", "$385.00"),
            ("Edge Appliance 16GB", "Edge-16M", "Per appliance", "Y", "$100.00"),
            ("Real-Time Penetration Testing", "RTPT", "Per IP", "M", "$7.00"),
            ("Vulnerability Assessment", "PT", "Per device", "M", "$1.00"),
            ("Network Assessment", "NA", "Per device", "M", "$1.00"),
        ],
    )

    d.page_break()

    d.section_header("3.5 Email, Compliance, and Identity")
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("E-Mail Archiving", "EA", "Per user", "M", "$4.00"),
            ("DMARC/DKIM Management", "DKIM", "Per domain", "M", "$20.00"),
            ("Site Assessment", "SA", "Per domain", "M", "$50.00"),
            ("Veeam 365 Backup", "V365", "Per user", "M", "$2.50"),
            ("Anti-Spam Basic", "ASB", "Per user", "M", "$4.25"),
            ("Anti-Spam Standard (+ Encryption)", "ASA", "Per user", "M", "$6.25"),
            ("Anti-Spam Professional", "ASP", "Per user", "M", "$10.25"),
            ("Anti-Spam GSuite", "ASG", "Per user", "M", "$6.25"),
            ("Phishing Training", "PHT", "Per user", "M", "$6.00"),
            ("SSO / 2FA Login", "SSO-2FA", "Per user", "M", "$12.00"),
            ("SSO / Desktop", "SSO-DSK", "Per user", "M", "$6.00"),
            ("Credential Manager", "CRM", "Per user", "M", "$5.00"),
        ],
    )

    # ── SECTION 4: SIP TRUNK, HOSTING, LICENSING ──
    d.spacer()
    d.part_header("4. SIP TRUNK, HOSTING, AND LICENSING SERVICES")
    d.spacer()
    d.styled_table(
        ["Service", "Code", "Unit", "License", "Monthly Rate"],
        [
            ("SIP Trunk", "\u2014", "Per trunk", "M", "Per Service Order"),
            ("Voice Package", "\u2014", "Per package", "M", "Per Service Order"),
            ("DID Number", "\u2014", "Per number", "M", "Per Service Order"),
            ("Hosting \u2014 Web IIS", "HIIS", "Per domain", "M", "$10.00"),
            ("Hosting \u2014 Web WordPress", "HWP", "Per domain", "M", "$10.00"),
            ("Hosting \u2014 Web Magento", "HMG", "Per domain", "M", "$10.00"),
            ("My VDI", "MYVDI", "Per user", "M", "$10.00"),
            ("MDM (iOS)", "MDMIOS", "Per device", "M", "$4.00"),
            ("My Disk (File Sharing & Backup)", "MDU", "Per user", "M", "$16.00"),
            ("Veeam VBR (Backup & Replication)", "VBR", "Per VM", "M", "$13.00"),
            ("Veeam One (Backup Management)", "VONE", "Per VM", "M", "$3.00"),
            ("365 Azure Directory P2", "MS-ADP2", "Per user", "M", "$9.00"),
            ("MSFT \u2014 RDP SAL", "MS-WRDS", "Per license", "M", "$9.20"),
            ("MSFT \u2014 SQL Std SAL", "MS-SSSE", "Per license", "M", "$18.50"),
            ("MSFT \u2014 SQL Std Core", "MS-SSSC", "Per license", "M", "$175.00"),
            ("MSFT \u2014 Sharepoint SAL", "MS-SPSE", "Per license", "M", "$4.00"),
            ("MSFT \u2014 Exchange Std SAL", "MS-ESS", "Per license", "M", "$3.50"),
            ("MSFT \u2014 Exchange Ent SAL", "MS-EESS", "Per license", "M", "$4.00"),
            ("MSFT \u2014 Std Core 2 LIC", "MS-STD", "Per license", "M", "$5.25"),
            ("Foxit PDF Business", "FX-BSPDF", "Per license", "Y", "$14.00"),
        ],
    )

    d.page_break()

    # ── SECTION 5: TERMS ──
    d.part_header("5. TERMS")
    d.spacer()

    d.section_header("5.01 Rate Adjustments")
    d.body(
        "Technijian may adjust the rates in this Schedule upon sixty (60) days written notice to Client. "
        "Adjusted rates shall take effect at the start of the next Renewal Term of the Agreement."
    )

    d.section_header("5.02 Volume Discounts")
    d.body(
        "Volume-based pricing may be negotiated and documented in the applicable Service Order or Subscription Order."
    )

    d.section_header("5.03 Minimum Billing")
    d.body(
        "Contracted Virtual Staff hours are billed per the Cycle-Based Billing Model described in Schedule A, "
        "Part 3, at the Contracted Rate. Ad hoc (non-contracted) support is billed at the Hourly Rate in "
        "15-minute increments. Upon cancellation or termination, any unpaid hour balance shall be invoiced at "
        "the applicable Hourly Rate as set forth in Schedule A, Section 3.3(e)."
    )

    d.section_header("5.04 Relationship to Price List")
    d.body(
        "Technijian maintains a separate Services Price List that provides the current catalog of available "
        "services and list prices. The Price List may be updated by Technijian from time to time and shall be "
        "made available to Client upon request. In the event of a conflict between the Price List and this Rate "
        "Card, the rates documented in this Rate Card (Schedule C) shall govern for the duration of the "
        "then-current term. Any rate change shall be subject to the sixty (60) day written notice and Renewal "
        "Term effective date requirements set forth in Section 5.01 above and MSA Section 9.02."
    )

    # Rate Card Signatures
    d.signatures()
    d.footer_bar()

    d.save("MSA-CPM.docx")


if __name__ == "__main__":
    build_msa()
