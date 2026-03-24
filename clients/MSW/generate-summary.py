"""
MasonWest Executive Summary Generator
MSA Managed Services + Cybersecurity Proposal
Uses Technijian Brand Guide 2026 formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors (from Technijian_Brand_Guide_2026) ──
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex versions for shading
HEX_BLUE = "006DB6"
HEX_ORANGE = "F67D4B"
HEX_DARK = "1A1A2E"
HEX_OFF_WHITE = "F8F9FA"

FONT = "Open Sans"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "technijian-logo.png")

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = BRAND_GREY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# ── Helper Functions ──

def set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def accent_bar(color_hex, height_pt=4):
    """Full-width colored accent bar using a 1-row table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pf = p._element.get_or_add_pPr()
    spacing = pf.makeelement(qn("w:spacing"), {
        qn("w:before"): "0", qn("w:after"): "0", qn("w:line"): str(int(height_pt * 20))
    })
    pf.append(spacing)
    # Remove table borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def run_styled(paragraph, text, size=11, bold=False, color=None, italic=False, font=FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def section_header(title):
    """Section header with blue left-bar accent."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Blue accent bar cell (narrow)
    bar_cell = tbl.rows[0].cells[0]
    set_cell_shading(bar_cell, HEX_BLUE)
    bar_cell.width = Inches(0.08)
    bar_cell.text = ""
    bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Title cell
    title_cell = tbl.rows[0].cells[1]
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run_styled(p, title, size=14, bold=True, color=CORE_BLUE)

    # Remove table borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)
    return tbl


def styled_table(headers, rows, col_widths=None, total_row_indices=None):
    """Create a branded table with light grey borders and blue header row."""
    total_row_indices = total_row_indices or []
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEX_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        run_styled(p, h, size=9, bold=True, color=WHITE)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = tbl.add_row()
        is_total = row_idx in total_row_indices
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if is_total:
                set_cell_shading(cell, HEX_OFF_WHITE)
                run_styled(p, str(val), size=9, bold=True, color=DARK_CHARCOAL)
            else:
                run_styled(p, str(val), size=9, color=BRAND_GREY)

    # Apply light grey borders
    tbl_pr = tbl._element.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{LIGHT_GREY}"/>'
        "</w:tblBorders>"
    )
    tbl_pr.append(borders_xml)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return tbl


def bullet_item(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_styled(p, bold_prefix, size=10, bold=True, color=DARK_CHARCOAL)
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)
    return p


def body_text(text, space_after=8):
    p = doc.add_paragraph()
    run_styled(p, text, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ════════════════════════════════════════════════════════════════
# CLIENT DATA
# ════════════════════════════════════════════════════════════════

CLIENT_NAME = "MasonWest"
CLIENT_CONTACT = "Dustin Patel"
CLIENT_EMAIL = "DustinP@masonwest.com"

USERS = 74
ENDPOINTS = 80
LOCATIONS = 4

# ── Per-Unit Monthly Pricing (from PriceList_20260323) ──
PRICE_CROWDSTRIKE = 8.50       # AVD - CrowdStrike per desktop
PRICE_HUNTRESS = 6.00          # AVMH - Huntress per desktop
PRICE_UMBRELLA = 6.00          # SI - My Secure Internet (Cisco Umbrella DNS) per endpoint
PRICE_PATCH_MGMT = 4.00        # PMW - Patch Management per endpoint
PRICE_MY_REMOTE = 2.00         # MR - My Remote per endpoint
PRICE_INKY = 10.25             # ASP - Anti-Spam Professional (Inky) per user
PRICE_DMARC_DKIM = 20.00       # DKIM - DMARC/DKIM Monitoring per domain
PRICE_SITE_ASSESSMENT = 50.00  # SA - Site/365 Assessment per domain
PRICE_VEEAM_365 = 2.50         # V365 - Veeam 365 Backup per user

# Support hourly rates
RATE_US_TECH = 150.00
RATE_INDIA_NIGHT = 45.00    # India night = US business hours
RATE_INDIA_DAY = 35.00      # India day = US after-hours

# Support hours breakdown (74 total)
TOTAL_SUPPORT_HOURS = 74
DAY_HOURS = 37       # 50% day
NIGHT_HOURS = 37     # 50% night
US_DAY_HOURS = 18.5  # 50% of daytime = US
INDIA_NIGHT_HOURS = 18.5  # 50% of daytime = India (night shift, US biz hours)
INDIA_DAY_HOURS = 37      # all nighttime = India (day shift, US after-hours)

# ── Calculated Costs ──
# Endpoint Security
cost_crowdstrike = ENDPOINTS * PRICE_CROWDSTRIKE
cost_huntress = ENDPOINTS * PRICE_HUNTRESS
cost_umbrella = ENDPOINTS * PRICE_UMBRELLA  # DNS filtering is per endpoint
cost_patch = ENDPOINTS * PRICE_PATCH_MGMT
cost_remote = ENDPOINTS * PRICE_MY_REMOTE
subtotal_endpoint = cost_crowdstrike + cost_huntress + cost_umbrella + cost_patch + cost_remote

# Email Security
cost_inky = USERS * PRICE_INKY
cost_dmarc = 1 * PRICE_DMARC_DKIM  # 1 domain
cost_assessment = LOCATIONS * PRICE_SITE_ASSESSMENT
subtotal_email = cost_inky + cost_dmarc + cost_assessment

# Optional Add-On
cost_veeam_365 = USERS * PRICE_VEEAM_365

# Virtual Staff
cost_us_support = US_DAY_HOURS * RATE_US_TECH
cost_india_night = INDIA_NIGHT_HOURS * RATE_INDIA_NIGHT
cost_india_day = INDIA_DAY_HOURS * RATE_INDIA_DAY
subtotal_support = cost_us_support + cost_india_night + cost_india_day

# Totals
total_monthly = subtotal_endpoint + subtotal_email + subtotal_support
total_annual = total_monthly * 12


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

# Top blue accent bar
accent_bar(HEX_BLUE, height_pt=6)

# Spacer
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Centered logo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO_PATH):
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(3.0))

# Orange divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
div_tbl = doc.add_table(rows=1, cols=3)
div_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in div_tbl.rows[0].cells:
    cell.text = ""
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
mid_cell = div_tbl.rows[0].cells[1]
mid_cell.width = Inches(2.5)
set_cell_shading(mid_cell, HEX_ORANGE)
div_tbl.rows[0].cells[0].width = Inches(2.25)
div_tbl.rows[0].cells[2].width = Inches(2.25)
div_pr = div_tbl._element.tblPr
div_borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    "</w:tblBorders>"
)
div_pr.append(div_borders)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(8)
run_styled(p, "Executive Summary", size=26, bold=True, color=DARK_CHARCOAL)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run_styled(p, "Managed Services & Cybersecurity Solutions", size=14, color=BRAND_GREY)

# Prepared for
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run_styled(p, "Prepared for ", size=12, color=BRAND_GREY)
run_styled(p, CLIENT_NAME, size=12, bold=True, color=CORE_BLUE)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run_styled(p, "March 2026", size=11, color=BRAND_GREY)

# Spacer
for _ in range(6):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Bottom orange accent bar
accent_bar(HEX_ORANGE, height_pt=6)

# Confidential
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(p, "CONFIDENTIAL \u2014 For authorized use only", size=8, italic=True, color=BRAND_GREY)

# Contact
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run_styled(
    p,
    "Technijian  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)


# ════════════════════════════════════════════════════════════════
# PAGE 2 - ENVIRONMENT OVERVIEW + WHAT'S INCLUDED
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Overview")
body_text(
    f"Technijian proposes a comprehensive Managed Services Agreement for {CLIENT_NAME}, "
    "delivering enterprise-grade cybersecurity, endpoint protection, email security, and "
    "dedicated virtual staff support. This solution is designed to secure and manage "
    f"{CLIENT_NAME}\u2019s {USERS} users across {ENDPOINTS} endpoints and {LOCATIONS} office locations, "
    "providing 24/7 protection with predictable monthly costs."
)
body_text(
    f"Upon engagement, {CLIENT_NAME} will have a fully managed security and support infrastructure "
    "including next-generation endpoint detection and response, DNS-layer protection, advanced "
    "email threat prevention, and a blended US/offshore support team providing coverage across "
    "business hours and after-hours."
)

# Current Environment
section_header("Current Environment")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Details"],
    [
        ("Active Users", f"{USERS}"),
        ("Managed Endpoints", f"~{ENDPOINTS} (desktops/laptops)"),
        ("Office Locations", f"{LOCATIONS}"),
        ("Cloud Platform", "Microsoft 365 (Exchange Online, SharePoint, OneDrive, Teams)"),
        ("ERP / Finance", "Sage 100"),
        ("CRM", "Salesforce (integrated with Sage 100)"),
        ("Engineering Apps", "Autodesk AutoCAD, Revit, Navisworks; Bluebeam Revu"),
        ("Workload Profile", "Engineering and estimating-focused"),
    ],
    col_widths=[1.8, 4.2],
)

# Why Each Service
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Why Each Service")
body_text(
    f"Every component in this proposal addresses a specific risk or operational need for "
    f"{CLIENT_NAME}\u2019s environment. Below is the justification for each service."
)

services_justification = [
    ("CrowdStrike Falcon EDR",
     "Next-generation endpoint detection and response deployed on all managed desktops and laptops.",
     f"With {ENDPOINTS} endpoints spread across {LOCATIONS} locations, {CLIENT_NAME} needs real-time "
     "threat detection that goes beyond traditional antivirus. CrowdStrike uses AI-driven behavioral "
     "analysis to detect ransomware, fileless attacks, and zero-day exploits before they execute. "
     "Engineering workstations running AutoCAD, Revit, and Navisworks are high-value targets \u2014 "
     "a single ransomware incident could halt project delivery and expose client data."),

    ("Huntress Managed Detection",
     "24/7 SOC-backed threat hunting and persistent foothold detection with human-led analysis.",
     "CrowdStrike stops known and behavioral threats in real time, but sophisticated attackers "
     "establish persistence mechanisms that evade automated tools. Huntress provides a dedicated "
     "Security Operations Center that actively hunts for footholds, backdoors, and living-off-the-land "
     "techniques. This layered approach closes the gap between automated detection and human-led "
     "threat intelligence \u2014 critical for an environment with Sage 100 financial data and "
     "Salesforce integrations."),

    ("Cisco Umbrella (My Secure Internet)",
     "DNS-layer security that blocks malicious domains, phishing sites, and command-and-control "
     "callbacks before a connection is ever established.",
     "Most cyberattacks begin with a DNS query \u2014 whether a phishing link in an email, a "
     "malicious ad, or a compromised website. Cisco Umbrella intercepts these requests at the "
     "DNS layer, blocking threats before they reach the endpoint. This is especially important "
     f"for {CLIENT_NAME}\u2019s {LOCATIONS} office locations, where users access cloud applications "
     "(M365, Salesforce, Sage 100) from multiple networks. Umbrella provides consistent protection "
     "regardless of location."),

    ("Patch Management (ManageEngine Endpoint Central Plus)",
     "Automated OS and third-party application patching powered by ManageEngine Endpoint Central Plus, "
     "a unified endpoint management platform with software deployment, vulnerability scanning, and "
     "configuration management built in.",
     "Unpatched vulnerabilities are the #1 attack vector for ransomware and malware. With "
     f"{ENDPOINTS} endpoints running Windows plus engineering applications (AutoCAD, Revit, "
     "Navisworks, Bluebeam), keeping software current is both critical and complex. Endpoint "
     "Central Plus goes beyond basic patching \u2014 it provides automated vulnerability scanning "
     "to identify missing patches by severity, supports 850+ third-party application updates, "
     "and includes software deployment and configuration management from a single console. "
     "Scheduled deployment windows ensure security updates are applied consistently without "
     "disrupting engineering workflows."),

    ("My Remote (ConnectWise ScreenConnect)",
     "Secure remote access and support tool powered by ConnectWise ScreenConnect, providing "
     "on-demand and unattended access to all managed endpoints.",
     f"Technijian\u2019s support team needs secure, on-demand access to {CLIENT_NAME}\u2019s "
     "endpoints to deliver the contracted support hours, troubleshoot issues, deploy software, "
     "and perform maintenance. ScreenConnect provides lightning-fast remote sessions with role-based "
     "access controls, full session audit logging, and unattended access for after-hours maintenance "
     "\u2014 eliminating the need for users to install ad hoc remote tools that introduce security "
     "risk. Its lightweight agent has minimal impact on engineering workstation performance."),

    ("Inky Anti-Phishing Protection",
     "AI-powered email security that detects and blocks phishing, business email compromise (BEC), "
     "and social engineering attacks targeting Microsoft 365 mailboxes.",
     f"Email is the primary attack vector for business email compromise, and with {USERS} users "
     "on Microsoft 365, the attack surface is significant. Inky uses computer vision and machine "
     "learning to analyze every inbound email, detecting impersonation attempts, brand forgery, "
     "and zero-day phishing that bypasses Microsoft\u2019s built-in filters. For a firm handling "
     "engineering proposals, client contracts, and financial data through Sage 100, a single "
     "successful BEC attack could result in wire fraud or data exfiltration."),

    ("DMARC/DKIM Monitoring",
     "Email authentication monitoring and reporting to prevent unauthorized parties from spoofing "
     f"{CLIENT_NAME}\u2019s domain.",
     "Without DMARC and DKIM enforcement, attackers can send emails that appear to come from "
     f"{CLIENT_NAME}\u2019s domain \u2014 targeting your clients, vendors, and partners with "
     "fraudulent invoices or requests. Ongoing monitoring ensures your email authentication "
     "policies are correctly configured, alerts on spoofing attempts, and provides visibility "
     "into who is sending email on your behalf."),

    ("Monthly Site Assessments (Network Detective by RapidFire Tools)",
     f"Comprehensive security and compliance assessments for each of {CLIENT_NAME}\u2019s "
     f"{LOCATIONS} office locations and Microsoft 365 tenant, powered by Network Detective.",
     "Network Detective by RapidFire Tools performs automated, deep-scan assessments that go far "
     "beyond surface-level checks. Each monthly scan produces detailed reports covering network "
     "vulnerabilities, Active Directory health, M365 Secure Score analysis, external attack "
     "surface exposure, and compliance readiness (HIPAA, PCI, NIST, CMMC). The tool also "
     "generates executive-ready risk scorecards that track your security posture over time "
     f"\u2014 giving {CLIENT_NAME} leadership clear visibility into risk trends across all "
     f"{LOCATIONS} locations without requiring manual audit effort."),

    ("Virtual Staff Support",
     f"{TOTAL_SUPPORT_HOURS} hours/month of blended US and offshore technical support with "
     "50/50 day/night coverage.",
     f"With {USERS} users, engineering-heavy applications, and integrations between Sage 100, "
     "Salesforce, and Microsoft 365, the support demand is continuous. The blended model "
     "provides US-based technicians for complex escalations during business hours, India-based "
     "night-shift staff for cost-effective daytime coverage, and India-based day-shift staff "
     "for after-hours support \u2014 ensuring issues don\u2019t wait until the next business day."),
]

for svc_name, svc_desc, svc_why in services_justification:
    p = doc.add_paragraph()
    run_styled(p, svc_name, size=11, bold=True, color=CORE_BLUE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run_styled(p, svc_desc, size=10, italic=True, color=DARK_CHARCOAL)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run_styled(p, "Why: ", size=10, bold=True, color=CORE_ORANGE)
    run_styled(p, svc_why, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(6)


# ════════════════════════════════════════════════════════════════
# PAGE 3 - CYBERSECURITY SERVICES PRICING
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Endpoint & Device Security")
body_text(
    f"Enterprise cybersecurity stack deployed across all {ENDPOINTS} managed endpoints, "
    "providing layered protection from endpoint to network edge."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Unit", "Qty", "Unit Price", "Monthly"],
    [
        ("CrowdStrike Falcon EDR", "Per endpoint", str(ENDPOINTS), f"${PRICE_CROWDSTRIKE:.2f}", f"${cost_crowdstrike:,.2f}"),
        ("Huntress Managed Detection", "Per endpoint", str(ENDPOINTS), f"${PRICE_HUNTRESS:.2f}", f"${cost_huntress:,.2f}"),
        ("My Secure Internet (Cisco Umbrella)", "Per endpoint", str(ENDPOINTS), f"${PRICE_UMBRELLA:.2f}", f"${cost_umbrella:,.2f}"),
        ("Patch Mgmt (Endpoint Central Plus)", "Per endpoint", str(ENDPOINTS), f"${PRICE_PATCH_MGMT:.2f}", f"${cost_patch:,.2f}"),
        ("My Remote (ScreenConnect)", "Per endpoint", str(ENDPOINTS), f"${PRICE_MY_REMOTE:.2f}", f"${cost_remote:,.2f}"),
        ("SUBTOTAL \u2014 ENDPOINT SECURITY", "", "", "", f"${subtotal_endpoint:,.2f}"),
    ],
    col_widths=[2.2, 1.0, 0.5, 0.9, 1.0],
    total_row_indices=[5],
)

# Email Security
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Email Security & Compliance")
body_text(
    f"Advanced email protection for all {USERS} Microsoft 365 users with domain authentication "
    f"monitoring and monthly security assessments across all {LOCATIONS} sites."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Unit", "Qty", "Unit Price", "Monthly"],
    [
        ("Inky Anti-Phishing (Anti-Spam Pro)", "Per user", str(USERS), f"${PRICE_INKY:.2f}", f"${cost_inky:,.2f}"),
        ("DMARC/DKIM Monitoring", "Per domain", "1", f"${PRICE_DMARC_DKIM:.2f}", f"${cost_dmarc:,.2f}"),
        ("Site Assessment (Network Detective)", "Per site", str(LOCATIONS), f"${PRICE_SITE_ASSESSMENT:.2f}", f"${cost_assessment:,.2f}"),
        ("SUBTOTAL \u2014 EMAIL & COMPLIANCE", "", "", "", f"${subtotal_email:,.2f}"),
    ],
    col_widths=[2.2, 1.0, 0.5, 0.9, 1.0],
    total_row_indices=[3],
)

# Optional Add-On: Veeam 365 Backup
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Optional Add-On \u2014 Veeam 365 Backup")
body_text(
    "Microsoft\u2019s native retention policies provide limited protection against accidental deletion, "
    "ransomware, and compliance-driven data recovery needs. Veeam Backup for Microsoft 365 provides "
    "independent, automated backup of Exchange Online, SharePoint, OneDrive, and Teams data with "
    "granular point-in-time recovery."
)
body_text(
    f"Technijian recommends assessing {CLIENT_NAME}\u2019s current data volume, retention requirements, "
    "and compliance obligations before finalizing backup scope and storage tiers. The pricing below "
    "reflects the per-user backup license \u2014 storage costs will be scoped during the assessment."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service", "Unit", "Qty", "Unit Price", "Monthly"],
    [
        ("Veeam 365 Backup", "Per user", str(USERS), f"${PRICE_VEEAM_365:.2f}", f"${cost_veeam_365:,.2f}"),
    ],
    col_widths=[2.2, 1.0, 0.5, 0.9, 1.0],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run_styled(p, "Note: ", size=10, bold=True, color=CORE_ORANGE)
run_styled(
    p,
    "Final Veeam 365 pricing will be confirmed after a storage and retention assessment. "
    "The assessment will evaluate current mailbox sizes, SharePoint/OneDrive data volume, "
    "required retention periods, and any regulatory or compliance requirements. Additional "
    "backup storage costs (if applicable) will be quoted separately.",
    size=10, color=BRAND_GREY,
)


# ════════════════════════════════════════════════════════════════
# PAGE 4 - VIRTUAL STAFF + TOTAL INVESTMENT
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Virtual Staff \u2014 Contracted Support")
body_text(
    f"{TOTAL_SUPPORT_HOURS} hours per month of dedicated remote technical support, split 50/50 between "
    "daytime and nighttime coverage. Daytime hours are further split 50/50 between US-based "
    "and India-based (night shift) technicians for cost-optimized coverage during business hours."
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Role", "Location", "Shift", "Hours/Mo", "Rate", "Monthly"],
    [
        ("Tech Support", "United States", "Day (US Biz Hours)", f"{US_DAY_HOURS:.1f}", f"${RATE_US_TECH:.2f}/hr", f"${cost_us_support:,.2f}"),
        ("Tech Support", "India \u2014 Night", "Day (US Biz Hours)", f"{INDIA_NIGHT_HOURS:.1f}", f"${RATE_INDIA_NIGHT:.2f}/hr", f"${cost_india_night:,.2f}"),
        ("Tech Support", "India \u2014 Day", "Night (US After-Hours)", f"{INDIA_DAY_HOURS:.1f}", f"${RATE_INDIA_DAY:.2f}/hr", f"${cost_india_day:,.2f}"),
        ("SUBTOTAL \u2014 VIRTUAL STAFF", "", "", f"{TOTAL_SUPPORT_HOURS}", "", f"${subtotal_support:,.2f}"),
    ],
    col_widths=[1.2, 1.1, 1.3, 0.7, 0.9, 0.9],
    total_row_indices=[3],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run_styled(p, "Naming Convention:", size=10, bold=True, color=DARK_CHARCOAL)
naming_items = [
    ("\u201cIndia \u2014 Night\u201d ", "refers to India-based staff working during US business hours (nighttime in India)"),
    ("\u201cIndia \u2014 Day\u201d ", "refers to India-based staff working during US after-hours (daytime in India)"),
]
for bold_text, desc in naming_items:
    bullet_item(desc, bold_prefix=bold_text)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run_styled(p, "Billing Model: ", size=10, bold=True, color=DARK_CHARCOAL)
run_styled(
    p,
    "Cycle-based billing with reconciliation. Initial cycle based on estimated "
    f"{TOTAL_SUPPORT_HOURS} hours/month per the Service Order. Actual usage tracked from day one "
    "with reconciliation at end of each cycle period.",
    size=10, color=BRAND_GREY,
)

# Total Investment Summary
doc.add_paragraph().paragraph_format.space_after = Pt(8)
section_header("Total Monthly Investment Summary")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Category", "Monthly", "Annual"],
    [
        ("Endpoint & Device Security", f"${subtotal_endpoint:,.2f}", f"${subtotal_endpoint * 12:,.2f}"),
        ("Email Security & Compliance", f"${subtotal_email:,.2f}", f"${subtotal_email * 12:,.2f}"),
        ("Virtual Staff Support", f"${subtotal_support:,.2f}", f"${subtotal_support * 12:,.2f}"),
        ("TOTAL MONTHLY", f"${total_monthly:,.2f}", f"${total_annual:,.2f}"),
    ],
    col_widths=[2.8, 1.6, 1.6],
    total_row_indices=[3],
)


# ════════════════════════════════════════════════════════════════
# PAGE 5 - AGREEMENT STRUCTURE + NEXT STEPS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

section_header("Service Level Commitments")
doc.add_paragraph().paragraph_format.space_after = Pt(2)
styled_table(
    ["Service Level", "Target"],
    [
        ("Critical Incident Response", "Within 1 hour of notification"),
        ("Standard Support Response", "Within 4 business hours"),
        ("Scheduled Maintenance", "Tuesdays evenings and Saturdays (with advance notice)"),
        ("24/7 Monitoring", "All endpoints and email security continuously monitored"),
        ("Monthly Reporting", "Service reports summarizing incidents, uptime, and support activity"),
        ("Quarterly Reviews", f"Scheduled service reviews with {CLIENT_NAME}\u2019s designated representative"),
    ],
    col_widths=[2.2, 3.8],
)

# Agreement Structure
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Agreement Structure")
agreements = [
    ("Master Service Agreement (MSA): ",
     "Governs the overall relationship, payment terms, confidentiality, liability, and dispute "
     "resolution for a 12-month initial term with automatic annual renewal."),
    ("Schedule A \u2014 Monthly Services: ",
     "Details the managed security, endpoint protection, email security, and virtual staff "
     "services with SLA commitments."),
    ("Schedule B \u2014 Subscription Services: ",
     "Covers any third-party software subscriptions and licenses managed by Technijian on "
     f"{CLIENT_NAME}\u2019s behalf."),
    ("Schedule C \u2014 Rate Card: ",
     "Establishes hourly rates for any additional work, change orders, or ad hoc support "
     "beyond the contracted hours."),
]
for bold_text, desc in agreements:
    bullet_item(desc, bold_prefix=bold_text)

# Next Steps
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_header("Next Steps")
steps = [
    f"Review this summary and the attached agreements",
    f"Confirm any questions with {CLIENT_CONTACT}",
    "Sign the MSA, Schedule A, Schedule B, Schedule C",
    "Technijian begins onboarding and agent deployment within 5 business days of execution",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run_styled(p, f"{i}. ", size=10, bold=True, color=CORE_BLUE)
    run_styled(p, step, size=10, color=BRAND_GREY)
    p.paragraph_format.space_after = Pt(2)

# Footer divider
doc.add_paragraph().paragraph_format.space_after = Pt(8)
accent_bar(HEX_BLUE, height_pt=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run_styled(
    p,
    "Technijian, Inc.  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8499  |  technijian.com",
    size=8, color=BRAND_GREY,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(
    p,
    f"CONFIDENTIAL \u2014 Prepared exclusively for {CLIENT_NAME}",
    size=8, italic=True, color=BRAND_GREY,
)

OUTPUT = os.path.join(os.path.dirname(__file__), "MSW-Executive-Summary.docx")
doc.save(OUTPUT)
print(f"Created {OUTPUT}")
